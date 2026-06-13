from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_PATH = ROOT / "exports" / "ml" / "FestFlow_current_ai_implementation_report.docx"

FONT = "Malgun Gothic"
BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "ECFDF5"
YELLOW = "FFFBEB"
RED = "FEF2F2"


def set_korean_font(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    set_korean_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, header_fill=LIGHT_GRAY):
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.05)
                for run in paragraph.runs:
                    set_korean_font(run)
                    run.font.size = Pt(9)
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(BLUE)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    style_table(table, widths, header_fill)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.15)
    p.clear()
    add_run(p, text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.15)
    p.clear()
    add_run(p, text)


def add_callout(doc, title, body, fill=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    add_run(p, title + "\n", bold=True, color=BLUE)
    add_run(p, body)
    set_paragraph_spacing(p, after=0, line=1.15)
    style_table(table, [9360], fill)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=4)
    add_run(p, "FestFlow", bold=True, color=BLUE, size=14)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=20, after=8)
    add_run(title, "현재 실제 구현된 AI 기능 전체 설명서", bold=True, color="0B2545", size=24)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=18, line=1.2)
    add_run(
        subtitle,
        "데이터 속성, 모델 학습, 서버 실시간 추론, API 응답, 프론트 표시, 한계와 확장 방향 정리",
        color="555555",
        size=11,
    )

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 목적", "현재 프로젝트에 실제로 들어간 AI 구현 범위를 발표/질의응답용으로 설명"],
            ["작성 기준", "현재 저장소의 scripts/ml, exports/ml, backend, frontend 구현 기준"],
            ["핵심 AI", "RandomForest 기반 혼잡도 분류 모델"],
            ["실제 노출 화면", "프론트엔드 /analytics 페이지의 30분 혼잡도 예측 카드"],
            ["주의", "XGBoost는 비교 실험용이며 현재 서버 실시간 추론에는 RandomForest 모델만 연결"],
        ],
        [2200, 7160],
        LIGHT_BLUE,
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_current_scope(doc):
    doc.add_heading("1. 현재 실제로 들어간 AI 범위", level=1)
    add_callout(
        doc,
        "핵심 결론",
        "현재 FestFlow에는 단순 규칙 기반 표시만 있는 것이 아니라, 학습된 RandomForest 모델 파일을 Spring Boot 서버가 Python 추론 스크립트를 통해 직접 불러와 API 응답에 포함하는 구조가 들어가 있습니다.",
        GREEN,
    )
    add_table(
        doc,
        ["구분", "현재 상태", "설명"],
        [
            ["RandomForest 혼잡도 예측", "실제 적용", "저장된 .pkl 모델을 서버가 호출하여 부스별 혼잡도를 예측"],
            ["XGBoost", "비교 실험", "학습/평가 결과는 있지만 현재 운영 API의 실시간 추론 모델은 아님"],
            ["규칙 기반 혼잡도", "비교 기준", "기존 휴리스틱 방식과 AI 모델 성능을 비교하기 위한 baseline"],
            ["시간 변화 feature", "실제 적용", "최근 GPS 변화량, 예약 변화량, 체크인 변화량, 대기시간 변화량을 feature로 사용"],
            ["Drift 감지", "실제 적용", "현재 입력값이 학습 데이터 분포에서 얼마나 벗어났는지 점수와 경고로 표시"],
            ["LSTM/GNN/온라인 러닝", "미구현", "문서화된 향후 확장 방향이며 현재 기능에는 포함되지 않음"],
        ],
        [2200, 1700, 5460],
    )


def add_dataset_section(doc):
    doc.add_heading("2. 학습 데이터셋 개요", level=1)
    p = doc.add_paragraph()
    add_run(
        p,
        "학습 데이터는 사용자가 설명한 축제 운영 가정을 바탕으로 만든 HYBRID_SIMULATED 데이터입니다. "
        "즉, 실제 서비스 DB에서 완전히 자동 수집한 데이터는 아니지만, 공연 시간대, 인기 가수 여부, 야간 부스 쏠림, GPS/예약/체크인/대기시간 같은 운영 신호를 반영해 AI 모델 학습용 표 형태로 만든 데이터입니다.",
    )

    add_table(
        doc,
        ["항목", "값"],
        [
            ["파일", "exports/ml/congestion_training_dataset.csv"],
            ["행 수", "2,520 rows"],
            ["열 수", "28 columns"],
            ["데이터 출처 표기", "HYBRID_SIMULATED"],
            ["예측 대상", "target_congestion: LOW, NORMAL, BUSY, VERY_BUSY"],
            ["라벨 분포", "NORMAL 1131 / BUSY 619 / VERY_BUSY 546 / LOW 224"],
        ],
        [2600, 6760],
        LIGHT_BLUE,
    )

    doc.add_heading("2.1 데이터 속성 전체 목록", level=2)
    rows = [
        ["scenario_day", "숫자", "가상의 축제 운영 날짜 번호. 학습 분포와 현재 입력의 차이를 drift 판단에 사용"],
        ["hour", "숫자", "시간대. 18-22시 공연 집중 시간과 야간 부스 혼잡 패턴을 반영"],
        ["is_peak_time", "0/1", "18-22시 또는 핵심 혼잡 시간대 여부"],
        ["zone_type", "범주", "STAGE, FOOD, BOOTH 등 혼잡이 발생하는 공간 유형"],
        ["booth_id", "숫자", "부스 식별자. 실제 API에서는 부스별 예측 결과를 연결하는 기준"],
        ["artist_popularity", "범주", "LOW, MEDIUM, HIGH 등 공연자 인기 수준"],
        ["artist_popularity_score", "숫자", "인기 수준을 모델이 처리할 수 있도록 숫자로 바꾼 값"],
        ["stage_capacity", "숫자", "노천극장 최대 수용 가능 인원. 현재 기준은 4,000명"],
        ["expected_stage_crowd", "숫자", "공연 시간과 인기 가수 여부를 반영한 예상 무대 인파"],
        ["stage_load_ratio", "숫자", "expected_stage_crowd / stage_capacity. 무대 포화 정도"],
        ["is_night_booth", "0/1", "공연 외 시간대에 야간 부스 쪽으로 사람이 몰리는 상황 여부"],
        ["event_soon", "0/1", "곧 공연이나 이벤트가 시작되는지 여부"],
        ["minutes_to_next_event", "숫자", "다음 공연/이벤트까지 남은 시간"],
        ["gps_count_nearby", "숫자", "현재 부스 또는 구역 주변 GPS 로그 수"],
        ["gps_delta_5m", "숫자", "최근 5분 기준 주변 GPS 수 변화량"],
        ["gps_delta_15m", "숫자", "최근 15분 기준 주변 GPS 수 변화량"],
        ["reservation_count", "숫자", "해당 부스 예약 수"],
        ["reservation_delta_15m", "숫자", "최근 15분 예약 증가량"],
        ["checked_in_count", "숫자", "실제 체크인한 인원 수"],
        ["checked_in_delta_15m", "숫자", "최근 15분 체크인 증가량"],
        ["available_seats", "숫자", "남은 좌석 또는 수용 가능 여유량"],
        ["wait_minutes", "숫자", "예상 대기 시간"],
        ["wait_delta_15m", "숫자", "최근 15분 대기시간 변화량"],
        ["remaining_stock", "숫자", "부스 재고. 주문/판매 흐름을 간접적으로 반영"],
        ["event_count_context", "숫자", "주변 이벤트 개수 또는 시간대 이벤트 밀도"],
        ["data_source", "범주", "데이터 생성/출처 구분. 현재는 HYBRID_SIMULATED"],
        ["rule_based_level", "범주", "기존 규칙 기반 방식으로 산출한 혼잡도"],
        ["target_congestion", "범주", "AI가 맞춰야 하는 정답 라벨"],
    ]
    add_table(doc, ["컬럼", "타입", "의미"], rows, [2300, 1300, 5760])


def add_model_section(doc):
    doc.add_heading("3. 모델 학습과 성능", level=1)
    p = doc.add_paragraph()
    add_run(
        p,
        "학습은 scripts/ml/train_congestion_models.py에서 수행됩니다. 같은 데이터셋에 대해 규칙 기반 baseline, RandomForest, XGBoost를 비교했고, 현재 서버에 실제 연결된 모델은 RandomForest입니다.",
    )

    add_table(
        doc,
        ["모델", "Accuracy", "Macro F1", "현재 역할"],
        [
            ["규칙 기반 baseline", "0.6810", "0.6436", "기존 방식 비교 기준"],
            ["RandomForest", "0.7873", "0.7708", "현재 서버 실시간 추론에 실제 적용"],
            ["XGBoost", "0.8127", "0.7869", "성능 비교용. 현재 운영 API에는 미연결"],
        ],
        [2600, 1600, 1600, 3960],
        LIGHT_BLUE,
    )

    doc.add_heading("3.1 RandomForest를 실제 적용 모델로 둔 이유", level=2)
    for text in [
        "트리 여러 개를 앙상블로 묶는 방식이라 작은 규모의 tabular 데이터에서도 안정적으로 동작합니다.",
        "feature 중요도와 예측 확률을 해석하기 쉬워 발표에서 설명하기 좋습니다.",
        "XGBoost보다 성능은 조금 낮지만, 설치/배포 의존성과 설명 난이도가 상대적으로 낮습니다.",
        "현재 프로젝트에서는 '실제 동작하는 AI 기능'을 안정적으로 보여주는 것이 우선이므로 RandomForest를 운영 추론 모델로 선택했습니다.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("3.2 XGBoost의 위치", level=2)
    add_callout(
        doc,
        "XGBoost는 버린 것이 아니라 비교 대상으로 유지",
        "XGBoost는 실험 결과 가장 높은 accuracy를 보였지만 현재 백엔드 실시간 추론 구조에는 연결하지 않았습니다. 발표에서는 '추가 비교 실험까지 수행했고, 운영 안정성을 고려해 RandomForest를 먼저 적용했다'고 설명할 수 있습니다.",
        YELLOW,
    )


def add_inference_section(doc):
    doc.add_heading("4. 서버 실시간 추론 구조", level=1)
    add_number(doc, "프론트엔드가 /api/ai/congestion/predictions API를 요청합니다.")
    add_number(doc, "Spring Boot의 AiCongestionService가 부스, GPS, 예약, 체크인, 이벤트 데이터를 모아 모델 입력 feature를 구성합니다.")
    add_number(doc, "PythonCongestionModelService가 임시 JSON 파일을 만들고 scripts/ml/predict_congestion.py를 실행합니다.")
    add_number(doc, "Python 스크립트가 exports/ml/models/random_forest_congestion_model.pkl 모델 파일을 로드합니다.")
    add_number(doc, "모델이 부스별 predictedLevel, confidence, driftStatus, driftScore, driftWarnings를 계산합니다.")
    add_number(doc, "백엔드는 이 결과를 AiModelPredictionDto와 AiBoothRecommendationDto에 담아 프론트엔드로 반환합니다.")

    add_table(
        doc,
        ["구성요소", "역할"],
        [
            ["PythonCongestionModelService.java", "Spring Boot에서 Python 추론 프로세스를 실행하고 결과 JSON을 읽음"],
            ["AiCongestionService.java", "현재 운영 데이터를 모델 feature로 변환하고 AI 예측 응답을 조립"],
            ["predict_congestion.py", "저장된 RandomForest 모델을 로드해 실제 예측 수행"],
            ["random_forest_congestion_model.pkl", "학습이 완료된 실제 모델 파일"],
            ["congestion_training_profile.json", "drift 감지를 위한 학습 데이터 분포 정보"],
        ],
        [3600, 5760],
    )

    doc.add_heading("4.1 서버 배포 시 필요한 설정", level=2)
    add_table(
        doc,
        ["항목", "현재 기본값 또는 설명"],
        [
            ["Python 실행 파일", "../.venv-ml/Scripts/python.exe"],
            ["추론 스크립트", "../scripts/ml/predict_congestion.py"],
            ["모델 파일", "../exports/ml/models/random_forest_congestion_model.pkl"],
            ["타임아웃", "20,000 ms"],
            ["주의점", "서버 배포 환경에도 Python, scikit-learn, joblib 등 requirements-ml.txt 의존성이 필요"],
        ],
        [2600, 6760],
        LIGHT_BLUE,
    )


def add_api_front_section(doc):
    doc.add_heading("5. API 응답과 프론트엔드 표시", level=1)
    add_table(
        doc,
        ["응답 필드", "의미"],
        [
            ["modelType", "현재 예측에 사용된 모델명. 실제 적용 시 RandomForest"],
            ["rawPredictedLevel", "모델이 직접 예측한 원본 혼잡도 라벨"],
            ["displayPredictedLevel", "화면 표시용으로 변환된 혼잡도 라벨"],
            ["confidence", "모델 예측 확률 기반 신뢰도"],
            ["modelBased", "실제 모델 기반 예측인지 여부. true이면 모델 파일 추론 결과"],
            ["driftStatus", "NORMAL, CAUTION, WARNING 중 하나"],
            ["driftScore", "현재 입력이 학습 분포에서 벗어난 정도"],
            ["driftWarnings", "분포 이탈 원인 설명 목록"],
            ["factors", "예측에 영향을 준 주요 요인 설명"],
            ["error", "Python 추론 실패 시 오류 메시지"],
        ],
        [2800, 6560],
    )

    p = doc.add_paragraph()
    add_run(
        p,
        "프론트엔드에서는 frontend/src/pages/AnalyticsPage.jsx의 '30분 혼잡도 예측' 카드에서 확인할 수 있습니다. 화면에는 현재 혼잡도, AI 예측 혼잡도, risk score, RandomForest 모델명, confidence, drift 배지, 주요 판단 요인이 표시됩니다.",
    )


def add_temporal_drift_section(doc):
    doc.add_heading("6. 시간 변화 feature와 drift 감지", level=1)
    doc.add_heading("6.1 시간 변화 feature", level=2)
    for text in [
        "gps_delta_5m / gps_delta_15m: 단순 현재 인원 수가 아니라 최근 몇 분 동안 인원이 늘고 있는지 줄고 있는지를 반영합니다.",
        "reservation_delta_15m: 예약이 갑자기 증가하는 부스는 곧 혼잡해질 가능성이 높다고 봅니다.",
        "checked_in_delta_15m: 실제 방문 전환이 빠르게 늘어나는지를 반영합니다.",
        "wait_delta_15m: 대기시간이 증가하는 방향인지 감소하는 방향인지 반영합니다.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("6.2 Drift 감지", level=2)
    p = doc.add_paragraph()
    add_run(
        p,
        "drift는 현재 운영 입력값이 학습 데이터에서 봤던 범위를 벗어나는지 확인하는 안전장치입니다. 예를 들어 학습 데이터가 scenario_day 1-28 범위였는데 실제 입력이 훨씬 큰 값이면, 모델이 익숙하지 않은 상황으로 판단해 drift warning을 제공합니다.",
    )
    add_table(
        doc,
        ["driftStatus", "의미", "해석"],
        [
            ["NORMAL", "학습 데이터 분포와 큰 차이 없음", "예측을 일반적으로 신뢰 가능"],
            ["CAUTION", "일부 feature가 학습 분포에서 벗어남", "예측은 참고하되 운영자가 함께 판단"],
            ["WARNING", "여러 feature가 크게 벗어남", "모델 재학습 또는 데이터 보강 필요"],
        ],
        [1800, 3300, 4260],
        LIGHT_BLUE,
    )


def add_limitations_section(doc):
    doc.add_heading("7. 현재 한계와 발표에서의 정확한 표현", level=1)
    add_callout(
        doc,
        "정확한 표현",
        "현재 구현은 '학습된 AI 모델을 서버가 실제로 호출하여 혼잡도 예측 API와 화면에 연결한 프로토타입'입니다. 다만 학습 데이터가 실제 장기간 운영 로그가 아니라 가정 기반 synthetic/hybrid 데이터이므로, 상용 수준의 정확도 검증이 끝난 AI라고 말하면 과장입니다.",
        RED,
    )
    for text in [
        "현재 데이터는 실제 축제 로그 전체가 아니라 사용자가 제공한 운영 가정과 데모 데이터를 바탕으로 구성했습니다.",
        "RandomForest는 현재 시점의 feature를 보고 분류하는 모델이라 LSTM이나 TFT처럼 긴 시간 흐름 자체를 직접 학습하는 구조는 아닙니다.",
        "GPS 로그, 예약 로그, 체크인 로그가 실제로 충분히 쌓일수록 모델은 더 현실적인 데이터로 재학습할 수 있습니다.",
        "온라인 러닝, GNN, 강화학습, 수요 예측 기반 운영 최적화는 현재 미구현이며 향후 확장 주제로 분리하는 것이 정확합니다.",
    ]:
        add_bullet(doc, text)


def add_runbook_section(doc):
    doc.add_heading("8. 실행 및 재학습 방법", level=1)
    add_table(
        doc,
        ["작업", "명령 또는 파일"],
        [
            ["데이터셋 생성", ".venv-ml\\Scripts\\python.exe scripts\\ml\\build_congestion_dataset.py"],
            ["모델 학습/비교", ".venv-ml\\Scripts\\python.exe scripts\\ml\\train_congestion_models.py"],
            ["단일/배치 예측", ".venv-ml\\Scripts\\python.exe scripts\\ml\\predict_congestion.py"],
            ["백엔드 컴파일 확인", "backend\\gradlew.bat compileJava"],
            ["프론트엔드 빌드 확인", "cd frontend && npm run build"],
        ],
        [2600, 6760],
    )

    doc.add_heading("9. 주요 파일 위치", level=1)
    add_table(
        doc,
        ["파일", "역할"],
        [
            ["scripts/ml/build_congestion_dataset.py", "AI 학습용 혼잡도 데이터셋 생성"],
            ["scripts/ml/train_congestion_models.py", "RandomForest/XGBoost 학습과 비교 평가"],
            ["scripts/ml/predict_congestion.py", "저장된 RandomForest 모델을 로드해 예측 실행"],
            ["exports/ml/congestion_training_dataset.csv", "학습 데이터셋"],
            ["exports/ml/model_comparison.csv", "모델별 성능 비교 결과"],
            ["exports/ml/models/random_forest_congestion_model.pkl", "실제 서버 추론에 사용되는 모델 파일"],
            ["exports/ml/models/congestion_training_profile.json", "drift 감지용 학습 분포 프로파일"],
            ["backend/src/main/java/com/festflow/backend/service/PythonCongestionModelService.java", "Spring Boot에서 Python 모델 추론 호출"],
            ["backend/src/main/java/com/festflow/backend/service/AiCongestionService.java", "AI feature 구성과 예측 응답 생성"],
            ["frontend/src/pages/AnalyticsPage.jsx", "AI 예측 카드와 drift 정보 화면 표시"],
        ],
        [4700, 4660],
    )


def add_presentation_script(doc):
    doc.add_heading("10. 발표에서 말할 수 있는 요약 문장", level=1)
    add_callout(
        doc,
        "발표용 설명",
        "처음에는 대기 인원, 예약 수, 체크인 수, 시간대 등을 단순 규칙으로 계산했지만, AI 활용도가 부족하다고 판단해서 같은 정보를 feature로 정리하고 RandomForest 모델을 학습시켰습니다. 현재는 학습된 모델 파일을 백엔드가 직접 호출해 부스별 30분 혼잡도를 예측하고, 프론트엔드 분석 화면에 모델명, 신뢰도, drift 상태까지 함께 표시합니다. XGBoost도 비교 실험을 했지만, 운영 안정성과 설명 가능성을 고려해 현재 실시간 추론에는 RandomForest를 먼저 연결했습니다.",
        GREEN,
    )


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(footer, "FestFlow AI Implementation Report", color="777777", size=9)


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_current_scope(doc)
    add_dataset_section(doc)
    add_model_section(doc)
    add_inference_section(doc)
    add_api_front_section(doc)
    add_temporal_drift_section(doc)
    add_limitations_section(doc)
    add_runbook_section(doc)
    add_presentation_script(doc)
    add_footer(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(f"Report written to {OUTPUT_PATH}")
