from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_TITLE = "FestFlow 혼잡도 AI 예측 프로토타입 보고서"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(PROJECT_TITLE)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("규칙 기반 혼잡도 산정 방식과 RandomForest/XGBoost 모델 비교 실험")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    paragraph.add_run(text)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_code_block(doc: Document, commands: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(commands))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("111111")
    doc.add_paragraph()


def add_dataframe_table(doc: Document, dataframe: pd.DataFrame, widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    for index, column in enumerate(dataframe.columns):
        cell = table.cell(0, index)
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(str(column))
        run.bold = True

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(dataframe.columns):
            value = row[column]
            if isinstance(value, float):
                value = f"{value:.4f}"
            cells[index].paragraphs[0].add_run(str(value))
            set_cell_width(cells[index], widths[index])
    doc.add_paragraph()


def add_image_if_exists(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor.from_string("555555")


def build_report(output_dir: Path, docx_path: Path) -> None:
    figures_dir = output_dir / "figures"
    comparison = pd.read_csv(output_dir / "model_comparison.csv")
    dataset = pd.read_csv(output_dir / "congestion_training_dataset.csv")
    rf_importance = pd.read_csv(output_dir / "feature_importance_random_forest.csv").head(5)
    xgb_importance = pd.read_csv(output_dir / "feature_importance_xgboost.csv").head(5)

    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "핵심 요약",
        "현재 앱의 혼잡도 기능은 규칙 기반으로도 설명 가능하지만, AI 활용도를 보강하기 위해 운영 경험 기반 시뮬레이션 데이터로 RandomForest와 XGBoost를 비교 실험했다. 이 결과는 실제 운영 검증 모델이 아니라, 향후 실제 로그가 쌓였을 때 확장 가능한 AI 프로토타입으로 제시하는 것이 안전하다.",
    )

    doc.add_heading("1. 직접 실행 방법", level=1)
    doc.add_paragraph("아래 명령은 프로젝트 루트에서 실행한다. Windows PowerShell 기준으로 한 번에 전체 파이프라인을 다시 실행할 수 있다.")
    add_code_block(
        doc,
        [
            "cd \"c:\\Users\\jk636\\OneDrive\\문서\\GitHub\\Whims-of-Wonder\\Whims-of-Wonder\\MetaverseProgramming\\My project\\Assets\\FestFlow\"",
            "powershell -ExecutionPolicy Bypass -File scripts\\ml\\run_congestion_ml_pipeline.ps1",
        ],
    )
    doc.add_paragraph("단계별로 실행하고 싶다면 아래 순서를 사용한다.")
    add_code_block(
        doc,
        [
            "py -m venv .venv-ml",
            ".\\.venv-ml\\Scripts\\python.exe -m pip install -r requirements-ml.txt",
            ".\\.venv-ml\\Scripts\\python.exe scripts\\ml\\build_congestion_dataset.py",
            ".\\.venv-ml\\Scripts\\python.exe scripts\\ml\\train_congestion_models.py",
            ".\\.venv-ml\\Scripts\\python.exe scripts\\ml\\plot_congestion_results.py",
            ".\\.venv-ml\\Scripts\\python.exe scripts\\ml\\create_congestion_report_docx.py",
        ],
    )

    doc.add_heading("2. 데이터셋 구성 방식", level=1)
    doc.add_paragraph(
        f"이번 실험 데이터셋은 총 {len(dataset):,}행으로 구성했다. 현재 앱에서 추출한 CSV 구조를 참고하고, 사용자가 제공한 축제 운영 경험 가정을 반영해 HYBRID_SIMULATED 데이터로 생성했다."
    )
    add_bullets(
        doc,
        [
            "18시부터 22시 사이에는 무대 관람 수요가 급격히 증가한다.",
            "인기 가수가 오는 날에는 무대 구역의 혼잡도가 더 크게 증가한다.",
            "인기가 낮은 공연이거나 무대 피크 시간이 아니면 야간 부스, 음식, 주점 구역으로 수요가 이동한다.",
            "무대 수용량은 3,000명에서 4,000명 수준으로 가정한다.",
            "GPS 추정 인원, 예약 수, 체크인 수, 대기 시간, 잔여 재고, 공연 임박 여부 등을 혼잡도 feature로 사용한다.",
        ],
    )

    label_counts = dataset["target_congestion"].value_counts().reindex(["LOW", "NORMAL", "BUSY", "VERY_BUSY"], fill_value=0)
    add_dataframe_table(
        doc,
        label_counts.rename_axis("혼잡도 라벨").reset_index(name="행 수"),
        [3500, 2500],
    )
    add_image_if_exists(doc, figures_dir / "label_distribution.png", "그림 1. 학습 데이터의 혼잡도 라벨 분포")

    doc.add_heading("3. 모델 비교 구조", level=1)
    doc.add_paragraph("비교 대상은 기존 규칙 기반 방식, RandomForest, XGBoost 세 가지다. 같은 test split을 사용해 정확도와 macro F1을 비교했다.")
    add_dataframe_table(doc, comparison[["model", "accuracy", "macro_f1", "notes"]], [2200, 1500, 1500, 4160])
    add_image_if_exists(doc, figures_dir / "model_performance_comparison.png", "그림 2. 규칙 기반, RandomForest, XGBoost 성능 비교")

    doc.add_heading("4. RandomForest와 XGBoost 해석", level=1)
    doc.add_paragraph(
        "RandomForest는 여러 개의 결정트리를 만들고 다수결로 분류하는 모델이다. 데이터가 아주 크지 않고 feature 간 관계를 빠르게 확인해야 할 때 설명이 쉽다. XGBoost는 이전 트리의 오차를 다음 트리가 보완하는 boosting 계열 모델로, 성능이 강하지만 발표에서는 RandomForest보다 설명 난도가 조금 높다."
    )
    doc.add_heading("RandomForest 상위 feature", level=2)
    add_dataframe_table(doc, rf_importance, [5200, 2200])
    add_image_if_exists(doc, figures_dir / "feature_importance_random_forest.png", "그림 3. RandomForest feature importance")

    doc.add_heading("XGBoost 상위 feature", level=2)
    add_dataframe_table(doc, xgb_importance, [5200, 2200])
    add_image_if_exists(doc, figures_dir / "feature_importance_xgboost.png", "그림 4. XGBoost feature importance")

    doc.add_heading("5. Confusion Matrix 해석", level=1)
    doc.add_paragraph(
        "Confusion matrix는 실제 라벨과 예측 라벨이 얼마나 일치했는지 보여준다. 대각선 값이 클수록 모델이 해당 혼잡도 단계를 잘 맞힌 것이다. 발표에서는 전체 수치를 모두 설명하기보다, 규칙 기반보다 ML 모델의 분류 정확도가 개선되었다는 점을 중심으로 설명하면 된다."
    )
    add_image_if_exists(doc, figures_dir / "confusion_matrix_rule_based_baseline.png", "그림 5. 규칙 기반 confusion matrix", width=5.2)
    add_image_if_exists(doc, figures_dir / "confusion_matrix_random_forest.png", "그림 6. RandomForest confusion matrix", width=5.2)
    add_image_if_exists(doc, figures_dir / "confusion_matrix_xgboost.png", "그림 7. XGBoost confusion matrix", width=5.2)

    doc.add_heading("6. 발표에서 사용할 수 있는 설명", level=1)
    add_callout(
        doc,
        "발표 문장",
        "현재 혼잡도는 대기 시간, 예약 수, 체크인 수 같은 값을 기준으로 규칙 기반으로 산정하고 있습니다. 다만 AI 활용도를 높이기 위해 GPS 추정 인원, 공연 시간대, 가수 인기 여부, 야간 부스 여부 등을 feature로 추가한 시뮬레이션 데이터셋을 만들고 RandomForest와 XGBoost 모델을 실험했습니다. 아직 실제 운영 로그로 검증된 모델은 아니지만, 기존 규칙 기반보다 높은 분류 성능을 보여 향후 실제 데이터가 쌓이면 AI 기반 혼잡도 예측으로 확장할 수 있습니다.",
    )

    doc.add_heading("7. 주의점과 확장 방향", level=1)
    add_bullets(
        doc,
        [
            "이번 결과는 실제 축제 운영 로그로 검증한 최종 모델이 아니라, 운영 경험 기반 시뮬레이션 데이터로 만든 AI 프로토타입이다.",
            "실제 앱에 적용하려면 GPS 로그, 예약/체크인 로그, 주문량, 웨이팅 로그, 공연 스케줄을 시간 단위로 저장해야 한다.",
            "초기에는 RandomForest를 서비스 설명용 모델로 두고, XGBoost는 성능 비교용으로 함께 제시하는 구성이 적절하다.",
            "실제 운영 데이터가 충분히 쌓이면 시간대별 예측, 구역별 혼잡도 예측, 부스 재고 부족 예측, 안전 인력 배치 추천으로 확장할 수 있다.",
        ],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("부록. 생성되는 주요 파일", level=1)
    add_dataframe_table(
        doc,
        pd.DataFrame(
            [
                ["exports/ml/congestion_training_dataset.csv", "학습용 혼잡도 데이터셋"],
                ["exports/ml/model_comparison.csv", "규칙 기반/RF/XGBoost 성능 비교"],
                ["exports/ml/feature_importance_random_forest.csv", "RandomForest feature importance"],
                ["exports/ml/feature_importance_xgboost.csv", "XGBoost feature importance"],
                ["exports/ml/figures/*.png", "발표와 보고서에 사용할 그래프 이미지"],
                ["scripts/ml/run_congestion_ml_pipeline.ps1", "전체 파이프라인 재실행 스크립트"],
            ],
            columns=["파일", "설명"],
        ),
        [5200, 4160],
    )

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a DOCX report for FestFlow congestion ML results.")
    parser.add_argument("--output-dir", type=Path, default=Path("exports/ml"))
    parser.add_argument("--docx", type=Path, default=Path("exports/ml/페스트플로우_혼잡도_인공지능_예측_보고서.docx"))
    args = parser.parse_args()
    build_report(args.output_dir, args.docx)
    print(f"Report written to {args.docx}")


if __name__ == "__main__":
    main()
