from __future__ import annotations

import ast
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SCRIPTS_DIR = ROOT / "scripts" / "ml"
EXPORT_DIR = ROOT / "exports" / "ml"
OUTPUT_DOCX = EXPORT_DIR / "페스트플로우_혼잡도_AI_Python_코드_분석서.docx"

PYTHON_FILES = [
    SCRIPTS_DIR / "build_congestion_dataset.py",
    SCRIPTS_DIR / "train_congestion_models.py",
    SCRIPTS_DIR / "predict_congestion.py",
]

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F8FAFC"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.18) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 10, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], *, font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        shade_cell(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.05)
        add_run(p, str(header), bold=True, size=font_size, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.08)
            add_run(p, str(value), size=font_size)
    doc.add_paragraph()


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p)
    add_run(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4, line=1.15)
    add_run(p, text)


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.15)
    add_run(p, title + " | ", bold=True, size=9, color=DARK)
    add_run(p, body, size=9, color=MUTED)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.5)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def function_rows(path: Path) -> list[list[object]]:
    tree = ast.parse(path.read_text(encoding="utf-8"))
    rows = []
    summaries = {
        "parse_int": "CSV 문자열 값을 안전하게 정수로 변환",
        "read_csv": "원본 CSV가 있으면 읽고, 없으면 빈 리스트 반환",
        "label_from_score": "0~100 점수를 LOW/NORMAL/BUSY/VERY_BUSY로 변환",
        "rule_based_level": "현재 규칙 기반 혼잡도 기준값 계산",
        "target_congestion": "AI 학습용 정답 라벨 생성",
        "popularity_for_day": "날짜/시간대별 가수 인기도 시나리오 생성",
        "build_rows": "학습 데이터 행 전체 생성",
        "write_rows": "생성된 행을 CSV로 저장",
        "make_preprocessor": "숫자/범주 feature 전처리 파이프라인 생성",
        "write_comparison": "모델별 성능 비교 CSV 저장",
        "write_confusion_matrix": "혼동 행렬 CSV 저장",
        "write_feature_importance": "feature importance CSV 저장",
        "build_training_profile": "drift 판단 기준이 되는 학습 데이터 통계 생성",
        "write_training_profile": "training profile JSON 저장",
        "write_random_forest_model": "RandomForest pipeline을 pkl로 저장",
        "metric_row": "accuracy/macro F1 성능 행 생성",
        "read_payload": "추론 입력 JSON을 인자/파일/stdin에서 읽기",
        "normalize_row": "누락 feature를 0으로 채워 모델 입력 형태 정리",
        "drift_check": "운영 입력이 학습 분포를 벗어나는지 검사",
        "predict_one": "단일 부스/구역 혼잡도 예측",
        "predict_many": "여러 부스/구역을 batch로 예측",
        "main": "CLI 인자 처리와 전체 실행 흐름",
    }
    for node in tree.body:
        if isinstance(node, ast.FunctionDef):
            args = ", ".join(arg.arg for arg in node.args.args)
            rows.append([node.name, f"{node.lineno}~{getattr(node, 'end_lineno', node.lineno)}", args, summaries.get(node.name, "파일 실행을 구성하는 보조 함수")])
    return rows


def read_source(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def metric_rows() -> list[list[object]]:
    path = EXPORT_DIR / "model_comparison.csv"
    if not path.exists():
        return [["결과 파일 없음", "", "", "train_congestion_models.py 실행 후 생성"]]
    df = pd.read_csv(path)
    return [[row["model"], row["accuracy"], row["macro_f1"], row["notes"]] for _, row in df.iterrows()]


def dataset_rows() -> list[list[object]]:
    path = EXPORT_DIR / "congestion_training_dataset.csv"
    if not path.exists():
        return [["데이터셋", "파일 없음"]]
    df = pd.read_csv(path)
    return [
        ["파일", "exports/ml/congestion_training_dataset.csv"],
        ["행/컬럼", f"{len(df):,}행 / {len(df.columns)}컬럼"],
        ["정답 컬럼", "target_congestion"],
        ["분류 단계", ", ".join(str(x) for x in sorted(df["target_congestion"].unique()))],
        ["노천극장 수용 인원", f"{int(df['stage_capacity'].max()):,}명"],
        ["데이터 출처 표기", str(df["data_source"].iloc[0]) if "data_source" in df.columns else "HYBRID_SIMULATED"],
    ]


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)
    title = doc.add_paragraph()
    set_spacing(title, before=16, after=8, line=1.1)
    add_run(title, "혼잡도 AI Python 코드 분석서", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=12)
    add_run(subtitle, "데이터 생성, 모델 학습, 실시간 추론 코드의 목적과 동작 흐름을 발표/질문 대응용으로 정리", color=MUTED, size=11)
    add_table(
        doc,
        ["대상 코드", "역할"],
        [
            ["scripts/ml/build_congestion_dataset.py", "혼잡도 학습 데이터 생성"],
            ["scripts/ml/train_congestion_models.py", "RandomForest/XGBoost 학습 및 성능 비교"],
            ["scripts/ml/predict_congestion.py", "저장된 RandomForest 모델을 불러와 실시간 추론"],
        ],
        [4300, 5060],
        font_size=9,
    )


def add_overview(doc: Document) -> None:
    doc.add_heading("1. AI Python 코드 전체 흐름", level=1)
    add_callout(
        doc,
        "핵심 요약",
        "이 Python 코드는 단순히 화면에 AI라고 표시하는 코드가 아니라, 학습 데이터 생성 -> 모델 학습 -> 모델 파일 저장 -> 운영 추론까지 이어지는 작은 ML 파이프라인입니다.",
        GREEN_FILL,
    )
    for item in [
        "첫 번째 파일은 축제 운영 가정과 기존 CSV를 바탕으로 학습 가능한 표 형태 데이터셋을 만듭니다.",
        "두 번째 파일은 같은 데이터로 규칙 기반 baseline, RandomForest, XGBoost를 비교하고 RandomForest 모델 파일을 저장합니다.",
        "세 번째 파일은 서버가 넘긴 현재 상태 feature를 받아 저장된 모델로 predictedLevel, confidence, driftStatus를 계산합니다.",
        "현재 방식은 시계열/LSTM/GNN이 아니라, 특정 시점의 상태값을 보고 30분 뒤 혼잡도를 분류하는 tabular ML입니다.",
    ]:
        add_bullet(doc, item)


def add_dataset_file_section(doc: Document) -> None:
    source = read_source(SCRIPTS_DIR / "build_congestion_dataset.py")
    doc.add_heading("2. build_congestion_dataset.py 분석", level=1)
    add_para(doc, "이 파일의 목적은 AI 학습용 CSV를 만드는 것입니다. 실제 운영 DB만으로는 학습 데이터가 충분하지 않기 때문에, 부스/예약/GPS/공연 정보를 바탕으로 축제 시나리오를 확장합니다.")
    add_table(doc, ["항목", "내용"], dataset_rows(), [2500, 6860], font_size=8.5)
    add_table(doc, ["함수", "라인", "입력", "역할"], function_rows(SCRIPTS_DIR / "build_congestion_dataset.py"), [2100, 1100, 2600, 3560], font_size=7.5)

    doc.add_heading("2.1 핵심 로직: 혼잡도 라벨 기준", level=2)
    add_code_block(
        doc,
        """
def label_from_score(score: float) -> str:
    if score >= 75:
        return "VERY_BUSY"
    if score >= 55:
        return "BUSY"
    if score >= 30:
        return "NORMAL"
    return "LOW"
""",
    )
    add_para(doc, "모든 혼잡도 판단은 최종적으로 네 단계 라벨로 정리됩니다. 발표에서는 이 기준을 기준표처럼 설명하면 이해가 쉽습니다.")

    doc.add_heading("2.2 핵심 로직: 무대 피크 시간과 가수 인기도 반영", level=2)
    add_code_block(
        doc,
        """
if popularity == "HIGH" and is_peak:
    stage_crowd = rng.randint(2600, 4200)
elif popularity == "MEDIUM" and is_peak:
    stage_crowd = rng.randint(1400, 3000)
elif is_peak:
    stage_crowd = rng.randint(450, 1600)
else:
    stage_crowd = rng.randint(120, 1100)
capacity = 4000
""",
    )
    add_para(doc, "사용자가 말한 '18~22시에 무대 줄이 많아지고, 인기 가수면 더 심하다'는 운영 가정이 이 부분에 들어갑니다. 노천극장 수용 인원은 4,000명으로 반영됩니다.")

    doc.add_heading("2.3 핵심 로직: 정답 라벨 생성", level=2)
    add_code_block(
        doc,
        """
score += min(28, gps * 1.8)
score += min(18, wait * 0.45)
score += min(16, reservations * 3.0)

if zone == "STAGE":
    score += stage_ratio * 48
    if peak:
        score += 10
    if popularity == "HIGH":
        score += 16
""",
    )
    add_para(doc, "target_congestion은 GPS, 대기 시간, 예약 수, 구역 종류, 공연 인기도를 조합해서 만듭니다. 따라서 모델은 단순 GPS 숫자만 보는 것이 아니라 운영 맥락을 함께 학습합니다.")
    add_callout(doc, "주의점", "이 데이터는 완전한 실측 데이터가 아니라 HYBRID_SIMULATED 데이터입니다. 발표에서는 '실제 앱 데이터와 축제 운영 가정을 결합해 만든 학습용 데이터'라고 표현하는 것이 정확합니다.", YELLOW_FILL)


def add_training_file_section(doc: Document) -> None:
    doc.add_heading("3. train_congestion_models.py 분석", level=1)
    add_para(doc, "이 파일은 학습 데이터 CSV를 읽고, 규칙 기반 baseline과 ML 모델을 같은 기준으로 비교합니다. 최종적으로 서버가 사용할 RandomForest 모델 파일을 저장합니다.")
    add_table(doc, ["함수", "라인", "입력", "역할"], function_rows(SCRIPTS_DIR / "train_congestion_models.py"), [2100, 1100, 2600, 3560], font_size=7.5)

    doc.add_heading("3.1 Feature 구성", level=2)
    add_code_block(
        doc,
        """
CATEGORICAL_FEATURES = ["zone_type", "artist_popularity"]
NUMERIC_FEATURES = [
    "scenario_day", "hour", "is_peak_time",
    "stage_capacity", "expected_stage_crowd", "stage_load_ratio",
    "gps_count_nearby", "gps_delta_5m", "gps_delta_15m",
    "reservation_count", "checked_in_count", "available_seats",
    "wait_minutes", "remaining_stock",
]
""",
    )
    add_para(doc, "feature는 모델이 보는 입력값입니다. 시간, 공연, GPS, 예약, 대기, 재고 정보가 섞여 있으므로 단순 규칙보다 더 많은 요인을 한 번에 고려할 수 있습니다.")

    doc.add_heading("3.2 전처리 Pipeline", level=2)
    add_code_block(
        doc,
        """
ColumnTransformer(
    transformers=[
        ("numeric", "passthrough", NUMERIC_FEATURES),
        ("category", OneHotEncoder(handle_unknown="ignore"), CATEGORICAL_FEATURES),
    ],
)
""",
    )
    add_para(doc, "숫자 feature는 그대로 통과시키고, 문자열 feature는 OneHotEncoder로 바꿉니다. 예를 들어 STAGE/PUB/FOOD 같은 구역명은 모델이 이해할 수 있는 숫자 벡터가 됩니다.")
    add_table(
        doc,
        ["코드 요소", "뜻", "이 프로젝트에서의 의미"],
        [
            ["ColumnTransformer", "컬럼 종류별로 서로 다른 전처리를 적용하는 scikit-learn 도구", "숫자 feature와 문자열 feature를 한 pipeline 안에서 같이 처리합니다."],
            ["(\"numeric\", \"passthrough\", NUMERIC_FEATURES)", "숫자 컬럼은 변환하지 않고 그대로 모델에 전달", "GPS 수, 대기 시간, 예약 수처럼 이미 숫자인 값은 그대로 사용합니다."],
            ["OneHotEncoder", "문자열 범주를 0/1 벡터로 바꾸는 인코더", "STAGE, PUB, FOOD 같은 구역명을 모델이 계산 가능한 값으로 바꿉니다."],
            ["handle_unknown=\"ignore\"", "학습 때 못 본 범주가 들어와도 에러를 내지 않음", "운영 중 새 구역명이나 예외 값이 들어와도 추론 API가 바로 죽지 않게 합니다."],
            ["verbose_feature_names_out=False", "전처리 후 feature 이름을 너무 길게 만들지 않는 설정", "feature_importance CSV를 발표자가 읽기 쉽게 만듭니다."],
        ],
        [2200, 3600, 3560],
        font_size=7.5,
    )

    doc.add_heading("3.2.1 학습/검증 데이터 분리", level=2)
    add_code_block(
        doc,
        """
x_train, x_test, y_train, y_test, train_idx, test_idx = train_test_split(
    x,
    y,
    df.index,
    test_size=0.25,
    random_state=args.seed,
    stratify=y,
)
""",
    )
    add_table(
        doc,
        ["옵션", "뜻", "이 프로젝트에서의 의미"],
        [
            ["x, y", "입력 feature와 정답 라벨", "x는 GPS/예약/공연 feature, y는 target_congestion입니다."],
            ["df.index", "원본 데이터의 행 번호", "나중에 prediction_samples.csv를 만들 때 어떤 행이 test였는지 추적합니다."],
            ["test_size=0.25", "전체 데이터의 25%를 검증용으로 분리", "학습에 쓰지 않은 데이터로 모델 성능을 확인합니다."],
            ["random_state=args.seed", "데이터 분리 난수 고정", "실행할 때마다 결과가 크게 바뀌지 않도록 재현성을 확보합니다."],
            ["stratify=y", "정답 라벨 비율을 유지하면서 분리", "LOW/NORMAL/BUSY/VERY_BUSY 비율이 train/test에서 비슷하게 유지됩니다."],
        ],
        [2200, 3500, 3660],
        font_size=7.5,
    )

    doc.add_heading("3.3 RandomForest 학습", level=2)
    add_code_block(
        doc,
        """
RandomForestClassifier(
    n_estimators=350,
    max_depth=12,
    min_samples_leaf=3,
    class_weight="balanced",
    random_state=args.seed,
    n_jobs=-1,
)
""",
    )
    add_para(doc, "RandomForest는 여러 결정트리를 학습하고 그 결과를 종합해 분류합니다. 이 프로젝트에서는 운영 안정성을 위해 RandomForest를 실제 추론 모델로 저장합니다.")
    add_table(
        doc,
        ["파라미터", "뜻", "현재 값의 의미"],
        [
            ["n_estimators=350", "생성할 결정트리 개수", "350개의 트리를 만들어 여러 판단을 평균/투표합니다. 너무 적으면 불안정하고, 너무 많으면 느려지므로 중간 이상으로 잡았습니다."],
            ["max_depth=12", "각 트리가 최대 몇 단계까지 질문을 나눌 수 있는지", "모델이 너무 단순하지 않게 하면서도, 학습 데이터를 외워버리는 과적합을 줄이기 위한 제한입니다."],
            ["min_samples_leaf=3", "트리의 마지막 노드에 최소 몇 개 샘플이 있어야 하는지", "1개 샘플만 보고 결론 내리는 극단적 규칙을 막아 예측을 안정화합니다."],
            ["class_weight=\"balanced\"", "라벨 개수가 불균형할 때 적은 클래스에 더 큰 가중치 부여", "LOW처럼 상대적으로 적은 라벨이 무시되지 않게 합니다."],
            ["random_state=args.seed", "모델 학습 난수 고정", "발표/검증 때 같은 데이터로 같은 결과가 나오게 합니다."],
            ["n_jobs=-1", "사용 가능한 CPU 코어를 모두 사용", "여러 트리를 병렬로 학습해 학습 시간을 줄입니다."],
        ],
        [2300, 3300, 3760],
        font_size=7.5,
    )
    add_callout(
        doc,
        "발표용 해석",
        "RandomForest 설정은 성능만 높이기 위한 값이라기보다, 작은 프로젝트 데이터에서 안정적으로 동작하고 설명 가능한 결과를 얻기 위한 값입니다. 특히 class_weight, max_depth, min_samples_leaf는 과적합과 라벨 불균형을 줄이기 위한 안전장치입니다.",
        GREEN_FILL,
    )

    doc.add_heading("3.4 XGBoost 비교", level=2)
    add_code_block(
        doc,
        """
XGBClassifier(
    n_estimators=260,
    max_depth=4,
    learning_rate=0.06,
    objective="multi:softmax",
    eval_metric="mlogloss",
)
""",
    )
    add_para(doc, "XGBoost는 오답을 보완해가며 성능을 올리는 boosting 모델입니다. 현재 운영 서버에 직접 연결된 모델은 RandomForest이고, XGBoost는 비교 실험 결과로 사용합니다.")
    add_table(
        doc,
        ["파라미터", "뜻", "현재 값의 의미"],
        [
            ["n_estimators=260", "순차적으로 만들 boosting tree 개수", "260번에 걸쳐 이전 모델의 부족한 부분을 보완합니다."],
            ["max_depth=4", "각 트리의 최대 깊이", "XGBoost는 트리를 순차적으로 쌓기 때문에 RandomForest보다 얕은 트리로 과적합을 줄였습니다."],
            ["learning_rate=0.06", "각 트리가 최종 예측에 반영되는 비율", "한 번에 크게 보정하지 않고 조금씩 학습해 안정성을 높입니다."],
            ["subsample=0.9", "각 트리 학습에 사용할 행 비율", "전체 데이터의 90%만 샘플링해 과적합을 완화합니다."],
            ["colsample_bytree=0.9", "각 트리 학습에 사용할 컬럼 비율", "feature 일부만 사용하게 해서 특정 feature에 과도하게 의존하는 것을 줄입니다."],
            ["objective=\"multi:softmax\"", "다중 분류에서 최종 클래스를 직접 출력하는 목적 함수", "LOW/NORMAL/BUSY/VERY_BUSY 중 하나를 바로 예측합니다."],
            ["eval_metric=\"mlogloss\"", "다중 분류 확률 예측 품질 평가 지표", "모델이 얼마나 확신 있게 맞추는지도 고려합니다."],
            ["random_state=args.seed", "학습 난수 고정", "실험 결과 재현성을 확보합니다."],
        ],
        [2300, 3300, 3760],
        font_size=7.3,
    )
    add_callout(
        doc,
        "RandomForest와 XGBoost 차이",
        "RandomForest는 여러 트리를 독립적으로 만든 뒤 투표하는 방식이고, XGBoost는 이전 트리의 오답을 다음 트리가 보완하는 방식입니다. 그래서 XGBoost가 성능은 더 나올 수 있지만, 현재 운영 연결은 구조가 더 단순하고 안정적인 RandomForest로 두었습니다.",
        YELLOW_FILL,
    )

    doc.add_heading("3.5 학습 결과", level=2)
    add_table(doc, ["모델", "Accuracy", "Macro F1", "비고"], metric_rows(), [2100, 1300, 1300, 4660], font_size=8)
    add_callout(doc, "해석", "accuracy는 전체 정답률이고, macro F1은 LOW/NORMAL/BUSY/VERY_BUSY 각 라벨을 균형 있게 맞췄는지 보는 지표입니다. 클래스 불균형이 있으므로 macro F1을 함께 보는 것이 좋습니다.", GREEN_FILL)
    add_table(
        doc,
        ["지표/파일", "뜻", "왜 보는가"],
        [
            ["accuracy", "전체 샘플 중 맞춘 비율", "직관적으로 모델이 얼마나 맞는지 설명하기 좋습니다."],
            ["macro_f1", "각 라벨의 F1을 동일 비중으로 평균낸 값", "NORMAL 데이터가 많아도 LOW/VERY_BUSY 성능이 묻히지 않게 확인합니다."],
            ["confusion_matrix_*.csv", "실제 라벨과 예측 라벨의 교차표", "어떤 혼잡 단계를 자주 헷갈리는지 볼 수 있습니다."],
            ["feature_importance_*.csv", "모델이 중요하게 사용한 feature 순위", "GPS, 대기 시간, 예약 수 등이 실제 판단 근거로 쓰였는지 설명할 수 있습니다."],
            ["prediction_samples.csv", "검증 데이터 일부의 실제값과 예측값", "발표에서 모델이 실제로 어떤 행을 어떻게 예측했는지 예시로 보여줄 수 있습니다."],
        ],
        [2600, 3300, 3460],
        font_size=7.5,
    )


def add_prediction_file_section(doc: Document) -> None:
    doc.add_heading("4. predict_congestion.py 분석", level=1)
    add_para(doc, "이 파일은 학습 코드가 아니라 운영 추론 코드입니다. Java 백엔드가 현재 부스 상태를 JSON으로 넘기면, 저장된 RandomForest 모델을 로드해서 혼잡도를 예측합니다.")
    add_table(doc, ["함수", "라인", "입력", "역할"], function_rows(SCRIPTS_DIR / "predict_congestion.py"), [2100, 1100, 2600, 3560], font_size=7.5)

    doc.add_heading("4.1 입력 JSON 처리", level=2)
    add_code_block(
        doc,
        """
def read_payload(input_arg: str | None, input_file: Path | None) -> dict:
    if input_arg:
        return json.loads(input_arg)
    if input_file:
        return json.loads(input_file.read_text(encoding="utf-8"))
    raw = sys.stdin.read().strip()
    return json.loads(raw)
""",
    )
    add_para(doc, "추론 입력은 명령행 JSON, 입력 파일, 표준입력 중 하나로 받을 수 있습니다. 서버에서는 보통 임시 JSON 파일을 만들고 이 스크립트에 넘기는 구조입니다.")

    doc.add_heading("4.2 모델 로드와 예측", level=2)
    add_code_block(
        doc,
        """
model_payload = joblib.load(args.model)
features = model_payload["features"]
pipeline = model_payload["pipeline"]

frame = pd.DataFrame([row], columns=model_payload["features"])
prediction = str(pipeline.predict(frame)[0])
""",
    )
    add_para(doc, "여기서 실제 저장된 `random_forest_congestion_model.pkl`이 사용됩니다. 즉, 프론트에 AI 문구만 붙인 것이 아니라 Python 모델 파일을 불러와 예측합니다.")

    doc.add_heading("4.3 신뢰도 계산", level=2)
    add_code_block(
        doc,
        """
if hasattr(pipeline, "predict_proba"):
    classes = [str(item) for item in pipeline.classes_]
    values = pipeline.predict_proba(frame)[0]
    probabilities = {label: round(float(value), 4) for label, value in zip(classes, values)}
    confidence = max(probabilities.values())
""",
    )
    add_para(doc, "RandomForest는 각 라벨에 대한 확률을 반환할 수 있습니다. 화면에 보이는 신뢰도는 가장 높은 라벨 확률을 confidence로 변환한 값입니다.")
    add_table(
        doc,
        ["코드 요소", "뜻", "예시"],
        [
            ["hasattr(pipeline, \"predict_proba\")", "모델이 라벨별 확률을 제공할 수 있는지 확인", "RandomForest는 가능하므로 confidence 계산에 사용됩니다."],
            ["pipeline.classes_", "모델이 예측할 수 있는 클래스 목록", "LOW, NORMAL, BUSY, VERY_BUSY"],
            ["predict_proba(frame)", "입력 1건에 대한 라벨별 확률 계산", "NORMAL 0.12, BUSY 0.71, VERY_BUSY 0.17 같은 형태"],
            ["probabilities", "라벨별 확률을 JSON으로 정리한 값", "프론트/디버깅에서 모델 판단 분포를 확인할 수 있습니다."],
            ["confidence=max(...)", "가장 높은 라벨 확률", "BUSY가 0.71이면 신뢰도 71%로 표시할 수 있습니다."],
        ],
        [2700, 3500, 3160],
        font_size=7.5,
    )

    doc.add_heading("4.4 Drift 검사", level=2)
    add_code_block(
        doc,
        """
if value < min_value or value > max_value:
    drift_points += 1.0
elif value < p05 or value > p95:
    drift_points += 0.45
""",
    )
    add_para(doc, "실제 운영 입력이 학습 데이터 범위와 너무 다르면 모델 결과를 그대로 믿기 어렵습니다. 이 코드는 학습 분포를 벗어난 정도를 driftScore로 계산하고 NORMAL/CAUTION/WARNING으로 표시합니다.")
    add_table(
        doc,
        ["기준", "뜻", "판단 결과"],
        [
            ["min/max 범위 밖", "학습 데이터에서 한 번도 보지 못한 수준의 값", "drift_points += 1.0으로 강한 경고"],
            ["p05/p95 범위 밖", "학습 데이터의 일반적인 90% 구간 밖", "drift_points += 0.45로 주의 경고"],
            ["normal_max_score=0.15", "driftScore가 이 값 이하이면 정상", "AI 데이터 신뢰도 안정"],
            ["caution_max_score=0.35", "0.15 초과 0.35 이하이면 주의", "AI 데이터 신뢰도 주의"],
            ["0.35 초과", "학습 분포와 차이가 큼", "AI 데이터 신뢰도 낮음"],
        ],
        [2700, 3500, 3160],
        font_size=7.5,
    )

    doc.add_heading("4.5 최종 출력 JSON", level=2)
    add_code_block(
        doc,
        """
return {
    "modelType": model_payload.get("model_type", "RandomForest"),
    "predictedLevel": prediction,
    "confidence": confidence,
    "probabilities": probabilities,
    **drift_check(row, model_payload),
}
""",
    )
    add_para(doc, "이 JSON이 Java 백엔드로 돌아가고, 백엔드는 API 응답에 포함합니다. 프론트는 predictedLevel, confidence, driftStatus 등을 카드로 표시합니다.")


def add_artifact_section(doc: Document) -> None:
    doc.add_heading("5. 생성되는 산출물", level=1)
    add_table(
        doc,
        ["파일", "생성 주체", "의미"],
        [
            ["exports/ml/congestion_training_dataset.csv", "build_congestion_dataset.py", "AI 학습용 표 형태 데이터"],
            ["exports/ml/model_comparison.csv", "train_congestion_models.py", "규칙 기반, RandomForest, XGBoost 성능 비교"],
            ["exports/ml/feature_importance_random_forest.csv", "train_congestion_models.py", "RandomForest가 중요하게 본 feature"],
            ["exports/ml/confusion_matrix_*.csv", "train_congestion_models.py", "실제 라벨과 예측 라벨 비교"],
            ["exports/ml/models/random_forest_congestion_model.pkl", "train_congestion_models.py", "운영 추론에 쓰는 저장 모델"],
            ["exports/ml/models/congestion_training_profile.json", "train_congestion_models.py", "drift 체크 기준 통계"],
        ],
        [3800, 2500, 3060],
        font_size=8,
    )


def add_run_section(doc: Document) -> None:
    doc.add_heading("6. 직접 실행 방법", level=1)
    add_code_block(
        doc,
        """
# 1. 학습 데이터 생성
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\build_congestion_dataset.py

# 2. 모델 학습 및 결과 파일 생성
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\train_congestion_models.py

# 3. 저장된 RandomForest 모델로 단일 추론 테스트
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\predict_congestion.py --input-json "{\\"features\\":{\\"hour\\":20,\\"zone_type\\":\\"STAGE\\",\\"artist_popularity\\":\\"HIGH\\",\\"stage_capacity\\":4000,\\"expected_stage_crowd\\":3200,\\"gps_count_nearby\\":70,\\"wait_minutes\\":45}}"
""",
    )


def add_final_explanation(doc: Document) -> None:
    doc.add_heading("7. 발표에서 말할 수 있는 정확한 표현", level=1)
    add_table(
        doc,
        ["질문", "답변"],
        [
            ["이게 실제 AI인가?", "RandomForest 모델을 학습해 pkl로 저장하고, 서버가 Python 추론 스크립트를 호출해 예측 결과를 받는 구조이므로 실제 모델 기반 AI 기능이라고 말할 수 있습니다."],
            ["어떤 AI인가?", "특정 시점의 GPS, 예약, 대기, 공연, 구역 정보를 feature로 받아 30분 뒤 혼잡 단계를 분류하는 tabular ML입니다."],
            ["데이터 한계는?", "완전 실측 데이터가 아니라 실제 앱 데이터와 축제 운영 가정을 결합한 HYBRID_SIMULATED 데이터입니다."],
            ["왜 RandomForest인가?", "작은 표 형태 데이터에서도 안정적이고, feature importance를 통해 판단 근거를 설명하기 쉬워 발표와 운영 프로토타입에 적합합니다."],
            ["XGBoost는?", "비교 실험 모델입니다. 성능 비교에는 포함되지만 현재 서버 추론 모델은 RandomForest입니다."],
        ],
        [2600, 6760],
        font_size=8,
    )


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_overview(doc)
    add_dataset_file_section(doc)
    add_training_file_section(doc)
    add_prediction_file_section(doc)
    add_artifact_section(doc)
    add_run_section(doc)
    add_final_explanation(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_doc()
