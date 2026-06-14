from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EXPORT_DIR = ROOT / "exports" / "ml"
FIGURE_DIR = EXPORT_DIR / "figures" / "ai_code_deep_dive"
OUTPUT_DOCX = EXPORT_DIR / "페스트플로우_현재_인공지능_구현_코드상세해설서.docx"
FALLBACK_OUTPUT_DOCX = EXPORT_DIR / "페스트플로우_현재_인공지능_구현_코드상세해설서_적정분량.docx"

DATASET_PATH = EXPORT_DIR / "congestion_training_dataset.csv"
MODEL_COMPARISON_PATH = EXPORT_DIR / "model_comparison.csv"
RF_IMPORTANCE_PATH = EXPORT_DIR / "feature_importance_random_forest.csv"
XGB_IMPORTANCE_PATH = EXPORT_DIR / "feature_importance_xgboost.csv"
SAMPLES_PATH = EXPORT_DIR / "prediction_samples.csv"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F4F6F9"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: int = 10, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=6, line=1.2)
    if bold_label:
        add_run(p, bold_label, bold=True, color=DARK)
        add_run(p, text)
    else:
        add_run(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4, line=1.15)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    add_run(p, text)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for cell, header, width in zip(header_cells, headers, widths):
        set_cell_shading(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.05)
        add_run(p, str(header), bold=True, size=font_size, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.08)
            add_run(p, str(value), size=font_size)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=2, line=1.15)
    add_run(p, title + " | ", bold=True, color=DARK, size=9)
    add_run(p, body, color=MUTED, size=9)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.5)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def create_figures() -> list[tuple[str, Path, str]]:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    figures: list[tuple[str, Path, str]] = []

    dataset = pd.read_csv(DATASET_PATH)
    comparison = pd.read_csv(MODEL_COMPARISON_PATH)
    rf_importance = pd.read_csv(RF_IMPORTANCE_PATH)
    samples = pd.read_csv(SAMPLES_PATH)

    path = FIGURE_DIR / "01_target_label_distribution.png"
    dataset["target_congestion"].value_counts().reindex(["LOW", "NORMAL", "BUSY", "VERY_BUSY"]).plot(
        kind="bar", color=["#94A3B8", "#60A5FA", "#F59E0B", "#EF4444"], figsize=(7, 4)
    )
    plt.title("Target Label Distribution")
    plt.xlabel("target_congestion")
    plt.ylabel("rows")
    plt.tight_layout()
    plt.savefig(path, dpi=170)
    plt.close()
    figures.append(("정답 라벨 분포", path, "학습 데이터가 LOW/NORMAL/BUSY/VERY_BUSY 네 단계로 어떻게 나뉘는지 보여줍니다."))

    path = FIGURE_DIR / "02_model_performance_comparison.png"
    comparison.set_index("model")[["accuracy", "macro_f1"]].plot(kind="bar", figsize=(7, 4), color=["#2563EB", "#14B8A6"])
    plt.title("Model Performance Comparison")
    plt.ylim(0, 1)
    plt.ylabel("score")
    plt.xticks(rotation=15, ha="right")
    plt.tight_layout()
    plt.savefig(path, dpi=170)
    plt.close()
    figures.append(("모델 성능 비교", path, "규칙 기반 baseline과 RandomForest/XGBoost의 정확도, macro F1을 비교합니다."))

    path = FIGURE_DIR / "03_random_forest_feature_importance.png"
    top_rf = rf_importance.head(10).sort_values("importance")
    plt.figure(figsize=(7, 4.5))
    plt.barh(top_rf["feature"], top_rf["importance"], color="#16A34A")
    plt.title("RandomForest Feature Importance")
    plt.xlabel("importance")
    plt.tight_layout()
    plt.savefig(path, dpi=170)
    plt.close()
    figures.append(("RandomForest 특성 중요도 / Feature Importance", path, "모델이 혼잡도 판단에 크게 사용한 입력값을 순서대로 보여줍니다."))

    path = FIGURE_DIR / "04_prediction_sample_correctness.png"
    sample_view = samples.head(60).copy()
    sample_view["correct"] = sample_view["target_congestion"] == sample_view["random_forest_prediction"]
    counts = sample_view["correct"].map({True: "correct", False: "wrong"}).value_counts()
    counts.reindex(["correct", "wrong"]).fillna(0).plot(kind="bar", color=["#22C55E", "#EF4444"], figsize=(6, 3.5))
    plt.title("Prediction Sample Correctness")
    plt.xlabel("sample result")
    plt.ylabel("rows")
    plt.tight_layout()
    plt.savefig(path, dpi=170)
    plt.close()
    figures.append(("샘플 예측 일치 여부", path, "prediction_samples.csv 일부에서 RandomForest 예측이 정답 라벨과 맞았는지 요약합니다."))

    return figures


def dataset_summary() -> list[list[object]]:
    dataset = pd.read_csv(DATASET_PATH)
    return [
        ["행 수", f"{len(dataset):,}행"],
        ["컬럼 수", f"{len(dataset.columns)}개"],
        ["정답 컬럼", "target_congestion"],
        ["분류 단계", ", ".join(["LOW", "NORMAL", "BUSY", "VERY_BUSY"])],
        ["데이터 성격", str(dataset["data_source"].iloc[0]) if "data_source" in dataset.columns else "HYBRID_SIMULATED"],
        ["노천극장 수용 인원", f"{int(dataset['stage_capacity'].max()):,}명 기준"],
    ]


def top_features(path: Path, label: str) -> list[list[object]]:
    if not path.exists():
        return [[label, "파일 없음", ""]]
    df = pd.read_csv(path).head(8)
    return [[label, row["feature"], round(float(row["importance"]), 4)] for _, row in df.iterrows()]


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)
    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "현재 인공지능 구현 코드 해설서", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph()
    add_run(
        subtitle,
        "데이터 생성, 모델 학습, 실시간 추론, 백엔드 연동, 프론트 표시를 발표용으로 설명하기 위한 적정 분량 문서",
        color=MUTED,
        size=11,
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["핵심 목적", "코드 전체를 한 줄씩 설명하지 않고, AI 기능의 동작 원리를 블록 단위로 설명"],
            ["현재 AI 방식", "특정 시점의 상태값을 보고 30분 뒤 혼잡도를 분류하는 tabular ML"],
            ["운영 모델", "Python에서 저장한 RandomForest .pkl 모델을 서버가 호출해서 추론"],
            ["비교 모델", "XGBoost는 성능 비교용으로 학습 결과를 함께 생성"],
            ["주의점", "시계열/LSTM/GNN 수준은 아직 아니며, 현재는 표 형태 feature 기반 분류 모델"],
        ],
        [2200, 7160],
        font_size=9,
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_overview(doc: Document) -> None:
    doc.add_heading("1. 전체 구조 요약", level=1)
    add_callout(
        doc,
        "한 문장 요약",
        "현재 AI 기능은 축제장의 현재 상태 feature를 만들고, RandomForest 모델이 30분 뒤 혼잡 단계를 예측한 뒤, 백엔드 API와 프론트 화면에 표시하는 구조입니다.",
        GREEN_FILL,
    )
    for item in [
        "데이터 생성: 부스, 예약, GPS, 공연 흐름을 바탕으로 학습용 CSV를 생성합니다.",
        "모델 학습: 규칙 기반 baseline, RandomForest, XGBoost를 같은 데이터로 비교합니다.",
        "모델 저장: 실제 서버 추론에는 RandomForest 모델을 `random_forest_congestion_model.pkl`로 저장해 사용합니다.",
        "실시간 추론: Java 백엔드가 현재 부스 상태를 feature로 만들고 Python 추론 스크립트를 호출합니다.",
        "화면 표시: 프론트엔드는 예측 단계, 신뢰도, drift 상태, 위험 점수를 카드로 보여줍니다.",
    ]:
        add_bullet(doc, item)


def add_dataset_section(doc: Document) -> None:
    doc.add_heading("2. 학습 데이터 구성", level=1)
    add_table(doc, ["항목", "값"], dataset_summary(), [2600, 6760], font_size=9)
    add_para(
        doc,
        "데이터는 실제 서비스 DB에서 뽑은 부스/예약/GPS/공연 정보를 그대로 쓰는 것이 아니라, 사용자가 알고 있는 축제 운영 패턴을 반영해 학습 가능한 형태로 확장한 HYBRID_SIMULATED 데이터입니다.",
    )
    add_table(
        doc,
        ["feature 그룹", "주요 컬럼", "의미"],
        [
            ["시간", "scenario_day, hour, is_peak_time", "몇 번째 시나리오 날짜인지, 몇 시인지, 18~22시 피크인지"],
            ["공연", "artist_popularity, expected_stage_crowd, stage_capacity, stage_load_ratio", "인기 가수 여부와 노천극장 수용 대비 예상 관객 규모"],
            ["위치/구역", "zone_type, booth_id, is_night_booth", "무대, 주점, 푸드, 체험, 굿즈, 안전 구역의 차이"],
            ["현장 신호", "gps_count_nearby, gps_delta_5m, gps_delta_15m", "주변 추정 인원과 최근 증가/감소 흐름"],
            ["예약/대기", "reservation_count, checked_in_count, available_seats, wait_minutes", "예약, 체크인, 잔여 좌석, 예상 대기 시간"],
            ["정답", "target_congestion", "30분 뒤 혼잡도를 LOW/NORMAL/BUSY/VERY_BUSY 중 하나로 지정"],
        ],
        [1700, 3300, 4360],
        font_size=8,
    )


def add_code_sections(doc: Document) -> None:
    doc.add_heading("3. 핵심 Python 코드 블록 해설", level=1)
    add_callout(
        doc,
        "설명 방식",
        "이 문서는 모든 줄을 해설하지 않습니다. 발표에서 설명 가치가 큰 블록만 골라 입력, 처리, 출력 관점으로 정리합니다.",
        YELLOW_FILL,
    )

    doc.add_heading("3.1 데이터 생성: 축제 운영 가정을 feature로 변환", level=2)
    add_code_block(
        doc,
        """
if popularity == "HIGH" and is_peak:
    stage_crowd = rng.randint(2600, 4200)
elif popularity == "MEDIUM" and is_peak:
    stage_crowd = rng.randint(1400, 3000)
elif is_peak:
    stage_crowd = rng.randint(450, 1600)
else:
    stage_crowd = rng.randint(120, 1100)
capacity = 4000
""",
    )
    add_para(
        doc,
        "18~22시 피크 시간대와 가수 인기도를 이용해 무대 예상 관객 수를 다르게 생성합니다. 노천극장 수용 인원은 4,000명으로 잡고, `stage_load_ratio`를 만들어 모델이 무대 과밀 상황을 학습할 수 있게 합니다.",
    )

    doc.add_heading("3.2 정답 라벨 생성: 단순 규칙보다 조금 더 현실적인 목표값", level=2)
    add_code_block(
        doc,
        """
score += min(28, gps * 1.8)
score += min(18, wait * 0.45)
score += min(16, reservations * 3.0)

if zone == "STAGE":
    score += stage_ratio * 48
    if peak:
        score += 10
    if popularity == "HIGH":
        score += 16
""",
    )
    add_para(
        doc,
        "target_congestion은 GPS, 대기 시간, 예약 수, 구역 종류, 공연 인기도를 조합해서 만듭니다. 즉, 현재 규칙 기반 점수와 비슷하지만 공연/구역 맥락을 더 반영한 정답 라벨을 만들어 모델이 학습하도록 했습니다.",
    )

    doc.add_heading("3.3 학습 전처리: 숫자 feature와 범주 feature 분리", level=2)
    add_code_block(
        doc,
        """
ColumnTransformer(
    transformers=[
        ("numeric", "passthrough", NUMERIC_FEATURES),
        ("category", OneHotEncoder(handle_unknown="ignore"), CATEGORICAL_FEATURES),
    ],
)
""",
    )
    add_para(
        doc,
        "숫자값은 그대로 모델에 넣고, `zone_type`, `artist_popularity` 같은 문자열 범주는 OneHotEncoder로 숫자 벡터로 바꿉니다. 이 과정이 있어야 RandomForest가 구역/가수 인기도 같은 범주 정보를 사용할 수 있습니다.",
    )

    doc.add_heading("3.4 RandomForest 학습: 현재 운영에 연결된 모델", level=2)
    add_code_block(
        doc,
        """
RandomForestClassifier(
    n_estimators=350,
    max_depth=12,
    min_samples_leaf=3,
    class_weight="balanced",
    random_state=args.seed,
    n_jobs=-1,
)
""",
    )
    add_para(
        doc,
        "RandomForest는 여러 개의 결정트리를 학습해서 다수결로 혼잡 단계를 예측합니다. `class_weight='balanced'`는 LOW처럼 적은 라벨이 무시되지 않도록 보정하는 설정입니다.",
    )

    doc.add_heading("3.5 XGBoost 학습: 비교 실험용 모델", level=2)
    add_code_block(
        doc,
        """
XGBClassifier(
    n_estimators=260,
    max_depth=4,
    learning_rate=0.06,
    objective="multi:softmax",
    eval_metric="mlogloss",
)
""",
    )
    add_para(
        doc,
        "XGBoost는 이전 트리의 오답을 다음 트리가 보완하는 gradient boosting 계열 모델입니다. 현재 서버에 직접 연결된 운영 모델은 RandomForest이고, XGBoost는 발표에서 성능 비교 근거로 사용합니다.",
    )

    doc.add_heading("3.6 실시간 추론: 저장된 .pkl 모델을 로드", level=2)
    add_code_block(
        doc,
        """
model_payload = joblib.load(args.model)
pipeline = model_payload["pipeline"]
frame = pd.DataFrame([row], columns=model_payload["features"])
prediction = str(pipeline.predict(frame)[0])
probabilities = pipeline.predict_proba(frame)[0]
""",
    )
    add_para(
        doc,
        "`predict_congestion.py`는 Java 서버가 넘긴 현재 feature를 DataFrame으로 바꾼 뒤, 저장된 RandomForest pipeline으로 예측합니다. 결과는 predictedLevel, confidence, probabilities, driftStatus 형태의 JSON으로 돌아갑니다.",
    )

    doc.add_heading("3.7 Drift 체크: 학습 데이터와 너무 다른 입력인지 검사", level=2)
    add_code_block(
        doc,
        """
if value < min_value or value > max_value:
    drift_points += 1.0
elif value < p05 or value > p95:
    drift_points += 0.45
""",
    )
    add_para(
        doc,
        "모델이 학습 때 보지 못한 범위의 값이 들어오면 예측을 완전히 믿기 어렵습니다. 그래서 입력값이 학습 데이터의 min/max 또는 p05/p95 범위를 벗어나는지 보고 NORMAL/CAUTION/WARNING 상태를 함께 반환합니다.",
    )


def add_results_section(doc: Document, figures: list[tuple[str, Path, str]]) -> None:
    doc.add_heading("4. 모델 결과와 해석", level=1)
    comparison = pd.read_csv(MODEL_COMPARISON_PATH)
    add_table(
        doc,
        ["모델", "Accuracy", "Macro F1", "해석"],
        [
            [row["model"], row["accuracy"], row["macro_f1"], row["notes"]]
            for _, row in comparison.iterrows()
        ],
        [2200, 1400, 1400, 4360],
        font_size=8,
    )
    add_table(
        doc,
        ["모델", "중요 feature", "중요도"],
        top_features(RF_IMPORTANCE_PATH, "RandomForest") + top_features(XGB_IMPORTANCE_PATH, "XGBoost"),
        [1800, 5000, 2560],
        font_size=8,
    )
    for title, path, explanation in figures:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(5.8))
        add_para(doc, explanation)


def add_backend_frontend_section(doc: Document) -> None:
    doc.add_heading("5. 백엔드와 프론트 연결", level=1)

    doc.add_heading("5.1 Java 백엔드: 현재 상태를 모델 feature로 변환", level=2)
    add_code_block(
        doc,
        """
Map<Long, AiModelPredictionDto> modelPredictions = modelPredictions(snapshot, eventSoon);
AiModelPredictionDto aiModel = modelPrediction != null
        ? modelPrediction
        : AiModelPredictionDto.fallback(fallbackPredictedLevel, modelFactors, "MODEL_UNAVAILABLE");
""",
    )
    add_para(
        doc,
        "`AiCongestionService`는 부스별 GPS, 예약, 대기 시간, 재고, 공연 여부를 모아 모델 입력값을 만들고, 모델 결과가 있으면 AI 예측을 사용합니다. 모델 호출이 실패하면 서비스가 멈추지 않도록 fallback 규칙 결과를 사용합니다.",
    )

    doc.add_heading("5.2 Python 호출 서비스: 서버에서 실제 모델 파일 사용", level=2)
    add_code_block(
        doc,
        """
ProcessBuilder builder = new ProcessBuilder(
    pythonCommand,
    script.toString(),
    "--model", model.toString(),
    "--input-file", inputFile.toString(),
    "--output-file", outputFile.toString()
);
""",
    )
    add_para(
        doc,
        "`PythonCongestionModelService`는 Java 내부에서 Python 프로세스를 실행합니다. 이때 `random_forest_congestion_model.pkl` 파일과 `predict_congestion.py`를 직접 사용하므로, 단순 표시용 mock이 아니라 저장된 모델 기반 추론 구조입니다.",
    )

    doc.add_heading("5.3 프론트엔드: 사용자가 보는 AI 카드", level=2)
    add_code_block(
        doc,
        """
<b>{item.aiModel.modelBased ? item.aiModel.modelType || "RandomForest" : "Fallback"}</b>
<span>신뢰도 {Math.round(Number(item.aiModel.confidence) * 100)}%</span>
<em>{driftStatusLabel(item.aiModel.driftStatus)}</em>
<small>AI 위험 점수 {Number(item.riskScore) || 0}점</small>
""",
    )
    add_para(
        doc,
        "프론트에서는 모델 종류, 신뢰도, drift 상태, 위험 점수를 보여줍니다. 그래서 배포 환경에서 `Fallback`으로 표시되면 Python 모델 호출이 실패했거나 모델 파일/경로/환경변수 문제가 있다는 신호입니다.",
    )


def add_runbook_section(doc: Document) -> None:
    doc.add_heading("6. 직접 실행 방법", level=1)
    add_code_block(
        doc,
        """
# 1. 학습 데이터 생성
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\build_congestion_dataset.py

# 2. 모델 학습 및 결과 파일 생성
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\train_congestion_models.py

# 3. 저장된 RandomForest 모델로 추론 테스트
.\\.venv-ml\\Scripts\\python.exe scripts\\ml\\predict_congestion.py --input-json "{\\"features\\":{\\"hour\\":20,\\"zone_type\\":\\"STAGE\\",\\"artist_popularity\\":\\"HIGH\\",\\"stage_capacity\\":4000,\\"expected_stage_crowd\\":3200,\\"gps_count_nearby\\":70,\\"wait_minutes\\":45}}"
""",
    )
    add_table(
        doc,
        ["확인할 파일", "의미"],
        [
            ["exports/ml/congestion_training_dataset.csv", "학습 데이터"],
            ["exports/ml/model_comparison.csv", "모델 성능 비교"],
            ["exports/ml/models/random_forest_congestion_model.pkl", "실제 추론에 쓰는 모델 파일"],
            ["exports/ml/models/congestion_training_profile.json", "drift 판단 기준"],
            ["exports/ml/prediction_samples.csv", "예측 샘플 결과"],
        ],
        [4300, 5060],
        font_size=8,
    )


def add_limitations_section(doc: Document) -> None:
    doc.add_heading("7. 발표 때 정확히 말해야 할 범위", level=1)
    add_table(
        doc,
        ["질문", "답변 방향"],
        [
            ["진짜 AI가 들어간 것인가?", "네. RandomForest 모델을 학습해 .pkl로 저장했고, 서버가 Python 추론 스크립트를 호출해 결과를 API 응답과 화면에 반영합니다."],
            ["시계열 모델인가?", "아닙니다. 현재는 특정 시점의 feature를 보고 30분 뒤 혼잡도를 분류하는 tabular ML입니다."],
            ["XGBoost도 운영에 쓰는가?", "현재 운영 연결은 RandomForest입니다. XGBoost는 같은 데이터셋에서 비교 실험한 모델입니다."],
            ["데이터가 완전한 실제 데이터인가?", "완전한 실측 데이터는 아닙니다. 실제 앱 데이터와 축제 운영 가정을 결합해 만든 HYBRID_SIMULATED 학습 데이터입니다."],
            ["앞으로 고도화하면?", "실제 GPS 로그가 충분히 쌓이면 시계열 예측, 온라인 재학습, 공간-시간 그래프 모델로 확장할 수 있습니다."],
        ],
        [2600, 6760],
        font_size=8,
    )


def build_doc() -> None:
    figures = create_figures()
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_overview(doc)
    add_dataset_section(doc)
    add_code_sections(doc)
    add_results_section(doc, figures)
    add_backend_frontend_section(doc)
    add_runbook_section(doc)
    add_limitations_section(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT_DOCX)
        print(f"written: {OUTPUT_DOCX}")
    except PermissionError:
        doc.save(FALLBACK_OUTPUT_DOCX)
        print(f"target locked, written fallback: {FALLBACK_OUTPUT_DOCX}")


if __name__ == "__main__":
    build_doc()
