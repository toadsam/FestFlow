from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "festflow" / "페스트플로우_발표_질의응답_대비문서.docx"

FONT = "Malgun Gothic"
BLUE = "1F4D78"
DARK = "0B2545"
MUTED = "555555"
HEADER_FILL = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
NOTE_FILL = "FFFBEB"
OK_FILL = "ECFDF5"


def set_font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, size=None, bold=False, color=None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, header_fill=LIGHT_GRAY):
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_font(run, size=9, bold=row_idx == 0, color=BLUE if row_idx == 0 else None)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths, header_fill)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=OK_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    p = cell.paragraphs[0]
    add_run(p, f"{title}\n", bold=True, color=BLUE)
    add_run(p, body)
    set_spacing(p, after=0, line=1.2)
    style_table(table, [9360], fill)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc):
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)

    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "발표 질의응답 대비 문서", bold=True, color=DARK, size=24)

    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=14)
    add_run(
        subtitle,
        "기획, 운영, 기술, 실시간 처리, AI 혼잡도 예측, 데이터 한계, 배포, 보안, 확장성 질문과 답변",
        color=MUTED,
        size=11,
    )

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 목적", "발표 중 나올 수 있는 질문에 바로 답변하기 위한 예상 Q&A 정리"],
            ["답변 톤", "과장하지 않고 현재 구현 범위와 한계를 명확히 인정하는 방식"],
            ["현재 AI 범위", "RandomForest 기반 tabular ML 혼잡도 예측. 시계열/LSTM/GNN은 미구현"],
            ["주의", "현재 브랜치 기준 Java JSON 직접 추론 구조는 문서 답변에 포함하지 않음"],
            ["추천 사용법", "질문을 외우기보다 핵심 논리와 방어 문장을 익히는 용도"],
        ],
        [2200, 7160],
        HEADER_FILL,
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


SECTIONS = [
    (
        "1. 기획 의도와 문제 정의",
        [
            ("왜 FestFlow를 만들었나요?", "대학교 축제에서는 공연, 부스, 예약, 지도, 분실물, 공지, 운영자 대응이 여러 채널로 흩어져 있습니다. FestFlow는 방문자에게는 필요한 정보를 한 화면에서 제공하고, 운영자에게는 현장 상태를 빠르게 갱신할 수 있는 통합 운영 도구를 제공하기 위해 기획했습니다."),
            ("기존 축제 안내 방식과 가장 큰 차이는 무엇인가요?", "기존 방식은 포스터, 단체 채팅방, SNS 공지처럼 정보 전달 중심입니다. FestFlow는 정보 제공뿐 아니라 예약, 혼잡도, 실시간 공지, 관리자/운영자 화면, AI 안내까지 연결해서 실제 운영 흐름을 서비스 안에서 처리한다는 점이 다릅니다."),
            ("사용자 입장에서 해결되는 문제는 무엇인가요?", "사용자는 어떤 부스가 어디 있는지, 얼마나 기다려야 하는지, 공연이 언제 시작되는지, 지금 어디로 가면 좋은지 빠르게 판단할 수 있습니다. 특히 모바일 기준으로 축제장에서 바로 쓰는 상황을 고려했습니다."),
            ("운영자 입장에서 해결되는 문제는 무엇인가요?", "운영자는 부스 상태, 예약 가능 여부, 대기 시간, 공지, 스태프 상태를 관리자/운영 화면에서 관리할 수 있습니다. 데이터가 바뀌면 프론트에 실시간으로 반영되어 현장 안내 부담을 줄일 수 있습니다."),
            ("서비스의 핵심 가치는 무엇인가요?", "핵심 가치는 축제 현장의 분산된 정보를 하나의 운영형 웹앱으로 묶는 것입니다. 단순 안내 사이트가 아니라 방문자 화면, 운영자 화면, 관리자 기능, 실시간 데이터, AI 추천을 연결한 현장 운영 플랫폼을 목표로 했습니다."),
            ("왜 축제라는 도메인을 선택했나요?", "축제는 짧은 기간에 많은 사람이 몰리고, 정보 변화가 빠르며, 대기열과 혼잡 문제가 뚜렷합니다. 그래서 실시간 데이터 처리, 예약, 혼잡도 예측, 운영자 관리 같은 웹 서비스 기술을 종합적으로 보여주기 좋은 도메인입니다."),
            ("기획에서 가장 중요하게 본 사용자 경험은 무엇인가요?", "축제 현장에서는 사용자가 오래 탐색하지 않습니다. 그래서 홈 화면에서 공연, 부스, 지도, AI 가이드, 혼잡도 같은 핵심 정보로 바로 이동할 수 있도록 구성했습니다."),
            ("서비스가 실제 축제에 적용된다면 가장 먼저 효과가 나는 부분은 어디인가요?", "가장 먼저 효과가 나는 부분은 부스/공연 정보 통합과 실시간 공지입니다. 그 다음으로 예약, 대기 시간, 혼잡도 예측이 운영 효율을 높이는 역할을 할 수 있습니다."),
            ("이 프로젝트가 단순 CRUD와 다른 점은 무엇인가요?", "부스와 공연 CRUD만 있는 것이 아니라, 예약 상태, SSE 실시간 반영, 관리자/운영자 권한, AI 가이드, AI 혼잡도 예측, 분실물과 스태프 기능까지 연결되어 있습니다. 여러 기능이 실제 운영 흐름 안에서 상호작용합니다."),
            ("기획상 가장 어려웠던 점은 무엇인가요?", "방문자용 화면과 운영자용 화면의 요구가 다르기 때문에, 같은 데이터를 서로 다른 관점으로 보여주는 구조를 잡는 것이 어려웠습니다. 또한 AI 기능은 과장하지 않고 현재 데이터 수준에서 설득력 있게 설명하는 것이 중요했습니다."),
        ],
    ),
    (
        "2. 서비스 기능과 사용자 흐름",
        [
            ("주요 기능을 한 문장으로 설명하면?", "방문자는 축제 정보를 보고 이동 결정을 하고, 운영자는 현장 데이터를 관리하며, 시스템은 실시간 데이터와 AI 예측을 바탕으로 혼잡도와 추천 정보를 제공하는 서비스입니다."),
            ("홈 화면에서 사용자는 무엇을 할 수 있나요?", "홈 화면에서는 오늘의 축제 정보, AI 축제 가이드, 추천 부스/공연 카드, 전체 혼잡도 요약을 확인하고 부스, 공연, 지도, 분석 화면으로 이동할 수 있습니다."),
            ("AI 축제 가이드는 어떤 역할인가요?", "AI 축제 가이드는 현재 부스, 공연, 혼잡도 데이터를 바탕으로 사용자가 지금 어떤 동선을 선택하면 좋을지 안내하는 역할입니다. OpenAI API가 없거나 실패하면 기본 안내로 fallback됩니다."),
            ("부스 기능은 어떻게 동작하나요?", "부스 목록과 상세 화면에서 위치, 설명, 이미지, 예약 가능 여부, 대기 시간 등을 보여줍니다. 관리자나 운영자가 값을 수정하면 실시간 스트림을 통해 사용자 화면에 반영될 수 있습니다."),
            ("공연 기능은 어떻게 동작하나요?", "공연 목록과 시간 정보를 제공하고, 홈 화면에서는 다음 공연을 추천 카드로 보여줍니다. 공연 일정은 혼잡도 예측 feature에도 영향을 줄 수 있는 데이터로 사용됩니다."),
            ("예약 기능은 왜 필요한가요?", "축제 부스에서는 대기열이 길어지는 문제가 생깁니다. 예약 기능은 사용자가 대기 상태를 예측하고, 운영자가 좌석/수용 가능 상태를 관리하는 데 도움을 줍니다."),
            ("지도 기능의 목적은 무엇인가요?", "부스 위치를 빠르게 찾고, 혼잡하거나 추천되는 구역으로 이동할 수 있게 돕는 기능입니다. 축제 현장에서는 정보 자체보다 위치 기반 이동성이 중요합니다."),
            ("분실물 기능은 왜 포함했나요?", "축제 현장에서 자주 발생하는 운영 문제 중 하나가 분실물입니다. 사용자가 분실물을 조회하고, 운영자가 등록/관리할 수 있게 해 서비스의 실제 현장성을 높였습니다."),
            ("스태프 기능은 어떤 목적이 있나요?", "스태프 상태와 역할을 관리해 운영자가 현장 대응을 더 체계적으로 할 수 있게 하는 기능입니다. 혼잡하거나 문제가 있는 구역에 인력을 배치하는 운영 시나리오와 연결됩니다."),
            ("AI Match 기능은 전체 서비스에서 어떤 역할인가요?", "AI Match는 축제 부가 이벤트 기능입니다. 핵심 혼잡도 AI와는 별개의 기능이며, 이미지 변환/프로필/매칭 흐름을 통해 축제 참여 경험을 확장하는 역할입니다."),
            ("관리자 화면과 운영자 화면은 어떻게 다르나요?", "관리자 화면은 전체 데이터 관리와 설정에 가깝고, 운영자 화면은 현장에서 부스 상태나 실시간 운영 값을 빠르게 수정하는 데 초점을 둡니다."),
            ("사용자 화면과 관리자 화면이 같은 데이터를 쓰나요?", "네, 기본적으로 같은 백엔드 데이터를 사용합니다. 다만 권한과 목적이 다르기 때문에 방문자 화면은 조회 중심, 관리자/운영 화면은 수정과 운영 중심으로 분리했습니다."),
        ],
    ),
    (
        "3. 프론트엔드 구조",
        [
            ("프론트엔드는 어떤 기술로 만들었나요?", "React와 Vite 기반으로 구성했습니다. React는 컴포넌트 단위로 화면을 나누기 좋고, Vite는 개발 서버와 빌드 속도가 빠르기 때문에 프로젝트 규모에 적합했습니다."),
            ("홈 화면의 핵심 파일은 무엇인가요?", "홈 화면은 frontend/src/pages/HomePage.jsx가 담당합니다. 이 파일에서 부스, 공연, 방문량, AI 가이드 데이터를 불러오고 추천 카드와 실시간 혼잡도 요약을 렌더링합니다."),
            ("API 호출은 어디에 모아두었나요?", "frontend/src/api.js에 모아두었습니다. 화면 컴포넌트가 직접 fetch URL을 흩뿌리지 않고, API 함수로 감싸서 유지보수성을 높였습니다."),
            ("라우팅은 어떻게 처리하나요?", "React Router를 사용합니다. App.jsx가 공통 레이아웃과 하단 네비게이션 노출 조건을 관리하고, 각 페이지 컴포넌트가 라우트에 연결됩니다."),
            ("fallback 데이터는 왜 있나요?", "백엔드가 꺼져 있거나 네트워크 요청이 실패해도 화면이 완전히 깨지지 않게 하기 위해 fallback 데이터를 둡니다. 발표나 데모 상황에서도 최소한의 화면 확인이 가능합니다."),
            ("홈 화면에서 useMemo를 쓰는 이유는 무엇인가요?", "부스와 공연 데이터에서 추천 카드나 혼잡도 비율을 계산할 때, 관련 데이터가 바뀔 때만 다시 계산하도록 하기 위해 사용합니다. 복잡한 계산은 아니지만 구조를 명확하게 유지하는 데 도움이 됩니다."),
            ("홈 화면은 실시간으로 어떻게 바뀌나요?", "홈 화면은 createBoothStream과 createEventStream으로 SSE를 구독합니다. 부스나 공연 데이터가 서버에서 방송되면 해당 데이터를 받아 state를 갱신합니다."),
            ("분석 화면은 어떤 역할인가요?", "AnalyticsPage.jsx는 전체 혼잡도, 구역별 혼잡도, 30분 뒤 AI 혼잡도 예측, AI 추천 근거를 보여주는 화면입니다. AI 혼잡도 기능은 주로 이 화면에서 확인합니다."),
            ("프론트엔드에서 AI 모델을 직접 실행하나요?", "아닙니다. 프론트엔드는 백엔드 API가 반환한 AI 예측 결과를 표시합니다. 모델 실행은 서버 또는 Python 추론 스크립트 쪽에서 처리합니다."),
            ("프론트에서 AI 신뢰도와 drift를 보여주는 이유는 무엇인가요?", "예측 결과만 보여주면 사용자가 무조건 정답처럼 받아들일 수 있습니다. 신뢰도와 drift 상태를 함께 보여주면 현재 예측이 얼마나 안정적인지 판단할 수 있습니다."),
        ],
    ),
    (
        "4. 백엔드 구조와 API",
        [
            ("백엔드는 어떤 기술로 만들었나요?", "Spring Boot 기반입니다. Controller, Service, Repository, DTO 계층으로 나누어 HTTP 요청 처리, 비즈니스 로직, DB 접근, 응답 구조를 분리했습니다."),
            ("Controller와 Service를 나눈 이유는 무엇인가요?", "Controller는 요청과 응답의 입구 역할만 하고, 실제 규칙과 계산은 Service에 두기 위해서입니다. 이렇게 나누면 기능이 커져도 테스트와 유지보수가 쉬워집니다."),
            ("DTO를 사용하는 이유는 무엇인가요?", "Entity를 그대로 외부에 노출하지 않고, 화면에 필요한 응답 형태만 전달하기 위해 DTO를 사용합니다. 보안, 응답 안정성, 프론트엔드 타입 안정성 측면에서 유리합니다."),
            ("Repository는 어떤 역할인가요?", "DB와 직접 통신하는 계층입니다. Service는 Repository를 통해 Entity를 조회하거나 저장하고, 그 결과를 DTO로 변환해 Controller에 전달합니다."),
            ("API는 어떤 기준으로 나누었나요?", "방문자 API, 관리자 API, 운영자 API, 분석/AI API, 실시간 스트림 API로 성격을 나누었습니다. URL prefix와 인증 조건을 통해 접근 범위를 구분합니다."),
            ("관리자 API는 어떻게 보호하나요?", "관리자 로그인과 JWT 기반 인증을 통해 보호합니다. 관리자 토큰이 필요한 API는 인증 헤더를 확인하고 권한이 없으면 접근을 막습니다."),
            ("운영자 API는 어떤 방식으로 접근하나요?", "운영자 페이지는 운영 키나 전용 접근 조건을 통해 관리됩니다. 현장 운영자가 복잡한 관리자 로그인 없이 부스 상태를 빠르게 수정할 수 있도록 고려했습니다."),
            ("백엔드에서 AI 혼잡도 예측은 어디서 처리되나요?", "AiCongestionService가 예측에 필요한 feature를 만들고, PythonCongestionModelService가 Python 추론 스크립트를 호출해 RandomForest 모델 예측 결과를 받아옵니다."),
            ("AI 예측 API는 프론트에 어떤 값을 주나요?", "예측 혼잡도, 표시용 혼잡도, 신뢰도, 모델 기반 여부, drift 상태, drift 점수, 판단 요인, 위험 점수 등을 포함해 프론트가 카드 형태로 보여줄 수 있게 합니다."),
            ("백엔드가 실패하면 프론트는 어떻게 되나요?", "프론트는 일부 fallback 데이터를 사용하거나 오류 메시지를 보여줍니다. AI 예측이 실패해도 전체 서비스가 멈추지 않도록 기본 안내와 규칙 기반 정보가 남도록 설계했습니다."),
        ],
    ),
    (
        "5. 실시간 처리와 SSE",
        [
            ("SSE는 무엇인가요?", "SSE(Server-Sent Events)는 서버가 브라우저로 이벤트를 계속 보내는 단방향 실시간 통신 방식입니다. 브라우저에서는 EventSource로 구독합니다."),
            ("왜 WebSocket이 아니라 SSE를 사용했나요?", "이 프로젝트의 실시간 요구는 대부분 서버에서 클라이언트로 상태 변경을 전달하는 단방향 구조입니다. 양방향 채팅처럼 복잡한 상호작용이 핵심이 아니므로 SSE가 더 단순하고 구현 비용이 낮습니다."),
            ("SSE는 어떤 데이터에 사용되나요?", "부스, 공연, 공지, 혼잡도, 스태프, 분실물, 예약 등 변경 이벤트를 브라우저에 전달하는 데 사용됩니다."),
            ("SSE가 주기적으로 계속 발동되는 구조인가요?", "SSE 연결 자체는 계속 유지되지만, 실제 이벤트는 서버에서 데이터 변경이나 방송 로직이 발생할 때 전달됩니다. 클라이언트는 연결을 유지하다가 이벤트가 오면 state를 갱신합니다."),
            ("예약이나 부스 정보가 바뀌면 어떻게 반영되나요?", "관리자나 운영자가 데이터를 수정하면 백엔드 Service가 변경을 처리하고, StreamService가 해당 이벤트를 구독 중인 클라이언트에게 방송합니다."),
            ("SSE 연결이 많아지면 문제는 없나요?", "동시 접속자가 매우 많아지면 서버 메모리와 연결 관리 부담이 생길 수 있습니다. 현재 프로젝트 규모에서는 SSE가 적절하지만, 실제 대규모 운영에서는 Redis Pub/Sub, 메시지 브로커, 로드밸런서 sticky session 등을 고려해야 합니다."),
            ("SSE 연결이 끊기면 어떻게 되나요?", "브라우저 EventSource는 기본적으로 재연결을 시도합니다. 다만 프론트는 연결 실패 시에도 기존 데이터나 fallback 데이터로 화면을 유지하도록 구성할 수 있습니다."),
            ("SSE와 HTTP API는 어떤 관계인가요?", "초기 데이터는 HTTP API로 가져오고, 이후 변경분은 SSE로 받는 구조입니다. 즉 HTTP는 초기 조회, SSE는 실시간 갱신에 가깝습니다."),
            ("실시간 갱신이 모든 페이지에 필요한가요?", "아닙니다. 홈 화면처럼 부스/공연 변경이 중요한 페이지와 운영 화면처럼 즉시 반영이 필요한 곳에 우선 적용하는 것이 효율적입니다."),
            ("SSE를 WebSocket으로 바꿔야 할 가능성은 있나요?", "실시간 채팅, 사용자별 양방향 제어, 매우 빈번한 상호작용이 핵심이 되면 WebSocket이 더 적합할 수 있습니다. 현재 기능 범위에서는 SSE로 충분하다고 판단했습니다."),
        ],
    ),
    (
        "6. AI 혼잡도 예측",
        [
            ("현재 실제로 들어간 AI 기능은 무엇인가요?", "현재 실제 AI 기능은 RandomForest 기반 혼잡도 예측입니다. 특정 시점의 부스/구역 상태값을 feature로 넣어 30분 뒤 혼잡도 등급을 예측하는 tabular ML 구조입니다."),
            ("이 기능은 시계열 모델인가요?", "아닙니다. Prophet, LSTM, Transformer처럼 시간 순서 데이터를 길게 학습하는 시계열 모델은 아닙니다. 현재는 특정 시점의 상태값과 최근 변화량 feature를 이용하는 tabular ML입니다."),
            ("tabular ML이란 무엇인가요?", "엑셀 표처럼 행과 열로 정리된 데이터를 학습하는 머신러닝입니다. 시간대, 구역 유형, GPS 추정 인원, 대기 시간, 예약 수, 체크인 수, 최근 변화량 같은 컬럼을 보고 혼잡도 등급을 분류합니다."),
            ("30분 뒤 혼잡도는 어떻게 예측하나요?", "현재 시점의 feature를 만들고, 학습된 RandomForest 모델에 넣습니다. 모델은 여러 결정트리의 판단을 종합해 30분 뒤 혼잡도 등급과 confidence를 반환합니다."),
            ("왜 RandomForest를 선택했나요?", "데이터 규모가 크지 않고, feature가 표 형태이기 때문에 RandomForest가 안정적으로 동작합니다. 또한 feature 중요도와 예측 확률을 설명하기 쉬워 발표와 운영 관점에서 적합합니다."),
            ("XGBoost도 학습했는데 왜 운영 모델은 RandomForest인가요?", "XGBoost의 성능이 조금 더 높게 나올 수 있지만, RandomForest는 설치/운영 부담이 상대적으로 낮고 설명이 쉽습니다. 현재 프로젝트에서는 안정성과 설명 가능성을 우선해 RandomForest를 운영 연결 기준으로 두었습니다."),
            ("모델 성능은 어느 정도인가요?", "현재 비교 기준으로 규칙 기반 accuracy는 약 0.7270, RandomForest는 약 0.7984, XGBoost는 약 0.8143입니다. 발표에서는 RandomForest가 규칙 기반보다 개선됐다는 점을 강조하면 됩니다."),
            ("AI 모델이 예측하는 값은 정확히 무엇인가요?", "부스 또는 구역의 30분 뒤 혼잡도 등급입니다. 예를 들어 여유, 보통, 혼잡, 매우 혼잡 같은 분류 결과와 함께 confidence, drift 상태, 위험 점수를 제공합니다."),
            ("AI 위험 점수는 무엇인가요?", "예측 혼잡도, 현재 대기 시간, GPS 추정 인원, 최근 변화량 등을 종합해 운영상 주의가 필요한 정도를 점수화한 값입니다. 사용자에게는 방문 추천/회피 판단을 돕고 운영자에게는 대응 우선순위를 줄 수 있습니다."),
            ("drift는 무엇인가요?", "현재 입력 데이터가 학습 데이터 분포에서 얼마나 벗어났는지 보는 지표입니다. 예를 들어 학습 데이터보다 훨씬 큰 인원이나 대기 시간이 들어오면 모델 예측을 조심해야 하므로 drift 경고를 표시합니다."),
            ("AI 예측 결과가 틀리면 어떻게 하나요?", "혼잡도 예측은 의사결정 보조 정보로 제공해야 합니다. 그래서 신뢰도와 drift 상태를 함께 보여주고, 운영자는 실제 현장 상황과 함께 판단하도록 설계하는 것이 맞습니다."),
            ("AI가 실제로 서버에서 동작하나요?", "현재 구조는 백엔드가 Python 추론 스크립트를 호출하고, 해당 스크립트가 저장된 RandomForest .pkl 모델을 로드해 예측합니다. 프론트는 그 결과를 API로 받아 표시합니다."),
            ("모델 파일은 어디에 있나요?", "RandomForest 모델은 exports/ml/models/random_forest_congestion_model.pkl에 저장됩니다. drift 감지용 학습 분포 정보는 congestion_training_profile.json에 저장됩니다."),
            ("Java가 모델을 직접 읽어 추론하나요?", "현재 브랜치 기준으로는 아닙니다. 현재는 Python 스크립트가 .pkl 모델을 로드하는 구조입니다. Java JSON 직접 추론은 향후 배포 안정화 방향으로 고려할 수 있지만 현재 구현 범위는 아닙니다."),
            ("AI 혼잡도 기능은 홈페이지에서 바로 보이나요?", "홈 화면에는 AI 축제 가이드와 혼잡도 요약이 보이고, 구체적인 30분 뒤 AI 혼잡도 예측 카드는 주로 /analytics 분석 화면에서 확인할 수 있습니다."),
        ],
    ),
    (
        "7. 데이터셋과 모델 학습",
        [
            ("학습 데이터는 실제 축제 데이터인가요?", "완전한 실제 운영 로그는 아닙니다. 사용자가 제공한 축제 운영 경험 가정과 현재 앱의 데이터 구조를 바탕으로 만든 HYBRID_SIMULATED 데이터입니다."),
            ("시뮬레이션 데이터로 학습한 모델도 의미가 있나요?", "실제 운영 성능을 보장한다고 말할 수는 없습니다. 다만 프로젝트 단계에서는 운영 가정을 feature로 정리하고, 규칙 기반보다 나은 예측 모델 흐름을 구현했다는 점에서 의미가 있습니다."),
            ("데이터셋 규모는 어느 정도인가요?", "현재 혼잡도 학습 데이터셋은 약 2,520행, 28개 속성입니다. 발표에서는 소규모 프로토타입 데이터셋이라고 설명하는 것이 정확합니다."),
            ("어떤 feature를 사용하나요?", "시간대, 피크 시간 여부, 구역 유형, 부스 ID, 공연 인기도, 무대 예상 인원, 무대 부하율, 야간 부스 여부, 예약 수, 체크인 수, 잔여좌석, 대기시간, GPS 추정 인원, 최근 변화량 등을 사용합니다."),
            ("노천극장 4000명 기준은 어떻게 반영됐나요?", "무대 수용량 feature를 4000명 기준으로 두고, 공연 인기도와 시간대에 따라 무대 예상 인원과 부하율이 달라지도록 데이터셋을 구성했습니다."),
            ("데이터 라벨은 어떻게 만들었나요?", "현재 feature 조합을 바탕으로 30분 뒤 혼잡도 등급을 생성했습니다. 실제 운영 로그 기반 정답이 아니라 운영 가정 기반 라벨이므로 한계를 명확히 설명해야 합니다."),
            ("실제 운영 데이터가 쌓이면 무엇이 달라지나요?", "실제 예약, 체크인, GPS, 매출, 대기열 로그가 쌓이면 모델을 재학습해 예측 신뢰도를 높일 수 있습니다. 이때 시뮬레이션 데이터에서 실제 데이터 기반 모델로 발전할 수 있습니다."),
            ("GPS 데이터가 없으면 예측이 불가능한가요?", "불가능하지는 않습니다. 예약 수, 체크인 수, 대기시간, 주문량, 부스 판매량, 공연 일정, 앱 웨이팅 로그 같은 간접 지표로도 혼잡도를 예측할 수 있습니다."),
            ("데이터 편향 문제는 없나요?", "있을 수 있습니다. 시뮬레이션 데이터는 우리가 설정한 가정에 영향을 받기 때문에 실제 축제와 다를 수 있습니다. 그래서 drift 표시와 실제 운영 데이터 기반 재학습이 중요합니다."),
            ("모델을 얼마나 자주 재학습해야 하나요?", "현재는 오프라인 학습 모델입니다. 실제 운영에서는 축제 전날/당일 운영 로그를 반영해 주기적으로 재학습하거나, 운영 종료 후 다음 축제 모델을 개선하는 방식이 현실적입니다."),
            ("학습 데이터가 적은데 과적합은 없나요?", "가능성이 있습니다. 그래서 발표에서는 성능 수치를 절대적인 정확도로 말하기보다, 규칙 기반 대비 개선된 프로토타입 결과라고 설명하는 것이 안전합니다."),
            ("성능 지표는 accuracy만 보면 되나요?", "아닙니다. 혼잡도 등급은 클래스 불균형이 있을 수 있으므로 macro F1도 함께 보는 것이 좋습니다. 다만 발표에서는 accuracy와 규칙 기반 대비 개선을 중심으로 설명해도 충분합니다."),
        ],
    ),
    (
        "8. 배포와 운영 안정성",
        [
            ("배포 환경에서 AI 모델은 어떻게 실행되나요?", "Spring Boot 서버가 Python 명령을 실행해 scripts/ml/predict_congestion.py를 호출하고, Python 스크립트가 random_forest_congestion_model.pkl을 로드해 예측합니다."),
            ("배포 시 필요한 파일은 무엇인가요?", "predict_congestion.py, random_forest_congestion_model.pkl, congestion_training_profile.json, 그리고 필요한 Python 패키지가 필요합니다."),
            ("배포 시 필요한 환경변수는 무엇인가요?", "APP_ML_CONGESTION_ENABLED, APP_ML_PYTHON_COMMAND, APP_ML_CONGESTION_PREDICT_SCRIPT, APP_ML_CONGESTION_MODEL_PATH, APP_ML_CONGESTION_TIMEOUT_MS가 핵심입니다."),
            ("Python 의존성 문제란 무엇인가요?", "서버에 Python이 없거나, scikit-learn/joblib/pandas 같은 패키지가 설치되지 않았거나, 모델 파일 경로가 다르면 추론 스크립트 실행이 실패할 수 있다는 뜻입니다."),
            ("AI 모델 실행이 실패하면 서비스가 멈추나요?", "멈추지 않도록 설계해야 합니다. 현재 구조에서는 모델 예측이 실패하면 AI 모델 결과는 비거나 fallback으로 표시되고, 기본 혼잡도/규칙 기반 정보는 계속 보여줄 수 있습니다."),
            ("배포 환경에서 fallback만 뜨면 무엇을 확인해야 하나요?", "Python 실행 경로, 모델 파일 경로, predict_congestion.py 경로, Python 패키지 설치 여부, 환경변수 값, 서버 로그의 추론 실패 메시지를 확인해야 합니다."),
            ("AI 기능을 끌 수 있나요?", "APP_ML_CONGESTION_ENABLED 값을 통해 모델 추론 사용 여부를 제어할 수 있습니다. 문제가 생기면 AI 예측을 끄고 기본 기능을 유지하는 운영 전략이 가능합니다."),
            ("Python을 서버에서 실행하는 방식의 단점은 무엇인가요?", "Java 서버와 Python 런타임을 함께 관리해야 하므로 배포 복잡도가 올라갑니다. 장기적으로는 모델 서버 분리, ONNX 변환, Java 직접 추론, 또는 별도 ML API 서버를 고려할 수 있습니다."),
            ("현재 배포 구조의 장점은 무엇인가요?", "기존 Spring Boot 서버를 크게 바꾸지 않고 Python ML 모델을 붙일 수 있습니다. 프로토타입 단계에서 빠르게 실제 API와 연결하기 좋은 구조입니다."),
            ("실제 운영 규모가 커지면 어떻게 개선해야 하나요?", "AI 추론을 별도 서비스로 분리하고, 캐싱, 큐, 모니터링, 로드밸런싱을 추가하는 것이 좋습니다. SSE도 메시지 브로커 기반으로 확장할 수 있습니다."),
        ],
    ),
    (
        "9. 보안, 개인정보, 권한",
        [
            ("관리자와 일반 사용자는 어떻게 구분하나요?", "관리자 API는 로그인과 JWT 기반 인증으로 보호하고, 일반 방문자 API는 공개 조회 중심으로 구성합니다. 기능별로 접근 권한을 다르게 둡니다."),
            ("개인정보는 어떤 부분에서 발생하나요?", "예약, AI Match, SMS 인증, 연락처 관리 등에서 개인정보가 발생할 수 있습니다. 특히 전화번호와 프로필 정보는 최소 수집과 접근 제한이 필요합니다."),
            ("개인정보 보호를 위해 무엇을 고려했나요?", "관리자/운영자 권한 분리, 필요한 정보만 응답하는 DTO 사용, 인증이 필요한 API 분리, 삭제/정정 기능 등을 고려했습니다."),
            ("AI Match의 사진 정보는 어떻게 조심해야 하나요?", "사진은 민감한 개인정보가 될 수 있으므로 저장 위치, 접근 권한, 삭제 기능, 외부 AI API 전송 여부를 명확히 관리해야 합니다."),
            ("운영자 키가 유출되면 어떻게 하나요?", "운영자 키는 환경변수로 관리하고, 유출 시 즉시 교체할 수 있어야 합니다. 실제 운영에서는 권한별 계정 관리와 로그 추적을 강화하는 것이 좋습니다."),
            ("관리자 API를 프론트에서 숨기면 안전한가요?", "아닙니다. 프론트에서 숨기는 것은 UX일 뿐 보안이 아닙니다. 반드시 백엔드에서 인증과 권한 검사를 해야 합니다."),
            ("AI 답변에서 잘못된 안내가 나올 위험은 없나요?", "있습니다. 그래서 AI 답변은 운영 데이터 기반의 보조 안내로 제공하고, 중요한 공지나 안전 관련 내용은 관리자 공지를 우선해야 합니다."),
            ("로그에는 무엇을 남겨야 하나요?", "관리자 변경, 운영자 상태 변경, AI 판단 근거, 공지 생성, 예약 상태 변경 같은 주요 행위는 감사 로그로 남기면 운영 추적에 도움이 됩니다."),
        ],
    ),
    (
        "10. 한계와 향후 확장성",
        [
            ("현재 프로젝트의 가장 큰 한계는 무엇인가요?", "실제 운영 로그가 충분하지 않다는 점입니다. 기능 구조는 구현했지만, AI 예측 성능은 실제 축제 데이터로 검증해야 합니다."),
            ("시계열 모델로 확장한다면 어떻게 하나요?", "시간 순서대로 쌓이는 GPS, 예약, 대기열, 공연 시작/종료 데이터를 이용해 Prophet, LSTM, Temporal Fusion Transformer 같은 모델로 미래 혼잡도 흐름을 예측할 수 있습니다."),
            ("공간-시간 모델로 확장할 수 있나요?", "가능합니다. 부스와 공연장을 노드로 보고 이동 경로를 그래프로 구성하면, 공연 종료 후 푸드존/주점으로 이동하는 흐름을 Spatio-Temporal GNN으로 모델링할 수 있습니다."),
            ("온라인 러닝은 필요한가요?", "축제 당일 데이터가 계속 들어오므로 장기적으로는 주기적 재학습이나 온라인 러닝이 의미 있습니다. 다만 현재 프로젝트에서는 구현 범위가 크기 때문에 향후 확장 주제로 두는 것이 현실적입니다."),
            ("강화학습을 적용할 수 있나요?", "운영 정책 추천에는 가능성이 있습니다. 예를 들어 공지 발송, 스태프 배치, 대기열 분산 전략을 action으로 보고 혼잡 완화 효과를 reward로 둘 수 있습니다. 다만 실제 운영 실험과 안전 장치가 필요합니다."),
            ("AI를 더 고도화한다면 가장 현실적인 다음 단계는 무엇인가요?", "가장 현실적인 단계는 실제 로그 수집 후 재학습, drift 모니터링 강화, 시계열 feature 확대입니다. 바로 LSTM/GNN으로 가기보다 데이터 품질을 먼저 높이는 것이 중요합니다."),
            ("대규모 축제에도 적용할 수 있나요?", "기본 구조는 확장 가능하지만, 대규모 적용에는 서버 확장, 실시간 메시징 구조, 위치 데이터 정확도, 개인정보 보호, 운영자 권한 관리가 추가로 필요합니다."),
            ("모바일 앱으로 확장할 수 있나요?", "현재는 웹앱이지만 PWA 또는 React Native 앱으로 확장할 수 있습니다. 다만 축제 현장에서는 설치 부담이 적은 모바일 웹/PWA가 현실적입니다."),
            ("다른 축제에도 재사용할 수 있나요?", "부스, 공연, 지도, 예약, 공지, 혼잡도라는 구조는 다른 축제에도 적용 가능합니다. 다만 장소 구조, 수용 인원, 운영 방식, 데이터 수집 방식은 축제마다 설정해야 합니다."),
            ("실제 상용화하려면 무엇이 더 필요하나요?", "실제 운영 데이터, 안정적인 배포/모니터링, 개인정보 처리 정책, 관리자 교육, 장애 대응 매뉴얼, 현장 네트워크 환경 검증이 필요합니다."),
        ],
    ),
    (
        "11. 날카로운 질문 대비",
        [
            ("이게 진짜 AI라고 할 수 있나요?", "네, 현재는 학습된 RandomForest 모델을 사용해 feature 기반 예측을 수행하므로 AI/머신러닝 기능이라고 볼 수 있습니다. 다만 LLM이나 딥러닝 수준의 고도 AI는 아니며, tabular ML 기반 프로토타입이라고 정확히 설명하는 것이 맞습니다."),
            ("데이터가 실제가 아닌데 의미가 있나요?", "실제 운영 성능을 보장한다는 의미는 아닙니다. 대신 운영 가정을 데이터 구조로 만들고, 규칙 기반과 머신러닝 모델을 비교해 실제 서비스에 AI를 연결하는 흐름을 구현했다는 데 의미가 있습니다."),
            ("왜 성능이 가장 좋은 XGBoost를 쓰지 않았나요?", "XGBoost는 비교 실험에서 좋은 성능을 보였지만, 현재 운영 연결은 설명 가능성과 안정성을 우선해 RandomForest를 선택했습니다. 발표에서는 XGBoost는 향후 비교/개선 후보로 유지한다고 답하면 됩니다."),
            ("모델 성능 0.7984가 충분한가요?", "실제 서비스 정확도를 보장하기에는 부족할 수 있습니다. 하지만 규칙 기반 0.7270보다 개선됐고, 프로토타입 단계에서는 AI 예측 흐름을 검증하는 기준으로 충분합니다."),
            ("혼잡도 예측이 틀리면 사용자가 손해를 보지 않나요?", "그래서 예측을 절대적인 지시가 아니라 참고 정보로 제공해야 합니다. 신뢰도와 drift 상태를 같이 보여주고, 최종 안내는 운영자 공지와 현장 정보를 우선하도록 설계합니다."),
            ("SSE가 많아지면 서버가 버티나요?", "현재 프로젝트 규모에서는 충분하지만, 대규모 운영에서는 연결 수 관리가 필요합니다. Redis Pub/Sub, 메시지 브로커, 서버 스케일아웃, 로드밸런서 설정을 추가해야 합니다."),
            ("WebSocket이 더 좋아 보이는데 왜 안 썼나요?", "현재 요구는 서버가 변경 정보를 클라이언트에 보내는 단방향 실시간 갱신이 중심입니다. WebSocket은 양방향성이 필요한 경우에 더 적합하고, 현재는 SSE가 단순하고 충분합니다."),
            ("OpenAI API가 없으면 AI 기능이 다 죽나요?", "아닙니다. OpenAI 기반 AI 가이드는 fallback 답변이 가능하고, 혼잡도 ML은 별도의 RandomForest 모델 구조입니다. 다만 각 기능별로 필요한 환경변수가 다르므로 운영 설정은 확인해야 합니다."),
            ("현재 AI는 시계열이 아니라면 미래 예측이라고 말해도 되나요?", "30분 뒤 상태를 예측하므로 미래 예측이라고 말할 수는 있습니다. 다만 시계열 모델은 아니고, 현재 시점 feature와 최근 변화량을 이용한 tabular ML 예측이라고 정확히 설명해야 합니다."),
            ("프로젝트 범위가 너무 넓은 것 아닌가요?", "축제 운영 서비스라는 도메인을 기준으로 방문자, 관리자, 운영자 기능을 묶었기 때문에 범위가 넓어 보입니다. 하지만 핵심 흐름은 부스/공연/혼잡도/예약/실시간 반영으로 연결됩니다."),
            ("AI Match는 핵심 서비스와 너무 동떨어진 것 아닌가요?", "AI Match는 혼잡도 예측 AI와 별개의 축제 참여형 부가 기능입니다. 핵심 운영 기능은 FestFlow 본체이고, AI Match는 축제 경험을 확장하는 이벤트성 기능으로 설명하면 됩니다."),
            ("실제 GPS를 쓰면 개인정보 문제가 생기지 않나요?", "맞습니다. 실제 적용 시에는 개인 식별이 불가능한 집계 데이터로 처리하고, 위치 권한 동의와 보관 기간, 익명화 정책이 필요합니다. 현재 프로젝트에서는 GPS 추정/집계 관점으로 설명하는 것이 안전합니다."),
            ("이 프로젝트에서 가장 기술적으로 보여줄 만한 부분은 무엇인가요?", "React/Spring Boot 분리 구조, SSE 실시간 갱신, 운영자/관리자 권한 분리, RandomForest 기반 AI 혼잡도 예측, drift 표시, API와 프론트 화면 연결을 함께 보여줄 수 있다는 점입니다."),
            ("지금 당장 실제 축제에 적용할 수 있나요?", "데모 수준의 적용은 가능하지만, 실제 운영에는 현장 네트워크, 데이터 정확도, 관리자 교육, 장애 대응, 개인정보 정책, 실제 로그 기반 모델 검증이 필요합니다."),
            ("가장 먼저 개선해야 할 부분은 무엇인가요?", "AI 측면에서는 실제 운영 로그 수집과 재학습이 가장 중요합니다. 운영 측면에서는 배포 안정성, 환경변수 정리, 관리자/운영자 UX 검증이 우선입니다."),
        ],
    ),
    (
        "12. 짧은 답변 암기용",
        [
            ("현재 AI를 한 문장으로 설명하면?", "현재 AI는 특정 시점의 축제 운영 데이터를 feature로 넣어 30분 뒤 혼잡도를 예측하는 RandomForest 기반 tabular ML 모델입니다."),
            ("시계열 모델인가요?", "아직은 아닙니다. 최근 변화량 feature는 쓰지만, LSTM이나 Prophet처럼 시간 순서 전체를 학습하는 구조는 아닙니다."),
            ("데이터는 실제인가요?", "완전한 실제 로그는 아니고, 운영 가정 기반 HYBRID_SIMULATED 데이터입니다. 실제 적용 전에는 운영 로그로 검증과 재학습이 필요합니다."),
            ("왜 RandomForest인가요?", "작은 표 형태 데이터에서 안정적이고 설명이 쉬워서 현재 운영 연결 모델로 선택했습니다."),
            ("XGBoost는 왜 있나요?", "비교 실험용입니다. 성능 비교와 향후 개선 후보로 유지합니다."),
            ("SSE를 쓴 이유는?", "현재 실시간 요구가 서버에서 클라이언트로 상태 변경을 보내는 단방향 구조라 WebSocket보다 단순하고 적합하기 때문입니다."),
            ("배포에서 AI가 실패하면?", "Python 실행, 모델 경로, 패키지 문제가 있으면 AI 모델 결과는 fallback되고 기본 기능은 유지되도록 설계합니다."),
            ("가장 큰 한계는?", "실제 축제 운영 로그가 부족하다는 점입니다. 그래서 프로토타입 AI 예측 모델이라고 설명하는 것이 정확합니다."),
            ("향후 가장 좋은 확장은?", "실제 로그 수집 후 재학습, 시계열 feature 확대, 이후 Prophet/LSTM/GNN 같은 모델로 확장하는 것입니다."),
            ("이 프로젝트의 핵심 차별점은?", "방문자 화면, 운영자 화면, 관리자 기능, 실시간 데이터, AI 혼잡도 예측을 하나의 축제 운영 흐름으로 연결했다는 점입니다."),
        ],
    ),
]


def add_qa(doc, qno, question, answer):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2, line=1.2)
    add_run(p, f"Q{qno:02d}. {question}", bold=True, color=DARK)

    a = doc.add_paragraph()
    set_spacing(a, before=0, after=8, line=1.25)
    add_run(a, f"답변: {answer}")


def add_sections(doc):
    qno = 1
    for section_title, qas in SECTIONS:
        doc.add_heading(section_title, level=1)
        for question, answer in qas:
            add_qa(doc, qno, question, answer)
            qno += 1


def add_summary(doc):
    doc.add_heading("발표 답변 원칙", level=1)
    add_callout(
        doc,
        "가장 중요한 태도",
        "현재 구현 범위를 과장하지 않고, 구현된 것은 명확히 말하고, 미구현 영역은 향후 확장으로 분리해서 답변합니다. 특히 AI는 '시계열 딥러닝'이 아니라 'RandomForest 기반 tabular ML 프로토타입'이라고 설명하는 것이 정확합니다.",
        NOTE_FILL,
    )
    add_table(
        doc,
        ["상황", "답변 방향"],
        [
            ["기획 질문", "축제 현장의 정보 분산, 대기열, 운영자 대응 문제를 해결하기 위한 서비스라고 설명"],
            ["기술 질문", "React/Vite, Spring Boot, DTO/Service/Repository, SSE 구조를 기능 흐름과 연결해 설명"],
            ["AI 질문", "RandomForest 기반 30분 뒤 혼잡도 분류 모델이며 시계열 모델은 아니라고 명확히 설명"],
            ["데이터 질문", "실제 로그가 아닌 운영 가정 기반 시뮬레이션 데이터라는 한계를 인정"],
            ["확장 질문", "실제 로그 수집, 재학습, 시계열 모델, GNN, 운영 최적화 순서로 확장 가능하다고 설명"],
        ],
        [1900, 7460],
        HEADER_FILL,
    )


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(footer, "페스트플로우 발표 질의응답 대비 문서", color="777777", size=9)


def build():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_summary(doc)
    add_sections(doc)
    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"written: {OUTPUT}")


if __name__ == "__main__":
    build()
