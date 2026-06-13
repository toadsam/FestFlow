from __future__ import annotations

from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "exports" / "ml" / "페스트플로우_현재_인공지능_구현_전체_설명서.docx"
FALLBACK_OUTPUT = ROOT / "exports" / "ml" / "페스트플로우_현재_인공지능_구현_전체_설명서_상세본.docx"
DATASET = ROOT / "exports" / "ml" / "congestion_training_dataset.csv"
MODEL_COMPARISON = ROOT / "exports" / "ml" / "model_comparison.csv"
RF_IMPORTANCE = ROOT / "exports" / "ml" / "feature_importance_random_forest.csv"
XGB_IMPORTANCE = ROOT / "exports" / "ml" / "feature_importance_xgboost.csv"

FONT = "Malgun Gothic"
BLUE = "1F4D78"
DARK = "0B2545"
MUTED = "555555"
HEADER_FILL = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "ECFDF5"
YELLOW = "FFFBEB"
RED = "FEF2F2"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def dataset_summary() -> dict:
    rows = read_csv_rows(DATASET)
    if not rows:
        return {"row_count": 0, "columns": [], "target_counts": {}, "rule_counts": {}}
    columns = list(rows[0].keys())
    target_counts: dict[str, int] = {}
    rule_counts: dict[str, int] = {}
    for row in rows:
        target_counts[row.get("target_congestion", "")] = target_counts.get(row.get("target_congestion", ""), 0) + 1
        rule_counts[row.get("rule_based_level", "")] = rule_counts.get(row.get("rule_based_level", ""), 0) + 1
    return {
        "row_count": len(rows),
        "columns": columns,
        "target_counts": target_counts,
        "rule_counts": rule_counts,
    }


def set_font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, size=None, bold=False, color=None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_para(doc, text, before=0, after=6, line=1.25, bold=False, color=None):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=after, line=line)
    add_run(p, text, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4, line=1.2)
    p.clear()
    add_run(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=4, line=1.2)
    p.clear()
    add_run(p, text)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, header_fill=LIGHT_GRAY):
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_font(run, size=9, bold=row_idx == 0, color=BLUE if row_idx == 0 else None)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    style_table(table, widths, header_fill)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    p = cell.paragraphs[0]
    add_run(p, f"{title}\n", bold=True, color=BLUE)
    add_run(p, body)
    set_spacing(p, after=0, line=1.2)
    style_table(table, [9360], fill)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc, summary):
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)

    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "현재 인공지능 구현 전체 설명서", bold=True, color=DARK, size=24)

    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=14, line=1.2)
    add_run(
        subtitle,
        "데이터셋, feature 설계, RandomForest/XGBoost 학습, Python 추론, 백엔드 API 연결, 프론트 표시, fallback, drift, 한계와 발표용 설명까지 상세 정리",
        color=MUTED,
        size=11,
    )

    add_table(
        doc,
        ["항목", "현재 코드 기준 내용"],
        [
            ["문서 목적", "현재 프로젝트에 실제로 들어간 AI 기능을 기술적으로 설명"],
            ["핵심 AI", "RandomForest 기반 30분 뒤 혼잡도 분류 모델"],
            ["모델 유형", "시계열 모델이 아니라 tabular ML. 최근 변화량 feature는 사용하지만 LSTM/Prophet/TFT는 미구현"],
            ["운영 연결", "Spring Boot가 Python 스크립트 predict_congestion.py를 실행하고 .pkl 모델을 로드"],
            ["중요 주의", "현재 브랜치 기준 Java가 JSON 모델을 직접 읽어 추론하는 구조는 없음"],
            ["데이터셋 규모", f"{summary['row_count']}행, {len(summary['columns'])}개 컬럼"],
        ],
        [2300, 7060],
        HEADER_FILL,
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_scope_section(doc):
    doc.add_heading("1. 현재 실제로 구현된 AI 범위", level=1)
    add_callout(
        doc,
        "핵심 결론",
        "현재 FestFlow의 AI 혼잡도 기능은 단순 화면 문구가 아니라, 학습된 RandomForest .pkl 모델을 Python 스크립트가 로드하고, Spring Boot 백엔드가 그 결과를 API 응답에 포함하며, 프론트엔드 분석 화면이 모델명, 신뢰도, drift 상태, 위험 점수를 표시하는 구조입니다. 다만 실제 운영 로그가 아니라 운영 가정 기반 시뮬레이션 데이터로 학습한 프로토타입입니다.",
        GREEN,
    )
    add_table(
        doc,
        ["기능", "현재 상태", "상세 설명"],
        [
            ["AI 혼잡도 예측", "실제 구현", "부스/구역의 현재 상태 feature를 기반으로 30분 뒤 혼잡도를 LOW/NORMAL/BUSY/VERY_BUSY로 분류"],
            ["RandomForest 모델", "운영 연결", "exports/ml/models/random_forest_congestion_model.pkl에 저장된 scikit-learn Pipeline을 Python이 로드"],
            ["XGBoost 모델", "비교 실험", "성능 비교용으로 학습/평가하지만 현재 서버 실시간 추론의 운영 모델은 아님"],
            ["규칙 기반 fallback", "실제 구현", "모델이 실패하거나 비활성화되면 riskScore 기반 fallback 예측을 표시"],
            ["Drift 감지", "실제 구현", "현재 입력값이 학습 데이터 분포에서 벗어났는지 Python 추론 스크립트가 계산"],
            ["AI 축제 가이드", "별도 AI/추천 기능", "홈 화면과 분석 화면에서 현재 데이터 기반 추천/안내를 제공. OpenAI 기반 가이드와 fallback 안내가 함께 존재"],
            ["시계열 딥러닝", "미구현", "Prophet, LSTM, Temporal Fusion Transformer, GNN, 온라인 러닝은 현재 구현이 아니라 향후 확장 주제"],
            ["Java JSON 직접 추론", "미구현", "현재 브랜치에는 random_forest_congestion_model.json 및 Java tree traversal 구조가 없음"],
        ],
        [1900, 1500, 5960],
        HEADER_FILL,
    )
    add_para(doc, "발표에서는 'AI를 도입했다'고 말할 수 있지만, 표현은 정확해야 합니다. 가장 안전한 표현은 '운영 가정 기반 데이터를 이용해 RandomForest 분류 모델을 학습하고, 이를 백엔드 API와 프론트 분석 화면에 연결한 프로토타입 AI 예측 기능'입니다.")


def add_not_section(doc):
    doc.add_heading("2. 현재 AI가 아닌 것과 오해 방지", level=1)
    add_table(
        doc,
        ["질문", "정확한 답변"],
        [
            ["시계열 모델인가?", "아님. 시간 흐름 전체를 sequence로 학습하지 않고, 특정 시점 feature와 최근 변화량 컬럼을 사용"],
            ["딥러닝인가?", "아님. RandomForest와 XGBoost는 전통적인 tree 기반 머신러닝 모델"],
            ["LLM이 혼잡도를 예측하나?", "아님. 혼잡도 예측은 저장된 ML 모델이 수행하고, LLM/OpenAI는 가이드/챗봇 성격의 기능과 구분됨"],
            ["실제 축제 로그로 검증됐나?", "아님. 현재는 운영 가정 기반 HYBRID_SIMULATED 데이터셋으로 학습"],
            ["XGBoost가 실제 API 모델인가?", "아님. 비교 실험 모델이며 현재 운영 연결 기준은 RandomForest"],
            ["Java가 모델을 직접 읽나?", "아님. 현재는 Java 서버가 Python 프로세스를 실행하고 Python이 .pkl 모델을 읽음"],
        ],
        [2600, 6760],
        HEADER_FILL,
    )
    add_callout(
        doc,
        "방어 문장",
        "현재 모델은 고도 딥러닝이나 시계열 예측 모델은 아닙니다. 대신 축제 운영 데이터처럼 행과 열로 정리되는 tabular 데이터에 적합한 RandomForest를 사용했고, 실제 API/화면에 연결해 AI 활용 흐름을 완성했습니다.",
        YELLOW,
    )


def add_architecture_section(doc):
    doc.add_heading("3. 전체 AI 아키텍처", level=1)
    add_para(doc, "혼잡도 AI는 데이터 생성, 모델 학습, 모델 저장, 서버 추론, API 응답, 프론트 표시의 6단계로 연결됩니다. 각 단계가 분리되어 있어 모델을 재학습하거나 feature를 추가할 때 어느 파일을 수정해야 하는지 추적할 수 있습니다.")
    for step in [
        "1단계: scripts/ml/build_congestion_dataset.py가 학습용 CSV를 생성합니다.",
        "2단계: scripts/ml/train_congestion_models.py가 RandomForest/XGBoost를 학습하고 성능 CSV와 모델 파일을 저장합니다.",
        "3단계: exports/ml/models/random_forest_congestion_model.pkl에 RandomForest Pipeline, feature 목록, training profile이 저장됩니다.",
        "4단계: 백엔드 AiCongestionService가 현재 부스/예약/GPS/공연 데이터를 모아 모델 입력 feature를 만듭니다.",
        "5단계: PythonCongestionModelService가 predict_congestion.py를 별도 프로세스로 실행해 예측 결과 JSON을 받습니다.",
        "6단계: AnalyticsPage.jsx가 /ai/congestion/predictions 응답을 카드로 표시합니다.",
    ]:
        add_number(doc, step)
    add_table(
        doc,
        ["계층", "파일", "역할"],
        [
            ["데이터", "exports/ml/congestion_training_dataset.csv", "학습용 feature와 target_congestion 저장"],
            ["학습", "scripts/ml/train_congestion_models.py", "RandomForest/XGBoost 학습, 성능 비교, 모델 저장"],
            ["추론", "scripts/ml/predict_congestion.py", ".pkl 모델 로드, 예측, confidence, drift 계산"],
            ["백엔드 feature 생성", "AiCongestionService.java", "현재 운영 데이터를 모델 입력 feature로 변환"],
            ["백엔드 Python 호출", "PythonCongestionModelService.java", "Python 프로세스 실행, input/output JSON 파일로 batch 예측 처리"],
            ["DTO", "AiModelPredictionDto.java", "모델명, 예측 등급, confidence, drift, fallback 여부를 응답 구조로 표현"],
            ["프론트 API", "frontend/src/api.js", "fetchAiCongestionPredictions로 AI 예측 API 호출"],
            ["프론트 화면", "AnalyticsPage.jsx", "30분 뒤 혼잡 예측 카드, model/fallback, confidence, drift, 위험 점수 표시"],
        ],
        [1700, 3200, 4460],
        HEADER_FILL,
    )


def add_dataset_section(doc, summary):
    doc.add_heading("4. 학습 데이터셋 상세", level=1)
    add_para(doc, "학습 데이터셋은 사용자가 설명한 축제 운영 경험을 기반으로 만든 HYBRID_SIMULATED 데이터입니다. 실제 운영 로그를 그대로 수집한 것은 아니지만, 18시 이후 무대 인파 증가, 인기 가수 공연일의 무대 집중, 비피크 시간대 야간 부스/주점/푸드존 쏠림 같은 도메인 가정을 반영했습니다.")
    add_table(
        doc,
        ["항목", "값"],
        [
            ["파일", "exports/ml/congestion_training_dataset.csv"],
            ["행 수", str(summary["row_count"])],
            ["컬럼 수", str(len(summary["columns"]))],
            ["데이터 성격", "HYBRID_SIMULATED. 실제 DB 로그가 아니라 운영 가정 기반 시뮬레이션"],
            ["target", "target_congestion"],
            ["라벨", "LOW, NORMAL, BUSY, VERY_BUSY"],
            ["무대 수용량", "stage_capacity = 4000 기준"],
        ],
        [2600, 6760],
        HEADER_FILL,
    )
    target_rows = [[label, count] for label, count in sorted(summary["target_counts"].items()) if label]
    rule_rows = [[label, count] for label, count in sorted(summary["rule_counts"].items()) if label]
    add_table(doc, ["target_congestion", "행 수"], target_rows, [3800, 5560], LIGHT_GRAY)
    add_table(doc, ["rule_based_level", "행 수"], rule_rows, [3800, 5560], LIGHT_GRAY)
    add_para(doc, "target_congestion은 모델이 맞혀야 하는 정답 라벨이고, rule_based_level은 기존 규칙 기반 baseline의 판단입니다. 모델 성능 비교는 target_congestion을 기준으로 rule_based_level과 RandomForest/XGBoost 예측을 비교하는 방식입니다.")


FEATURE_DESCRIPTIONS = {
    "scenario_day": "시나리오 날짜. 실제 운영에서는 축제 날짜/요일/일차를 나타낼 수 있음",
    "hour": "현재 시간대. 18-22시 공연 피크와 야간 부스 수요를 반영",
    "is_peak_time": "피크 시간 여부. 공연 전후 또는 저녁 시간대 집중을 이진값으로 표현",
    "zone_type": "구역 유형. STAGE, FOOD, PUB, GOODS, EXPERIENCE, SAFETY 등 공간 성격",
    "booth_id": "부스 식별자. 학습 데이터에는 있지만 현재 모델 feature 목록에는 직접 사용하지 않음",
    "artist_popularity": "공연 인기도 범주. LOW/MEDIUM/HIGH 등",
    "artist_popularity_score": "공연 인기도를 숫자로 변환한 값",
    "stage_capacity": "노천극장 수용 가능 인원. 현재 기준 4000명",
    "expected_stage_crowd": "공연 시간대와 인기도에 따른 예상 무대 관람 인원",
    "stage_load_ratio": "expected_stage_crowd / stage_capacity 비율",
    "is_night_booth": "야간에 수요가 커질 가능성이 있는 부스/구역 여부",
    "event_soon": "가까운 시간 안에 공연이 시작되는지 여부",
    "minutes_to_next_event": "다음 공연까지 남은 시간",
    "gps_count_nearby": "해당 부스 주변 GPS/위치 기반 추정 인원",
    "gps_delta_5m": "최근 5분 주변 인원 변화량",
    "gps_delta_15m": "최근 15분 주변 인원 변화량",
    "reservation_count": "현재 활성 예약 수",
    "reservation_delta_15m": "최근 15분 예약 증가량",
    "checked_in_count": "체크인된 예약 수",
    "checked_in_delta_15m": "최근 15분 체크인 증가량",
    "available_seats": "예약 가능 좌석 또는 수용 여유",
    "wait_minutes": "현재 예상 대기 시간",
    "wait_delta_15m": "최근 15분 대기 시간 변화 추정",
    "remaining_stock": "남은 재고. 품절/재고 부족은 혼잡도와 운영 위험에 영향",
    "event_count_context": "현재 맥락에서 고려되는 공연/이벤트 수",
    "data_source": "데이터 생성 방식 표시. HYBRID_SIMULATED 등",
    "rule_based_level": "기존 규칙 기반 방식의 혼잡도 판단",
    "target_congestion": "모델이 학습할 정답 혼잡도 라벨",
}


def add_feature_section(doc, summary):
    doc.add_heading("5. 데이터 컬럼과 feature 의미", level=1)
    add_para(doc, "데이터셋에는 총 28개 컬럼이 있으며, 이 중 모델 입력으로 쓰는 feature와 학습/비교용 보조 컬럼이 섞여 있습니다. 현재 train_congestion_models.py의 FEATURES는 22개 numeric feature와 2개 categorical feature로 구성됩니다.")
    rows = []
    for column in summary["columns"]:
        role = "target" if column == "target_congestion" else "baseline" if column == "rule_based_level" else "metadata" if column in {"booth_id", "data_source"} else "model feature"
        rows.append([column, role, FEATURE_DESCRIPTIONS.get(column, "프로젝트 데이터셋 컬럼")])
    add_table(doc, ["컬럼", "역할", "설명"], rows, [2400, 1500, 5460], HEADER_FILL)
    add_callout(
        doc,
        "feature 설계 의도",
        "혼잡도는 한 가지 값으로 결정되지 않습니다. 시간대, 공연 인기도, 무대 수용량, GPS 추정 인원, 예약/체크인, 대기 시간, 재고, 최근 변화량을 함께 보도록 feature를 설계했습니다. 이것이 단순 if문 규칙보다 ML 모델을 붙이는 이유입니다.",
        GREEN,
    )


def add_model_section(doc):
    doc.add_heading("6. 모델 학습 구조", level=1)
    add_para(doc, "학습은 scripts/ml/train_congestion_models.py에서 수행됩니다. 같은 데이터셋을 train/test로 나누고, 기존 규칙 기반 baseline, RandomForest, XGBoost를 비교합니다.")
    add_table(
        doc,
        ["구성 요소", "현재 구현"],
        [
            ["전처리", "ColumnTransformer로 numeric feature는 passthrough, categorical feature는 OneHotEncoder(handle_unknown='ignore') 적용"],
            ["RandomForest", "n_estimators=350, max_depth=12, min_samples_leaf=3, class_weight='balanced', random_state=42"],
            ["XGBoost", "n_estimators=260, max_depth=4, learning_rate=0.06, subsample=0.9, colsample_bytree=0.9"],
            ["평가 방식", "test_size=0.25, stratify=y, accuracy와 macro_f1 산출"],
            ["저장 모델", "RandomForest Pipeline과 feature 목록, training_profile을 joblib .pkl로 저장"],
            ["비교 산출물", "model_comparison.csv, confusion_matrix_*.csv, feature_importance_*.csv, prediction_samples.csv"],
        ],
        [2600, 6760],
        HEADER_FILL,
    )
    comparison = read_csv_rows(MODEL_COMPARISON)
    add_table(
        doc,
        ["모델", "Accuracy", "Macro F1", "해석"],
        [[row["model"], row["accuracy"], row["macro_f1"], row["notes"]] for row in comparison],
        [2200, 1400, 1400, 4360],
        HEADER_FILL,
    )
    add_para(doc, "해석상 중요한 지점은 RandomForest가 규칙 기반 baseline보다 높은 성능을 보였다는 점입니다. XGBoost가 더 높은 수치를 보이지만, 현재 운영 연결은 설명 가능성과 배포 단순성을 고려해 RandomForest를 기준으로 둡니다.")


def add_importance_section(doc):
    doc.add_heading("7. feature importance 해석", level=1)
    rf_rows = read_csv_rows(RF_IMPORTANCE)[:10]
    xgb_rows = read_csv_rows(XGB_IMPORTANCE)[:10]
    if rf_rows:
        add_table(
            doc,
            ["RandomForest 상위 feature", "중요도", "발표 해석"],
            [[row["feature"], row["importance"], feature_importance_comment(row["feature"])] for row in rf_rows],
            [3200, 1500, 4660],
            HEADER_FILL,
        )
    if xgb_rows:
        add_table(
            doc,
            ["XGBoost 상위 feature", "중요도", "발표 해석"],
            [[row["feature"], row["importance"], feature_importance_comment(row["feature"])] for row in xgb_rows],
            [3200, 1500, 4660],
            HEADER_FILL,
        )
    add_para(doc, "두 모델 모두 gps_count_nearby, wait_minutes, reservation_count, is_night_booth, event_soon, zone_type 같은 운영적으로 설명 가능한 feature를 중요하게 사용합니다. 이 점은 모델이 무작위로 예측하는 것이 아니라 혼잡도와 관련 있는 운영 신호를 학습했다는 발표 근거가 됩니다.")


def feature_importance_comment(feature: str) -> str:
    if "gps_count" in feature:
        return "주변 인원 추정값은 혼잡도와 직접 연결되는 핵심 신호"
    if "wait" in feature:
        return "대기 시간은 체감 혼잡도를 가장 쉽게 반영하는 운영 지표"
    if "reservation" in feature:
        return "예약 수 증가는 향후 방문 수요를 예고하는 선행 지표"
    if "checked" in feature:
        return "체크인은 예약이 실제 방문으로 전환됐음을 의미"
    if "night" in feature:
        return "야간 부스/주점/푸드존 쏠림을 반영"
    if "stage" in feature:
        return "공연장 인파와 무대 수용량이 주변 혼잡에 영향"
    if "zone_type" in feature:
        return "구역 성격에 따라 혼잡 패턴이 다름"
    if "event" in feature:
        return "공연 임박 여부는 사람 이동을 유발"
    return "운영 맥락에서 혼잡도를 설명하는 보조 신호"


def add_prediction_section(doc):
    doc.add_heading("8. 실시간 예측 실행 흐름", level=1)
    add_para(doc, "실시간 예측은 Java 안에서 모델을 직접 실행하는 것이 아니라, 백엔드가 Python 프로세스를 실행해 추론 결과를 받는 방식입니다. 이 구조는 프로토타입 단계에서 Python ML 생태계를 그대로 사용할 수 있다는 장점이 있습니다.")
    for step in [
        "프론트가 /api/ai/congestion/predictions를 호출합니다.",
        "백엔드 AiCongestionService가 현재 FestivalSnapshot을 구성합니다.",
        "각 부스별로 GPS, 예약, 체크인, 대기 시간, 공연 맥락, 재고, 최근 변화량을 계산합니다.",
        "PythonCongestionModelService.ModelPredictionRequest 목록을 만듭니다.",
        "PythonCongestionModelService가 임시 input JSON 파일을 만들고 predict_congestion.py를 실행합니다.",
        "predict_congestion.py가 random_forest_congestion_model.pkl을 joblib.load로 읽고 batch prediction을 수행합니다.",
        "Python은 predictedLevel, confidence, driftStatus, driftScore, driftWarnings를 output JSON 파일에 씁니다.",
        "Java는 output JSON을 읽어 AiModelPredictionDto로 변환합니다.",
        "모델 결과가 있으면 해당 값을 사용하고, 없으면 RULE_FALLBACK으로 기본 예측을 표시합니다.",
    ]:
        add_number(doc, step)
    add_table(
        doc,
        ["실패 지점", "현재 처리", "화면/발표 설명"],
        [
            ["Python 명령 없음", "빈 예측 Map 반환 가능", "AI 모델 결과 대신 fallback 표시"],
            ["predict_congestion.py 경로 오류", "로그 경고 후 모델 예측 생략", "배포 환경변수 확인 필요"],
            [".pkl 모델 누락", "모델 존재 여부 확인 실패", "모델 파일이 배포 산출물에 포함되어야 함"],
            ["Python 패키지 누락", "프로세스 exitCode != 0", "scikit-learn, pandas, joblib 설치 필요"],
            ["timeout", "프로세스 강제 종료", "APP_ML_CONGESTION_TIMEOUT_MS 조정 또는 추론 최적화 필요"],
        ],
        [2200, 3000, 4160],
        HEADER_FILL,
    )


def add_backend_section(doc):
    doc.add_heading("9. 백엔드 구현 상세", level=1)
    add_para(doc, "백엔드의 핵심 클래스는 AiCongestionService와 PythonCongestionModelService입니다. AiCongestionService는 현재 운영 데이터를 모델 feature로 만드는 역할이고, PythonCongestionModelService는 실제 Python 추론 스크립트를 호출하는 역할입니다.")
    add_table(
        doc,
        ["메서드/구조", "역할"],
        [
            ["buildFestivalGuide", "추천/회피/나중 방문 추천을 만들고 AI 판단 로그 기록"],
            ["predictCongestion", "부스별 AI 혼잡도 예측 목록 생성"],
            ["analyzeBooth", "riskScore, fallbackPredictedLevel, finalPredictedLevel, reasons, aiModel을 조립"],
            ["modelPredictions", "부스별 ModelPredictionRequest 목록을 만들어 Python 추론 서비스에 전달"],
            ["temporalFeatures", "최근 5분/15분 GPS 변화량, 예약/체크인 증가량, 대기 변화량 추정"],
            ["modelFeatures", "Python 모델 입력 feature Map 구성"],
            ["modelFactors", "프론트에서 보여줄 판단 근거 문자열 구성"],
            ["AiModelPredictionDto.fallback", "모델이 없을 때 RULE_FALLBACK 응답 생성"],
        ],
        [2800, 6560],
        HEADER_FILL,
    )
    add_heading_para(doc, "riskScore 계산", "riskScore는 모델 입력과 별개로 추천/회피 판단에 쓰이는 운영 위험 점수입니다. 주변 인원, 대기 시간, 활성 예약 수, 테이블 대비 예약 비율, 예약 가능 좌석, 재고, 공연 임박 여부를 가중 합산해 0-100 범위로 제한합니다.")
    add_heading_para(doc, "fallbackPredictedLevel 계산", "모델 예측이 실패해도 화면이 비지 않도록 riskScore에 공연 임박/예약 증가 가중치를 더해 fallbackPredictedLevel을 계산합니다. 이 값은 AI 모델 자체 결과가 아니라 운영 규칙 기반 안전망입니다.")
    add_heading_para(doc, "최종 predictedLevel 결정", "모델 결과가 있으면 aiModel.displayPredictedLevel을 사용하고, 모델 결과가 없으면 fallbackPredictedLevel을 사용합니다. 따라서 프론트의 'Fallback' 표시는 모델 파일/환경 문제가 있을 때 매우 중요한 진단 신호입니다.")


def add_heading_para(doc, heading, body):
    doc.add_heading(heading, level=2)
    add_para(doc, body)


def add_frontend_section(doc):
    doc.add_heading("10. 프론트엔드 표시 상세", level=1)
    add_para(doc, "프론트엔드에서는 AnalyticsPage.jsx가 AI 혼잡도 예측을 가장 직접적으로 표시합니다. 홈 화면은 AI 축제 가이드와 전체 혼잡도 요약을 보여주고, 자세한 AI 예측 카드는 분석 화면에서 확인합니다.")
    add_table(
        doc,
        ["프론트 요소", "현재 표시 내용"],
        [
            ["AI 혼잡도 예측 헤더", "AI 혼잡도 예측 제목과 설명, 분석 페이지 히어로 영역"],
            ["AI 현재 요약", "전체 혼잡도와 추천 행동 요약"],
            ["30분 뒤 혼잡 예측 카드", "현재 혼잡도 -> 30분 뒤 예측 혼잡도"],
            ["모델 표시", "modelBased가 true면 modelType, false면 Fallback 표시"],
            ["신뢰도", "confidence가 있으면 백분율로 표시"],
            ["drift 상태", "NORMAL/CAUTION/WARNING을 한국어 상태 문구로 변환"],
            ["판단 근거", "aiModel.factors 중 일부를 작은 설명으로 표시"],
            ["AI 위험 점수", "riskScore를 점수로 표시"],
            ["fallback 카드", "aiPredictions가 없으면 구역별 기본 혼잡도 카드 표시"],
        ],
        [2800, 6560],
        HEADER_FILL,
    )
    add_para(doc, "사용자가 보는 핵심은 '현재 매우혼잡 -> 30분 뒤 혼잡' 같은 변화 문장입니다. 그 아래에 RandomForest/Fallback, 신뢰도, drift 상태, 판단 요인이 붙기 때문에 모델 기반 예측인지 아닌지 확인할 수 있습니다.")


def add_drift_section(doc):
    doc.add_heading("11. Drift 감지 상세", level=1)
    add_para(doc, "drift는 현재 입력값이 학습 데이터 분포와 얼마나 다른지를 나타냅니다. 모델이 학습하지 못한 범위의 값이 들어오면 예측을 그대로 믿기 어렵기 때문에, drift 상태를 함께 표시합니다.")
    add_table(
        doc,
        ["구성", "설명"],
        [
            ["training_profile", "학습 데이터의 numeric p05, p95, min, max와 categorical 허용값 저장"],
            ["numeric drift", "현재 값이 min/max 밖이면 큰 drift, p05/p95 밖이면 약한 drift로 계산"],
            ["categorical drift", "학습 중 보지 못한 category가 들어오면 drift 경고"],
            ["driftScore", "drift point를 검사 feature 수로 나눈 값"],
            ["NORMAL", "driftScore <= 0.15"],
            ["CAUTION", "0.15 < driftScore <= 0.35"],
            ["WARNING", "driftScore > 0.35"],
            ["driftWarnings", "최대 4개까지 구체적인 경고 문구 반환"],
        ],
        [2500, 6860],
        HEADER_FILL,
    )
    add_callout(
        doc,
        "왜 중요한가",
        "AI 예측값만 보여주면 사용자가 정답처럼 받아들일 수 있습니다. drift 상태를 같이 보여주면 현재 입력이 학습 데이터와 비슷한지, 아니면 모델이 낯선 상황을 보고 있는지 설명할 수 있습니다.",
        YELLOW,
    )


def add_deploy_section(doc):
    doc.add_heading("12. 배포와 환경변수", level=1)
    add_para(doc, "현재 AI 혼잡도 추론은 Java 애플리케이션과 Python 런타임이 함께 필요합니다. 따라서 배포에서는 모델 파일뿐 아니라 Python 실행 경로와 패키지 의존성까지 확인해야 합니다.")
    add_table(
        doc,
        ["환경변수/파일", "역할", "현재 기본값 또는 설명"],
        [
            ["APP_ML_CONGESTION_ENABLED", "AI 혼잡도 모델 사용 여부", "기본 true"],
            ["APP_ML_PYTHON_COMMAND", "Python 실행 파일 경로", "기본 ../.venv-ml/Scripts/python.exe"],
            ["APP_ML_CONGESTION_PREDICT_SCRIPT", "예측 스크립트 경로", "../scripts/ml/predict_congestion.py"],
            ["APP_ML_CONGESTION_MODEL_PATH", "RandomForest .pkl 모델 경로", "../exports/ml/models/random_forest_congestion_model.pkl"],
            ["APP_ML_CONGESTION_TIMEOUT_MS", "Python 추론 timeout", "기본 20000ms"],
            ["random_forest_congestion_model.pkl", "운영 예측 모델", "배포 산출물에 포함 필요"],
            ["congestion_training_profile.json", "drift 판단 기준", "모델 payload에도 포함되지만 별도 산출물로 관리"],
            ["requirements", "Python 패키지", "pandas, scikit-learn, joblib 필요"],
        ],
        [2600, 2600, 4160],
        HEADER_FILL,
    )
    add_callout(
        doc,
        "배포에서 Fallback만 보일 때",
        "Python 경로, predict_congestion.py 경로, .pkl 모델 경로, Python 패키지 설치 여부, 서버 로그의 exitCode/stderr를 먼저 확인해야 합니다. 현재 구조는 Python 실패 시 Java JSON 직접 추론으로 복구하지 않습니다.",
        RED,
    )


def add_limit_section(doc):
    doc.add_heading("13. 현재 한계", level=1)
    for item in [
        "학습 데이터가 실제 축제 운영 로그가 아니라 운영 가정 기반 시뮬레이션 데이터입니다.",
        "모델은 시계열 sequence를 학습하지 않습니다. 최근 5분/15분 변화량을 feature로 사용할 뿐입니다.",
        "현재 운영 연결 모델은 RandomForest이고, XGBoost는 비교 실험용입니다.",
        "Python 런타임과 패키지 의존성이 있어 배포 환경 설정이 중요합니다.",
        "실제 GPS/예약/체크인 로그가 충분하지 않으면 모델 예측 신뢰도가 제한됩니다.",
        "혼잡도 예측은 사용자 행동을 보조하는 정보이며 안전/운영 공지를 대체해서는 안 됩니다.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "발표에서의 정확한 표현",
        "현재 구현은 고도 시계열 AI가 아니라, 축제 운영 데이터를 표 형태 feature로 구성해 RandomForest 모델을 학습하고 실제 API/프론트에 연결한 프로토타입 AI 예측 기능입니다.",
        YELLOW,
    )


def add_runbook_section(doc):
    doc.add_heading("14. 실행, 재학습, 검증 방법", level=1)
    add_table(
        doc,
        ["목적", "명령/파일", "설명"],
        [
            ["데이터셋 확인", "exports/ml/congestion_training_dataset.csv", "학습 데이터 2520행, 28컬럼 확인"],
            ["모델 학습", ".venv-ml\\Scripts\\python.exe scripts\\ml\\train_congestion_models.py", "RandomForest/XGBoost 재학습 및 산출물 생성"],
            ["단일/배치 예측", ".venv-ml\\Scripts\\python.exe scripts\\ml\\predict_congestion.py", ".pkl 모델 로드 후 JSON 입력 예측"],
            ["성능 비교", "exports/ml/model_comparison.csv", "baseline/RF/XGBoost accuracy, macro_f1"],
            ["feature 중요도", "exports/ml/feature_importance_random_forest.csv", "RandomForest가 어떤 feature를 봤는지 설명"],
            ["백엔드 컴파일", "backend\\gradlew.bat compileJava", "Java 코드 컴파일 검증"],
            ["프론트 빌드", "cd frontend && npm run build", "AnalyticsPage 렌더링/타입 오류 검증"],
        ],
        [1900, 3900, 3560],
        HEADER_FILL,
    )


def add_glossary_section(doc):
    doc.add_heading("15. 용어 정리", level=1)
    add_table(
        doc,
        ["용어", "설명"],
        [
            ["Tabular ML", "행과 열로 구성된 표 형태 데이터를 학습하는 머신러닝"],
            ["RandomForest", "여러 결정트리를 학습해 다수결/확률로 분류하는 앙상블 모델"],
            ["XGBoost", "오차를 순차적으로 줄이는 gradient boosting 기반 tree 모델"],
            ["Feature", "모델 입력 컬럼. 예: hour, gps_count_nearby, wait_minutes"],
            ["Target", "모델이 맞혀야 하는 정답. 여기서는 target_congestion"],
            ["Confidence", "모델이 선택한 클래스에 부여한 최대 확률값"],
            ["Drift", "현재 입력값이 학습 데이터 분포와 달라지는 현상"],
            ["Fallback", "모델 실패 시 화면/API를 유지하기 위한 규칙 기반 대체 결과"],
            ["Macro F1", "각 클래스 F1을 평균낸 지표. 클래스 불균형 상황에서 accuracy보다 균형 잡힌 해석 가능"],
            ["HYBRID_SIMULATED", "실제 운영 로그가 아니라 운영 가정과 앱 구조를 섞어 만든 시뮬레이션 데이터"],
        ],
        [2300, 7060],
        HEADER_FILL,
    )


def add_qa_section(doc):
    doc.add_heading("16. 발표 질의응답용 핵심 답변", level=1)
    qas = [
        ("이게 진짜 AI인가요?", "네. 학습된 RandomForest 모델을 사용해 feature 기반 분류 예측을 수행하므로 머신러닝 기반 AI 기능입니다. 다만 딥러닝/시계열 AI는 아니고 tabular ML 프로토타입입니다."),
        ("데이터가 실제가 아닌데 의미가 있나요?", "실제 운영 성능을 보장하는 단계는 아닙니다. 그러나 운영 가정을 데이터 구조로 만들고, 규칙 기반 대비 ML 모델을 학습해 실제 API와 화면에 연결했다는 점에서 프로젝트 AI 활용도를 보여줍니다."),
        ("왜 XGBoost가 더 좋은데 RandomForest를 쓰나요?", "XGBoost는 비교 실험에서 좋은 성능을 보였지만, 현재 운영 연결은 설명 가능성, 안정성, 단순성을 고려해 RandomForest를 기준으로 했습니다."),
        ("시계열 예측이라고 말해도 되나요?", "그렇게 말하면 부정확합니다. 30분 뒤를 예측하지만 시계열 모델은 아니고, 특정 시점 feature와 최근 변화량을 이용한 tabular ML 예측입니다."),
        ("AI가 실패하면 서비스가 멈추나요?", "아닙니다. 모델 예측이 실패하면 RULE_FALLBACK으로 기본 예측을 표시하고, 기본 혼잡도/추천 정보는 유지됩니다."),
        ("가장 먼저 개선할 부분은 무엇인가요?", "실제 축제 운영 로그를 수집해 재학습하고, 이후 시간 흐름을 더 잘 반영하는 시계열 feature 또는 시계열 모델로 확장하는 것입니다."),
    ]
    for index, (q, a) in enumerate(qas, start=1):
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=2, line=1.2)
        add_run(p, f"Q{index}. {q}", bold=True, color=DARK)
        add_para(doc, f"답변: {a}", after=8)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(footer, "페스트플로우 현재 인공지능 구현 전체 설명서", color="777777", size=9)


def build():
    summary = dataset_summary()
    doc = Document()
    configure_styles(doc)
    add_cover(doc, summary)
    add_scope_section(doc)
    add_not_section(doc)
    add_architecture_section(doc)
    add_dataset_section(doc, summary)
    add_feature_section(doc, summary)
    add_model_section(doc)
    add_importance_section(doc)
    add_prediction_section(doc)
    add_backend_section(doc)
    add_frontend_section(doc)
    add_drift_section(doc)
    add_deploy_section(doc)
    add_limit_section(doc)
    add_runbook_section(doc)
    add_glossary_section(doc)
    add_qa_section(doc)
    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "페스트플로우 현재 인공지능 구현 전체 설명서"
    try:
        doc.save(OUTPUT)
        print(f"written: {OUTPUT}")
    except PermissionError:
        doc.save(FALLBACK_OUTPUT)
        print(f"written: {FALLBACK_OUTPUT}")


if __name__ == "__main__":
    build()
