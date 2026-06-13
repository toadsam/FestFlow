from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("exports/ml/FestFlow_ai_future_extension_report.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("FestFlow AI 고도화 확장 제안서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("현재 RandomForest 혼잡도 예측 이후의 심화 AI 확장 방향")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc: Document, label: str, text: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    paragraph.add_run(text)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].paragraphs[0].add_run(value)
            set_cell_width(cells[index], widths[index])
    doc.add_paragraph()


def add_question_box(doc: Document, question: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("교수님께 질문: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    paragraph.add_run(question)
    doc.add_paragraph()


def add_extension_section(
    doc: Document,
    title: str,
    concept: str,
    why_deep: list[str],
    required_data: list[str],
    architecture: list[str],
    difficulty: str,
    question: str,
) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(concept)

    doc.add_heading("기술적으로 깊어 보이는 이유", level=2)
    add_bullets(doc, why_deep)

    doc.add_heading("필요한 데이터", level=2)
    add_bullets(doc, required_data)

    doc.add_heading("가능한 시스템 구조", level=2)
    add_bullets(doc, architecture)

    add_callout(doc, "현실적 판단", difficulty, "FFF7ED")
    add_question_box(doc, question)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "문서 목적",
        "현재 FestFlow에는 RandomForest 모델 기반 혼잡도 예측과 drift 감지 기능이 들어가 있다. 이 문서는 당장 구현하기에는 범위가 크지만, 기술적으로 더 심화된 AI 활용으로 확장할 수 있는 방향을 정리한다.",
    )

    doc.add_heading("1. 현재 구현 상태와 한계", level=1)
    doc.add_paragraph(
        "현재 구현은 운영 경험 기반 시뮬레이션 데이터로 RandomForest 모델을 학습하고, 저장된 모델 파일을 Spring Boot 백엔드가 Python 추론 스크립트로 호출해 실제 API 응답에 AI 예측 결과를 포함하는 구조다. 또한 학습 데이터 분포와 현재 입력 데이터의 차이를 계산하는 drift 감지 기능까지 포함한다."
    )
    add_bullets(
        doc,
        [
            "현재 강점: 실제 모델 파일을 서버 예측 흐름에 연결했기 때문에 단순 규칙 기반보다 AI 기능으로 설명하기 쉽다.",
            "현재 한계: 학습 데이터가 실제 축제 운영 로그가 아니라 시뮬레이션 기반이므로 실제 운영 성능을 보장할 수 없다.",
            "현재 한계: RandomForest는 특정 시점의 feature를 보고 분류하므로 시간 흐름과 공간 이동 관계를 직접 모델링하지 않는다.",
            "현재 한계: 모델이 예측은 하지만, 운영 정책을 스스로 최적화하거나 지속 학습하는 단계는 아니다.",
        ],
    )

    doc.add_heading("2. 확장 후보 우선순위", level=1)
    add_table(
        doc,
        ["확장 방향", "핵심 아이디어", "현실성", "발표/질문 가치"],
        [
            ["시계열 예측", "시간 흐름에 따른 혼잡 변화 예측", "중간", "높음"],
            ["공간-시간 그래프 모델", "부스 간 이동 관계를 그래프로 모델링", "낮음", "매우 높음"],
            ["온라인 러닝/주기적 재학습", "축제 당일 새 데이터를 반영", "중간", "높음"],
            ["불확실성 추정", "예측값의 신뢰 구간과 위험도 제공", "중간", "높음"],
            ["강화학습 기반 운영 최적화", "공지/스태프 배치 액션을 학습", "낮음", "높음"],
        ],
        [2100, 3300, 1500, 2460],
    )

    add_extension_section(
        doc,
        "3. 시계열 예측 확장: Prophet, LSTM, Temporal Fusion Transformer",
        "현재 모델은 한 시점의 GPS 추정 인원, 예약 수, 대기 시간 등을 사용해 혼잡도를 분류한다. 그러나 실제 축제 혼잡도는 시간의 흐름에 따라 변한다. 예를 들어 18시 이후 공연 시작 전에는 무대 쪽으로 인파가 몰리고, 공연 종료 후에는 주점과 푸드존으로 이동한다. 시계열 모델은 이런 시간적 패턴을 직접 학습하는 방향이다.",
        [
            "단순 현재 상태 분류가 아니라 미래 시점의 혼잡도를 예측하는 문제로 확장된다.",
            "최근 5분, 15분, 30분의 변화량과 공연 스케줄의 영향을 함께 고려할 수 있다.",
            "Temporal Fusion Transformer는 시간대별 feature importance와 장기/단기 패턴을 함께 해석할 수 있어 설명력도 확보할 수 있다.",
        ],
        [
            "부스별 5분 또는 10분 단위 GPS 추정 인원",
            "부스별 대기 시간 변화 로그",
            "예약 등록, 체크인, 취소 이벤트의 시간 로그",
            "공연 시작/종료 시각과 공연 인기 지표",
            "날씨, 요일, 시간대 같은 외부 변수",
        ],
        [
            "현재 API 입력 feature를 시간 단위 로그 테이블로 저장한다.",
            "부스별로 최근 N개 시점의 sequence를 구성한다.",
            "초기에는 Prophet이나 LightGBM/XGBoost + lag feature로 시작한다.",
            "데이터가 충분해지면 LSTM 또는 Temporal Fusion Transformer로 확장한다.",
            "예측 결과는 '30분 뒤 혼잡도' 또는 '다음 3개 시간 구간 혼잡도' 형태로 제공한다.",
        ],
        "당장 LSTM/TFT까지 구현하기는 어렵지만, 최근 5분/15분 변화량을 feature로 추가하는 방식은 바로 적용 가능하다. 실제 딥러닝 시계열 모델은 충분한 시간 로그가 쌓인 뒤 적용하는 것이 맞다.",
        "현재는 RandomForest/XGBoost로 특정 시점의 혼잡도를 분류하고 있는데, 실제 축제 환경에서는 시간에 따른 인파 이동 패턴이 중요하다고 생각합니다. 향후 Prophet, LSTM, Temporal Fusion Transformer 같은 시계열 모델로 확장하는 것이 더 적절할까요?",
    )

    add_extension_section(
        doc,
        "4. 공간-시간 그래프 모델: GNN, Spatio-Temporal GNN",
        "축제장의 부스와 무대는 독립된 점이 아니라 이동 경로로 연결된 공간이다. 공연장 혼잡이 줄면 주변 푸드존과 주점 혼잡이 증가할 수 있고, 특정 통로가 막히면 인접 구역의 혼잡도도 같이 변한다. Graph Neural Network는 이런 공간적 연결 관계를 모델에 반영하는 방법이다.",
        [
            "부스 간 거리, 이동 경로, 인접 구역 영향을 모델 구조에 직접 포함한다.",
            "각 부스를 그래프의 노드로 보고, 길이나 이동 가능성을 edge로 정의할 수 있다.",
            "시간 축까지 결합하면 '어느 구역의 혼잡이 다음 시점에 어느 구역으로 전파되는가'를 모델링할 수 있다.",
        ],
        [
            "부스/무대/출입구의 좌표 데이터",
            "이동 가능한 길과 통로 연결 정보",
            "구역별 시간대별 혼잡도 로그",
            "공연 종료 후 이동 흐름 또는 GPS 이동 궤적",
            "통로 폭, 병목 지점, 안전 구역 같은 공간 제약 정보",
        ],
        [
            "부스와 주요 지점을 node로 구성한다.",
            "거리 또는 실제 이동 가능 경로를 edge로 구성한다.",
            "각 node에 시간별 혼잡도, 대기 시간, 예약 수, 이벤트 여부를 feature로 넣는다.",
            "ST-GCN, DCRNN, Graph WaveNet, Spatio-Temporal GNN 계열 모델을 검토한다.",
            "출력은 구역별 미래 혼잡도와 혼잡 전파 경로로 구성한다.",
        ],
        "기술적으로 가장 고급스럽지만, 현재 프로젝트 범위에서 실제 구현하기에는 데이터와 시간이 부족하다. 발표에서는 '향후 연구 방향' 또는 '실제 운영 데이터가 쌓였을 때 가능한 고도화 방향'으로 제시하는 것이 적절하다.",
        "부스 간 거리와 이동 경로를 그래프로 보고, 각 구역의 혼잡도가 이웃 구역에 영향을 준다고 가정하면 Graph Neural Network나 Spatio-Temporal GNN을 적용하는 것이 의미 있을까요?",
    )

    add_extension_section(
        doc,
        "5. 온라인 러닝과 주기적 재학습",
        "축제 당일에는 GPS 로그, 예약 로그, 웨이팅 데이터가 계속 쌓인다. 사전에 학습된 모델만 사용하면 당일 상황 변화, 인기 가수, 날씨, 돌발 이벤트를 충분히 반영하지 못할 수 있다. 온라인 러닝 또는 주기적 재학습은 새로 들어온 데이터를 모델에 반영하는 운영 구조다.",
        [
            "AI 모델을 한 번 학습하고 끝내는 것이 아니라, 운영 중 계속 개선되는 시스템으로 확장한다.",
            "실시간 데이터 파이프라인, 모델 버전 관리, 재학습 주기, 배포 안정성까지 고려해야 한다.",
            "모델 성능이 떨어지는 시점을 감지하고 새 모델로 교체하는 MLOps 관점이 포함된다.",
        ],
        [
            "실제 운영 중 수집되는 GPS 로그",
            "예약 상태 변경 이벤트 로그",
            "부스 운영자가 입력한 대기 시간과 재고 변화 로그",
            "실제 혼잡도 라벨 또는 사후 검증 데이터",
            "모델 예측 결과와 실제 결과의 오차 기록",
        ],
        [
            "초기에는 하루 또는 몇 시간 단위의 주기적 재학습을 사용한다.",
            "새 데이터가 일정량 이상 쌓이면 학습 파이프라인을 실행해 새 모델을 생성한다.",
            "기존 모델과 새 모델을 검증 데이터로 비교한 뒤 성능이 좋을 때만 교체한다.",
            "모델 파일에 버전을 붙이고, API 응답에 사용 모델 버전을 포함한다.",
            "완전한 온라인 러닝은 안정성 검증 후 별도 단계로 검토한다.",
        ],
        "현재는 수동 재학습 파이프라인 수준이 현실적이다. 자동 온라인 러닝은 오류가 생기면 서비스 품질에 직접 영향을 주므로, 실제 운영에서는 검증 단계와 rollback 전략이 필요하다.",
        "축제 당일 GPS 로그와 예약 로그가 실시간으로 쌓이면, 사전에 학습된 모델을 고정해서 쓰는 것보다 온라인 러닝이나 주기적 재학습 구조가 더 적합할까요?",
    )

    add_extension_section(
        doc,
        "6. 모델 Drift와 Concept Drift 고도화",
        "현재 구현에는 학습 데이터 분포와 실제 입력값의 차이를 계산하는 기본 drift 감지가 들어가 있다. 향후에는 단순 범위 비교를 넘어서 통계적 drift 검정, feature별 drift 원인 분석, 실제 예측 오차 기반 concept drift 감지까지 확장할 수 있다.",
        [
            "시뮬레이션 데이터 기반 모델의 가장 큰 약점인 실제 데이터와의 분포 차이를 정면으로 다룬다.",
            "data drift는 입력 데이터 분포 변화, concept drift는 입력과 정답의 관계 변화까지 다룬다.",
            "모델 성능 저하를 사전에 감지하고 재학습 또는 fallback 전략을 설계할 수 있다.",
        ],
        [
            "학습 데이터 feature 분포",
            "실시간 입력 feature 분포",
            "실제 혼잡도 사후 라벨",
            "모델 예측과 실제 결과의 오차 로그",
            "시간대/공연/부스 유형별 drift 통계",
        ],
        [
            "현재 구현된 p05/p95 기반 drift score를 유지한다.",
            "feature별 PSI(Population Stability Index) 또는 KL divergence를 추가한다.",
            "드리프트가 큰 feature를 운영자 화면에 표시한다.",
            "예측 오차가 반복적으로 커지면 concept drift로 판단한다.",
            "drift 상태가 WARNING이면 모델 예측보다 규칙 기반 fallback 또는 보수적 안내를 우선한다.",
        ],
        "이 방향은 지금 프로젝트와 가장 직접적으로 연결된다. 기본 drift 감지는 이미 구현했으므로, 발표에서는 현재 구현된 부분과 향후 고도화 가능성을 함께 설명하기 좋다.",
        "시뮬레이션 데이터로 학습한 모델을 실제 축제에 적용하면 data drift나 concept drift가 발생할 수 있는데, 이를 감지하고 보정하려면 어떤 지표와 재학습 전략을 설계하는 것이 좋을까요?",
    )

    add_extension_section(
        doc,
        "7. 불확실성 추정과 Conformal Prediction",
        "혼잡도 예측은 운영 의사결정과 연결되기 때문에 단순 예측값만으로는 부족할 수 있다. 예측값이 '혼잡'이라도 모델이 얼마나 확신하는지, 또는 예측이 틀릴 가능성이 얼마나 되는지를 함께 제공하는 것이 더 안전하다.",
        [
            "AI가 단정적으로 답하는 것이 아니라 예측의 불확실성을 함께 제공한다.",
            "운영자는 확신이 낮은 예측을 참고용으로 보고, 확신이 높은 예측에는 적극 대응할 수 있다.",
            "Conformal prediction은 모델 종류와 무관하게 예측 집합이나 보장 수준을 제공할 수 있다.",
        ],
        [
            "모델의 class probability",
            "검증 데이터에서의 예측 오차",
            "라벨별 오분류 패턴",
            "시간대/구역별 모델 정확도",
            "실제 혼잡도 라벨",
        ],
        [
            "현재 RandomForest의 probability를 confidence로 제공한다.",
            "검증 데이터 기반으로 confidence calibration을 수행한다.",
            "불확실성이 높으면 '혼잡 또는 매우 혼잡 가능성'처럼 예측 집합을 제공한다.",
            "운영 화면에는 confidence, drift score, uncertainty level을 함께 표시한다.",
            "최종적으로 conformal prediction으로 일정 신뢰수준의 예측 범위를 제공한다.",
        ],
        "현재 confidence는 이미 제공되지만, calibration이나 conformal prediction은 아직 없다. 추가하면 안전한 AI 운영 관점에서 기술적으로 더 깊어 보인다.",
        "혼잡도 예측을 운영에 활용하려면 예측값뿐 아니라 예측 불확실성도 중요하다고 생각합니다. RandomForest의 확률값만으로 충분한지, conformal prediction 같은 방식까지 고려해야 할까요?",
    )

    add_extension_section(
        doc,
        "8. 강화학습 기반 혼잡 완화 정책",
        "예측 모델은 혼잡을 '맞히는' 데 초점이 있다. 강화학습은 한 단계 더 나아가 어떤 조치를 해야 전체 혼잡이 줄어드는지 학습하는 접근이다. 예를 들어 공지 발행, 스태프 배치, 동선 우회 안내, 예약 제한 같은 액션을 선택하고 그 결과로 혼잡도가 줄어드는지를 보상으로 학습할 수 있다.",
        [
            "예측 문제를 운영 정책 최적화 문제로 확장한다.",
            "AI가 단순 분석이 아니라 action을 선택하는 의사결정 시스템이 된다.",
            "시뮬레이터가 있으면 실제 운영 전에 다양한 정책을 가상 실험할 수 있다.",
        ],
        [
            "혼잡도 변화 시뮬레이터",
            "운영 액션 로그",
            "공지 발행 후 방문자 이동 변화",
            "스태프 배치 후 대기열 변화",
            "혼잡 완화 효과를 측정할 reward 정의",
        ],
        [
            "현재 운영 시뮬레이션 기능을 강화학습 환경으로 확장한다.",
            "상태는 구역별 혼잡도, 대기 시간, 공연 시간, 예약 수로 정의한다.",
            "액션은 공지 발행, 우회 추천, 스태프 배치, 예약 제한으로 정의한다.",
            "보상은 전체 혼잡도 감소, 위험 구역 감소, 대기 시간 감소로 설계한다.",
            "초기에는 실제 강화학습보다 rule-based policy simulation으로 시작한다.",
        ],
        "가장 심화된 방향이지만, 현재 프로젝트 발표 전 구현 대상으로는 과하다. 대신 '장기 확장 방향'으로 제시하면 기술적 깊이를 보여줄 수 있다.",
        "혼잡도 예측을 넘어서 공지 발행이나 스태프 배치 같은 액션을 통해 전체 혼잡도를 줄이는 문제로 보면, 강화학습으로 접근할 수 있을까요?",
    )

    doc.add_heading("9. 발표에서 안전하게 말하는 방식", level=1)
    add_callout(
        doc,
        "추천 표현",
        "현재는 RandomForest 모델 기반 혼잡도 예측과 drift 감지까지 구현했습니다. 향후에는 시간 흐름을 직접 반영하는 시계열 모델, 부스 간 이동 관계를 반영하는 공간-시간 그래프 모델, 실제 운영 데이터가 쌓였을 때의 주기적 재학습 및 drift 보정 구조로 확장할 수 있습니다.",
        "ECFDF5",
    )
    add_bullets(
        doc,
        [
            "현재 구현: 모델 파일 기반 RandomForest 추론, confidence 제공, drift score 제공",
            "단기 확장: 최근 5분/15분 변화량 feature 추가, drift 원인 표시 강화",
            "중기 확장: 주기적 재학습, 모델 버전 관리, confidence calibration",
            "장기 확장: LSTM/TFT, Spatio-Temporal GNN, 강화학습 기반 운영 정책 최적화",
        ],
    )

    doc.add_heading("10. 최종 질문 후보", level=1)
    add_bullets(
        doc,
        [
            "현재는 특정 시점 feature 기반 RandomForest 모델을 사용하지만, 실제 축제 혼잡도는 시간 흐름의 영향을 크게 받습니다. 시계열 모델로 확장하는 것이 더 적절할까요?",
            "부스와 공연장을 그래프 구조로 보고 혼잡이 인접 구역으로 전파된다고 가정하면 Spatio-Temporal GNN을 적용하는 것이 의미 있을까요?",
            "시뮬레이션 데이터 기반 모델을 실제 운영 데이터에 적용할 때 data drift와 concept drift를 어떻게 감지하고 보정하는 것이 좋을까요?",
            "예측값뿐 아니라 불확실성이나 confidence calibration을 함께 제공하는 것이 운영 의사결정에 더 안전할까요?",
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(f"Report written to {OUTPUT_PATH}")
