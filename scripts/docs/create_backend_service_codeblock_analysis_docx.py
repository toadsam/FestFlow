from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SERVICE_DIR = ROOT / "backend" / "src" / "main" / "java" / "com" / "festflow" / "backend" / "service"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "페스트플로우_백엔드_Service_코드블록_상세해설서.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F8FAFC"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"


SERVICE_NOTES = {
    "AdminActionService": {
        "method": "publishCongestionReliefNotice",
        "role": "관리자가 혼잡 완화 공지를 자동으로 발행할 때 쓰는 서비스입니다.",
        "why": "Controller가 공지 문구를 직접 만들지 않고, 가장 혼잡한 부스 조회와 공지 생성을 서비스 계층에서 묶어 처리합니다.",
        "flow": "관리자 버튼 클릭 -> AdminActionService -> BoothService로 혼잡 부스 조회 -> NoticeService로 공지 저장 -> 프론트 공지 갱신",
        "point": "여기서는 DB 저장을 직접 하지 않고 NoticeService를 호출합니다. 이미 있는 도메인 로직을 재사용하는 구조입니다.",
    },
    "AdminDashboardService": {
        "method": "getKpis",
        "role": "관리자 대시보드에 표시할 KPI 숫자를 계산합니다.",
        "why": "방문자 수, 부스 수, 공연 수, 혼잡도 같은 여러 도메인의 데이터를 한 화면용 DTO로 묶어야 하기 때문입니다.",
        "flow": "관리자 대시보드 요청 -> GPS 로그 집계 -> 부스/공연 현황 조회 -> AdminDashboardKpiDto 반환",
        "point": "KPI 서비스는 데이터를 새로 만드는 것이 아니라 여러 서비스와 Repository에서 읽어온 값을 조합합니다.",
    },
    "AdminImportService": {
        "method": "importBoothsCsv",
        "role": "관리자가 업로드한 CSV 파일을 읽어 부스를 대량 등록합니다.",
        "why": "파일 업로드는 JSON이 아니라 MultipartFile로 들어옵니다. 서비스는 파일 내용을 읽고 한 줄씩 DTO로 변환합니다.",
        "flow": "CSV 업로드 -> MultipartFile -> BufferedReader로 읽기 -> BoothUpsertRequestDto 생성 -> BoothService.createBooth 호출",
        "point": "파일 자체를 DB에 넣는 것이 아니라 CSV 내용을 해석해서 기존 부스 생성 로직을 여러 번 호출합니다.",
    },
    "AiCongestionService": {
        "method": "analyzeCurrent",
        "role": "현재 축제 상태를 보고 AI 혼잡도 예측과 추천 부스를 만듭니다.",
        "why": "AI 모델 결과만 믿지 않고, 모델 실패 시 fallback 규칙 결과도 함께 준비해서 운영 화면이 끊기지 않게 합니다.",
        "flow": "프론트 분석 페이지 -> /ai/congestion/predictions -> AiCongestionService -> Python 모델 추론 -> 추천 DTO 반환",
        "point": "이 서비스는 AI 결과, GPS, 예약, 대기 시간, 공연 임박 여부를 한 번에 조합하는 중심 서비스입니다.",
    },
    "AiDecisionLogService": {
        "method": "record",
        "role": "AI가 왜 그런 추천을 했는지 최근 판단 기록을 남깁니다.",
        "why": "AI 결과만 보여주면 근거가 부족하므로, 제목/요약/이유를 로그로 남겨 화면에서 설명할 수 있게 합니다.",
        "flow": "AI 분석 완료 -> record 호출 -> 메모리 로그에 추가 -> /ai/decisions에서 최근 로그 조회",
        "point": "DB 로그가 아니라 메모리 기반 최근 로그입니다. 서버 재시작 시 사라질 수 있습니다.",
    },
    "AiImageGenerationService": {
        "method": "generateFestivalProfileImage",
        "role": "AI 매칭용 프로필 이미지를 웹툰풍 이미지로 생성합니다.",
        "why": "원본 이미지를 그대로 쓰지 않고, AI 이미지 생성 API와 저장 서비스를 연결해 사용자 프로필 이미지를 만듭니다.",
        "flow": "이미지 업로드 -> 원본 저장 -> OpenAI 이미지 생성 -> 생성 이미지 저장 -> URL 반환",
        "point": "외부 API 호출은 실패할 수 있으므로 예외 처리와 fallback 흐름을 반드시 같이 봐야 합니다.",
    },
    "AiMatchService": {
        "method": "createProfile",
        "role": "AI 매칭 프로필 생성, 이미지 처리, 전화번호 사용 기록 저장을 처리합니다.",
        "why": "프로필 생성은 단순 DB insert가 아니라 이미지 저장, AI 이미지 생성, 중복 확인, 개인정보 동의가 함께 묶인 업무입니다.",
        "flow": "프론트 프로필 등록 -> 파일/폼 데이터 전송 -> AiMatchService -> 이미지 저장/AI 생성 -> 프로필 Entity 저장",
        "point": "@Transactional이 붙어 있어 프로필 생성 중 DB 작업을 하나의 단위로 처리합니다. 파일 저장은 DB와 완전히 같은 트랜잭션은 아닙니다.",
    },
    "AuditLogService": {
        "method": "log",
        "role": "관리자가 어떤 작업을 했는지 로그로 저장합니다.",
        "why": "운영자가 공지, 부스, 공연을 수정했을 때 나중에 누가 무엇을 했는지 추적하기 위해 필요합니다.",
        "flow": "관리자 작업 완료 -> AuditLogService.log -> AuditLogRepository.save -> 관리자 로그 화면 조회",
        "point": "로그 서비스는 기능 자체보다 운영 추적성과 문제 분석을 위한 보조 기능입니다.",
    },
    "AnalyticsService": {
        "method": "dashboard",
        "role": "혼잡도 분석 화면의 대시보드 데이터를 만듭니다.",
        "why": "프론트가 직접 GPS 로그를 계산하지 않도록, 백엔드에서 시간대/구역별 집계를 끝낸 DTO를 내려줍니다.",
        "flow": "프론트 /analytics -> fetchAnalyticsDashboard -> AnalyticsService.dashboard -> GPS/부스 집계 -> AnalyticsDashboardDto",
        "point": "분석 서비스는 DB row를 그대로 반환하지 않고 화면에 바로 쓸 수 있는 통계 형태로 가공합니다.",
    },
    "AuthService": {
        "method": "login",
        "role": "관리자 로그인과 JWT 토큰 발급을 처리합니다.",
        "why": "관리자 API는 아무나 호출하면 안 되므로 비밀번호 검증 후 토큰을 발급합니다.",
        "flow": "로그인 요청 -> AdminUserRepository 조회 -> PasswordEncoder 검증 -> JwtService 토큰 발급",
        "point": "비밀번호는 평문 비교가 아니라 PasswordEncoder.matches로 해시 검증을 해야 합니다.",
    },
    "BoothService": {
        "method": "createBooth",
        "role": "부스 생성, 수정, 조회, 삭제, 혼잡도 계산을 담당하는 핵심 서비스입니다.",
        "why": "부스는 지도, 운영자 콘솔, 예약, 혼잡도, 이미지 업로드 등 여러 기능의 중심 데이터이기 때문입니다.",
        "flow": "관리자 부스 등록 -> BoothService.createBooth -> Booth Entity 저장 -> StreamService로 부스 목록 갱신",
        "point": "부스 변경 후에는 단순히 DB 저장으로 끝내지 않고 SSE 이벤트를 발행해 프론트를 갱신합니다.",
    },
    "ChatService": {
        "method": "answer",
        "role": "축제 안내 챗봇 답변을 생성합니다.",
        "why": "OpenAI 설정이 있으면 AI 답변을 쓰고, 실패하거나 설정이 없으면 fallback 답변으로 서비스가 계속 동작하게 합니다.",
        "flow": "사용자 질문 -> ChatService.answer -> AI API 또는 fallback -> ChatResponseDto 반환",
        "point": "AI 기능은 외부 API 의존성이 있으므로 항상 실패 대비 코드가 있는지 보는 것이 중요합니다.",
    },
    "EventService": {
        "method": "createEvent",
        "role": "공연 생성, 수정, 삭제, 목록 조회와 실시간 갱신을 담당합니다.",
        "why": "공연 일정 변경은 사용자 화면과 운영 화면에 즉시 반영되어야 하므로 StreamService와 연결됩니다.",
        "flow": "관리자 공연 등록 -> EventService -> EventRepository.save -> StreamService.publishEvents",
        "point": "공연 데이터가 바뀌면 SSE로 events 스트림에 갱신 이벤트를 보냅니다.",
    },
    "FestivalSnapshotService": {
        "method": "current",
        "role": "AI와 운영 분석에 필요한 현재 축제 상태를 한 번에 모읍니다.",
        "why": "AI 예측은 부스, 예약, GPS, 공연 데이터를 동시에 봐야 하므로 snapshot 객체로 묶어 전달합니다.",
        "flow": "AiCongestionService -> FestivalSnapshotService.current -> 여러 Repository 조회 -> FestivalSnapshot 반환",
        "point": "snapshot은 어느 한 테이블이 아니라 현재 축제 상태를 한 번에 담은 읽기 전용 묶음입니다.",
    },
    "GpsService": {
        "method": "saveGpsLog",
        "role": "사용자 위치 로그를 저장하고 혼잡도 갱신 이벤트를 발행합니다.",
        "why": "GPS 로그는 혼잡도 계산의 핵심 입력이므로 저장 직후 분석 화면이 갱신되어야 합니다.",
        "flow": "프론트 위치 전송 -> /gps -> GpsService.saveGpsLog -> DB 저장 -> StreamService.publishCongestion",
        "point": "GPS 저장은 단순 기록이 아니라 실시간 혼잡도 SSE 갱신의 트리거입니다.",
    },
    "LostItemService": {
        "method": "create",
        "role": "분실물 등록, 수정, 상태 변경, 삭제를 처리합니다.",
        "why": "분실물은 개인정보 연락처를 포함할 수 있어, 조회 상황에 따라 마스킹 처리도 같이 필요합니다.",
        "flow": "스태프 분실물 등록 -> LostItemService.create -> LostItemRepository.save -> 목록 조회",
        "point": "연락처 같은 민감 정보는 응답 DTO에서 그대로 보여줄지 마스킹할지 분리해서 봐야 합니다.",
    },
    "NoticeService": {
        "method": "createNotice",
        "role": "공지 생성, 수정, 삭제와 활성 공지 실시간 갱신을 담당합니다.",
        "why": "공지 변경은 사용자 화면에 즉시 반영되어야 해서 저장 후 SSE 발행까지 한 서비스에서 처리합니다.",
        "flow": "관리자 공지 등록 -> NoticeService.createNotice -> DB 저장 -> StreamService.publishNotices",
        "point": "공지 저장 후 broadcastActiveNotices 같은 갱신 호출이 붙어 있는지 보면 실시간 구조를 이해할 수 있습니다.",
    },
    "OpsAiService": {
        "method": "masterBriefing",
        "role": "운영자용 AI 브리핑, 공지 초안, 스태프 현장 체크리스트를 생성합니다.",
        "why": "운영자는 여러 지표를 빠르게 읽어야 하므로, 데이터를 AI 프롬프트로 요약해 운영 판단 문구를 만듭니다.",
        "flow": "운영 콘솔 요청 -> OpsAiService -> 운영 데이터 수집 -> OpenAI 호출 또는 fallback -> AiAssistResponseDto",
        "point": "운영 AI는 실제 명령 실행이 아니라, 사람이 판단할 수 있도록 문구와 체크리스트를 보조합니다.",
    },
    "PublicAiGuideService": {
        "method": "guide",
        "role": "일반 방문자에게 페이지별 AI 방문 가이드를 제공합니다.",
        "why": "analytics, events, stage-map 같은 scope에 따라 필요한 안내가 달라서 scope를 받아 분기합니다.",
        "flow": "프론트 fetchAiVisitorGuide(scope) -> PublicAiGuideService.guide -> 데이터 수집/AI 호출 -> AiVisitorGuideDto",
        "point": "scope는 같은 AI 가이드 API를 여러 화면에서 재사용하기 위한 구분값입니다.",
    },
    "PythonCongestionModelService": {
        "method": "predictBatch",
        "role": "Java 서버에서 Python RandomForest 모델 추론 스크립트를 실행합니다.",
        "why": "Java가 pkl 모델을 직접 읽기보다 Python 스크립트를 별도 프로세스로 실행하는 구조가 구현이 단순합니다.",
        "flow": "AiCongestionService -> predictBatch -> input JSON 파일 생성 -> ProcessBuilder로 Python 실행 -> output JSON 파싱",
        "point": "여기서 모델 파일 경로, Python 실행 경로, timeout 설정이 맞지 않으면 프론트에 Fallback으로 보일 수 있습니다.",
    },
    "ReservationAuthService": {
        "method": "sendCode",
        "role": "예약 사용자 전화번호 인증번호를 발송하고 인증 세션을 관리합니다.",
        "why": "예약 조회/체크인 토큰 발급 같은 기능은 전화번호 인증을 거친 사용자만 접근해야 합니다.",
        "flow": "전화번호 입력 -> sendCode -> 인증번호 생성 -> SmsSender 발송 -> 세션 저장",
        "point": "SMS 발송 구현체는 SmsSender 인터페이스 뒤에 숨겨져 있어 업체를 바꿔도 인증 서비스 구조는 유지됩니다.",
    },
    "ReservationService": {
        "method": "createReservation",
        "role": "부스 예약 생성, 체크인, 완료, 테이블 해제를 처리합니다.",
        "why": "예약은 좌석 수, 사용자 인증, 중복 예약, 상태 변경이 얽혀 있어 트랜잭션으로 관리해야 합니다.",
        "flow": "예약 요청 -> 인증 토큰 확인 -> 부스/설정/테이블 확인 -> Reservation 저장 -> SSE 발행",
        "point": "예약 상태는 REGISTERED, CHECKED_IN, COMPLETED 같은 상태 전이가 핵심입니다.",
    },
    "SimulationService": {
        "method": "tick",
        "role": "운영 시뮬레이션의 주기적 변화를 실제 부스 상태에 반영합니다.",
        "why": "실제 축제 데이터가 없어도 운영 콘솔에서 혼잡 변화와 갱신 흐름을 테스트하기 위해 필요합니다.",
        "flow": "스케줄 tick -> SimulationStateService.tick -> BoothService.updateLiveStatus -> StreamService 이벤트 발행",
        "point": "시뮬레이션은 DB 원본을 영구적으로 바꾸기보다 현재 상태를 흉내 내는 운영 테스트 기능입니다.",
    },
    "SimulationStateService": {
        "method": "tick",
        "role": "시뮬레이션의 내부 상태와 변화량을 메모리에서 계산합니다.",
        "why": "여러 요청이 동시에 상태를 바꿀 수 있어 synchronized로 동시 접근을 제어합니다.",
        "flow": "SimulationService -> SimulationStateService.tick -> 메모리 상태 계산 -> 변화량 목록 반환",
        "point": "이 서비스는 Repository보다 메모리 Map과 record를 많이 씁니다. DB 서비스와 읽는 방식이 다릅니다.",
    },
    "StaffService": {
        "method": "login",
        "role": "스태프 로그인, 세션 발급, 상태 변경을 처리합니다.",
        "why": "스태프 화면은 관리자 JWT가 아니라 스태프용 토큰으로 인증되는 별도 흐름입니다.",
        "flow": "스태프 번호/PIN 입력 -> StaffService.login -> 세션 저장 -> StaffLoginResponseDto 반환",
        "point": "StaffSessionRepository가 스태프 토큰을 관리합니다. 관리자 로그인과 다른 인증 체계입니다.",
    },
    "StreamService": {
        "method": "subscribeCongestion",
        "role": "브라우저가 SSE로 혼잡도 변경 이벤트를 구독할 수 있게 합니다.",
        "why": "일반 fetch는 한 번 응답하면 끝나지만, SSE는 연결을 유지해 서버가 변경 이벤트를 계속 보낼 수 있습니다.",
        "flow": "프론트 EventSource('/stream/congestion') -> StreamService.subscribeCongestion -> emitter 저장 -> publishCongestion 시 이벤트 전송",
        "point": "SseEmitter 목록에서 완료/타임아웃/에러 연결을 제거해야 죽은 연결이 쌓이지 않습니다.",
    },
    "TranslateMetricsService": {
        "method": "recordSuccess",
        "role": "번역 API 성공/실패와 지연 시간을 누적합니다.",
        "why": "외부 번역 API 품질을 운영 화면에서 확인하려면 성공 수, 실패 수, 평균 지연 시간이 필요합니다.",
        "flow": "TranslateService 호출 완료 -> recordSuccess 또는 recordFailure -> snapshot에서 지표 조회",
        "point": "AtomicLong은 여러 요청이 동시에 값을 바꿔도 안전하게 누적하기 위한 도구입니다.",
    },
    "TranslateService": {
        "method": "translate",
        "role": "외부 번역 API를 호출해 입력 문장을 번역합니다.",
        "why": "번역은 서버 내부 계산이 아니라 외부 HTTP API 호출이므로 timeout, 예외 처리, 응답 파싱이 중요합니다.",
        "flow": "번역 요청 -> URI 생성 -> HttpClient.send -> 응답 JSON 파싱 -> TranslateResponseDto 반환",
        "point": "외부 API 응답 형식이 바뀌면 파싱 코드가 깨질 수 있어 방어 코드가 필요합니다.",
    },
    "UploadStorageService": {
        "method": "saveImage",
        "role": "업로드된 이미지를 로컬 디스크 또는 S3에 저장합니다.",
        "why": "프론트 파일 업로드는 JSON이 아니라 multipart/form-data로 오고, 백엔드에서는 MultipartFile을 받아 저장합니다.",
        "flow": "프론트 FormData 파일 업로드 -> Controller MultipartFile 수신 -> UploadStorageService.saveImage -> local 또는 S3 저장 -> URL 반환",
        "point": "파일 업로드에서는 파일 bytes와 contentType을 다뤄야 하므로 일반 JSON payload 저장과 완전히 다릅니다.",
    },
    "AiMatchSmsNotifier": {
        "method": "notifyRequestCreated",
        "role": "AI 매칭 신청이 생성되거나 수락될 때 문자 알림을 보냅니다.",
        "why": "매칭 당사자에게 앱 밖에서도 알림을 전달하기 위해 SMS 클라이언트를 사용합니다.",
        "flow": "매칭 요청 생성/수락 -> notifier 호출 -> SolapiMessageClient.sendText -> 문자 발송",
        "point": "enabled 플래그와 client 설정을 확인한 뒤 발송하므로, 설정이 없을 때는 조용히 건너뛸 수 있습니다.",
    },
    "AligoSmsSender": {
        "method": "sendVerificationCode",
        "role": "Aligo API로 예약 인증번호를 발송합니다.",
        "why": "SmsSender 인터페이스 구현체로 만들어 두면 다른 SMS 업체와 교체하기 쉽습니다.",
        "flow": "ReservationAuthService -> SmsSender -> AligoSmsSender -> Aligo HTTP API",
        "point": "문자 업체별 인증값과 요청 형식이 달라 구현체를 분리합니다.",
    },
    "NoopSmsSender": {
        "method": "sendVerificationCode",
        "role": "실제 문자를 보내지 않고 로그만 남기는 개발용 SMS 구현체입니다.",
        "why": "개발/데모 환경에서 문자 비용이나 외부 API 없이 인증 흐름을 테스트할 수 있습니다.",
        "flow": "인증번호 발송 요청 -> NoopSmsSender -> 로그 출력 -> 실제 SMS 없음",
        "point": "운영 환경에서 Noop 구현체가 선택되면 실제 사용자는 문자를 받지 못합니다.",
    },
    "SmsSender": {
        "method": "sendVerificationCode",
        "role": "인증번호 발송 기능의 공통 인터페이스입니다.",
        "why": "ReservationAuthService가 Aligo, Solapi, Twilio 중 어떤 업체인지 몰라도 같은 메서드로 발송하게 하기 위함입니다.",
        "flow": "서비스는 SmsSender만 호출 -> Spring 설정에 따라 실제 구현체 선택",
        "point": "인터페이스는 구현이 아니라 약속입니다. 실제 코드는 구현 클래스에 있습니다.",
    },
    "SolapiMessageClient": {
        "method": "sendText",
        "role": "Solapi SDK를 사용해 문자 한 건을 실제 발송합니다.",
        "why": "Solapi 전용 설정과 SDK 호출을 한 클래스에 가둬 두면 다른 서비스가 복잡한 SDK 코드를 몰라도 됩니다.",
        "flow": "Notifier 또는 SolapiSmsSender -> SolapiMessageClient.sendText -> Solapi SDK sendOne",
        "point": "isConfigured와 isEnabled를 확인해서 설정이 없을 때 발송을 막습니다.",
    },
    "SolapiSmsSender": {
        "method": "sendVerificationCode",
        "role": "SolapiMessageClient를 이용해 인증번호 문자를 보냅니다.",
        "why": "공통 SmsSender 인터페이스를 구현하면서 실제 발송은 SolapiMessageClient에 위임합니다.",
        "flow": "ReservationAuthService -> SolapiSmsSender -> SolapiMessageClient.sendText",
        "point": "얇은 adapter 역할입니다. 업체별 문구나 호출 방식을 여기서 맞춥니다.",
    },
    "TwilioSmsSender": {
        "method": "sendVerificationCode",
        "role": "Twilio API로 인증번호 문자를 보내는 대체 구현체입니다.",
        "why": "같은 SmsSender 인터페이스를 구현해 SMS 업체를 바꿀 수 있게 합니다.",
        "flow": "ReservationAuthService -> TwilioSmsSender -> Twilio Message.creator",
        "point": "Twilio.init에 필요한 SID/token/from-number 설정이 배포 환경에 있어야 합니다.",
    },
}


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.17) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 9.5, color: RGBColor | None = None, font: str = "Malgun Gothic"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], font_size: float = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        shade_cell(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.05)
        add_run(p, str(header), bold=True, size=font_size, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.08)
            add_run(p, str(value), size=font_size)
    doc.add_paragraph()


def add_para(doc: Document, text: str, label: str | None = None) -> None:
    p = doc.add_paragraph()
    set_spacing(p)
    if label:
        add_run(p, label + " ", bold=True, color=DARK)
    add_run(p, text)


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.15)
    add_run(p, title + " | ", bold=True, color=DARK, size=8.7)
    add_run(p, body, color=MUTED, size=8.7)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.3)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.17

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def service_files() -> list[Path]:
    return sorted(SERVICE_DIR.rglob("*.java"), key=lambda path: str(path.relative_to(SERVICE_DIR)).lower())


def source(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def extract_method(src: str, method_name: str, max_lines: int = 44) -> str:
    pattern = re.compile(
        r"(?:(?:@\w+(?:\([^)]*\))?\s*)*)"
        r"(?:public|private|protected)\s+[\w<>, ?\[\]]+\s+"
        + re.escape(method_name)
        + r"\s*\([^)]*\)\s*(?:throws [^{]+)?\{",
        re.M,
    )
    match = pattern.search(src)
    if not match:
        return extract_first_public_method(src, max_lines=max_lines)
    start = match.start()
    end = find_block_end(src, src.find("{", match.start()))
    snippet = src[start:end].strip()
    return limit_lines(snippet, max_lines)


def extract_first_public_method(src: str, max_lines: int = 34) -> str:
    pattern = re.compile(r"(?:(?:@\w+(?:\([^)]*\))?\s*)*)public\s+[\w<>, ?\[\]]+\s+\w+\s*\([^)]*\)\s*(?:throws [^{]+)?\{", re.M)
    match = pattern.search(src)
    if not match:
        return limit_lines(src, max_lines)
    end = find_block_end(src, src.find("{", match.start()))
    return limit_lines(src[match.start():end].strip(), max_lines)


def find_block_end(src: str, open_index: int) -> int:
    if open_index < 0:
        return min(len(src), open_index + 1500)
    depth = 0
    in_string = False
    escaped = False
    for index in range(open_index, len(src)):
        char = src[index]
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index + 1
    return len(src)


def limit_lines(snippet: str, max_lines: int) -> str:
    lines = snippet.splitlines()
    if len(lines) <= max_lines:
        return snippet
    head = lines[: max_lines - 4]
    tail = lines[-3:]
    return "\n".join(head + ["    // ... 중간 로직 생략: 문서에서는 핵심 흐름만 발췌했습니다."] + tail)


def class_name(path: Path) -> str:
    return path.stem


def method_list(src: str, class_name_value: str) -> list[str]:
    methods = []
    for match in re.finditer(r"public\s+[\w<>, ?\[\]]+\s+([A-Za-z0-9_]+)\s*\(", src):
        name = match.group(1)
        if name != class_name_value:
            methods.append(name)
    return methods


def add_cover(doc: Document, files: list[Path]) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)
    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "백엔드 Service 코드블록 상세해설서", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph()
    add_run(
        subtitle,
        "각 Service의 실제 핵심 코드 블록을 기준으로 기능, 구조 선택 이유, 요청 흐름, 초보자 주의점을 설명",
        color=MUTED,
        size=11,
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상", "backend/src/main/java/com/festflow/backend/service 하위 Java 서비스 파일"],
            ["포함 파일", f"{len(files)}개"],
            ["문서 방식", "파일별 대표 코드 블록 + 역할 설명 + 왜 이런 구조인지 + 전체 흐름 + 헷갈릴 포인트"],
            ["기존 문서와 차이", "기계적인 줄별 표가 아니라, 발표/공부에 필요한 코드 블록 중심 해설"],
        ],
        [2200, 7160],
        font_size=8.5,
    )


def add_intro(doc: Document) -> None:
    doc.add_heading("1. 이 문서를 읽는 방법", level=1)
    add_callout(
        doc,
        "핵심",
        "Service 코드를 볼 때는 코드 한 줄 자체보다 '이 요청이 어디서 들어와서, 어떤 데이터를 조회하고, 어떤 DTO로 나가는지'를 같이 봐야 합니다.",
        GREEN_FILL,
    )
    add_table(
        doc,
        ["볼 것", "설명"],
        [
            ["입력", "Controller에서 넘어오는 id, requestDto, file, token 같은 값입니다."],
            ["검증", "데이터가 있는지, 권한이 맞는지, 상태 변경이 가능한지 확인합니다."],
            ["처리", "Repository 조회/저장, 다른 Service 호출, 외부 API 호출, 파일 저장, AI 호출 등이 일어납니다."],
            ["출력", "프론트가 받을 DTO를 만들어 반환하거나, SSE/SMS 같은 알림을 발행합니다."],
            ["예외", "잘못된 요청이면 ResponseStatusException 등으로 중단합니다."],
        ],
        [1800, 7560],
        font_size=8,
    )
    add_para(
        doc,
        "예를 들어 파일 업로드는 JSON으로 처리하지 않습니다. 프론트는 FormData를 보내고, 백엔드는 MultipartFile로 받습니다. 반대로 일반 저장/수정 요청은 JSON body를 DTO로 받아 처리합니다. 이 차이를 코드 블록별 설명에서 계속 연결해 볼 수 있게 구성했습니다.",
    )


def add_service_section(doc: Document, path: Path, index: int) -> None:
    name = class_name(path)
    src = source(path)
    note = SERVICE_NOTES.get(name)
    if not note:
        methods = method_list(src, name)
        method = methods[0] if methods else ""
        note = {
            "method": method,
            "role": "서비스 계층에 포함된 보조 클래스입니다.",
            "why": "Controller나 다른 Service에서 필요한 업무 로직을 분리해 관리하기 위한 구조입니다.",
            "flow": "요청 또는 내부 호출 -> Service 메서드 -> Repository/외부 API/다른 Service -> DTO 또는 상태 반환",
            "point": "메서드 이름과 의존성 필드를 같이 보면 이 클래스가 맡은 책임을 파악할 수 있습니다.",
        }
    doc.add_heading(f"2.{index}. {name}", level=2)
    add_table(
        doc,
        ["항목", "설명"],
        [
            ["파일", str(path.relative_to(ROOT))],
            ["이 서비스의 역할", note["role"]],
            ["대표 메서드", note["method"] or "대표 public 메서드 없음"],
        ],
        [1800, 7560],
        font_size=8,
    )
    if note["method"]:
        add_code_block(doc, extract_method(src, note["method"]))
    else:
        add_code_block(doc, limit_lines(src, 32))
    add_para(doc, note["role"], "이 코드가 하는 일:")
    add_para(doc, note["why"], "왜 이렇게 되어 있는지:")
    add_para(doc, note["flow"], "흐름:")
    add_para(doc, note["point"], "헷갈릴 포인트:")
    add_keyword_notes(doc, src, name)


def add_keyword_notes(doc: Document, src: str, name: str) -> None:
    rows = []
    if "@Transactional" in src:
        rows.append(["@Transactional", "DB 조회/수정/저장을 하나의 작업 단위로 묶습니다. 예약, 매칭, 스태프 상태처럼 상태 변경이 있는 서비스에서 중요합니다."])
    if "MultipartFile" in src:
        rows.append(["MultipartFile", "프론트가 FormData로 보낸 파일을 백엔드에서 받는 타입입니다. JSON.stringify(payload)로 보내는 일반 데이터와 다릅니다."])
    if "ProcessBuilder" in src:
        rows.append(["ProcessBuilder", "Java 서버가 Python 추론 스크립트 같은 외부 프로그램을 실행할 때 사용합니다. AI 모델 pkl 추론 연결의 핵심입니다."])
    if "SseEmitter" in src or "publish" in src and "StreamService" in src:
        rows.append(["SSE / StreamService", "일반 fetch처럼 요청-응답으로 끝나는 구조가 아니라 연결을 유지하고 변경 이벤트를 밀어주는 실시간 구조입니다."])
    if "RestClient" in src or "HttpClient" in src:
        rows.append(["외부 API 호출", "OpenAI, 번역, SMS 같은 외부 시스템과 통신합니다. timeout, 예외 처리, fallback을 같이 봐야 합니다."])
    if "Repository" in src:
        rows.append(["Repository", "DB 접근 계층입니다. Service는 Repository를 통해 Entity를 조회, 저장, 삭제합니다."])
    if "orElseThrow" in src:
        rows.append(["orElseThrow", "조회 결과가 없을 때 null로 계속 진행하지 않고 즉시 예외를 던져 요청을 중단합니다."])
    if "stream()" in src:
        rows.append(["stream()", "목록을 필터링하거나 DTO로 변환할 때 쓰는 Java 컬렉션 처리 방식입니다."])
    if rows:
        add_table(doc, ["코드에서 볼 포인트", "설명"], rows, [2300, 7060], font_size=7.8)


def build_doc() -> None:
    files = service_files()
    doc = Document()
    configure_doc(doc)
    add_cover(doc, files)
    add_intro(doc)
    doc.add_heading("2. Service별 코드블록 상세해설", level=1)
    for index, path in enumerate(files, start=1):
        add_service_section(doc, path, index)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")
    print(f"services: {len(files)}")


if __name__ == "__main__":
    build_doc()
