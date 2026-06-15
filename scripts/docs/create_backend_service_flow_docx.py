from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
INPUT_MD = ROOT / "docs" / "festflow" / "백엔드_서비스_흐름_문서.md"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "백엔드_서비스_흐름_문서.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "E8EEF5"
CODE_FILL = "F4F6F9"
BORDER = "D6DEE8"


def read_markdown(path: Path) -> list[str]:
    for encoding in ("utf-8-sig", "utf-8", "cp949"):
        try:
            return path.read_text(encoding=encoding).splitlines()
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Could not decode {path}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_style_font(style, name: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_style_paragraph(style, before: float, after: float, line_spacing: float) -> None:
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_spacing


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    set_style_paragraph(normal, 0, 6, 1.25)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        set_style_paragraph(style, before, after, 1.25)

    list_bullet = styles["List Bullet"]
    set_style_font(list_bullet, "Calibri", 10.5, INK)
    set_style_paragraph(list_bullet, 0, 4, 1.25)
    list_bullet.paragraph_format.left_indent = Inches(0.375)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.188)

    list_number = styles["List Number"]
    set_style_font(list_number, "Calibri", 10.5, INK)
    set_style_paragraph(list_number, 0, 4, 1.25)
    list_number.paragraph_format.left_indent = Inches(0.375)
    list_number.paragraph_format.first_line_indent = Inches(-0.188)

    code_style = styles.add_style("Code Block", 1)
    set_style_font(code_style, "Consolas", 8.5, RGBColor(0x33, 0x33, 0x33))
    set_style_paragraph(code_style, 3, 5, 1.1)
    code_style.paragraph_format.left_indent = Inches(0.15)
    code_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    caption_style = styles.add_style("Muted Caption", 1)
    set_style_font(caption_style, "Calibri", 9, MUTED)
    set_style_paragraph(caption_style, 4, 4, 1.15)


def set_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FestFlow Backend Service Flow")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("FestFlow 백엔드 서비스 흐름 문서")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph(style="Muted Caption")
    subtitle.add_run(
        "Backend Service 역할, public 메서드, Controller/Frontend 연결 흐름 정리"
    )
    subtitle.paragraph_format.space_after = Pt(12)


def split_inline_code(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    current = ""
    in_code = False
    for char in text:
        if char == "`":
            if current:
                parts.append((current, in_code))
                current = ""
            in_code = not in_code
        else:
            current += char
    if current:
        parts.append((current, in_code))
    return parts


def add_text_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    for segment, is_code in split_inline_code(text):
        run = paragraph.add_run(segment)
        run.font.name = "Consolas" if is_code else "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10 if is_code else 11)
        run.font.color.rgb = RGBColor(0x2F, 0x3A, 0x45) if is_code else INK


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    set_paragraph_shading(paragraph, CODE_FILL)
    for idx, line in enumerate(lines):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def normalize_table_line(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def is_separator_row(line: str) -> bool:
    cells = normalize_table_line(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def table_widths(column_count: int) -> list[int]:
    patterns = {
        2: [2700, 6660],
        3: [2100, 4300, 2960],
        4: [1700, 2300, 3600, 1760],
        5: [1500, 1950, 2050, 2200, 1660],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base for _ in range(column_count)]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = [normalize_table_line(line) for line in table_lines if not is_separator_row(line)]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (max_cols - len(row)))

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(max_cols))

    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            for segment, is_code in split_inline_code(value):
                run = paragraph.add_run(segment)
                run.font.name = "Consolas" if is_code else "Calibri"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                run.font.size = Pt(7.8 if max_cols >= 5 else 8.5)
                run.font.color.rgb = RGBColor(0x2F, 0x3A, 0x45) if is_code else INK
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            else:
                set_cell_shading(cell, "FFFFFF")

    doc.add_paragraph(style="Muted Caption")


def markdown_to_docx(lines: list[str], output_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    set_footer(doc)
    add_title_block(doc)

    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_separator_row(lines[index + 1]):
            table_lines = [stripped, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue
        if stripped.startswith("## "):
            paragraph = doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue
        if stripped.startswith("#### "):
            paragraph = doc.add_paragraph(stripped[5:].strip(), style="Heading 3")
            paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue

        if stripped.startswith("- "):
            add_text_paragraph(doc, stripped[2:].strip(), style="List Bullet")
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            add_text_paragraph(doc, number_match.group(1), style="List Number")
            index += 1
            continue

        add_text_paragraph(doc, stripped)
        index += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "FestFlow 백엔드 서비스 흐름 문서"
    doc.core_properties.subject = "Backend service role and request flow reference"
    doc.core_properties.author = "Codex"
    doc.save(output_path)


def main() -> int:
    if not INPUT_MD.exists():
        print(f"Input markdown not found: {INPUT_MD}", file=sys.stderr)
        return 1
    markdown_to_docx(read_markdown(INPUT_MD), OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
