from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SERVICE_DIR = ROOT / "backend" / "src" / "main" / "java" / "com" / "festflow" / "backend" / "service"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "페스트플로우_백엔드_Service_코드_한줄해설서.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F8FAFC"
YELLOW_FILL = "FFFAEB"


SERVICE_ROLE_HINTS = {
    "AdminActionService": "관리자 버튼 액션으로 공지 발행을 자동화하는 서비스입니다.",
    "AdminDashboardService": "관리자 대시보드 KPI를 계산하는 서비스입니다.",
    "AdminImportService": "CSV 파일을 읽어 부스/공연 데이터를 대량 등록하는 서비스입니다.",
    "AiCongestionService": "AI 혼잡도 예측과 추천, 판단 로그를 만드는 서비스입니다.",
    "AiDecisionLogService": "AI 판단 로그를 메모리에 보관하는 서비스입니다.",
    "AiImageGenerationService": "AI 매칭 프로필 이미지를 검증하고 생성하는 서비스입니다.",
    "AiMatchService": "AI 매칭 프로필, 좋아요, 신청, 수락/거절, 만남 제안을 처리하는 서비스입니다.",
    "AuditLogService": "관리자 작업 로그를 저장하고 조회하는 서비스입니다.",
    "AnalyticsService": "GPS와 부스 데이터를 집계해 분석 대시보드를 만드는 서비스입니다.",
    "AuthService": "관리자 로그인과 JWT 발급을 담당하는 서비스입니다.",
    "BoothService": "부스 CRUD, 운영 상태, 혼잡도 계산을 담당하는 핵심 서비스입니다.",
    "ChatService": "축제 챗봇 답변을 생성하는 서비스입니다.",
    "EventService": "공연 CRUD와 공연 변경 알림을 담당하는 서비스입니다.",
    "FestivalSnapshotService": "AI/분석에 필요한 현재 축제 상태를 한 번에 모으는 서비스입니다.",
    "GpsService": "사용자 GPS 로그를 저장하고 혼잡도 변경 이벤트를 발행하는 서비스입니다.",
    "LostItemService": "분실물 등록, 수정, 상태 변경, 삭제를 처리하는 서비스입니다.",
    "NoticeService": "공지 CRUD와 공지 SSE 알림을 담당하는 서비스입니다.",
    "OpsAiService": "운영자/스태프용 AI 브리핑과 안내문을 생성하는 서비스입니다.",
    "PublicAiGuideService": "일반 방문자용 AI 방문 가이드를 생성하는 서비스입니다.",
    "PythonCongestionModelService": "Java에서 Python RandomForest 모델 추론 스크립트를 실행하는 서비스입니다.",
    "ReservationAuthService": "예약용 전화번호 인증번호 발송과 인증 토큰을 관리하는 서비스입니다.",
    "ReservationService": "예약 생성, 체크인, 완료, 테이블 해제 등 예약 기능을 처리하는 서비스입니다.",
    "SimulationService": "운영 시뮬레이션을 시작/정지하고 tick을 적용하는 서비스입니다.",
    "SimulationStateService": "시뮬레이션 상태를 메모리에서 관리하고 변화량을 계산하는 서비스입니다.",
    "StaffService": "스태프 로그인, 상태 변경, 세션 인증을 처리하는 서비스입니다.",
    "StreamService": "SSE 연결과 실시간 이벤트 발행을 담당하는 서비스입니다.",
    "TranslateMetricsService": "번역 성공/실패/지연시간 지표를 누적하는 서비스입니다.",
    "TranslateService": "외부 번역 API를 호출해 텍스트를 번역하는 서비스입니다.",
    "UploadStorageService": "이미지를 로컬 또는 S3 저장소에 저장/조회/삭제하는 서비스입니다.",
}


def set_spacing(paragraph, before: int = 0, after: int = 2, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 8, color: RGBColor | None = None, font: str = "Malgun Gothic"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], font_size: float = 6.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        shade_cell(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.0)
        add_run(p, str(header), bold=True, size=font_size + 0.4, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for index, (cell, value, width) in enumerate(zip(cells, row, widths)):
            set_cell_width(cell, width)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.0)
            if index == 1:
                add_run(p, str(value), size=font_size, font="Consolas")
            else:
                add_run(p, str(value), size=font_size)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [12960])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.12)
    add_run(p, title + " | ", bold=True, size=8.5, color=DARK)
    add_run(p, body, size=8.5, color=MUTED)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def service_files() -> list[Path]:
    return sorted(SERVICE_DIR.rglob("*.java"), key=lambda path: str(path.relative_to(SERVICE_DIR)).lower())


def class_name(path: Path) -> str:
    return path.stem


def explain_line(line: str, file_name: str, previous_nonblank: str = "") -> str:
    stripped = line.strip()
    name = file_name.replace(".java", "")

    if not stripped:
        return "빈 줄입니다. 코드 블록을 시각적으로 나누어 읽기 쉽게 합니다."
    if stripped == "{" or stripped.endswith(" {"):
        if " class " in stripped or stripped.startswith("public class") or stripped.startswith("public interface"):
            return "클래스나 인터페이스 본문을 시작합니다. 이 중괄호 안에 필드와 메서드가 들어갑니다."
        if re.search(r"\)\s*\{$", stripped):
            return "메서드나 조건문의 실행 블록을 시작합니다. 아래 줄들이 실제 처리 내용입니다."
        return "새 코드 블록을 시작하는 중괄호입니다."
    if stripped == "}" or stripped == "};" or stripped == "});":
        return "현재 코드 블록을 닫습니다. 메서드, 조건문, 람다식, 클래스 영역 중 하나가 끝나는 위치입니다."
    if stripped.startswith("package "):
        return "이 파일이 속한 Java 패키지를 선언합니다. 폴더 구조와 Java namespace를 맞추는 역할입니다."
    if stripped.startswith("import "):
        imported = stripped.replace("import ", "").replace(";", "")
        if ".dto." in imported:
            return f"`{imported}` DTO를 가져옵니다. API 요청이나 응답 형태를 표현할 때 사용합니다."
        if ".entity." in imported:
            return f"`{imported}` Entity를 가져옵니다. DB 테이블과 연결되는 도메인 객체입니다."
        if ".repository." in imported:
            return f"`{imported}` Repository를 가져옵니다. DB 조회/저장/삭제에 사용합니다."
        if "Service" in imported:
            return f"`{imported}` 서비스를 가져옵니다. 다른 업무 로직을 재사용하기 위한 의존성입니다."
        if "springframework" in imported:
            return f"Spring Framework 기능인 `{imported}`를 사용하기 위한 import입니다."
        if "java." in imported:
            return f"Java 표준 라이브러리 `{imported}`를 사용하기 위한 import입니다."
        return f"`{imported}` 타입을 이 파일에서 사용하기 위한 import입니다."
    if stripped.startswith("@Service"):
        return "이 클래스를 Spring Service Bean으로 등록합니다. Controller나 다른 Service에서 자동 주입해 사용할 수 있습니다."
    if stripped.startswith("@Component"):
        return "이 클래스를 Spring이 관리하는 일반 Bean으로 등록합니다. Service는 아니지만 주입 가능한 객체가 됩니다."
    if stripped.startswith("@Transactional(readOnly = true)"):
        return "읽기 전용 트랜잭션입니다. DB를 수정하지 않는 조회 메서드라는 의도를 나타냅니다."
    if stripped.startswith("@Transactional"):
        return "트랜잭션을 시작합니다. 이 메서드 안의 DB 작업은 하나의 작업 단위로 묶입니다."
    if stripped.startswith("@Value"):
        return "application.properties나 배포 환경변수의 설정값을 이 필드나 생성자 인자로 주입합니다."
    if stripped.startswith("@PostConstruct"):
        return "Bean 생성 후 초기화 메서드를 자동 실행하도록 표시합니다."
    if stripped.startswith("@Scheduled"):
        return "Spring이 정해진 주기마다 이 메서드를 자동 실행하도록 지정합니다."
    if stripped.startswith("@Override"):
        return "인터페이스나 부모 클래스의 메서드를 구현/재정의한다는 표시입니다."

    if re.match(r"public\s+(class|interface|record)\s+", stripped):
        if "class" in stripped:
            return f"`{name}` 클래스 선언입니다. {SERVICE_ROLE_HINTS.get(name, '이 서비스 파일의 핵심 로직이 이 클래스 안에 들어갑니다.')}"
        if "interface" in stripped:
            return "인터페이스 선언입니다. 구현 클래스들이 반드시 가져야 할 메서드 약속을 정의합니다."
        return "record 선언입니다. 여러 값을 묶어 전달하는 불변 데이터 구조를 정의합니다."

    if re.match(r"(private|public|protected)\s+static\s+final\s+", stripped):
        return "상수 선언입니다. 여러 곳에서 반복해서 쓰는 고정값을 이름 붙여 관리합니다."
    if re.match(r"private\s+final\s+.+;", stripped):
        type_and_name = stripped.replace("private final ", "").rstrip(";")
        if "Repository" in type_and_name:
            return "Repository 의존성 필드입니다. DB 접근을 위해 생성자에서 주입받아 보관합니다."
        if "Service" in type_and_name:
            return "다른 Service 의존성 필드입니다. 이미 만들어진 업무 로직을 이 서비스에서 재사용합니다."
        if "Client" in type_and_name or "Sender" in type_and_name:
            return "외부 API 호출이나 메시지 발송을 담당하는 협력 객체 필드입니다."
        return f"`{type_and_name}` 값을 생성자에서 주입받아 저장하는 final 필드입니다."
    if re.match(r"private\s+.+;", stripped):
        return "클래스 내부에서 사용할 필드입니다. 서비스의 설정값이나 상태를 저장합니다."

    if re.match(r"public\s+" + re.escape(name) + r"\s*\(", stripped):
        return "생성자 선언입니다. Spring이 필요한 Repository, Service, 설정값을 이 생성자로 주입합니다."
    if re.match(r"public\s+[\w<>, ?\[\]]+\s+\w+\s*\(", stripped):
        method = re.sub(r"\s+", " ", stripped)
        return f"public 메서드 선언입니다. Controller나 다른 Service에서 호출할 수 있는 기능입니다. 선언부: `{method}`"
    if re.match(r"private\s+[\w<>, ?\[\]]+\s+\w+\s*\(", stripped):
        return "private 보조 메서드 선언입니다. 이 클래스 내부에서만 쓰는 세부 로직을 분리한 것입니다."
    if re.match(r"protected\s+[\w<>, ?\[\]]+\s+\w+\s*\(", stripped):
        return "protected 메서드 선언입니다. 같은 패키지나 상속 관계에서 사용할 수 있는 기능입니다."

    if stripped.startswith("this."):
        if "=" in stripped:
            return "생성자나 메서드에서 받은 값을 현재 객체의 필드에 저장합니다."
        return "현재 객체의 필드나 메서드를 명시적으로 참조합니다."
    if ".save(" in stripped:
        return "Repository의 save를 호출해 Entity를 DB에 저장하거나 수정합니다."
    if ".delete(" in stripped or ".deleteBy" in stripped:
        return "Repository 삭제 기능을 호출해 DB 데이터를 제거합니다."
    if ".findById(" in stripped:
        return "ID로 DB 데이터를 조회합니다. 보통 뒤에서 없을 때 예외 처리와 연결됩니다."
    if ".findAll(" in stripped:
        return "DB 테이블의 전체 목록을 조회합니다."
    if ".findBy" in stripped or ".existsBy" in stripped or ".countBy" in stripped:
        return "Repository의 query method를 호출합니다. 메서드 이름 규칙으로 조건 조회/존재 확인/개수 계산을 수행합니다."
    if ".orElseThrow" in stripped:
        return "조회 결과가 없을 때 예외를 던집니다. API에서는 보통 404나 400 오류로 이어집니다."
    if "ResponseStatusException" in stripped:
        return "HTTP 상태 코드가 포함된 예외를 만듭니다. 잘못된 요청이나 찾을 수 없는 데이터 상황을 표현합니다."
    if stripped.startswith("throw "):
        return "조건을 만족하지 못해 예외를 발생시킵니다. 이후 코드는 실행되지 않습니다."

    if stripped.startswith("if " ) or stripped.startswith("if("):
        return "조건문입니다. 괄호 안 조건이 참일 때만 아래 블록을 실행합니다."
    if stripped.startswith("else if"):
        return "앞 조건이 거짓일 때 추가 조건을 검사합니다."
    if stripped.startswith("else"):
        return "앞의 if 조건들이 모두 거짓일 때 실행되는 대체 흐름입니다."
    if stripped.startswith("for " ) or stripped.startswith("for("):
        return "반복문입니다. 목록이나 범위를 순회하면서 같은 처리를 반복합니다."
    if stripped.startswith("while " ) or stripped.startswith("while("):
        return "조건이 참인 동안 반복 실행하는 반복문입니다."
    if stripped.startswith("try"):
        return "예외가 발생할 수 있는 코드를 실행하기 위한 try 블록입니다."
    if stripped.startswith("catch"):
        return "try 블록에서 발생한 예외를 잡아 처리합니다."
    if stripped.startswith("finally"):
        return "예외 발생 여부와 관계없이 마지막에 실행되는 정리 블록입니다."
    if stripped.startswith("return "):
        if "Dto" in stripped or "Response" in stripped:
            return "최종 응답 DTO를 반환합니다. Controller는 이 값을 JSON 응답으로 내려줍니다."
        if ".stream()" in stripped:
            return "stream 처리 결과를 바로 반환합니다. 보통 목록을 DTO 목록으로 변환합니다."
        return "메서드 실행 결과를 호출한 쪽으로 반환합니다."

    if ".stream()" in stripped:
        return "Java Stream API를 시작합니다. 컬렉션을 필터링하거나 변환하기 위한 체인 처리입니다."
    if ".map(" in stripped:
        return "Stream의 각 원소를 다른 형태로 변환합니다. Entity를 DTO로 바꿀 때 자주 사용합니다."
    if ".filter(" in stripped:
        return "조건에 맞는 원소만 남깁니다."
    if ".sorted(" in stripped:
        return "목록을 정렬합니다."
    if ".toList()" in stripped or "Collectors.toList" in stripped:
        return "Stream 처리 결과를 List로 모읍니다."
    if "Collectors.toMap" in stripped:
        return "Stream 처리 결과를 Map 형태로 모읍니다."
    if "Optional<" in stripped or "Optional." in stripped:
        return "값이 있을 수도 없을 수도 있음을 표현하는 Optional을 사용합니다."

    if "new " in stripped:
        if "Dto" in stripped:
            return "새 DTO 객체를 생성합니다. 화면/API로 내보낼 응답 데이터를 구성합니다."
        if "Entity" in stripped or re.search(r"new\s+[A-Z][A-Za-z0-9]+\(", stripped):
            return "새 객체를 생성합니다. Entity, DTO, 예외, 요청 객체 중 하나를 만들고 있습니다."
    if "RestClient" in stripped or "restClient" in stripped:
        return "Spring RestClient로 외부 HTTP API를 호출하는 코드입니다."
    if "HttpClient" in stripped or "httpClient" in stripped:
        return "Java HttpClient로 외부 HTTP API를 호출하는 코드입니다."
    if "SseEmitter" in stripped:
        return "SSE 연결 객체를 다룹니다. 서버가 브라우저에 실시간 이벤트를 보낼 때 사용합니다."
    if "ProcessBuilder" in stripped:
        return "외부 프로세스를 실행하기 위한 객체입니다. 이 프로젝트에서는 Python 추론 스크립트 실행에 사용됩니다."
    if "MultipartFile" in stripped:
        return "브라우저에서 업로드된 파일을 표현합니다. 이미지나 CSV 업로드 처리에 사용됩니다."
    if "LocalDateTime" in stripped or "Instant" in stripped:
        return "시간 값을 다룹니다. 로그 집계, 만료 시간, 예약 시간 계산에 사용됩니다."
    if "Duration" in stripped:
        return "시간 간격을 표현합니다. timeout이나 만료 기준 계산에 사용합니다."
    if "UUID" in stripped:
        return "중복 가능성이 낮은 고유 문자열을 생성합니다. 토큰이나 파일명에 자주 사용됩니다."
    if "log." in stripped:
        return "서버 로그를 남깁니다. 성공/실패/예외 상황을 운영자가 확인할 수 있게 합니다."
    if "builder" in stripped.lower():
        return "Builder 패턴을 사용해 객체 설정을 단계적으로 구성합니다."
    if stripped.startswith("List<") or stripped.startswith("Map<") or stripped.startswith("Set<"):
        return "컬렉션 변수를 선언합니다. 여러 데이터를 목록, 키-값, 집합 형태로 다룹니다."
    if "=" in stripped and stripped.endswith(";"):
        return "변수에 값을 계산해 저장합니다. 이후 조건 판단, DTO 생성, 저장 로직에서 사용됩니다."
    if stripped.endswith(";"):
        return "하나의 Java 실행문입니다. 메서드 호출, 값 저장, 객체 생성 중 하나를 수행합니다."
    if stripped.startswith("//"):
        return "코드 설명을 위한 주석입니다. 실행되지는 않습니다."

    return "서비스 로직의 일부입니다. 앞뒤 줄과 함께 읽으면 데이터 조회, 검증, 변환, 저장 흐름에 포함됩니다."


def file_summary(path: Path) -> str:
    name = class_name(path)
    lines = path.read_text(encoding="utf-8").splitlines()
    public_methods = []
    for line in lines:
        stripped = line.strip()
        match = re.match(r"public\s+[\w<>, ?\[\]]+\s+([A-Za-z0-9_]+)\s*\(", stripped)
        if match and match.group(1) != name:
            public_methods.append(match.group(1))
    method_text = ", ".join(public_methods[:8]) if public_methods else "public 메서드 없음 또는 인터페이스/구현 보조 클래스"
    if len(public_methods) > 8:
        method_text += " ..."
    return f"{SERVICE_ROLE_HINTS.get(name, '서비스 계층에 포함된 Java 파일입니다.')} 주요 public 메서드: {method_text}"


def line_rows(path: Path) -> list[list[object]]:
    rows = []
    previous_nonblank = ""
    for number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        explanation = explain_line(line, path.name, previous_nonblank)
        rows.append([number, line if line.strip() else "(blank)", explanation])
        if line.strip():
            previous_nonblank = line.strip()
    return rows


def add_cover(doc: Document, files: list[Path]) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=13)
    title = doc.add_paragraph()
    set_spacing(title, before=10, after=6, line=1.05)
    add_run(title, "백엔드 Service 코드 한줄해설서", bold=True, color=DARK, size=22)
    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=8, line=1.12)
    add_run(
        subtitle,
        "backend service 하위 Java 파일 전체를 줄 번호, 실제 코드, 초보자용 설명 형태로 정리한 학습 문서",
        color=MUTED,
        size=10,
    )
    total_lines = sum(len(path.read_text(encoding="utf-8").splitlines()) for path in files)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상 폴더", "backend/src/main/java/com/festflow/backend/service"],
            ["포함 파일", f"{len(files)}개"],
            ["총 코드 줄 수", f"{total_lines:,}줄"],
            ["문서 형식", "각 파일별로 줄 번호 / 실제 코드 / 한줄 설명 표 제공"],
            ["추천 사용법", "처음에는 파일 요약을 보고, 막히는 줄은 오른쪽 설명을 같이 읽는 방식"],
        ],
        [2200, 10760],
        font_size=8.2,
    )
    add_callout(
        doc,
        "읽기 기준",
        "설명은 Java와 Spring을 처음 공부하는 관점에 맞췄습니다. import, 어노테이션, 필드, 생성자, 조건문, Repository 호출, DTO 반환, 외부 API 호출, SSE, 트랜잭션을 각각 구분해서 설명합니다.",
        YELLOW_FILL,
    )


def add_foundation(doc: Document) -> None:
    doc.add_heading("0. 먼저 알아야 할 기본 표현", level=1)
    add_table(
        doc,
        ["표현", "뜻"],
        [
            ["@Service", "Spring이 이 클래스를 서비스 객체로 만들어 관리합니다."],
            ["@Component", "Spring이 관리하는 일반 객체입니다. SMS 발송 구현체 등에 쓰입니다."],
            ["@Transactional", "DB 작업을 하나의 작업 단위로 묶습니다. 실패하면 변경을 되돌릴 수 있습니다."],
            ["Repository", "DB에 접근하는 객체입니다. find, save, delete 같은 메서드를 사용합니다."],
            ["Entity", "DB 테이블과 연결되는 Java 객체입니다."],
            ["DTO", "API 요청/응답에 쓰는 데이터 모양입니다."],
            ["SseEmitter", "서버가 브라우저에 실시간 이벤트를 보내는 SSE 연결 객체입니다."],
            ["ProcessBuilder", "Java에서 Python 같은 외부 프로그램을 실행할 때 쓰는 객체입니다."],
            ["stream().map().toList()", "목록을 하나씩 변환해서 새 목록으로 만드는 Java 문법입니다."],
            ["orElseThrow()", "조회 결과가 없으면 예외를 던지는 Optional 처리 방식입니다."],
        ],
        [2600, 10360],
        font_size=7.8,
    )


def build_doc() -> None:
    files = service_files()
    doc = Document()
    configure_doc(doc)
    add_cover(doc, files)
    add_foundation(doc)

    doc.add_heading("1. Service 파일별 한줄해설", level=1)
    for index, path in enumerate(files, start=1):
        rel = path.relative_to(ROOT)
        doc.add_heading(f"1.{index}. {class_name(path)}", level=2)
        add_callout(doc, "파일 요약", f"{rel} | {file_summary(path)}")
        rows = line_rows(path)
        add_table(doc, ["줄", "실제 코드", "설명"], rows, [650, 6100, 6210], font_size=5.8)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")
    print(f"files: {len(files)}")
    print(f"lines: {sum(len(path.read_text(encoding='utf-8').splitlines()) for path in files)}")


if __name__ == "__main__":
    build_doc()
