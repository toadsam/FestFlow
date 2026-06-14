# -*- coding: utf-8 -*-
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SERVICE_DIR = ROOT / "backend" / "src" / "main" / "java" / "com" / "festflow" / "backend" / "service"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "페스트플로우_백엔드_Service_전체코드_완전해설서.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
RED = RGBColor(153, 27, 27)
GREEN = RGBColor(22, 101, 52)

HEADER_FILL = "E8EEF5"
NOTE_FILL = "F8FAFC"
CODE_FILL = "F3F4F6"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"


@dataclass
class CodeBlock:
    kind: str
    name: str
    signature: str
    body: str
    visibility: str
    return_type: str
    params: str
    start_line: int
    end_line: int


def service_files() -> list[Path]:
    return sorted(SERVICE_DIR.rglob("*.java"), key=lambda p: str(p.relative_to(SERVICE_DIR)).lower())


def read_source(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def class_name(path: Path) -> str:
    return path.stem


def line_number(src: str, index: int) -> int:
    return src.count("\n", 0, index) + 1


def find_block_end(src: str, open_index: int) -> int:
    depth = 0
    in_string = False
    escaped = False
    in_line_comment = False
    in_block_comment = False
    for i in range(open_index, len(src)):
        ch = src[i]
        nxt = src[i + 1] if i + 1 < len(src) else ""
        if in_line_comment:
            if ch == "\n":
                in_line_comment = False
            continue
        if in_block_comment:
            if ch == "*" and nxt == "/":
                in_block_comment = False
            continue
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == "/" and nxt == "/":
            in_line_comment = True
            continue
        if ch == "/" and nxt == "*":
            in_block_comment = True
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return i + 1
    return len(src)


def parse_methods(src: str, cls: str) -> list[CodeBlock]:
    blocks: list[CodeBlock] = []
    method_pattern = re.compile(
        r"(?ms)^"
        r"(?P<indent>\s*)"
        r"(?P<ann>(?:@\w+(?:\([^)]*\))?\s*)*)"
        r"(?P<vis>public|private|protected)\s+"
        r"(?P<mods>(?:(?:static|final|synchronized|abstract)\s+)*)"
        r"(?P<ret>[A-Za-z0-9_<>, ?\[\].]+?)\s+"
        r"(?P<name>[A-Za-z_][A-Za-z0-9_]*)\s*"
        r"\((?P<params>[^)]*)\)\s*"
        r"(?P<throws>throws\s+[^{]+)?\{",
    )
    ctor_pattern = re.compile(
        r"(?ms)^"
        r"(?P<indent>\s*)"
        r"(?P<ann>(?:@\w+(?:\([^)]*\))?\s*)*)"
        r"(?P<vis>public|private|protected)\s+"
        + re.escape(cls)
        + r"\s*\((?P<params>[^)]*)\)\s*(?P<throws>throws\s+[^{]+)?\{",
    )
    seen: set[tuple[int, int]] = set()
    for pattern, kind in [(ctor_pattern, "constructor"), (method_pattern, "method")]:
        for match in pattern.finditer(src):
            open_index = src.find("{", match.end() - 1)
            end = find_block_end(src, open_index)
            key = (match.start(), end)
            if key in seen:
                continue
            seen.add(key)
            raw = src[match.start():end].strip()
            first_line = raw.splitlines()[0].strip()
            if kind == "constructor":
                name = cls
                return_type = "생성자"
            else:
                name = match.group("name")
                return_type = " ".join(match.group("ret").split())
                if name == cls:
                    continue
            blocks.append(
                CodeBlock(
                    kind=kind,
                    name=name,
                    signature=first_line,
                    body=raw,
                    visibility=match.group("vis"),
                    return_type=return_type,
                    params=" ".join(match.group("params").split()),
                    start_line=line_number(src, match.start()),
                    end_line=line_number(src, end),
                )
            )
    blocks.sort(key=lambda block: block.start_line)
    return blocks


def top_level_lines(src: str, cls: str) -> tuple[list[str], list[str], list[str], str]:
    package = next((line.strip() for line in src.splitlines() if line.strip().startswith("package ")), "")
    imports = [line.strip() for line in src.splitlines() if line.strip().startswith("import ")]
    annotations = []
    class_match = re.search(r"(?ms)((?:^\s*@\w+(?:\([^)]*\))?\s*)*)\s*public\s+(?:class|interface|record)\s+" + re.escape(cls), src)
    if class_match:
        annotations = [line.strip() for line in class_match.group(1).splitlines() if line.strip()]
    fields = extract_fields(src, cls)
    return imports, annotations, fields, package


def extract_fields(src: str, cls: str) -> list[str]:
    class_match = re.search(r"\b(?:class|interface|record)\s+" + re.escape(cls) + r"\b[^{]*\{", src)
    if not class_match:
        return []
    start = src.find("{", class_match.start()) + 1
    end = find_block_end(src, start - 1) - 1
    body = src[start:end]
    lines = body.splitlines()
    fields: list[str] = []
    current: list[str] = []
    depth = 1
    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            continue
        opens = stripped.count("{")
        closes = stripped.count("}")
        if depth == 1 and (stripped.startswith("private ") or stripped.startswith("public ") or stripped.startswith("protected ")):
            if "(" not in stripped or stripped.endswith(";"):
                current.append(stripped)
                if stripped.endswith(";"):
                    fields.append(" ".join(current))
                    current = []
        elif current:
            current.append(stripped)
            if stripped.endswith(";"):
                fields.append(" ".join(current))
                current = []
        depth += opens - closes
    clean = []
    for field in fields:
        if "(" in field and not field.endswith(";"):
            continue
        if field not in clean:
            clean.append(field)
    return clean


def nested_records(src: str) -> list[str]:
    records = []
    for match in re.finditer(r"(?ms)^\s*(?:private|public|protected)?\s*record\s+([A-Za-z_][A-Za-z0-9_]*)\s*\([^)]*\)\s*\{", src):
        end = find_block_end(src, src.find("{", match.end() - 1))
        records.append(src[match.start():end].strip())
    return records


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 9.3, color: RGBColor | None = None, font: str = "Malgun Gothic"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "FestFlow Backend Service Full Code Guide", size=8, color=MUTED)


def add_paragraph(doc: Document, text: str, label: str | None = None, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    set_spacing(p)
    if label:
        add_run(p, label + " ", bold=True, color=DARK)
    add_run(p, text, color=color)


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL, color: RGBColor = DARK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.2)
    add_run(p, title + " | ", bold=True, color=color, size=8.7)
    add_run(p, body, color=MUTED, size=8.7)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], font_size: float = 7.8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        shade_cell(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.15)
        add_run(p, str(header), bold=True, size=font_size, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.16)
            add_run(p, str(value), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_FILL)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.rstrip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(6.25)
    doc.add_paragraph()


def add_cover(doc: Document, files: list[Path], total_methods: int) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)
    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "백엔드 Service 전체코드 완전해설서", bold=True, color=DARK, size=23)
    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=10, line=1.2)
    add_run(
        subtitle,
        "대표 코드만 뽑은 문서가 아니라, Service 파일의 package/import/필드/생성자/public 메서드/private 메서드/record까지 전체 구조를 따라가며 설명하는 공부용 문서입니다.",
        color=MUTED,
        size=10.5,
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상", "backend/src/main/java/com/festflow/backend/service 이하 전체 Java 파일"],
            ["Service 파일 수", f"{len(files)}개"],
            ["추출된 생성자/메서드 수", f"{total_methods}개"],
            ["설명 방식", "파일 구조 설명 + 모든 코드 블록 수록 + 메서드별 입력/반환/흐름/문법 포인트 설명"],
            ["주의", "소스 파일 안의 일부 한글 문자열은 기존 인코딩 문제로 깨져 보일 수 있으나, 코드 구조 해설에는 영향을 주지 않습니다."],
        ],
        [2300, 7060],
        font_size=8.2,
    )


def add_global_guide(doc: Document) -> None:
    doc.add_heading("1. 이 문서를 읽는 방식", level=1)
    add_callout(
        doc,
        "핵심",
        "이번 문서는 대표 메서드가 아니라 모든 메서드를 다룹니다. 단, 한 줄씩 기계적으로 번역하지 않고 메서드 단위로 코드 전체를 보여준 뒤 초보자가 이해해야 할 흐름을 설명합니다.",
        GREEN_FILL,
        GREEN,
    )
    add_table(
        doc,
        ["문서 항목", "무엇을 보면 되는가"],
        [
            ["package/import", "이 파일이 어느 패키지에 있고 어떤 DTO, Entity, Repository, Spring 기능을 쓰는지 봅니다."],
            ["필드/상수", "이 서비스가 의존하는 DB 저장소, 다른 서비스, 외부 API 클라이언트, 설정값을 봅니다."],
            ["생성자", "Spring이 어떤 의존성을 주입하는지 봅니다."],
            ["public 메서드", "Controller나 다른 서비스가 실제로 호출할 수 있는 기능 입구입니다."],
            ["private 메서드", "public 메서드 내부의 복잡한 로직을 나눠 놓은 보조 함수입니다."],
            ["record", "여러 값을 묶어서 전달하기 위한 내부 데이터 구조입니다."],
        ],
        [2100, 7260],
        font_size=8.1,
    )
    add_paragraph(
        doc,
        "공부할 때는 먼저 public 메서드를 보고, 그 public 메서드가 호출하는 private 메서드를 따라가면 됩니다. private 메서드를 먼저 보면 전체 기능 흐름이 잘 보이지 않습니다.",
        "추천 순서:",
        BLUE,
    )


def import_summary(imports: list[str]) -> list[list[str]]:
    groups = {
        "DTO": 0,
        "Entity": 0,
        "Repository": 0,
        "Service": 0,
        "Spring": 0,
        "Java 표준 라이브러리": 0,
        "외부 라이브러리": 0,
    }
    for imp in imports:
        if ".dto." in imp:
            groups["DTO"] += 1
        elif ".entity." in imp:
            groups["Entity"] += 1
        elif ".repository." in imp:
            groups["Repository"] += 1
        elif ".service." in imp:
            groups["Service"] += 1
        elif "org.springframework" in imp:
            groups["Spring"] += 1
        elif imp.startswith("import java."):
            groups["Java 표준 라이브러리"] += 1
        else:
            groups["외부 라이브러리"] += 1
    return [[k, v, import_group_explain(k)] for k, v in groups.items() if v]


def import_group_explain(group: str) -> str:
    return {
        "DTO": "프론트와 주고받는 데이터 모양입니다.",
        "Entity": "DB 테이블과 가까운 객체입니다.",
        "Repository": "DB 조회/저장 계층입니다.",
        "Service": "다른 업무 로직을 재사용합니다.",
        "Spring": "@Service, @Value, ResponseStatusException 같은 Spring 기능입니다.",
        "Java 표준 라이브러리": "List, Map, LocalDateTime, Optional 같은 기본 도구입니다.",
        "외부 라이브러리": "Jackson, SMS SDK, HTTP 클라이언트 등 외부 라이브러리입니다.",
    }[group]


def field_explain(field: str) -> str:
    lower = field.lower()
    if "static final" in lower:
        return "고정 상수입니다. 반경, 상태 목록, logger처럼 여러 곳에서 반복해서 쓰는 값을 한 곳에 둡니다."
    if "repository" in lower:
        return "DB 접근 객체입니다. 이 필드를 통해 Entity를 조회하거나 저장합니다."
    if "service" in lower:
        return "다른 서비스의 업무 로직을 재사용하기 위해 주입받습니다."
    if "@value" in lower or "enabled" in lower or "timeout" in lower or "command" in lower or "path" in lower:
        return "환경 설정값입니다. 로컬/배포 환경에 따라 값이 달라질 수 있습니다."
    if "client" in lower or "sender" in lower:
        return "외부 API 또는 SMS 발송을 담당하는 객체입니다."
    if "list<" in lower or "map<" in lower:
        return "여러 값을 메모리에 보관하는 컬렉션입니다."
    if "logger" in lower or " log " in lower:
        return "서버 로그를 남기기 위한 객체입니다."
    return "이 서비스가 기능을 처리하기 위해 내부에 보관하는 값 또는 의존성입니다."


def explain_block(block: CodeBlock, src: str) -> list[list[str]]:
    rows: list[list[str]] = [
        ["종류", "생성자" if block.kind == "constructor" else "메서드"],
        ["접근 제한자", visibility_explain(block.visibility)],
        ["입력값", params_explain(block.params)],
        ["반환값", return_explain(block.return_type)],
        ["코드 위치", f"{block.start_line}~{block.end_line}행"],
    ]
    rows.extend(pattern_rows(block.body, src))
    return rows


def visibility_explain(vis: str) -> str:
    if vis == "public":
        return "외부 Controller나 다른 서비스가 호출할 수 있는 공개 기능입니다."
    if vis == "private":
        return "이 클래스 내부에서만 쓰는 보조 로직입니다. public 메서드를 읽다가 필요할 때 따라가면 됩니다."
    if vis == "protected":
        return "상속 관계에서 접근 가능한 범위입니다. 일반 public보다는 제한적입니다."
    return "접근 범위를 나타냅니다."


def params_explain(params: str) -> str:
    if not params.strip():
        return "입력값이 없습니다. 내부 상태나 Repository 조회 결과를 기준으로 처리합니다."
    parts = [p.strip() for p in params.split(",") if p.strip()]
    explained = []
    for part in parts:
        tokens = part.split()
        if len(tokens) >= 2:
            name = tokens[-1]
            typ = " ".join(tokens[:-1])
            explained.append(f"{name}: {param_type_explain(typ)}")
        else:
            explained.append(part)
    return " / ".join(explained)


def param_type_explain(typ: str) -> str:
    if "Dto" in typ or "Request" in typ:
        return f"{typ} 형태의 요청 데이터입니다."
    if typ in {"Long", "long", "Integer", "int"} or "Id" in typ:
        return f"{typ} 숫자 값입니다. 보통 id나 개수를 의미합니다."
    if "String" in typ:
        return "문자열 값입니다."
    if "MultipartFile" in typ:
        return "프론트 FormData로 업로드된 파일입니다."
    if "List" in typ:
        return "여러 개의 값을 담은 목록입니다."
    if "Map" in typ:
        return "key-value 형태의 묶음 데이터입니다."
    if "LocalDateTime" in typ:
        return "날짜와 시간을 나타냅니다."
    return f"{typ} 타입의 입력입니다."


def return_explain(ret: str) -> str:
    if ret == "생성자":
        return "객체를 만들 때 의존성을 주입받습니다. 별도 값을 반환하지 않습니다."
    if ret == "void":
        return "반환값 없이 저장, 삭제, 발송, 이벤트 발행 같은 부수 효과를 수행합니다."
    if "Dto" in ret or "Response" in ret:
        return f"{ret}를 반환합니다. 프론트가 받기 좋은 응답 데이터입니다."
    if "List" in ret:
        return f"{ret}를 반환합니다. 여러 개의 결과를 목록으로 돌려줍니다."
    if "Map" in ret:
        return f"{ret}를 반환합니다. id나 이름을 key로 결과를 빠르게 찾기 위한 구조입니다."
    if "Optional" in ret:
        return f"{ret}를 반환합니다. 값이 있을 수도 있고 없을 수도 있음을 표현합니다."
    if ret in {"boolean", "Boolean"}:
        return "참/거짓 판단 결과를 반환합니다."
    if ret in {"int", "long", "double", "Integer", "Long", "Double"}:
        return "계산된 숫자 값을 반환합니다."
    if ret == "String":
        return "문자열 결과를 반환합니다."
    return f"{ret} 타입을 반환합니다."


def pattern_rows(body: str, src: str) -> list[list[str]]:
    rows: list[list[str]] = []
    checks = [
        ("findById", "id로 DB에서 하나의 Entity를 조회합니다. 보통 없으면 orElseThrow로 예외 처리합니다."),
        ("findAll", "DB에서 전체 목록을 가져옵니다. 이후 정렬하거나 DTO로 변환하는 경우가 많습니다."),
        ("save(", "DB에 새 데이터를 저장하거나 기존 데이터를 수정합니다."),
        ("deleteById", "id 기준으로 DB 데이터를 삭제합니다."),
        ("orElseThrow", "조회 결과가 없을 때 null로 진행하지 않고 즉시 예외를 던집니다."),
        ("ResponseStatusException", "서비스에서 HTTP 오류 상태를 명확히 반환하기 위한 예외입니다."),
        ("stream()", "목록을 필터링, 정렬, 변환, 집계하기 위한 Java Stream 문법입니다."),
        (".map(", "각 항목을 다른 형태로 바꿉니다. Entity를 DTO로 바꿀 때 자주 나옵니다."),
        (".filter(", "조건에 맞는 항목만 남깁니다."),
        (".sorted(", "목록을 정렬합니다."),
        (".toList()", "Stream 처리 결과를 List로 확정합니다."),
        ("@Transactional", "여러 DB 작업을 하나의 트랜잭션으로 묶습니다."),
        ("MultipartFile", "파일 업로드 입력입니다. 프론트 FormData와 연결됩니다."),
        ("SseEmitter", "SSE 실시간 연결 객체입니다. EventSource와 연결됩니다."),
        ("ProcessBuilder", "Java에서 Python 추론 스크립트 같은 외부 프로세스를 실행합니다."),
        ("ObjectMapper", "Java 객체와 JSON 문자열을 서로 변환합니다."),
        ("HttpClient", "외부 HTTP API를 호출합니다."),
        ("RestClient", "외부 HTTP API를 호출합니다."),
        ("SmsSender", "SMS 발송 인터페이스입니다. 실제 업체 구현체를 교체할 수 있습니다."),
        ("Files.", "파일 시스템에 임시 파일을 만들거나 읽고 쓰는 코드입니다."),
        ("Math.", "점수, 거리, 비율, 상한/하한 같은 숫자 계산입니다."),
        ("Duration.", "시간 차이를 계산합니다."),
        ("LocalDateTime", "현재 시각 또는 특정 시각을 다룹니다."),
        ("switch", "값에 따라 여러 분기 중 하나를 선택합니다."),
        ("try", "실패 가능성이 있는 코드를 감싸고 예외를 처리합니다."),
        ("catch", "예외가 발생했을 때 fallback, 로그, 빈 결과 반환 등을 수행합니다."),
        ("return Map.of()", "실패하거나 데이터가 없을 때 빈 Map을 반환해 호출자가 안전하게 처리하게 합니다."),
        ("return List.of()", "실패하거나 데이터가 없을 때 빈 List를 반환합니다."),
    ]
    for keyword, explanation in checks:
        if keyword in body:
            rows.append([keyword, explanation])
    if not rows:
        rows.append(["기본 흐름", "입력값을 받아 내부 계산 또는 단순 반환을 수행하는 보조 코드입니다."])
    return rows[:12]


def method_plain_summary(block: CodeBlock) -> str:
    name = block.name
    lower = name.lower()
    if block.kind == "constructor":
        return "이 생성자는 Spring이 이 서비스를 만들 때 필요한 의존성을 주입받는 부분입니다. Controller나 다른 서비스에서 직접 new로 만들기보다 Spring 컨테이너가 관리합니다."
    if lower.startswith("get") or lower.startswith("find") or lower.startswith("list"):
        return "이 메서드는 조회 기능입니다. DB나 내부 상태에서 데이터를 가져와 화면 또는 다른 서비스가 쓰기 좋은 형태로 반환합니다."
    if lower.startswith("create") or lower.startswith("save") or lower.startswith("add"):
        return "이 메서드는 생성/저장 기능입니다. 입력 DTO나 값을 Entity로 바꾸고 Repository save 흐름으로 이어지는 경우가 많습니다."
    if lower.startswith("update") or lower.startswith("change") or lower.startswith("set"):
        return "이 메서드는 수정 기능입니다. 기존 데이터를 조회한 뒤 입력된 값으로 상태를 바꾸고 저장합니다."
    if lower.startswith("delete") or lower.startswith("remove"):
        return "이 메서드는 삭제 기능입니다. 먼저 대상이 존재하는지 확인하고 삭제합니다."
    if lower.startswith("send") or lower.startswith("notify"):
        return "이 메서드는 알림/문자/외부 전송 기능입니다. 외부 API 실패 가능성을 함께 봐야 합니다."
    if lower.startswith("publish"):
        return "이 메서드는 실시간 이벤트 발행 기능입니다. 연결된 프론트 EventSource들에게 변경 내용을 전달합니다."
    if lower.startswith("subscribe"):
        return "이 메서드는 프론트가 SSE 연결을 구독하도록 emitter를 만들어 보관합니다."
    if lower.startswith("predict") or lower.startswith("analyze") or lower.startswith("guide"):
        return "이 메서드는 AI 또는 분석 결과를 만드는 기능입니다. 입력 feature, fallback, DTO 변환 흐름을 함께 봐야 합니다."
    if lower.startswith("login"):
        return "이 메서드는 인증 기능입니다. 사용자 정보를 검증하고 토큰이나 세션을 반환합니다."
    if block.visibility == "private":
        return "이 메서드는 외부에서 직접 호출하는 기능이 아니라, 같은 클래스 안의 public 메서드가 복잡한 로직을 나누기 위해 호출하는 보조 함수입니다."
    return "이 메서드는 서비스의 업무 흐름 중 하나를 담당합니다. 입력값, Repository 호출, 반환값을 순서대로 읽으면 됩니다."


def add_file_section(doc: Document, path: Path, index: int) -> int:
    src = read_source(path)
    cls = class_name(path)
    imports, anns, fields, package = top_level_lines(src, cls)
    blocks = parse_methods(src, cls)
    records = nested_records(src)

    doc.add_heading(f"2.{index}. {cls}", level=2)
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["파일 위치", str(path.relative_to(ROOT))],
            ["전체 줄 수", len(src.splitlines())],
            ["package", package or "package 선언 없음"],
            ["annotation", ", ".join(anns) if anns else "없음"],
            ["필드/상수 수", len(fields)],
            ["생성자/메서드 수", len(blocks)],
            ["내부 record 수", len(records)],
        ],
        [1900, 7460],
        font_size=7.8,
    )

    doc.add_heading("파일 맨 위 구조: package/import/annotation", level=3)
    add_paragraph(doc, "package는 이 클래스의 주소입니다. import는 이 파일이 사용하려는 외부 클래스 목록입니다. annotation은 Spring이 이 클래스를 어떻게 관리할지 알려주는 표시입니다.")
    if imports:
        add_table(doc, ["import 그룹", "개수", "쉽게 말하면"], import_summary(imports), [2100, 800, 6460], font_size=7.5)
    else:
        add_paragraph(doc, "import가 거의 없는 단순 파일입니다.", "import:")

    doc.add_heading("필드와 상수 전체 설명", level=3)
    if fields:
        add_table(doc, ["필드/상수 코드", "의미"], [[field, field_explain(field)] for field in fields], [4700, 4660], font_size=7.0)
    else:
        add_paragraph(doc, "이 파일은 별도 필드나 상수가 거의 없습니다. 인터페이스이거나 단순 구현체일 가능성이 있습니다.", "필드:")

    doc.add_heading("생성자와 모든 메서드", level=3)
    if not blocks:
        add_paragraph(doc, "추출된 메서드가 없습니다. 인터페이스나 단순 타입 정의 파일일 수 있습니다.")
    for method_index, block in enumerate(blocks, start=1):
        add_method_section(doc, cls, method_index, block, src)

    if records:
        doc.add_heading("내부 record 전체 설명", level=3)
        for idx, record_code in enumerate(records, start=1):
            add_code_block(doc, record_code)
            add_paragraph(
                doc,
                "record는 여러 값을 한 묶음으로 전달하기 위한 간단한 데이터 타입입니다. 보통 private helper 메서드들이 계산한 결과를 담아 같은 클래스 내부에서 사용합니다.",
                f"record {idx}:",
            )
    return len(blocks)


def add_method_section(doc: Document, cls: str, method_index: int, block: CodeBlock, src: str) -> None:
    doc.add_heading(f"{cls} - {method_index}. {block.name}", level=3)
    add_paragraph(doc, method_plain_summary(block), "이 코드의 역할:")
    add_table(doc, ["구분", "설명"], explain_block(block, src), [2200, 7160], font_size=7.4)
    add_code_block(doc, block.body)
    add_paragraph(
        doc,
        "위 코드는 먼저 입력값과 필요한 내부 필드를 사용해 조건을 확인하고, Repository나 다른 Service를 호출한 뒤, 결과를 반환하거나 상태를 변경합니다. private 메서드라면 public 메서드의 일부 계산을 분리한 것입니다.",
        "읽는 순서:",
    )


def add_final_index(doc: Document, files: list[Path]) -> None:
    doc.add_heading("3. 전체 Service 공부 순서", level=1)
    add_table(
        doc,
        ["순서", "먼저 볼 파일", "이유"],
        [
            [1, "BoothService", "부스 조회, DTO 변환, GPS 기반 혼잡도 계산, 예약 좌석 요약이 모두 들어 있습니다."],
            [2, "ReservationService / ReservationAuthService", "예약 상태 전이와 사용자 인증 토큰 흐름을 이해할 수 있습니다."],
            [3, "StreamService", "SSE 실시간 연결과 publish 구조를 이해할 수 있습니다."],
            [4, "AiCongestionService / PythonCongestionModelService", "현재 실제 AI 혼잡도 예측 기능의 핵심입니다."],
            [5, "UploadStorageService / AdminImportService", "FormData, MultipartFile, CSV 읽기처럼 JSON이 아닌 요청을 이해할 수 있습니다."],
            [6, "SMS/Chat/Translate/Ops AI 계열", "외부 API 호출, timeout, fallback, 설정값 주입을 이해할 수 있습니다."],
        ],
        [800, 3000, 5560],
        font_size=7.8,
    )
    rows = []
    for idx, path in enumerate(files, start=1):
        src = read_source(path)
        rows.append([idx, class_name(path), len(src.splitlines()), len(parse_methods(src, class_name(path)))])
    add_table(doc, ["번호", "파일", "줄 수", "메서드 수"], rows, [700, 4200, 1200, 1200], font_size=7.2)


def build_doc() -> None:
    files = service_files()
    total_methods = sum(len(parse_methods(read_source(path), class_name(path))) for path in files)
    doc = Document()
    configure_doc(doc)
    add_cover(doc, files, total_methods)
    add_global_guide(doc)
    doc.add_heading("2. Service 파일별 전체 코드 해설", level=1)
    counted = 0
    for idx, path in enumerate(files, start=1):
        counted += add_file_section(doc, path, idx)
    add_final_index(doc, files)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")
    print(f"services: {len(files)}")
    print(f"methods: {counted}")


if __name__ == "__main__":
    build_doc()
