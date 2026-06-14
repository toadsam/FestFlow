from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SERVICE_DIR = ROOT / "backend" / "src" / "main" / "java" / "com" / "festflow" / "backend" / "service"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "페스트플로우_백엔드_Service_코드_분석서.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F4F6F9"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"


@dataclass
class JavaServiceInfo:
    name: str
    package_group: str
    path: Path
    kind: str
    annotations: list[str]
    dependencies: list[str]
    values: list[str]
    methods: list[tuple[str, str, str, str]]
    line_count: int


GROUP_ORDER = [
    "core",
    "analytics",
    "ai",
    "reservation",
    "operations",
    "staff",
    "infra",
    "sms",
]


SERVICE_PROFILES = {
    "AdminActionService": {
        "group": "operations",
        "role": "관리자가 버튼 한 번으로 혼잡 완화 공지나 공연 시작 공지를 만들 수 있게 돕는 서비스입니다.",
        "beginner": "컨트롤러가 직접 공지 내용을 조립하지 않고, 이 서비스가 BoothService/EventService/NoticeService를 조합해서 운영 액션을 완성합니다.",
        "flow": "가장 혼잡한 부스를 찾거나 특정 공연 정보를 조회한 뒤, NoticeService.createNotice()로 실제 공지를 저장합니다.",
    },
    "AdminDashboardService": {
        "group": "operations",
        "role": "관리자 메인 대시보드의 KPI 숫자를 계산합니다.",
        "beginner": "KPI는 오늘 방문자 수, 현재 혼잡도, 등록된 부스/공연 개수처럼 운영자가 한눈에 보는 지표입니다.",
        "flow": "GpsLogRepository로 오늘 방문 로그를 세고, BoothService와 EventService에서 현황을 가져와 AdminDashboardKpiDto로 묶습니다.",
    },
    "AdminImportService": {
        "group": "operations",
        "role": "관리자가 업로드한 CSV 파일을 읽어 부스와 공연 데이터를 대량 등록합니다.",
        "beginner": "MultipartFile은 사용자가 업로드한 파일입니다. 이 서비스는 CSV 한 줄을 DTO로 바꿔 기존 생성 서비스를 호출합니다.",
        "flow": "CSV를 줄 단위로 읽고 BoothUpsertRequestDto 또는 EventUpsertRequestDto를 만들어 BoothService/EventService에 넘깁니다.",
    },
    "AiCongestionService": {
        "group": "ai",
        "role": "현재 축제 상태를 분석해 부스별 혼잡 예측, 추천 행동, AI 판단 로그를 만듭니다.",
        "beginner": "이 서비스는 AI 혼잡도 기능의 중심입니다. 실시간 스냅샷을 만들고 Python 모델 결과와 규칙 기반 fallback을 함께 사용합니다.",
        "flow": "FestivalSnapshotService에서 현재 상태를 가져오고, PythonCongestionModelService로 RandomForest 예측을 요청한 뒤, 추천/위험 점수/근거를 DTO로 반환합니다.",
    },
    "AiDecisionLogService": {
        "group": "ai",
        "role": "AI가 어떤 판단을 했는지 최근 기록을 메모리 목록에 저장합니다.",
        "beginner": "DB에 저장하는 로그가 아니라 서버 메모리에 최근 판단 기록을 보관하는 가벼운 로그 서비스입니다.",
        "flow": "record()로 새 로그를 앞에 넣고, recent()로 최근 로그를 화면에 보여줄 DTO 목록으로 바꿉니다.",
    },
    "AiImageGenerationService": {
        "group": "ai",
        "role": "AI 매칭 프로필 이미지를 검증하고 웹툰풍 이미지로 생성하는 OpenAI 연동 서비스입니다.",
        "beginner": "RestClient로 외부 OpenAI API를 호출하고, 결과 이미지 bytes를 UploadStorageService로 저장합니다.",
        "flow": "원본 이미지를 읽고, 필요하면 사진 검증을 수행한 뒤, 이미지 생성 API 응답의 base64 데이터를 파일/S3에 저장합니다.",
    },
    "AiMatchService": {
        "group": "ai",
        "role": "AI 매칭 기능의 프로필, 좋아요, 신청, 수락/거절, 만남 제안을 처리하는 가장 큰 도메인 서비스입니다.",
        "beginner": "프로필 저장, 이미지 생성, 전화번호 중복 확인, 요청 상태 변경처럼 AI 매칭 화면에서 일어나는 대부분의 업무 규칙이 여기에 있습니다.",
        "flow": "Repository로 프로필/신청/좋아요/전화번호 사용 기록을 조회하고, 상태 검증 후 Entity를 저장하며, 필요하면 이미지/SMS 서비스를 호출합니다.",
    },
    "AuditLogService": {
        "group": "operations",
        "role": "관리자 작업 로그를 저장하고 최근 로그를 조회합니다.",
        "beginner": "누가 어떤 관리자 작업을 했는지 남겨두는 서비스입니다. 운영 추적과 문제 확인에 필요합니다.",
        "flow": "log()에서 AuditLog Entity를 저장하고, getRecentLogs()에서 최근 로그를 DTO로 변환합니다.",
    },
    "AnalyticsService": {
        "group": "analytics",
        "role": "방문자 수, 인기 부스, 혼잡 히트맵, 무대 혼잡도, 분석 대시보드 데이터를 계산합니다.",
        "beginner": "프론트의 /analytics 화면에서 보는 차트와 수치 대부분이 이 서비스에서 계산됩니다.",
        "flow": "GPS 로그와 부스 정보를 시간/위치 기준으로 집계해서 TrafficHourlyDto, HeatPointDto, AnalyticsDashboardDto 같은 응답을 만듭니다.",
    },
    "AuthService": {
        "group": "core",
        "role": "관리자 로그인과 JWT 토큰 발급을 담당합니다.",
        "beginner": "사용자가 보낸 아이디/비밀번호를 확인하고, 맞으면 이후 요청에 사용할 토큰을 발급합니다.",
        "flow": "AdminUserRepository에서 계정을 찾고 PasswordEncoder로 비밀번호를 검증한 뒤 JwtService로 토큰을 만듭니다.",
    },
    "BoothService": {
        "group": "core",
        "role": "부스 목록, 상세, 생성, 수정, 삭제, 실시간 운영 상태, 혼잡도 계산을 담당합니다.",
        "beginner": "부스 도메인의 중심 서비스입니다. Repository에서 Booth를 가져오고 DTO로 바꿔 컨트롤러에 돌려줍니다.",
        "flow": "부스 Entity를 저장/수정하고, 위치 로그/대기 시간/운영 상태를 바탕으로 CongestionResponseDto를 계산합니다.",
    },
    "ChatService": {
        "group": "ai",
        "role": "축제 안내 챗봇 응답을 생성합니다.",
        "beginner": "외부 AI API가 설정돼 있으면 AI 응답을 만들고, 없거나 실패하면 로컬 fallback 응답을 제공합니다.",
        "flow": "질문을 정리하고 축제 맥락 prompt를 만든 뒤 RestClient로 AI API를 호출하거나 규칙 기반 답변으로 대체합니다.",
    },
    "EventService": {
        "group": "core",
        "role": "공연 목록, 상세, 생성, 수정, 삭제와 공연 변경 SSE 알림을 담당합니다.",
        "beginner": "공연 Entity를 관리하는 기본 CRUD 서비스입니다.",
        "flow": "EventRepository로 DB를 조작하고, 변경 후 StreamService.publishEvents()로 프론트 화면 갱신 이벤트를 보냅니다.",
    },
    "FestivalSnapshotService": {
        "group": "ai",
        "role": "AI와 운영 분석에 필요한 현재 축제 상태를 한 번에 모읍니다.",
        "beginner": "여러 Repository를 매번 따로 조회하면 복잡하므로, 부스/예약/공연/GPS 정보를 하나의 스냅샷 record로 묶습니다.",
        "flow": "현재 부스, 예약 통계, GPS 로그, 공연 정보를 읽어 FestivalSnapshot record에 담습니다.",
    },
    "GpsService": {
        "group": "core",
        "role": "사용자 GPS 위치 로그를 저장하고 혼잡도 SSE 이벤트를 발행합니다.",
        "beginner": "프론트가 위치를 보내면 이 서비스가 DB에 저장하고, 혼잡도 화면이 갱신되도록 이벤트를 보냅니다.",
        "flow": "GpsLogRepository.save()로 위치를 저장한 뒤 BoothService.getAllCongestions() 결과를 StreamService로 발행합니다.",
    },
    "LostItemService": {
        "group": "staff",
        "role": "분실물 목록, 등록, 상태 변경, 수정, 찾은 사람 청구, 삭제를 처리합니다.",
        "beginner": "분실물 게시판의 핵심 업무 규칙을 담당합니다. 연락처 마스킹 같은 개인정보 처리도 포함됩니다.",
        "flow": "LostItemRepository로 분실물 Entity를 조회/저장하고 상태에 맞게 LostItemResponseDto를 만듭니다.",
    },
    "NoticeService": {
        "group": "core",
        "role": "공지 목록, 생성, 수정, 삭제, 활성 공지 SSE 발행을 담당합니다.",
        "beginner": "관리자가 만든 공지를 사용자 화면에 보여주기 위한 서비스입니다.",
        "flow": "NoticeRepository로 공지를 관리하고, 변경 후 StreamService.publishNotices()로 실시간 갱신을 보냅니다.",
    },
    "OpsAiService": {
        "group": "ai",
        "role": "운영자와 스태프를 위한 AI 브리핑, 공지 초안, 현장 체크리스트, 분실물 응대 문구를 생성합니다.",
        "beginner": "운영 데이터와 축제 상황을 prompt로 정리해 AI에게 보내고, 실패 시 fallback 문구를 제공합니다.",
        "flow": "관리/스태프 서비스에서 현장 데이터를 읽고 RestClient로 AI API를 호출해 AiAssistResponseDto를 반환합니다.",
    },
    "PublicAiGuideService": {
        "group": "ai",
        "role": "일반 방문자용 AI 방문 가이드를 생성합니다.",
        "beginner": "혼잡도, 공연, 부스 현황을 바탕으로 방문자에게 어디를 피하고 어디로 가면 좋은지 안내합니다.",
        "flow": "AnalyticsService, AiCongestionService, EventService 데이터를 prompt로 묶어 AI 응답을 만들거나 fallback guide를 반환합니다.",
    },
    "PythonCongestionModelService": {
        "group": "ai",
        "role": "Java 서버에서 Python RandomForest 모델 추론 스크립트를 실행합니다.",
        "beginner": "Java가 직접 pkl 모델을 읽지 않고, Python 프로세스를 실행해 JSON 파일로 입력/출력을 주고받습니다.",
        "flow": "ProcessBuilder로 predict_congestion.py를 실행하고, output JSON을 AiModelPredictionDto로 변환합니다.",
    },
    "ReservationAuthService": {
        "group": "reservation",
        "role": "예약 사용자의 전화번호 인증번호 발송, 검증, 인증 토큰 관리를 담당합니다.",
        "beginner": "예약을 아무나 수정하지 못하게 전화번호 기반 인증을 제공합니다.",
        "flow": "인증 코드를 만들고 SmsSender로 보내며, 성공 시 ReservationAuthSession을 저장하고 토큰을 발급합니다.",
    },
    "ReservationService": {
        "group": "reservation",
        "role": "부스 예약 설정, 예약 생성, 체크인, 완료, 테이블 해제를 처리합니다.",
        "beginner": "예약 기능의 중심 서비스입니다. 좌석/테이블 수, 예약 상태, 체크인 토큰 같은 규칙을 관리합니다.",
        "flow": "BoothReservationRepository와 BoothReservationTableRepository를 사용해 예약 상태를 바꾸고, StreamService로 예약 갱신을 알립니다.",
    },
    "SimulationService": {
        "group": "operations",
        "role": "운영 시뮬레이션 시작/정지/리셋/시나리오 적용과 주기적 tick 처리를 담당합니다.",
        "beginner": "실제 데이터가 부족할 때 운영 화면에서 혼잡도 변화 시나리오를 흉내 내기 위한 서비스입니다.",
        "flow": "SimulationStateService가 계산한 변화량을 BoothService.updateLiveStatus()에 반영하고 SSE로 갱신을 보냅니다.",
    },
    "SimulationStateService": {
        "group": "operations",
        "role": "시뮬레이션의 현재 상태, 원본 상태, 시나리오별 변화량을 메모리에 관리합니다.",
        "beginner": "DB에 저장하기보다 서버 메모리에 시뮬레이션 상태를 들고 있으며, synchronized로 동시 접근을 막습니다.",
        "flow": "start/stop/clear/tick 메서드로 상태를 변경하고, boothSnapshots와 stageSnapshot으로 화면 표시용 상태를 제공합니다.",
    },
    "StaffService": {
        "group": "staff",
        "role": "스태프 로그인, 대시보드, 상태 업데이트, 관리자 수정, 로그아웃, 토큰 인증을 처리합니다.",
        "beginner": "스태프 앱에서 누가 로그인했는지, 현재 근무 상태가 무엇인지 관리합니다.",
        "flow": "StaffMemberRepository와 StaffSessionRepository로 직원과 세션을 관리하고 StreamService로 상태 변경을 알립니다.",
    },
    "TranslateMetricsService": {
        "group": "infra",
        "role": "번역 성공/실패 횟수와 지연 시간을 메모리에 누적합니다.",
        "beginner": "번역 API가 얼마나 자주 성공하고 실패했는지 운영 지표로 보여주기 위한 서비스입니다.",
        "flow": "AtomicLong으로 카운터를 증가시키고 snapshot()에서 TranslateMetricsDto를 반환합니다.",
    },
    "TranslateService": {
        "group": "infra",
        "role": "Google 번역 엔드포인트를 호출해 텍스트를 번역합니다.",
        "beginner": "Java HttpClient로 외부 번역 API에 요청을 보내고 응답 JSON에서 번역문을 꺼냅니다.",
        "flow": "TranslateRequestDto를 검증하고 URI를 만든 뒤 HttpClient.send()로 응답을 받아 TranslateResponseDto로 반환합니다.",
    },
    "UploadStorageService": {
        "group": "infra",
        "role": "이미지 파일을 로컬 디스크 또는 S3에 저장/조회/삭제합니다.",
        "beginner": "업로드 파일 저장 방식을 한곳에 모아두면 AI 이미지, 부스 이미지, 메뉴 이미지가 같은 저장 로직을 공유할 수 있습니다.",
        "flow": "storageType 설정에 따라 local 또는 s3 분기를 타고, 저장 후 프론트가 접근할 URL을 반환합니다.",
    },
    "StreamService": {
        "group": "infra",
        "role": "SSE 연결을 관리하고 혼잡도/공연/공지/부스/스태프/분실물/예약 이벤트를 발행합니다.",
        "beginner": "SSE는 서버가 브라우저에게 단방향으로 이벤트를 밀어주는 방식입니다. 프론트의 EventSource와 연결됩니다.",
        "flow": "subscribe 메서드는 SseEmitter를 목록에 등록하고, publish 메서드는 해당 목록의 모든 연결에 eventName과 payload를 전송합니다.",
    },
    "AiMatchSmsNotifier": {
        "group": "sms",
        "role": "AI 매칭 신청 생성/수락 시 SMS 알림을 보냅니다.",
        "beginner": "매칭 기능 전용 알림 서비스입니다. 실제 발송은 SolapiMessageClient에 위임합니다.",
        "flow": "기능이 켜져 있고 SMS 클라이언트가 설정돼 있으면 정해진 안내 문구를 전송합니다.",
    },
    "AligoSmsSender": {
        "group": "sms",
        "role": "Aligo API로 인증번호 SMS를 발송하는 구현체입니다.",
        "beginner": "SmsSender 인터페이스의 구현체 중 하나이며, HTTP form 요청으로 문자 발송 API를 호출합니다.",
        "flow": "apiKey/userId/sender 설정을 사용해 요청 body를 만들고 HttpClient로 Aligo 서버에 전송합니다.",
    },
    "NoopSmsSender": {
        "group": "sms",
        "role": "실제 SMS를 보내지 않고 로그만 남기는 대체 구현체입니다.",
        "beginner": "개발/데모 환경에서 문자 비용 없이 인증 흐름을 테스트할 때 쓰기 좋습니다.",
        "flow": "sendVerificationCode()가 호출되면 실제 발송 대신 로그에 인증번호를 기록합니다.",
    },
    "SmsSender": {
        "group": "sms",
        "role": "인증번호 SMS 발송 기능의 공통 인터페이스입니다.",
        "beginner": "인터페이스는 '이런 메서드를 반드시 가져야 한다'는 약속입니다. 구현체를 바꿔도 ReservationAuthService 코드는 그대로 유지됩니다.",
        "flow": "sendVerificationCode(phoneNumber, code) 메서드 하나를 정의합니다.",
    },
    "SolapiMessageClient": {
        "group": "sms",
        "role": "Solapi SDK를 직접 감싸는 낮은 수준의 문자 발송 클라이언트입니다.",
        "beginner": "Solapi 설정이 올바른지 확인하고, 실제 SDK 객체를 사용해 메시지를 보냅니다.",
        "flow": "provider/apiKey/apiSecret/fromNumber가 설정되면 DefaultMessageService를 만들고 sendText()에서 문자 한 건을 발송합니다.",
    },
    "SolapiSmsSender": {
        "group": "sms",
        "role": "SolapiMessageClient를 이용해 예약 인증번호 문자를 보내는 SmsSender 구현체입니다.",
        "beginner": "ReservationAuthService 입장에서는 SmsSender만 알면 되며, 실제 구현은 이 클래스가 담당할 수 있습니다.",
        "flow": "인증번호 문구를 만들고 SolapiMessageClient.sendText()에 넘깁니다.",
    },
    "TwilioSmsSender": {
        "group": "sms",
        "role": "Twilio API로 인증번호 SMS를 보내는 구현체입니다.",
        "beginner": "다른 SMS 업체를 쓰기 위한 대안 구현체입니다. 같은 SmsSender 인터페이스를 구현합니다.",
        "flow": "Twilio.init()으로 SDK를 설정하고 Message.creator()로 문자를 발송합니다.",
    },
}

GROUP_LABELS = {
    "core": "핵심 도메인 서비스",
    "analytics": "분석/혼잡도 서비스",
    "ai": "AI 기능 서비스",
    "reservation": "예약 서비스",
    "operations": "운영/관리 서비스",
    "staff": "스태프/분실물 서비스",
    "infra": "인프라/공통 서비스",
    "sms": "SMS/알림 서비스",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.2) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 10, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=6, line=1.22)
    add_run(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4, line=1.18)
    add_run(p, text)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], *, font_size: float = 8.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        set_cell_shading(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.05)
        add_run(p, str(header), bold=True, color=DARK, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.08)
            add_run(p, str(value), size=font_size)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.16)
    add_run(p, title + " | ", bold=True, color=DARK, size=9)
    add_run(p, body, color=MUTED, size=9)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.4)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def clean_type(type_name: str) -> str:
    type_name = type_name.strip()
    type_name = re.sub(r"\s+", " ", type_name)
    return type_name


def strip_comments(source: str) -> str:
    source = re.sub(r"/\*.*?\*/", "", source, flags=re.S)
    source = re.sub(r"//.*", "", source)
    return source


def parse_java_service(path: Path) -> JavaServiceInfo:
    source = path.read_text(encoding="utf-8")
    no_comments = strip_comments(source)
    rel = path.relative_to(SERVICE_DIR)
    name = path.stem
    profile = SERVICE_PROFILES.get(name, {})
    package_group = profile.get("group") or infer_group(rel, no_comments, name)

    if re.search(r"\binterface\s+" + re.escape(name) + r"\b", no_comments):
        kind = "interface"
    elif re.search(r"\brecord\s+" + re.escape(name) + r"\b", no_comments):
        kind = "record"
    else:
        kind = "class"

    annotations = sorted(set(re.findall(r"@(Service|Component|Transactional|PostConstruct|PreDestroy|Scheduled)\b(?:\([^)]*\))?", no_comments)))
    values = re.findall(r'@Value\("([^"]+)"\)', no_comments)
    dependencies = []
    for match in re.finditer(r"private\s+final\s+([A-Za-z0-9_<>, ?]+)\s+([A-Za-z0-9_]+)\s*;", no_comments):
        dependencies.append(clean_type(match.group(1)))
    for match in re.finditer(r"private\s+(?:static\s+final\s+)?([A-Za-z0-9_<>, ?]+)\s+([A-Za-z0-9_]+)\s*;", no_comments):
        type_name = clean_type(match.group(1))
        if type_name.endswith("Repository") or type_name.endswith("Service") or type_name.endswith("Sender") or type_name.endswith("Client"):
            dependencies.append(type_name)
    dependencies = sorted(set(dependencies))

    methods = []
    method_pattern = re.compile(
        r"(?:@Transactional(?:\([^)]*\))?\s*)?(public|private|protected)\s+"
        r"([A-Za-z0-9_<>, ?\[\]]+)\s+([A-Za-z0-9_]+)\s*\(([^)]*)\)",
        re.M,
    )
    for match in method_pattern.finditer(no_comments):
        visibility, return_type, method_name, params = match.groups()
        if method_name == name:
            continue
        if method_name in {"if", "for", "while", "switch"}:
            continue
        line = no_comments[: match.start()].count("\n") + 1
        params_short = summarize_params(params)
        methods.append((method_name, clean_type(return_type), params_short, str(line)))
    return JavaServiceInfo(
        name=name,
        package_group=package_group,
        path=path,
        kind=kind,
        annotations=annotations,
        dependencies=dependencies,
        values=values,
        methods=methods,
        line_count=len(source.splitlines()),
    )


def infer_group(rel: Path, source: str, name: str) -> str:
    rel_text = str(rel).replace("\\", "/")
    if rel_text.startswith("analytics/"):
        return "analytics"
    if rel_text.startswith("sms/") or rel_text.startswith("notification/"):
        return "sms"
    if rel_text.startswith("stream/"):
        return "infra"
    if name.startswith("Ai") or name in {"ChatService", "OpsAiService", "PublicAiGuideService", "PythonCongestionModelService"}:
        return "ai"
    if "Reservation" in name:
        return "reservation"
    if name.startswith("Staff") or name.startswith("LostItem"):
        return "staff"
    if name.startswith("Admin") or name.startswith("Simulation") or name.startswith("Audit"):
        return "operations"
    if name in {"TranslateService", "TranslateMetricsService", "UploadStorageService", "StreamService"}:
        return "infra"
    return "core"


def summarize_params(params: str) -> str:
    params = " ".join(params.strip().split())
    if not params:
        return "없음"
    parts = []
    for raw in params.split(","):
        raw = raw.strip()
        raw = re.sub(r"@\w+(?:\([^)]*\))?\s*", "", raw)
        tokens = raw.split()
        if len(tokens) >= 2:
            parts.append(f"{tokens[-2]} {tokens[-1]}")
        else:
            parts.append(raw)
    return ", ".join(parts[:5]) + (" ..." if len(parts) > 5 else "")


def method_explanation(method_name: str) -> str:
    lower = method_name.lower()
    if lower.startswith("get") or lower.startswith("find") or lower in {"recent", "status", "snapshot", "guide", "dashboard"}:
        return "조회 계열 메서드입니다. DB나 메모리 상태를 읽고 화면 응답 DTO로 변환합니다."
    if lower.startswith("create") or lower.startswith("save") or lower.startswith("import"):
        return "생성 계열 메서드입니다. 요청 DTO를 검증하고 Entity를 만들어 저장합니다."
    if lower.startswith("update") or lower.startswith("patch") or lower.startswith("apply"):
        return "수정 계열 메서드입니다. 기존 Entity나 상태를 찾아 필요한 필드를 바꾼 뒤 저장합니다."
    if lower.startswith("delete") or lower.startswith("purge"):
        return "삭제 계열 메서드입니다. 권한/상태를 확인한 뒤 데이터를 삭제하거나 비활성화합니다."
    if lower.startswith("publish") or lower.startswith("broadcast") or lower.startswith("notify"):
        return "알림 계열 메서드입니다. 공지, SSE, SMS 같은 외부 전달을 담당합니다."
    if lower.startswith("subscribe"):
        return "SSE 구독 메서드입니다. 브라우저가 받을 실시간 연결을 생성합니다."
    if lower.startswith("send"):
        return "외부 전송 메서드입니다. SMS나 HTTP 요청을 실제로 보냅니다."
    if lower.startswith("login") or lower.startswith("authenticate") or lower.startswith("verify") or lower.startswith("require"):
        return "인증 계열 메서드입니다. 토큰, 비밀번호, 인증번호 같은 접근 조건을 확인합니다."
    if lower.startswith("analyze") or lower.startswith("predict") or lower.startswith("translate") or lower.startswith("answer"):
        return "계산/AI 계열 메서드입니다. 입력을 분석하고 모델/API/규칙을 통해 결과를 만듭니다."
    if lower.startswith("start") or lower.startswith("stop") or lower.startswith("reset") or lower.startswith("tick"):
        return "상태 진행 메서드입니다. 시뮬레이션이나 운영 상태의 흐름을 변경합니다."
    return "도메인 내부 로직을 보조하는 메서드입니다. 이름과 반환 타입을 기준으로 처리 목적을 읽으면 됩니다."


def code_snippet_for(info: JavaServiceInfo) -> str | None:
    snippets = {
        "BoothService": """
public BoothResponseDto getBoothById(Long boothId) {
    Booth booth = boothRepository.findById(boothId)
            .orElseThrow(() -> new ResponseStatusException(HttpStatus.NOT_FOUND, "Booth not found"));
    return toResponseDto(booth);
}
""",
        "AiCongestionService": """
Map<Long, AiModelPredictionDto> modelPredictions = modelPredictions(snapshot, eventSoon);
AiModelPredictionDto aiModel = modelPrediction != null
        ? modelPrediction
        : AiModelPredictionDto.fallback(fallbackPredictedLevel, modelFactors, "MODEL_UNAVAILABLE");
""",
        "PythonCongestionModelService": """
ProcessBuilder builder = new ProcessBuilder(
        pythonCommand,
        script.toString(),
        "--model", model.toString(),
        "--input-file", inputFile.toString(),
        "--output-file", outputFile.toString()
);
""",
        "StreamService": """
SseEmitter emitter = new SseEmitter(0L);
emitters.add(emitter);
emitter.onCompletion(() -> emitters.remove(emitter));
emitter.onTimeout(() -> emitters.remove(emitter));
""",
        "ReservationService": """
BoothReservation reservation = reservationRepository.findById(reservationId)
        .orElseThrow(() -> new ResponseStatusException(HttpStatus.NOT_FOUND, "Reservation not found"));
reservation.checkIn();
return toDto(reservation);
""",
        "UploadStorageService": """
if ("s3".equalsIgnoreCase(storageType)) {
    return saveToS3(fileBytes, prefix, extension, contentType);
}
return saveToLocal(fileBytes, prefix, extension);
""",
        "AiMatchService": """
@Transactional
public AiMatchRequestResponseDto acceptRequest(Long requestId, AiMatchProfileAccessRequestDto requestDto) {
    AiMatchRequest request = findRequest(requestId);
    request.accept();
    return toRequestResponseDto(request);
}
""",
        "AnalyticsService": """
public AnalyticsDashboardDto dashboard(int minutesWindow) {
    LocalDateTime since = LocalDateTime.now().minusMinutes(minutesWindow);
    List<GpsLog> recentLogs = gpsLogRepository.findByCreatedAtAfter(since);
    return buildDashboard(recentLogs, minutesWindow);
}
""",
    }
    return snippets.get(info.name)


def parse_all_services() -> list[JavaServiceInfo]:
    files = sorted(SERVICE_DIR.rglob("*.java"))
    infos = [parse_java_service(path) for path in files]
    return sorted(
        infos,
        key=lambda info: (GROUP_ORDER.index(info.package_group) if info.package_group in GROUP_ORDER else 99, info.name),
    )


def add_cover(doc: Document, infos: list[JavaServiceInfo]) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)
    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "백엔드 Service 코드 분석서", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=12, line=1.18)
    add_run(
        subtitle,
        "Spring Boot 서비스 계층을 기초 개념부터 각 클래스 역할, 의존성, 메서드 흐름까지 공부하기 위한 문서",
        color=MUTED,
        size=11,
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["분석 대상", "backend/src/main/java/com/festflow/backend/service 하위 Java 파일 전체"],
            ["포함 파일 수", f"{len(infos)}개"],
            ["문서 관점", "초보자 기준: Service가 무엇인지, Controller/Repository/DTO/Entity와 어떻게 연결되는지부터 설명"],
            ["읽는 방법", "1장 기초 개념 -> 2장 전체 목록 -> 3장 그룹별 상세 분석 -> 4장 자주 나오는 코드 패턴 순서 추천"],
        ],
        [2200, 7160],
        font_size=8.8,
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_foundation(doc: Document) -> None:
    doc.add_heading("1. Service 계층을 이해하기 위한 기초", level=1)
    add_callout(
        doc,
        "한 문장 요약",
        "Service는 Controller가 받은 요청을 실제 업무 규칙으로 처리하는 계층입니다. Repository에서 데이터를 읽고, Entity를 바꾸고, DTO로 응답을 만들어 Controller에 돌려줍니다.",
        GREEN_FILL,
    )
    add_table(
        doc,
        ["개념", "기초 설명", "이 프로젝트에서의 예"],
        [
            ["Controller", "HTTP 요청을 받는 입구입니다.", "AnalyticsController가 /analytics/dashboard 요청을 받음"],
            ["Service", "실제 업무 규칙을 처리합니다.", "AnalyticsService가 GPS 로그를 집계해 대시보드 DTO 생성"],
            ["Repository", "DB 접근을 담당합니다.", "GpsLogRepository가 GPS 로그를 조회"],
            ["Entity", "DB 테이블과 연결되는 객체입니다.", "Booth, Event, GpsLog, Reservation 등"],
            ["DTO", "API 요청/응답 모양입니다.", "BoothResponseDto, AnalyticsDashboardDto 등"],
            ["StreamService", "SSE 실시간 이벤트를 관리합니다.", "부스/공지/혼잡도 변경 시 프론트에 이벤트 발행"],
        ],
        [1700, 3700, 3960],
        font_size=8,
    )
    add_para(doc, "전형적인 흐름은 Controller -> Service -> Repository -> Entity/DB -> DTO -> Controller -> Frontend 순서입니다. Service는 이 흐름의 가운데에서 데이터 조회, 검증, 계산, 저장, 외부 API 호출을 조합합니다.")

    doc.add_heading("1.1 자주 나오는 Spring/Java 문법", level=2)
    add_table(
        doc,
        ["코드 표현", "뜻", "초보자용 해석"],
        [
            ["@Service", "Spring Bean으로 등록", "Spring이 이 클래스를 자동으로 만들어 필요한 곳에 주입할 수 있게 합니다."],
            ["@Component", "일반 Spring Bean으로 등록", "Service라는 이름은 아니지만 Spring이 관리하는 객체입니다. SMS 발송 구현체 등에 쓰입니다."],
            ["@Transactional", "메서드 실행을 하나의 DB 작업 단위로 묶음", "중간에 실패하면 저장 변경을 되돌릴 수 있어 데이터 꼬임을 줄입니다."],
            ["@Transactional(readOnly = true)", "읽기 전용 트랜잭션", "조회만 할 때 성능과 의도를 명확히 합니다."],
            ["private final Repository", "생성자 주입 대상 필드", "이 서비스가 DB 접근 객체를 필요로 한다는 뜻입니다."],
            ["@Value(\"${...}\")", "설정값 주입", "application.properties나 환경변수 값을 코드에 넣습니다."],
            ["orElseThrow()", "Optional 값이 없으면 예외 발생", "ID로 찾은 데이터가 없을 때 404 같은 에러를 내는 패턴입니다."],
            ["stream().map(...).toList()", "목록 변환", "Entity 목록을 DTO 목록으로 바꿀 때 자주 씁니다."],
            ["ResponseStatusException", "HTTP 상태 코드가 있는 예외", "서비스에서 NOT_FOUND, BAD_REQUEST 같은 API 오류를 직접 표현합니다."],
            ["SseEmitter", "SSE 연결 객체", "서버가 브라우저에 실시간 이벤트를 보낼 때 사용하는 Spring 객체입니다."],
            ["RestClient / HttpClient", "외부 API 호출 도구", "OpenAI, 번역, SMS API를 호출할 때 사용합니다."],
            ["synchronized", "동시 접근 제어", "시뮬레이션 상태처럼 여러 요청이 동시에 바꾸면 안 되는 값에 사용합니다."],
        ],
        [2300, 3000, 4060],
        font_size=7.7,
    )
    add_code_block(
        doc,
        """
@Service
public class BoothService {
    private final BoothRepository boothRepository;

    public BoothService(BoothRepository boothRepository) {
        this.boothRepository = boothRepository;
    }
}
""",
    )
    add_para(doc, "위 코드는 가장 기본적인 Service 구조입니다. @Service가 붙으면 Spring이 BoothService 객체를 만들고, 생성자에 필요한 BoothRepository도 자동으로 넣어줍니다. 그래서 Controller는 직접 new BoothService()를 하지 않습니다.")


def add_inventory(doc: Document, infos: list[JavaServiceInfo]) -> None:
    doc.add_heading("2. Service 전체 목록", level=1)
    rows = []
    for info in infos:
        profile = SERVICE_PROFILES.get(info.name, {})
        rel = info.path.relative_to(ROOT)
        rows.append([
            GROUP_LABELS.get(info.package_group, info.package_group),
            info.name,
            info.kind,
            f"{info.line_count}줄",
            profile.get("role", "서비스 계층의 보조 클래스입니다."),
            str(rel),
        ])
    add_table(
        doc,
        ["그룹", "클래스", "종류", "크기", "역할", "파일 위치"],
        rows,
        [1300, 1900, 700, 700, 3300, 1460],
        font_size=6.6,
    )


def add_group_summary(doc: Document, infos: list[JavaServiceInfo]) -> None:
    doc.add_heading("3. 그룹별로 보는 서비스 구조", level=1)
    grouped: dict[str, list[JavaServiceInfo]] = {}
    for info in infos:
        grouped.setdefault(info.package_group, []).append(info)

    for group in GROUP_ORDER:
        items = grouped.get(group, [])
        if not items:
            continue
        doc.add_heading(GROUP_LABELS.get(group, group), level=2)
        summary = group_description(group)
        add_para(doc, summary)
        rows = []
        for info in items:
            profile = SERVICE_PROFILES.get(info.name, {})
            deps = ", ".join(info.dependencies[:4]) if info.dependencies else "없음/직접 외부 API 사용"
            if len(info.dependencies) > 4:
                deps += " ..."
            rows.append([info.name, profile.get("role", "서비스 보조 클래스"), deps])
        add_table(doc, ["서비스", "역할", "주요 의존성"], rows, [2200, 4300, 2860], font_size=7.3)


def group_description(group: str) -> str:
    descriptions = {
        "core": "핵심 도메인 서비스는 부스, 공연, 공지, GPS, 로그인처럼 서비스의 기본 기능을 담당합니다. 대부분 Controller 요청을 받아 Repository를 사용하고 DTO로 응답합니다.",
        "analytics": "분석 서비스는 저장된 로그와 부스 상태를 통계로 바꿉니다. 프론트의 혼잡도 화면과 무대 인원 화면이 이 계층을 사용합니다.",
        "ai": "AI 서비스는 모델 추론, OpenAI 호출, AI 안내문 생성, AI 매칭처럼 외부 모델 또는 학습 모델과 연결된 기능을 담당합니다.",
        "reservation": "예약 서비스는 전화번호 인증, 예약 생성, 체크인, 테이블 상태 관리를 담당합니다. 데이터 일관성이 중요해서 @Transactional이 많이 쓰입니다.",
        "operations": "운영/관리 서비스는 관리자 대시보드, CSV 업로드, 시뮬레이션, 운영 액션처럼 축제를 운영하는 사람을 위한 기능입니다.",
        "staff": "스태프/분실물 서비스는 현장 직원 로그인, 근무 상태, 분실물 등록/처리 같은 현장 운영 기능을 담당합니다.",
        "infra": "인프라/공통 서비스는 파일 저장, 번역, SSE, 메트릭처럼 여러 기능에서 공통으로 쓰이는 기반 기능입니다.",
        "sms": "SMS/알림 서비스는 인증번호와 AI 매칭 알림을 외부 문자 발송 업체에 연결합니다. 인터페이스와 구현체를 분리해 업체를 바꿀 수 있게 했습니다.",
    }
    return descriptions.get(group, "서비스 그룹입니다.")


def add_service_details(doc: Document, infos: list[JavaServiceInfo]) -> None:
    doc.add_heading("4. 서비스별 상세 코드 분석", level=1)
    for index, info in enumerate(infos, start=1):
        profile = SERVICE_PROFILES.get(info.name, {})
        doc.add_heading(f"4.{index}. {info.name}", level=2)
        add_table(
            doc,
            ["항목", "내용"],
            [
                ["파일", str(info.path.relative_to(ROOT))],
                ["그룹", GROUP_LABELS.get(info.package_group, info.package_group)],
                ["종류", info.kind],
                ["라인 수", f"{info.line_count}줄"],
                ["역할", profile.get("role", "서비스 계층의 보조 클래스입니다.")],
                ["초보자 설명", profile.get("beginner", "이 클래스의 public 메서드가 외부에서 호출되는 주요 기능입니다.")],
                ["처리 흐름", profile.get("flow", "의존성을 사용해 데이터를 조회/가공/저장하고 DTO나 상태값을 반환합니다.")],
            ],
            [1900, 7460],
            font_size=7.7,
        )
        if info.dependencies:
            add_table(
                doc,
                ["의존성", "무슨 뜻인가"],
                [[dep, dependency_explanation(dep)] for dep in info.dependencies[:10]],
                [3000, 6360],
                font_size=7.5,
            )
        if info.values:
            add_table(
                doc,
                ["설정값", "의미"],
                [[value, value_explanation(value)] for value in info.values[:12]],
                [4300, 5060],
                font_size=7.3,
            )
        public_methods = [m for m in info.methods if not m[0].startswith("lambda")]
        if public_methods:
            rows = []
            for method_name, return_type, params, line in public_methods[:14]:
                rows.append([method_name, return_type, params, line, method_explanation(method_name)])
            add_table(
                doc,
                ["메서드", "반환", "입력", "라인", "기초 설명"],
                rows,
                [1650, 1550, 2600, 650, 2910],
                font_size=6.8,
            )
        snippet = code_snippet_for(info)
        if snippet:
            doc.add_heading("핵심 코드 예시", level=3)
            add_code_block(doc, snippet)
            add_para(doc, snippet_explanation(info.name))


def dependency_explanation(dep: str) -> str:
    if dep.endswith("Repository"):
        return "DB 테이블에 접근하는 객체입니다. Entity 조회, 저장, 삭제를 담당합니다."
    if dep.endswith("Service"):
        return "다른 업무 로직을 재사용하기 위해 주입된 서비스입니다."
    if dep.endswith("Sender") or dep.endswith("Notifier"):
        return "문자나 알림을 외부로 보내는 역할입니다."
    if dep.endswith("Client") or dep in {"RestClient", "HttpClient", "S3Client"}:
        return "외부 API나 저장소와 통신하는 클라이언트 객체입니다."
    if "PasswordEncoder" in dep:
        return "비밀번호를 평문으로 비교하지 않고 해시 검증하는 Spring Security 도구입니다."
    return "서비스 내부 처리를 위해 주입되거나 보관되는 협력 객체입니다."


def value_explanation(value: str) -> str:
    if "api-key" in value or "secret" in value or "token" in value:
        return "외부 API 인증에 필요한 민감 설정값입니다. 배포 환경변수로 관리해야 합니다."
    if "enabled" in value:
        return "기능을 켜고 끄는 플래그입니다. 개발/배포 환경별로 동작을 바꿀 수 있습니다."
    if "model" in value:
        return "AI 모델명 또는 ML 모델 경로 설정입니다."
    if "upload" in value or "storage" in value or "s3" in value:
        return "파일 저장 위치나 S3 저장소 설정입니다."
    if "sms" in value:
        return "문자 발송 업체 설정입니다."
    return "application.properties 또는 환경변수에서 주입되는 설정값입니다."


def snippet_explanation(name: str) -> str:
    explanations = {
        "BoothService": "ID로 Entity를 찾고, 없으면 예외를 던진 뒤 DTO로 변환하는 가장 기본적인 조회 패턴입니다.",
        "AiCongestionService": "AI 모델 결과가 있으면 그 결과를 쓰고, 모델 호출이 실패하면 fallback 결과를 써서 API가 끊기지 않게 하는 구조입니다.",
        "PythonCongestionModelService": "Java 프로세스 안에서 Python 스크립트를 별도 프로세스로 실행해 .pkl 모델 추론 결과를 JSON으로 받는 부분입니다.",
        "StreamService": "브라우저가 SSE를 구독하면 SseEmitter를 목록에 보관하고, 연결이 끝나면 목록에서 제거합니다.",
        "ReservationService": "예약을 찾고 상태를 바꾼 뒤 DTO로 돌려주는 전형적인 트랜잭션 처리 흐름입니다.",
        "UploadStorageService": "설정에 따라 로컬 저장소와 S3 저장소를 분기하는 공통 파일 저장 패턴입니다.",
        "AiMatchService": "매칭 요청의 상태를 변경하는 도메인 메서드 예시입니다. @Transactional 안에서 상태 변경과 저장이 함께 처리됩니다.",
        "AnalyticsService": "시간 기준으로 GPS 로그를 조회하고 대시보드 응답을 만드는 분석 서비스 흐름입니다.",
    }
    return explanations.get(name, "해당 서비스에서 자주 쓰이는 핵심 처리 흐름입니다.")


def add_pattern_section(doc: Document) -> None:
    doc.add_heading("5. 자주 나오는 코드 패턴 풀이", level=1)
    patterns = [
        [
            "생성자 주입",
            "public BoothService(BoothRepository boothRepository) { this.boothRepository = boothRepository; }",
            "Spring이 BoothRepository 객체를 만들어 생성자에 넣어줍니다. 테스트하기 쉽고, 필요한 의존성이 명확합니다.",
        ],
        [
            "Entity -> DTO 변환",
            "return booths.stream().map(this::toResponseDto).toList();",
            "DB Entity를 그대로 API로 내보내지 않고, 화면에 필요한 응답 모양으로 바꿉니다.",
        ],
        [
            "없으면 404",
            "repository.findById(id).orElseThrow(() -> new ResponseStatusException(HttpStatus.NOT_FOUND));",
            "ID로 찾은 데이터가 없으면 null을 계속 넘기지 않고 즉시 HTTP 에러로 중단합니다.",
        ],
        [
            "트랜잭션",
            "@Transactional public Dto update(...) { entity.change(...); return toDto(entity); }",
            "조회, 수정, 저장이 하나의 작업으로 묶입니다. 중간 실패 시 DB 변경을 되돌릴 수 있습니다.",
        ],
        [
            "SSE 발행",
            "streamService.publishBooths(boothService.getAllBooths());",
            "데이터가 바뀐 직후 프론트가 새 데이터를 받도록 실시간 이벤트를 보냅니다.",
        ],
        [
            "외부 API fallback",
            "try { return callAi(); } catch (...) { return fallback(); }",
            "OpenAI, Python, SMS 같은 외부 시스템이 실패해도 서비스 전체가 멈추지 않도록 대체 응답을 준비합니다.",
        ],
    ]
    add_table(doc, ["패턴", "예시", "의미"], patterns, [1800, 3600, 3960], font_size=7.3)


def add_study_order(doc: Document) -> None:
    doc.add_heading("6. 공부 추천 순서", level=1)
    add_para(doc, "모든 Service를 한 번에 외우려고 하면 어렵습니다. 화면에서 바로 확인되는 흐름부터 보고, 그다음 복잡한 AI/예약/시뮬레이션으로 넘어가는 순서가 좋습니다.")
    for item in [
        "1단계: BoothService, EventService, NoticeService로 기본 CRUD 흐름 이해",
        "2단계: GpsService, AnalyticsService, StreamService로 혼잡도와 SSE 흐름 이해",
        "3단계: ReservationAuthService, ReservationService로 트랜잭션과 상태 변경 이해",
        "4단계: AiCongestionService, PythonCongestionModelService로 AI 모델 추론 흐름 이해",
        "5단계: AiMatchService, AiImageGenerationService, OpsAiService로 큰 서비스의 구조 읽기",
        "6단계: UploadStorageService, SMS 서비스, TranslateService로 외부 시스템 연동 이해",
        "7단계: SimulationService, SimulationStateService로 메모리 상태와 synchronized 이해",
    ]:
        add_bullet(doc, item)


def build_doc() -> None:
    infos = parse_all_services()
    doc = Document()
    configure_doc(doc)
    add_cover(doc, infos)
    add_foundation(doc)
    add_inventory(doc, infos)
    add_group_summary(doc, infos)
    add_service_details(doc, infos)
    add_pattern_section(doc)
    add_study_order(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")
    print(f"services: {len(infos)}")


if __name__ == "__main__":
    build_doc()
