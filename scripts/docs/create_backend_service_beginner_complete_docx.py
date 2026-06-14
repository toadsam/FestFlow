# -*- coding: utf-8 -*-
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SERVICE_DIR = ROOT / "backend" / "src" / "main" / "java" / "com" / "festflow" / "backend" / "service"
OUTPUT_DOCX = ROOT / "docs" / "festflow" / "페스트플로우_백엔드_Service_초보자_완전해설서.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
RED = RGBColor(153, 27, 27)
GREEN = RGBColor(22, 101, 52)

HEADER_FILL = "E8EEF5"
NOTE_FILL = "F8FAFC"
GREEN_FILL = "ECFDF3"
YELLOW_FILL = "FFFAEB"
CODE_FILL = "F3F4F6"


SERVICE_GUIDE: dict[str, dict[str, object]] = {
    "AdminActionService": {
        "main_method": "publishCongestionReliefNotice",
        "purpose": "관리자 버튼 한 번으로 운영 공지나 혼잡 완화 안내를 만들기 위한 서비스입니다.",
        "beginner": "이 서비스는 직접 모든 일을 처리하지 않고, 필요한 일을 다른 서비스에 맡기는 조율자 역할을 합니다. 예를 들어 혼잡한 부스를 찾는 일은 BoothService가 더 잘 알고, 공지를 저장하는 일은 NoticeService가 더 잘 알기 때문에 여기서는 둘을 연결합니다.",
        "flow": "관리자 화면 버튼 클릭 -> Controller -> AdminActionService -> BoothService로 혼잡 부스 확인 -> NoticeService로 공지 생성 -> StreamService를 통해 프론트 실시간 갱신",
        "frontend": "관리자/운영 콘솔의 빠른 조치 버튼과 연결됩니다.",
        "study": "Service가 반드시 DB를 직접 다루는 것은 아닙니다. 여러 서비스의 기능을 묶어 하나의 운영 기능으로 만드는 서비스도 있습니다.",
    },
    "AdminDashboardService": {
        "main_method": "getKpis",
        "purpose": "관리자 대시보드에 표시할 핵심 지표를 계산합니다.",
        "beginner": "대시보드는 원본 DB 데이터를 그대로 보여주지 않습니다. 방문자 수, 부스 수, 공연 수, 예약 상태처럼 여러 테이블의 값을 모아서 화면용 숫자로 정리해야 합니다.",
        "flow": "관리자 대시보드 접속 -> Controller -> AdminDashboardService -> 여러 Repository 조회 -> KPI DTO 생성 -> 프론트 카드에 표시",
        "frontend": "관리자 대시보드의 KPI 카드, 요약 통계 영역과 연결됩니다.",
        "study": "Repository에서 가져온 Entity를 화면에 바로 주지 않고 DTO로 바꾸는 이유를 이해하기 좋은 서비스입니다.",
    },
    "AdminImportService": {
        "main_method": "importBoothsCsv",
        "purpose": "관리자가 CSV 파일로 부스 데이터를 한 번에 등록하거나 갱신할 수 있게 합니다.",
        "beginner": "파일 업로드는 JSON이 아닙니다. 프론트에서는 FormData로 파일을 보내고, 백엔드 Controller는 MultipartFile로 받습니다. 이 서비스는 파일 안의 줄을 읽어서 부스 생성 요청 DTO로 바꾼 뒤 BoothService를 호출합니다.",
        "flow": "CSV 파일 선택 -> FormData 업로드 -> Controller MultipartFile 수신 -> AdminImportService가 줄 단위로 읽음 -> BoothService로 부스 생성/수정",
        "frontend": "관리자 부스 가져오기, 대량 등록 기능과 연결됩니다.",
        "study": "파일 자체를 DB에 넣는 것이 아니라 파일 내용을 해석해서 기존 서비스 로직을 반복 호출하는 구조입니다.",
    },
    "AiCongestionService": {
        "main_method": "analyzeCurrent",
        "purpose": "현재 축제 상태를 보고 AI 혼잡도 예측과 추천 부스를 만들어냅니다.",
        "beginner": "이 서비스가 현재 AI 기능의 중심입니다. GPS, 예약, 체크인, 대기시간, 재고, 공연 임박 여부 같은 현재 상태값을 모아 Python 모델에 넘기고, 모델 결과가 없으면 규칙 기반 fallback 결과를 사용합니다.",
        "flow": "분석 페이지 -> /ai/congestion/predictions -> AiCongestionService -> FestivalSnapshotService로 현재 상태 수집 -> PythonCongestionModelService로 모델 추론 -> 추천 DTO 반환",
        "frontend": "혼잡도 분석 화면의 AI 예측 카드, 추천/회피 부스 목록과 연결됩니다.",
        "study": "현재 구현은 시계열 딥러닝이 아니라 특정 시점의 상태값을 보고 30분 뒤 혼잡도를 분류하는 tabular ML 구조입니다.",
        "extra": "AI_FEATURE",
    },
    "AiDecisionLogService": {
        "main_method": "record",
        "purpose": "AI가 어떤 판단을 했는지 최근 결정 로그로 남깁니다.",
        "beginner": "AI 결과만 보여주면 왜 그런 결론이 나왔는지 설명하기 어렵습니다. 이 서비스는 제목, 요약, 이유, 운영자 경고를 함께 저장해서 발표나 운영 화면에서 'AI 판단 근거'를 보여줄 수 있게 합니다.",
        "flow": "AI 분석 완료 -> AiDecisionLogService.record 호출 -> 메모리 로그에 저장 -> 최근 AI 판단 목록 API에서 조회",
        "frontend": "AI 분석 결과의 근거/로그 표시 영역과 연결될 수 있습니다.",
        "study": "DB가 아니라 메모리에 저장되는 로그라 서버를 재시작하면 사라질 수 있습니다. 운영 로그와 영구 저장 로그의 차이를 이해해야 합니다.",
    },
    "AiImageGenerationService": {
        "main_method": "generateFestivalProfileImage",
        "purpose": "AI 매칭 프로필 이미지를 생성하거나 이미지 생성 API를 호출합니다.",
        "beginner": "이미지 생성은 내부 계산이 아니라 외부 AI API 호출입니다. 그래서 API 키, 네트워크 실패, 응답 포맷, fallback 처리가 중요합니다.",
        "flow": "프로필 이미지 요청 -> 이미지 저장/프롬프트 구성 -> 외부 AI 이미지 API 호출 -> 결과 URL 저장/반환",
        "frontend": "AI 매칭 프로필 이미지 생성 기능과 연결됩니다.",
        "study": "외부 AI 기능은 성공만 가정하면 안 됩니다. API 키가 없거나 요청이 실패해도 서비스가 완전히 죽지 않게 fallback이 있어야 합니다.",
        "extra": "EXTERNAL_AI",
    },
    "AiMatchService": {
        "main_method": "createProfile",
        "purpose": "AI 매칭 프로필 생성, 이미지 처리, 매칭 요청 상태를 관리합니다.",
        "beginner": "단순히 사용자 정보를 저장하는 서비스가 아닙니다. 프로필 정보 저장, 이미지 저장, AI 이미지 생성, 중복 확인, 개인정보 동의 같은 여러 처리가 한 요청 안에 묶입니다.",
        "flow": "사용자 프로필 입력 -> AiMatchService -> 파일/이미지 처리 -> 프로필 Entity 저장 -> 매칭 요청/상태 관리",
        "frontend": "AI 매칭 페이지, 프로필 등록 화면, 관리자 매칭 관리 화면과 연결됩니다.",
        "study": "여러 DB 저장이 묶인 기능에서는 @Transactional이 왜 필요한지 보기 좋습니다.",
    },
    "AnalyticsService": {
        "main_method": "dashboard",
        "purpose": "혼잡도 분석 화면에 필요한 통계 데이터를 만듭니다.",
        "beginner": "분석 화면은 GPS 로그 한 줄 한 줄을 그대로 보여주지 않습니다. 시간대별, 부스별, 구역별로 묶어서 프론트가 바로 그릴 수 있는 DTO로 가공합니다.",
        "flow": "분석 페이지 접속 -> fetchAnalyticsDashboard -> AnalyticsService.dashboard -> GPS/부스/예약 집계 -> Dashboard DTO 반환",
        "frontend": "분석 페이지의 그래프, 카드, 혼잡도 요약 영역과 연결됩니다.",
        "study": "원본 데이터와 화면용 집계 데이터의 차이를 이해하기 좋은 서비스입니다.",
    },
    "AuditLogService": {
        "main_method": "log",
        "purpose": "관리자나 운영자가 수행한 중요한 작업을 기록합니다.",
        "beginner": "운영 시스템에서는 누가 언제 무엇을 바꿨는지 추적할 수 있어야 합니다. 이 서비스는 기능 자체보다 운영 추적성과 문제 분석을 위한 보조 기능입니다.",
        "flow": "관리자 작업 완료 -> AuditLogService.log -> AuditLogRepository.save -> 로그 조회 화면에서 확인",
        "frontend": "관리자 작업 이력, 운영 로그 화면과 연결될 수 있습니다.",
        "study": "서비스가 사용자 화면 기능만 담당하는 것이 아니라 운영 안정성을 위한 기록도 담당한다는 점을 봐야 합니다.",
    },
    "AuthService": {
        "main_method": "login",
        "purpose": "관리자 로그인과 JWT 토큰 발급을 처리합니다.",
        "beginner": "로그인은 아이디와 비밀번호가 맞는지 확인한 뒤, 이후 요청에서 사용할 토큰을 발급하는 과정입니다. 비밀번호는 평문 비교가 아니라 PasswordEncoder로 검증합니다.",
        "flow": "로그인 폼 제출 -> AuthService.login -> 관리자 조회 -> 비밀번호 검증 -> JwtService로 토큰 생성 -> 프론트가 Authorization 헤더에 저장",
        "frontend": "관리자 로그인 화면과 보호된 관리자 API 호출에 연결됩니다.",
        "study": "로그인 이후 요청마다 비밀번호를 다시 보내는 것이 아니라 Bearer 토큰을 보내는 이유를 이해해야 합니다.",
    },
    "BoothService": {
        "main_method": "getCongestionByBoothId",
        "purpose": "부스 생성, 수정, 삭제, 조회, 이미지 갱신, 실시간 상태, 혼잡도 계산을 담당하는 핵심 서비스입니다.",
        "beginner": "부스는 홈페이지에서 가장 많이 쓰이는 데이터입니다. 지도, 예약, 혼잡도, 관리자 수정, 운영 상태가 모두 부스를 기준으로 연결됩니다.",
        "flow": "프론트 부스 요청 -> BoothService -> BoothRepository/GpsLogRepository/ReservationRepository 조회 -> BoothResponseDto 또는 CongestionResponseDto 반환",
        "frontend": "홈, 지도, 부스 상세, 관리자 부스 관리, 혼잡도 화면과 연결됩니다.",
        "study": "Entity를 DTO로 변환하는 toDto, GPS 거리 계산, 예약 좌석 요약, fallback 시뮬레이션을 함께 볼 수 있는 가장 중요한 서비스입니다.",
    },
    "ChatService": {
        "main_method": "answer",
        "purpose": "축제 안내 챗봇 답변을 생성합니다.",
        "beginner": "챗봇은 외부 AI API를 사용할 수도 있고, 설정이 없거나 실패하면 미리 준비된 fallback 답변을 줄 수도 있습니다. 사용자는 중간 실패를 몰라도 됩니다.",
        "flow": "사용자 질문 입력 -> ChatService.answer -> OpenAI 또는 fallback 답변 생성 -> ChatResponseDto 반환",
        "frontend": "챗봇 또는 AI 안내 UI와 연결됩니다.",
        "study": "AI 기능은 모델 호출 자체보다 실패했을 때 서비스가 계속 동작하게 만드는 구조가 중요합니다.",
        "extra": "EXTERNAL_AI",
    },
    "EventService": {
        "main_method": "createEvent",
        "purpose": "공연/이벤트 일정 생성, 수정, 삭제, 조회를 담당합니다.",
        "beginner": "공연 일정은 사용자 화면뿐 아니라 AI 혼잡도 예측에도 영향을 줍니다. 공연이 곧 시작되면 무대 주변이나 이동 동선의 혼잡 위험이 달라질 수 있습니다.",
        "flow": "관리자 공연 등록 -> EventService -> EventRepository.save -> StreamService.publishEvents -> 프론트 일정 갱신",
        "frontend": "공연 일정, 무대 지도, 운영 콘솔, AI 혼잡도 기능과 연결됩니다.",
        "study": "하나의 도메인 데이터가 여러 기능에 영향을 주는 예시입니다.",
    },
    "FestivalSnapshotService": {
        "main_method": "current",
        "purpose": "AI와 운영 분석에 필요한 현재 축제 상태를 한 번에 모읍니다.",
        "beginner": "AI가 판단하려면 부스만 봐서는 부족합니다. 현재 부스, 예약, 공연, 스태프, GPS 등 여러 데이터를 한 묶음으로 만들어야 합니다. 그 묶음이 snapshot입니다.",
        "flow": "AiCongestionService 호출 -> FestivalSnapshotService.current -> 여러 Repository/Service 조회 -> FestivalSnapshot 반환",
        "frontend": "직접 화면에 보이기보다는 AI/운영 분석 서비스의 내부 입력으로 사용됩니다.",
        "study": "snapshot은 DB 테이블이 아니라 특정 순간의 상태를 담은 읽기용 묶음입니다.",
    },
    "GpsService": {
        "main_method": "saveGpsLog",
        "purpose": "사용자 위치 로그를 저장하고 혼잡도 갱신을 유발합니다.",
        "beginner": "GPS 로그는 AI와 규칙 기반 혼잡도 계산의 핵심 입력입니다. 사용자의 좌표를 저장한 뒤, 혼잡도 스트림을 갱신해 프론트가 새로고침 없이 변화를 볼 수 있게 합니다.",
        "flow": "프론트 위치 전송 -> GpsService.saveGpsLog -> GpsLogRepository.save -> StreamService.publishCongestion",
        "frontend": "지도/혼잡도 분석/실시간 혼잡도 갱신 기능과 연결됩니다.",
        "study": "데이터 저장과 실시간 이벤트 발행이 한 서비스 흐름 안에 함께 들어갈 수 있습니다.",
    },
    "LostItemService": {
        "main_method": "create",
        "purpose": "분실물 등록, 수정, 상태 변경, 조회를 처리합니다.",
        "beginner": "분실물은 공개 화면과 운영자 화면에서 보여주는 정보가 다를 수 있습니다. 연락처 같은 민감 정보는 어떤 DTO에 얼마나 노출할지 신중해야 합니다.",
        "flow": "스태프/사용자 분실물 등록 -> LostItemService -> LostItemRepository.save -> StreamService로 목록 갱신",
        "frontend": "분실물 페이지, 스태프 페이지, 운영 콘솔과 연결됩니다.",
        "study": "같은 Entity라도 사용자용 DTO와 운영자용 DTO를 다르게 설계할 수 있다는 점을 봐야 합니다.",
    },
    "NoticeService": {
        "main_method": "createNotice",
        "purpose": "공지 생성, 수정, 삭제, 활성 공지 조회, 실시간 갱신을 처리합니다.",
        "beginner": "공지는 저장만 하면 끝이 아니라 사용자 화면에 즉시 반영되어야 합니다. 그래서 저장 후 StreamService로 공지 갱신 이벤트를 보냅니다.",
        "flow": "관리자 공지 등록 -> NoticeService.createNotice -> NoticeRepository.save -> StreamService.publishNotices",
        "frontend": "공지 목록, 홈 화면 알림, 운영 콘솔 공지 관리와 연결됩니다.",
        "study": "CRUD와 SSE 발행이 함께 묶인 구조를 보기 좋습니다.",
    },
    "OpsAiService": {
        "main_method": "masterBriefing",
        "purpose": "운영자용 AI 브리핑, 공지 초안, 체크리스트를 생성합니다.",
        "beginner": "운영 AI는 사용자를 대신해 명령을 실행하는 기능이 아니라, 운영자가 빠르게 판단하도록 문장과 체크리스트를 보조하는 기능입니다.",
        "flow": "운영 콘솔 요청 -> OpsAiService -> 현재 운영 데이터 요약 -> 외부 AI 또는 fallback -> 브리핑 DTO 반환",
        "frontend": "운영자 AI 브리핑, 마스터 콘솔, 공지 초안 기능과 연결됩니다.",
        "study": "AI가 최종 권한자가 아니라 운영자의 판단을 돕는 보조 도구로 설계되어 있다는 점이 중요합니다.",
        "extra": "EXTERNAL_AI",
    },
    "PublicAiGuideService": {
        "main_method": "guide",
        "purpose": "일반 방문자에게 페이지별 AI 방문 가이드를 제공합니다.",
        "beginner": "scope라는 값을 받아 analytics, events, stage-map처럼 화면별로 다른 안내를 만들 수 있습니다. 같은 API 구조를 여러 화면에서 재사용하는 방식입니다.",
        "flow": "프론트 fetchAiVisitorGuide(scope) -> PublicAiGuideService.guide -> 데이터 수집/AI 또는 fallback -> AiVisitorGuideDto 반환",
        "frontend": "분석, 공연, 지도 등 페이지별 AI 가이드 섹션과 연결됩니다.",
        "study": "하나의 서비스가 scope에 따라 다른 내용을 반환하는 구조를 이해하기 좋습니다.",
        "extra": "EXTERNAL_AI",
    },
    "PythonCongestionModelService": {
        "main_method": "predictBatch",
        "purpose": "Java 서버에서 Python RandomForest 혼잡도 모델 추론 스크립트를 실행합니다.",
        "beginner": "Java가 pkl 모델 파일을 직접 읽는 것이 아니라, Python 스크립트를 별도 프로세스로 실행합니다. Java는 입력 JSON 파일을 만들고, Python은 모델 추론 후 출력 JSON 파일을 만듭니다.",
        "flow": "AiCongestionService -> PythonCongestionModelService.predictBatch -> 임시 input JSON 생성 -> ProcessBuilder로 Python 실행 -> output JSON 읽기 -> AiModelPredictionDto 반환",
        "frontend": "프론트에 직접 연결되지는 않지만 AI 예측 카드의 실제 모델 결과를 만드는 핵심 내부 서비스입니다.",
        "study": "배포 환경에서 pythonCommand, predictScript, modelPath, timeout 설정이 맞지 않으면 fallback으로 보일 수 있습니다.",
        "extra": "AI_FEATURE",
    },
    "ReservationAuthService": {
        "main_method": "sendCode",
        "purpose": "예약 사용자 전화번호 인증번호 발송과 인증 세션을 관리합니다.",
        "beginner": "예약은 아무나 남의 예약을 조회하거나 체크인하면 안 됩니다. 그래서 전화번호 인증을 먼저 하고, 성공하면 예약용 토큰을 발급합니다.",
        "flow": "전화번호 입력 -> ReservationAuthService.sendCode -> 인증번호 생성 -> SmsSender로 발송 -> 인증 성공 시 토큰 발급",
        "frontend": "부스 예약 인증 화면, QR 체크인 전 인증 흐름과 연결됩니다.",
        "study": "관리자 JWT와 별개로 예약 사용자는 X-Reservation-Token 같은 별도 인증 흐름을 사용합니다.",
    },
    "ReservationService": {
        "main_method": "createReservation",
        "purpose": "부스 예약 생성, 체크인, 완료, 취소, 만료 처리를 담당합니다.",
        "beginner": "예약 기능은 단순 저장이 아닙니다. 좌석이 남았는지, 이미 예약 중인지, 상태가 어떤지, 만료 시간이 지났는지 확인해야 합니다.",
        "flow": "예약 요청 -> 예약 토큰 확인 -> 테이블/좌석 확인 -> Reservation 저장 -> StreamService.publishReservations",
        "frontend": "부스 예약 화면, 예약 확인, QR 체크인, 운영자 예약 관리 화면과 연결됩니다.",
        "study": "RESERVED, CHECKED_IN, COMPLETED, CANCELED 같은 상태 전이를 이해해야 합니다.",
    },
    "SimulationService": {
        "main_method": "tick",
        "purpose": "데모/운영 시뮬레이션 상태를 주기적으로 실제 부스 상태처럼 반영합니다.",
        "beginner": "실제 축제 데이터가 부족할 때도 화면에서 혼잡도 변화와 운영 흐름을 보여주기 위해 시뮬레이션이 필요합니다.",
        "flow": "스케줄/운영 요청 -> SimulationService.tick -> SimulationStateService 계산 -> BoothService 상태 반영 -> StreamService 갱신",
        "frontend": "운영 콘솔, 데모 혼잡도 화면, 실시간 갱신 테스트와 연결됩니다.",
        "study": "실제 데이터와 데모 데이터가 섞일 수 있으므로 어떤 값이 시뮬레이션인지 구분해야 합니다.",
    },
    "SimulationStateService": {
        "main_method": "tick",
        "purpose": "시뮬레이션의 내부 상태와 변화량을 메모리에서 계산합니다.",
        "beginner": "이 서비스는 DB보다 메모리 Map, record, synchronized 같은 개념이 더 중요합니다. 여러 요청이 동시에 들어와도 상태가 꼬이지 않게 제어합니다.",
        "flow": "SimulationService -> SimulationStateService.tick -> 메모리 상태 갱신 -> 변경량 반환",
        "frontend": "직접 노출보다는 시뮬레이션 결과가 부스/혼잡도 화면에 반영됩니다.",
        "study": "메모리 기반 상태는 서버 재시작 시 초기화될 수 있다는 점을 이해해야 합니다.",
    },
    "StaffService": {
        "main_method": "login",
        "purpose": "스태프 로그인, 세션 토큰, 스태프 상태 변경을 처리합니다.",
        "beginner": "스태프는 관리자와 다른 사용자 유형입니다. 관리자 JWT가 아니라 스태프용 세션 토큰으로 운영 기능에 접근합니다.",
        "flow": "스태프 번호/PIN 입력 -> StaffService.login -> StaffSession 저장 -> 프론트가 스태프 토큰 사용 -> 상태 변경 시 StreamService 갱신",
        "frontend": "스태프 로그인, 스태프 상태 관리, 운영 콘솔 스태프 현황과 연결됩니다.",
        "study": "사용자 유형마다 인증 방식이 다를 수 있다는 점을 봐야 합니다.",
    },
    "StreamService": {
        "main_method": "subscribeCongestion",
        "purpose": "SSE 연결을 유지하고 혼잡도/공지/부스/예약 등 변경 이벤트를 프론트로 보냅니다.",
        "beginner": "일반 fetch는 요청하고 응답을 받으면 끝입니다. SSE는 EventSource 연결을 열어두고, 서버가 변경이 생길 때마다 이벤트를 밀어줍니다.",
        "flow": "프론트 EventSource('/stream/congestion') -> StreamService.subscribeCongestion -> emitter 저장 -> publishCongestion 호출 시 모든 연결에 전송",
        "frontend": "혼잡도, 공지, 부스, 스태프, 분실물, 예약 실시간 갱신과 연결됩니다.",
        "study": "WebSocket처럼 양방향 채팅이 필요한 구조가 아니라 서버가 변경 알림만 보내면 되는 구조라 SSE가 적합합니다.",
    },
    "TranslateMetricsService": {
        "main_method": "recordSuccess",
        "purpose": "번역 API 성공/실패 횟수와 지연 시간을 기록합니다.",
        "beginner": "외부 API는 정상 동작 여부를 운영자가 볼 수 있어야 합니다. 이 서비스는 번역 기능 자체가 아니라 번역 기능의 상태를 측정합니다.",
        "flow": "TranslateService 호출 완료 -> 성공/실패 기록 -> 운영 화면에서 지표 조회",
        "frontend": "번역 운영 지표, 관리자 모니터링 카드와 연결될 수 있습니다.",
        "study": "AtomicLong 같은 동시성 안전 카운터를 왜 쓰는지 이해하기 좋은 서비스입니다.",
    },
    "TranslateService": {
        "main_method": "translate",
        "purpose": "외부 번역 API를 호출해 문장을 번역합니다.",
        "beginner": "번역은 내부 계산이 아니라 HTTP 요청입니다. URI 구성, 헤더, 요청 body, 응답 JSON 파싱, 오류 처리가 모두 필요합니다.",
        "flow": "프론트 번역 요청 -> TranslateService.translate -> 외부 API HTTP 호출 -> 응답 파싱 -> TranslateResponseDto 반환",
        "frontend": "다국어 안내, 번역 버튼, 방문자 안내 UI와 연결됩니다.",
        "study": "외부 API와 통신하는 서비스는 timeout과 실패 처리까지 봐야 합니다.",
    },
    "UploadStorageService": {
        "main_method": "saveImage",
        "purpose": "업로드된 이미지를 로컬 또는 S3 저장소에 저장하고 접근 URL을 반환합니다.",
        "beginner": "이미지 파일은 JSON으로 보내지 않습니다. 프론트에서는 FormData에 file을 담고, 백엔드는 MultipartFile로 파일 bytes, 파일명, contentType을 받습니다.",
        "flow": "프론트 FormData 업로드 -> Controller MultipartFile 수신 -> UploadStorageService.saveImage -> 로컬/S3 저장 -> URL 반환 -> BoothService가 imageUrl 저장",
        "frontend": "부스 이미지 업로드, 메뉴판 이미지 업로드, AI 매칭 이미지 업로드와 연결됩니다.",
        "study": "파일 업로드에서는 Content-Type: application/json을 직접 넣지 않고 브라우저가 multipart/form-data boundary를 설정하게 둡니다.",
    },
    "AiMatchSmsNotifier": {
        "main_method": "notifyRequestCreated",
        "purpose": "AI 매칭 요청 생성이나 수락 상황에서 문자 알림을 보냅니다.",
        "beginner": "매칭 상태 변화는 화면에만 표시하면 사용자가 놓칠 수 있습니다. 그래서 SMS 알림을 별도 서비스로 분리해 필요한 시점에 호출합니다.",
        "flow": "매칭 요청/수락 발생 -> AiMatchSmsNotifier -> SolapiMessageClient.sendText -> SMS 발송",
        "frontend": "AI 매칭 신청/수락 흐름의 외부 알림과 연결됩니다.",
        "study": "enabled 설정이 꺼져 있으면 실제 문자를 보내지 않도록 방어하는 구조를 봐야 합니다.",
    },
    "AligoSmsSender": {
        "main_method": "sendVerificationCode",
        "purpose": "Aligo SMS API로 인증번호를 발송하는 구현체입니다.",
        "beginner": "SmsSender라는 공통 인터페이스를 구현합니다. ReservationAuthService는 Aligo인지 Solapi인지 몰라도 같은 메서드로 문자를 보낼 수 있습니다.",
        "flow": "ReservationAuthService -> SmsSender -> AligoSmsSender -> Aligo HTTP API",
        "frontend": "예약 전화번호 인증 화면과 간접 연결됩니다.",
        "study": "업체별 API 차이를 구현체 안에 숨기는 방식입니다.",
    },
    "NoopSmsSender": {
        "main_method": "sendVerificationCode",
        "purpose": "실제 문자를 보내지 않고 로그만 남기는 개발/데모용 SMS 구현체입니다.",
        "beginner": "개발 중에는 문자 비용이나 외부 API 설정 없이 인증 흐름을 테스트해야 합니다. 이때 Noop 구현체가 사용됩니다.",
        "flow": "인증번호 발송 요청 -> NoopSmsSender -> 로그 출력 -> 실제 SMS 없음",
        "frontend": "예약 인증 화면 테스트와 연결됩니다.",
        "study": "운영 환경에서 Noop이 선택되면 사용자는 실제 인증번호를 받지 못하므로 프로필 설정을 확인해야 합니다.",
    },
    "SmsSender": {
        "main_method": "sendVerificationCode",
        "purpose": "SMS 인증번호 발송 기능의 공통 인터페이스입니다.",
        "beginner": "인터페이스는 실제 동작 코드가 아니라 약속입니다. 이 약속을 구현한 Aligo, Solapi, Twilio, Noop 중 하나가 실제 발송을 담당합니다.",
        "flow": "서비스는 SmsSender만 호출 -> Spring 설정에 따라 실제 구현체 선택 -> 업체별 발송 처리",
        "frontend": "예약 인증 화면과 간접 연결됩니다.",
        "study": "인터페이스를 쓰면 나중에 SMS 업체를 바꿔도 ReservationAuthService 코드를 크게 바꾸지 않아도 됩니다.",
    },
    "SolapiMessageClient": {
        "main_method": "sendText",
        "purpose": "Solapi SDK를 사용해 실제 문자 메시지를 전송합니다.",
        "beginner": "Solapi 관련 설정, SDK 객체 생성, sendOne 호출 같은 업체 전용 세부사항을 한 곳에 모아둡니다.",
        "flow": "Notifier 또는 SolapiSmsSender -> SolapiMessageClient.sendText -> Solapi SDK -> SMS 발송",
        "frontend": "AI 매칭 문자, 예약 인증 문자와 간접 연결됩니다.",
        "study": "isConfigured나 enabled 체크를 통해 설정이 없을 때 조용히 건너뛰는 방어 구조를 봐야 합니다.",
    },
    "SolapiSmsSender": {
        "main_method": "sendVerificationCode",
        "purpose": "SolapiMessageClient를 이용해 예약 인증번호 문자를 보냅니다.",
        "beginner": "SolapiSmsSender는 SmsSender 인터페이스와 SolapiMessageClient 사이를 연결하는 adapter 역할입니다.",
        "flow": "ReservationAuthService -> SolapiSmsSender -> SolapiMessageClient.sendText",
        "frontend": "예약 전화번호 인증 화면과 간접 연결됩니다.",
        "study": "공통 인터페이스와 업체 전용 클라이언트를 분리하는 구조를 이해해야 합니다.",
    },
    "TwilioSmsSender": {
        "main_method": "sendVerificationCode",
        "purpose": "Twilio API로 인증번호 문자를 보내는 대체 구현체입니다.",
        "beginner": "같은 SmsSender 인터페이스를 구현하므로 설정만 바꾸면 다른 SMS 업체로 전환할 수 있습니다.",
        "flow": "ReservationAuthService -> TwilioSmsSender -> Twilio Message.creator",
        "frontend": "예약 전화번호 인증 화면과 간접 연결됩니다.",
        "study": "SID, token, from number 같은 배포 환경변수가 맞아야 실제 발송됩니다.",
    },
}


def service_files() -> list[Path]:
    return sorted(SERVICE_DIR.rglob("*.java"), key=lambda p: str(p.relative_to(SERVICE_DIR)).lower())


def read_source(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def class_name(path: Path) -> str:
    return path.stem


def set_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 9.5,
    color: RGBColor | None = None,
    font: str = "Malgun Gothic",
):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "FestFlow Backend Service Study Guide", size=8, color=MUTED)


def add_paragraph(doc: Document, text: str, label: str | None = None, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    set_spacing(p)
    if label:
        add_run(p, label + " ", bold=True, color=DARK)
    add_run(p, text, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int], font_size: float = 8.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        shade_cell(cell, HEADER_FILL)
        set_cell_width(cell, width)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1.15)
        add_run(p, str(header), bold=True, size=font_size, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.18)
            add_run(p, str(value), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = NOTE_FILL, color: RGBColor = DARK) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.2)
    add_run(p, title + " | ", bold=True, color=color, size=8.8)
    add_run(p, body, color=MUTED, size=8.8)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_FILL)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.0)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(6.7)
    doc.add_paragraph()


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        set_spacing(p, after=4, line=1.25)
        add_run(p, step, size=9.3)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, after=4, line=1.25)
        add_run(p, item, size=9.3)


def public_methods(src: str, cls: str) -> list[str]:
    names = []
    for match in re.finditer(r"public\s+(?:static\s+)?[\w<>, ?\[\]]+\s+([A-Za-z0-9_]+)\s*\(", src):
        name = match.group(1)
        if name != cls and name not in names:
            names.append(name)
    return names


def dependency_names(src: str) -> list[str]:
    fields = []
    for match in re.finditer(r"private\s+final\s+([A-Za-z0-9_<>, ?]+)\s+([A-Za-z0-9_]+)\s*;", src):
        fields.append(f"{match.group(2)}: {match.group(1).strip()}")
    return fields


def annotations(src: str) -> list[str]:
    found = []
    for item in ["@Service", "@Component", "@Transactional", "@Value", "@Scheduled", "@Async", "@Profile"]:
        if item in src:
            found.append(item)
    return found


def extract_method(src: str, method_name: str | None, max_lines: int = 72) -> str:
    if method_name:
        pattern = re.compile(
            r"(?:(?:@\w+(?:\([^)]*\))?\s*)*)"
            r"(?:public|private|protected)\s+(?:static\s+)?[\w<>, ?\[\]]+\s+"
            + re.escape(method_name)
            + r"\s*\([^)]*\)\s*(?:throws [^{]+)?\{",
            re.M,
        )
        match = pattern.search(src)
        if match:
            return limit_lines(src[match.start():find_block_end(src, src.find("{", match.start()))].strip(), max_lines)
    return extract_first_public_method(src, max_lines=max_lines)


def extract_first_public_method(src: str, max_lines: int = 56) -> str:
    pattern = re.compile(
        r"(?:(?:@\w+(?:\([^)]*\))?\s*)*)public\s+(?:static\s+)?[\w<>, ?\[\]]+\s+\w+\s*\([^)]*\)\s*(?:throws [^{]+)?\{",
        re.M,
    )
    match = pattern.search(src)
    if not match:
        return limit_lines(src, max_lines)
    return limit_lines(src[match.start():find_block_end(src, src.find("{", match.start()))].strip(), max_lines)


def find_block_end(src: str, open_index: int) -> int:
    if open_index < 0:
        return min(len(src), 1800)
    depth = 0
    in_string = False
    escaped = False
    in_line_comment = False
    in_block_comment = False
    for index in range(open_index, len(src)):
        char = src[index]
        nxt = src[index + 1] if index + 1 < len(src) else ""
        if in_line_comment:
            if char == "\n":
                in_line_comment = False
            continue
        if in_block_comment:
            if char == "*" and nxt == "/":
                in_block_comment = False
            continue
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == "/" and nxt == "/":
            in_line_comment = True
            continue
        if char == "/" and nxt == "*":
            in_block_comment = True
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index + 1
    return len(src)


def limit_lines(snippet: str, max_lines: int) -> str:
    lines = snippet.splitlines()
    if len(lines) <= max_lines:
        return snippet
    head = lines[: max_lines - 6]
    tail = lines[-5:]
    return "\n".join(head + ["    // ... 중간 세부 로직 생략: 문서에서는 핵심 흐름 중심으로 설명합니다."] + tail)


def pattern_explanations(src: str, cls: str) -> list[list[str]]:
    rows: list[list[str]] = []
    if "@Service" in src:
        rows.append(["@Service", "Spring이 이 클래스를 서비스 객체로 등록합니다. Controller나 다른 Service가 생성자를 통해 이 객체를 주입받아 사용할 수 있습니다."])
    if "@Component" in src:
        rows.append(["@Component", "특정 계층 이름은 아니지만 Spring Bean으로 등록한다는 의미입니다. 공통 보조 객체나 클라이언트 클래스에서 자주 사용됩니다."])
    if "@Transactional" in src:
        rows.append(["@Transactional", "메서드 실행 중 여러 DB 작업을 하나의 작업 단위로 묶습니다. 중간에 예외가 나면 일부만 저장되는 상황을 줄입니다."])
    if "@Value" in src:
        rows.append(["@Value", "application.properties나 배포 환경변수 값을 코드 안으로 주입합니다. API 키, 모델 경로, timeout처럼 환경마다 달라지는 값에 사용합니다."])
    if "private final" in src:
        rows.append(["private final 필드", "이 서비스가 의존하는 Repository, 다른 Service, 외부 Client를 저장합니다. final은 생성 후 다른 객체로 바뀌지 않게 하는 의미입니다."])
    if "Repository" in src:
        rows.append(["Repository", "DB 접근 계층입니다. findById, findAll, save, deleteById처럼 Entity를 조회하거나 저장하는 역할을 합니다."])
    if "save(" in src:
        rows.append(["save(...)", "새 Entity면 INSERT, 이미 존재하는 Entity면 UPDATE에 해당하는 저장 작업입니다. JPA가 Entity 상태를 보고 처리합니다."])
    if "orElseThrow" in src:
        rows.append(["orElseThrow", "조회 결과가 없으면 null로 계속 진행하지 않고 즉시 예외를 던집니다. 보통 404 Not Found 같은 응답으로 이어집니다."])
    if "ResponseStatusException" in src:
        rows.append(["ResponseStatusException", "서비스 내부에서 HTTP 상태 코드를 가진 예외를 발생시킵니다. 예: 존재하지 않는 부스면 404를 반환합니다."])
    if "stream()" in src:
        rows.append(["stream()", "목록 데이터를 필터링, 정렬, 변환, 집계할 때 쓰는 Java 문법입니다. Entity 목록을 DTO 목록으로 바꿀 때 자주 보입니다."])
    if ".map(" in src and ".toList()" in src:
        rows.append(["map(...).toList()", "각 항목을 다른 형태로 바꾼 뒤 리스트로 모읍니다. 예를 들어 Booth Entity를 BoothResponseDto로 바꿉니다."])
    if "MultipartFile" in src:
        rows.append(["MultipartFile", "프론트가 FormData로 보낸 파일을 백엔드에서 받는 타입입니다. 파일 업로드는 JSON.stringify(payload)가 아니라 multipart/form-data 흐름입니다."])
    if "SseEmitter" in src:
        rows.append(["SseEmitter", "SSE 연결 객체입니다. 프론트 EventSource와 연결되어 서버가 변경 이벤트를 계속 보낼 수 있게 합니다."])
    if "ProcessBuilder" in src:
        rows.append(["ProcessBuilder", "Java 서버가 Python 스크립트 같은 외부 프로세스를 실행할 때 사용합니다. FestFlow에서는 혼잡도 ML 모델 추론을 위해 사용합니다."])
    if "ObjectMapper" in src:
        rows.append(["ObjectMapper", "Java 객체와 JSON 문자열을 서로 변환합니다. Python 모델 입력/출력이나 외부 API 응답 파싱에 자주 사용됩니다."])
    if "HttpClient" in src or "RestClient" in src or "WebClient" in src:
        rows.append(["HTTP Client", "외부 API로 요청을 보내는 도구입니다. 번역, OpenAI, SMS API처럼 서버 밖의 서비스를 호출할 때 필요합니다."])
    if "SmsSender" in src:
        rows.append(["SmsSender", "SMS 업체를 직접 고정하지 않기 위한 인터페이스입니다. Aligo, Solapi, Twilio, Noop 구현체를 바꿔 끼울 수 있습니다."])
    if "record " in src:
        rows.append(["record", "Java의 간단한 데이터 묶음 문법입니다. 읽기 전용 DTO나 내부 계산 결과를 표현할 때 코드가 짧아집니다."])
    if "synchronized" in src:
        rows.append(["synchronized", "여러 요청이 동시에 같은 메모리 상태를 바꿀 때 값이 꼬이지 않도록 한 번에 하나씩 접근하게 합니다."])
    if "Math." in src:
        rows.append(["Math 계산", "거리, 점수, 비율, 상한/하한 보정처럼 화면에 보여줄 숫자를 안정적으로 계산하기 위해 사용됩니다."])
    if not rows:
        rows.append(["핵심 구조", f"{cls}는 메서드 이름과 의존성 필드를 함께 읽으면 역할을 파악할 수 있습니다. 입력을 받고, 필요한 객체를 조회한 뒤, 결과를 반환하는 구조입니다."])
    return rows[:10]


def method_summary_rows(src: str, cls: str) -> list[list[str]]:
    rows = []
    for name in public_methods(src, cls)[:16]:
        rows.append([name, method_guess(name)])
    if not rows:
        rows.append(["public 메서드 없음", "인터페이스이거나 보조 타입일 수 있습니다."])
    return rows


def method_guess(name: str) -> str:
    lower = name.lower()
    if lower.startswith("get") or lower.startswith("find") or lower.startswith("list"):
        return "데이터를 조회해서 DTO 또는 Entity 목록으로 반환하는 메서드입니다."
    if lower.startswith("create") or lower.startswith("add") or lower.startswith("save"):
        return "새 데이터를 생성하거나 저장하는 메서드입니다."
    if lower.startswith("update") or lower.startswith("change"):
        return "기존 데이터를 수정하는 메서드입니다."
    if lower.startswith("delete") or lower.startswith("remove"):
        return "데이터를 삭제하거나 비활성화하는 메서드입니다."
    if lower.startswith("send") or lower.startswith("notify"):
        return "외부 알림 또는 메시지를 보내는 메서드입니다."
    if lower.startswith("publish"):
        return "SSE 같은 실시간 이벤트를 발행하는 메서드입니다."
    if lower.startswith("subscribe"):
        return "프론트가 실시간 이벤트를 구독하도록 연결을 만드는 메서드입니다."
    if lower.startswith("predict") or lower.startswith("analyze") or lower.startswith("guide"):
        return "AI/분석 판단을 수행하고 결과 DTO를 만드는 메서드입니다."
    if lower.startswith("login"):
        return "인증 정보를 검증하고 세션 또는 토큰을 발급하는 메서드입니다."
    return "서비스 내부의 업무 흐름을 처리하는 공개 메서드입니다."


def add_cover(doc: Document, files: list[Path]) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=4)
    add_run(p, "페스트플로우", bold=True, color=BLUE, size=14)

    title = doc.add_paragraph()
    set_spacing(title, before=18, after=8, line=1.1)
    add_run(title, "백엔드 Service 코드 초보자 완전해설서", bold=True, color=DARK, size=23)

    subtitle = doc.add_paragraph()
    set_spacing(subtitle, after=10, line=1.2)
    add_run(
        subtitle,
        "Spring Boot Service 계층을 처음 공부하는 사람도 흐름을 따라갈 수 있도록, 실제 FestFlow 서비스 코드를 기준으로 역할, 요청 흐름, 대표 코드, 문법 의미, 발표용 설명 포인트를 정리한 문서입니다.",
        color=MUTED,
        size=10.5,
    )

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상 코드", "backend/src/main/java/com/festflow/backend/service 이하 전체 Service Java 파일"],
            ["포함 파일 수", f"{len(files)}개"],
            ["문서 방식", "기초 개념 -> 전체 요청 흐름 -> 서비스별 코드 해설 -> AI/SSE/파일/SMS 같은 핵심 패턴 해설"],
            ["읽는 목적", "발표 답변, 코드 공부, 교수님 질문 대응, 기능 구조 설명에 바로 사용할 수 있게 정리"],
        ],
        [2100, 7260],
        font_size=8.3,
    )

    add_callout(
        doc,
        "읽기 전 주의",
        "일부 Java 파일의 한글 문자열 리터럴은 인코딩 문제로 깨져 보일 수 있습니다. 이 문서는 문자열 표현보다 서비스 구조, 데이터 흐름, 메서드 역할을 중심으로 설명합니다.",
        YELLOW_FILL,
        RED,
    )


def add_foundation(doc: Document) -> None:
    doc.add_heading("1. Service 코드를 보기 전 반드시 알아야 할 기초", level=1)
    add_callout(
        doc,
        "핵심",
        "Service는 Controller와 Repository 사이에서 실제 업무 규칙을 처리하는 계층입니다. 프론트가 API를 호출하면 Controller가 요청을 받고, Service가 판단하고, Repository가 DB와 통신합니다.",
        GREEN_FILL,
        GREEN,
    )

    add_table(
        doc,
        ["개념", "초보자 설명"],
        [
            ["Controller", "프론트의 HTTP 요청을 받는 입구입니다. URL, method(GET/POST/PUT/DELETE), request body, header를 처리합니다."],
            ["Service", "실제 기능 규칙을 처리합니다. 예: 예약 가능한지 판단, 혼잡도 계산, AI 예측 요청, 공지 생성."],
            ["Repository", "DB에 접근합니다. Entity를 조회, 저장, 삭제합니다."],
            ["Entity", "DB 테이블과 가까운 객체입니다. 예: Booth, Reservation, GpsLog."],
            ["DTO", "프론트와 주고받기 좋게 만든 데이터 모양입니다. Entity를 그대로 노출하지 않기 위해 사용합니다."],
            ["SSE", "프론트가 EventSource로 연결을 유지하고, 서버가 변경 이벤트를 계속 보내는 실시간 방식입니다."],
            ["Fallback", "AI 모델, 외부 API, SMS 발송 등이 실패해도 서비스가 완전히 멈추지 않게 대신 사용하는 결과입니다."],
        ],
        [1800, 7560],
        font_size=8.2,
    )

    doc.add_heading("1.1 일반 요청 흐름", level=2)
    add_numbered_steps(
        doc,
        [
            "프론트가 fetch 또는 axios로 API를 호출합니다.",
            "Controller가 URL, body, header를 받습니다.",
            "Controller는 복잡한 로직을 직접 처리하지 않고 Service 메서드를 호출합니다.",
            "Service는 Repository나 다른 Service를 이용해 업무 규칙을 처리합니다.",
            "Repository는 DB에서 Entity를 조회하거나 저장합니다.",
            "Service는 Entity를 DTO로 바꾸고 Controller에 반환합니다.",
            "Controller는 JSON 응답을 프론트에 돌려줍니다.",
        ],
    )

    doc.add_heading("1.2 JSON 요청과 파일 업로드 요청은 다르다", level=2)
    add_code_block(
        doc,
        """
// 일반 데이터 예시: JSON body
body: JSON.stringify(payload)

// 파일 업로드 예시: FormData
const formData = new FormData();
formData.append("file", file);
        """,
    )
    add_paragraph(
        doc,
        "일반 부스 정보, 예약 정보, 로그인 정보는 보통 JSON으로 보냅니다. 그래서 백엔드에서는 RequestDto로 받습니다. 하지만 이미지나 CSV 같은 파일은 JSON 문자열로 보내지 않습니다. 프론트는 FormData를 만들고, 백엔드는 MultipartFile로 받습니다.",
    )
    add_paragraph(
        doc,
        "파일 업로드에서 Content-Type: application/json을 직접 넣으면 안 되는 경우가 많습니다. 브라우저가 multipart/form-data와 boundary 값을 자동으로 설정해야 파일이 정상 전송됩니다.",
        "중요:",
        RED,
    )

    doc.add_heading("1.3 인증 헤더는 사용자 종류마다 다를 수 있다", level=2)
    add_table(
        doc,
        ["헤더/토큰", "어디에 쓰는가"],
        [
            ["Authorization: Bearer ...", "관리자 로그인 후 보호된 관리자 API를 호출할 때 사용합니다."],
            ["X-OPS-KEY", "운영 콘솔처럼 별도 운영자 키로 접근하는 기능에서 사용할 수 있습니다."],
            ["X-Reservation-Token", "예약 생성/조회/체크인에서 전화번호 인증을 마친 사용자임을 확인할 때 사용합니다."],
            ["Staff Session Token", "스태프 로그인 후 스태프 전용 기능에 접근할 때 사용합니다."],
        ],
        [2300, 7060],
        font_size=8.2,
    )

    doc.add_heading("1.4 FestFlow에서 특히 중요한 흐름", level=2)
    add_table(
        doc,
        ["흐름", "설명"],
        [
            ["AI 혼잡도", "AiCongestionService가 현재 상태값을 모으고 PythonCongestionModelService가 Python 모델을 실행합니다."],
            ["SSE 실시간 갱신", "StreamService가 EventSource 연결을 관리하고 변경 이벤트를 프론트에 보냅니다."],
            ["부스/예약/혼잡도", "BoothService와 ReservationService가 가장 핵심적인 도메인 흐름을 담당합니다."],
            ["외부 API", "ChatService, OpsAiService, TranslateService, SMS 서비스는 외부 API 실패 가능성을 항상 고려해야 합니다."],
        ],
        [2200, 7160],
        font_size=8.2,
    )


def add_ai_deep_dive(doc: Document) -> None:
    doc.add_heading("2. 현재 실제로 들어간 AI 기능 구조", level=1)
    add_callout(
        doc,
        "정리",
        "현재 혼잡도 AI는 시계열 LSTM이나 GNN이 아니라, 특정 시점의 상태값을 feature로 만들어 RandomForest 모델이 30분 뒤 혼잡도 등급을 분류하는 tabular ML 구조입니다.",
        GREEN_FILL,
        GREEN,
    )
    add_table(
        doc,
        ["구성 요소", "역할"],
        [
            ["AiCongestionService", "현재 부스 상태, GPS, 예약, 공연 정보를 feature로 만들고 모델 결과와 규칙 기반 fallback을 합쳐 최종 추천을 만듭니다."],
            ["PythonCongestionModelService", "Java 서버에서 Python 추론 스크립트와 RandomForest 모델 파일을 실행합니다."],
            ["predict_congestion.py", "Python 쪽에서 model.pkl을 읽고 입력 JSON에 대해 예측 결과 JSON을 생성합니다."],
            ["RandomForest 모델 파일", "학습된 모델입니다. 서버는 이 파일을 직접 계산하지 않고 Python 프로세스에 맡깁니다."],
            ["Fallback", "모델 파일이나 Python 환경이 없거나 timeout이 나면 최소한의 규칙 기반 결과를 반환합니다."],
        ],
        [2500, 6860],
        font_size=8.2,
    )
    add_paragraph(
        doc,
        "프론트에서 보이는 RandomForest, 신뢰도, drift 상태, 위험 점수는 단순 문구가 아니라 백엔드가 현재 상태값을 만들고 Python 모델 추론 결과를 DTO에 담아 내려주는 구조입니다. 다만 배포 환경에서 Python 경로, 모델 파일 경로, 스크립트 경로가 맞지 않으면 Fallback으로 보입니다.",
    )
    add_paragraph(
        doc,
        "이 기능을 발표할 때는 '완전한 시계열 예측 시스템'이라고 말하면 안 됩니다. 정확한 표현은 '현재 시점의 상태값을 이용해 30분 뒤 혼잡도 등급을 분류하는 tabular ML 기반 예측 프로토타입'입니다.",
        "발표 표현:",
        BLUE,
    )


def add_service_section(doc: Document, path: Path, index: int) -> None:
    src = read_source(path)
    cls = class_name(path)
    guide = SERVICE_GUIDE.get(cls, {})
    main_method = str(guide.get("main_method") or (public_methods(src, cls)[0] if public_methods(src, cls) else ""))

    doc.add_heading(f"3.{index}. {cls}", level=2)
    add_table(
        doc,
        ["항목", "설명"],
        [
            ["파일 위치", str(path.relative_to(ROOT))],
            ["이 서비스의 목적", guide.get("purpose", "서비스 계층의 업무 로직을 처리합니다.")],
            ["프론트 연결", guide.get("frontend", "Controller를 통해 프론트 기능과 연결됩니다.")],
            ["대표 메서드", main_method or "대표 public 메서드 없음"],
        ],
        [1900, 7460],
        font_size=8.0,
    )

    add_paragraph(doc, str(guide.get("beginner", "이 서비스는 Controller에서 전달된 요청을 받아 필요한 Repository나 다른 Service를 호출하고 결과를 DTO로 반환합니다.")), "초보자 설명:")
    add_paragraph(doc, str(guide.get("flow", "프론트 요청 -> Controller -> Service -> Repository/다른 Service -> DTO 반환")), "요청 흐름:")
    add_paragraph(doc, str(guide.get("study", "메서드 이름, 의존성 필드, 반환 DTO를 함께 읽으면 역할을 이해할 수 있습니다.")), "공부 포인트:")

    deps = dependency_names(src)
    if deps:
        add_table(
            doc,
            ["의존성", "쉽게 말하면"],
            [[dep, dependency_explain(dep)] for dep in deps[:14]],
            [3100, 6260],
            font_size=7.8,
        )
    else:
        add_paragraph(doc, "주입받는 private final 의존성이 거의 없거나 인터페이스 성격의 파일입니다. 다른 객체를 많이 호출하지 않는 단순 구조일 가능성이 큽니다.", "의존성:")

    add_table(doc, ["public 메서드", "무슨 일을 하는가"], method_summary_rows(src, cls), [3000, 6360], font_size=7.8)

    doc.add_heading("대표 코드", level=3)
    add_code_block(doc, extract_method(src, main_method, max_lines=76))

    doc.add_heading("대표 코드를 읽는 법", level=3)
    add_table(doc, ["코드 요소", "의미"], pattern_explanations(src, cls), [2500, 6860], font_size=7.8)

    add_beginner_walkthrough(doc, cls, src, guide)

    if guide.get("extra") == "AI_FEATURE":
        add_callout(
            doc,
            "AI 관련 발표 포인트",
            "이 서비스는 단순 화면 보조 문구가 아니라 실제 모델 입력 feature, Python 추론, 모델 결과 DTO, fallback 판단이 연결되는 부분입니다. 단, 현재 방식은 시계열 딥러닝이 아니라 tabular ML 분류입니다.",
            YELLOW_FILL,
            RED,
        )
    elif guide.get("extra") == "EXTERNAL_AI":
        add_callout(
            doc,
            "외부 AI/API 포인트",
            "이 서비스의 AI는 외부 API 의존성이 있으므로 API 키, 네트워크, 응답 포맷, timeout, fallback 처리가 함께 설명되어야 합니다.",
            YELLOW_FILL,
            RED,
        )


def dependency_explain(dep: str) -> str:
    lower = dep.lower()
    if "repository" in lower:
        return "DB에 접근하기 위한 객체입니다. Entity를 조회하거나 저장합니다."
    if "service" in lower:
        return "다른 업무 로직을 재사용하기 위해 주입받은 서비스입니다."
    if "objectmapper" in lower:
        return "Java 객체와 JSON 문자열을 서로 바꾸는 도구입니다."
    if "sender" in lower or "client" in lower:
        return "외부 API나 문자 발송 같은 외부 시스템 호출을 담당하는 객체입니다."
    if "mapper" in lower:
        return "데이터 형태 변환을 담당하는 객체입니다."
    return "이 서비스가 일을 처리할 때 필요한 외부 객체입니다."


def add_beginner_walkthrough(doc: Document, cls: str, src: str, guide: dict[str, object]) -> None:
    doc.add_heading("처음 공부할 때 이렇게 읽으면 된다", level=3)
    steps = [
        f"먼저 클래스 이름 {cls}를 봅니다. 이름 끝의 Service는 '업무 규칙을 처리하는 계층'이라는 뜻입니다.",
        "그다음 private final 필드를 봅니다. 이 필드들이 이 서비스가 의존하는 DB, 다른 서비스, 외부 API입니다.",
        "public 메서드를 봅니다. public 메서드는 Controller나 다른 서비스가 실제로 호출할 수 있는 입구입니다.",
        "메서드 안에서 findById, findAll, save, deleteById가 나오면 DB 조회/저장 흐름이라고 보면 됩니다.",
        "마지막으로 반환 타입을 봅니다. ResponseDto, GuideDto, PredictionDto처럼 끝나면 프론트가 받기 좋은 응답 모양으로 바꿔서 내보내는 것입니다.",
    ]
    add_numbered_steps(doc, steps)

    notes = []
    if "FormData" in str(guide.get("study", "")) or "MultipartFile" in src:
        notes.append("파일은 JSON으로 보내지 않습니다. 프론트 FormData -> 백엔드 MultipartFile -> 저장소 URL -> DB imageUrl 저장 순서로 이해하면 됩니다.")
    if "SseEmitter" in src or cls == "StreamService":
        notes.append("SSE는 프론트가 서버에 계속 연결되어 있는 구조입니다. 서버는 publish 메서드가 호출될 때마다 연결된 브라우저들에게 이벤트를 보냅니다.")
    if "ProcessBuilder" in src or cls == "PythonCongestionModelService":
        notes.append("Python 모델 추론은 Java 내부 함수 호출이 아니라 외부 Python 프로세스 실행입니다. 그래서 배포 환경의 Python 경로와 모델 파일 경로가 매우 중요합니다.")
    if "SmsSender" in src or "Sms" in cls:
        notes.append("SMS 코드는 인터페이스와 구현체로 나뉩니다. 예약 서비스는 SmsSender만 알고, 실제 업체 구현은 설정에 따라 바뀔 수 있습니다.")
    if "Jwt" in src or cls == "AuthService":
        notes.append("관리자 인증은 로그인 후 JWT를 받고, 이후 요청에 Authorization: Bearer 토큰을 붙이는 방식입니다.")
    if "@Transactional" in src:
        notes.append("DB 상태가 여러 번 바뀌는 기능은 중간 실패에 대비해 트랜잭션 단위로 묶는 것이 안전합니다.")
    if not notes:
        notes.append("이 서비스는 메서드 이름을 기준으로 입력, DB 접근, DTO 반환 순서를 따라가면 구조를 이해할 수 있습니다.")
    add_bullets(doc, notes)


def add_appendix(doc: Document, files: list[Path]) -> None:
    doc.add_heading("4. 발표에서 바로 쓸 수 있는 Service 계층 설명", level=1)
    add_paragraph(
        doc,
        "FestFlow 백엔드는 Controller가 요청을 받고 Service가 실제 업무 로직을 처리하며 Repository가 DB와 통신하는 구조입니다. 부스, 예약, 공지, GPS, AI 예측, SSE 실시간 갱신, SMS 발송 같은 기능이 각각 Service로 분리되어 있어 기능별 책임이 비교적 명확합니다.",
    )
    add_paragraph(
        doc,
        "특히 AI 혼잡도 기능은 프론트에 단순 문구를 보여주는 수준이 아니라, 현재 시점의 GPS/예약/공연/대기시간/재고 데이터를 feature로 만들고 Python RandomForest 모델에 전달해 혼잡도 등급을 예측하는 구조입니다. 모델을 사용할 수 없을 때는 fallback 결과를 사용해 화면이 멈추지 않게 설계했습니다.",
    )
    add_paragraph(
        doc,
        "실시간 기능은 WebSocket이 아니라 SSE를 사용합니다. 사용자가 서버로 계속 메시지를 보내야 하는 채팅형 기능이 아니라, 서버가 혼잡도나 공지 변경을 프론트로 밀어주면 되는 구조이기 때문에 SSE가 충분히 적합합니다.",
    )

    doc.add_heading("4.1 Service 파일 전체 목록", level=2)
    rows = []
    for idx, file in enumerate(files, start=1):
        cls = class_name(file)
        guide = SERVICE_GUIDE.get(cls, {})
        rows.append([idx, cls, guide.get("purpose", "서비스 계층 파일")])
    add_table(doc, ["번호", "Service", "역할"], rows, [700, 2500, 6160], font_size=7.4)

    doc.add_heading("4.2 공부 순서 추천", level=2)
    add_numbered_steps(
        doc,
        [
            "BoothService를 먼저 봅니다. 부스 조회, DTO 변환, GPS 혼잡도 계산, 예약 좌석 요약이 모두 들어 있습니다.",
            "ReservationService와 ReservationAuthService를 봅니다. 상태 전이와 사용자 인증 토큰 흐름을 이해할 수 있습니다.",
            "StreamService를 봅니다. SSE가 어떻게 연결되고 publish되는지 이해할 수 있습니다.",
            "AiCongestionService와 PythonCongestionModelService를 봅니다. 현재 AI 기능의 실제 핵심입니다.",
            "UploadStorageService를 봅니다. JSON 요청과 파일 업로드 요청의 차이를 이해할 수 있습니다.",
            "SMS/Translate/Chat/Ops AI 서비스를 봅니다. 외부 API 연동과 fallback의 필요성을 이해할 수 있습니다.",
        ],
    )


def build_doc() -> None:
    files = service_files()
    doc = Document()
    configure_doc(doc)
    add_cover(doc, files)
    add_foundation(doc)
    add_ai_deep_dive(doc)
    doc.add_heading("3. Service별 코드 해설", level=1)
    for index, path in enumerate(files, start=1):
        add_service_section(doc, path, index)
    add_appendix(doc, files)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"written: {OUTPUT_DOCX}")
    print(f"services: {len(files)}")


if __name__ == "__main__":
    build_doc()
