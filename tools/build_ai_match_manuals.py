from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]


def ensure_manual_assets():
    out_dir = ROOT / "manual-assets"
    out_dir.mkdir(exist_ok=True)
    source = Image.open(ROOT / "admin-final-desktop.png")
    crops = {
        "ai-match-admin-dashboard-desktop.png": (140, 0, 1140, 980),
        "ai-match-admin-profiles-desktop.png": (140, 760, 760, 2680),
        "ai-match-admin-matches-desktop.png": (730, 620, 1140, 2300),
        "ai-match-admin-requests-desktop.png": (730, 1900, 1140, 3280),
    }
    for name, box in crops.items():
        source.crop(box).save(out_dir / name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(document):
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = document.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, "111827"),
        ("Heading 1", 17, "4F46E5"),
        ("Heading 2", 13, "111827"),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(document, title, subtitle):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(79, 70, 229)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)

    document.add_paragraph()


def add_note(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF2FF")
    set_cell_text(cell, text)
    document.add_paragraph()


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "DBEAFE")
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    document.add_paragraph()


def image_width_cm(path, max_width_cm, max_height_cm):
    with Image.open(path) as image:
        width, height = image.size
    width_limited = max_width_cm
    height_limited = max_height_cm * (width / height)
    return min(width_limited, height_limited)


def add_image(document, path, caption, max_width_cm=15.8, max_height_cm=19.0):
    image_path = ROOT / path
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(image_width_cm(image_path, max_width_cm, max_height_cm)))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.font.name = "Malgun Gothic"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(107, 114, 128)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def build_user_manual():
    document = Document()
    style_document(document)
    add_title(document, "AI Match 사용자 사용설명서", "대상 화면: /ai-match · 화면 기준: 모바일 모드")
    add_note(document, "이 문서는 실제 모바일 화면 캡처를 기준으로 AI 소개팅 부스 사용 순서를 설명합니다.")

    add_heading(document, "1. AI 소개팅 부스 접속하기")
    add_image(document, "qa-screenshots/ai-match-ref-intro.png", "그림 1. AI Match 모바일 소개 화면", 10.0)
    add_numbered(document, [
        "모바일 브라우저에서 서비스 주소로 접속합니다.",
        "첫 화면에서 AI 소개팅 부스 안내와 QR 영역을 확인합니다.",
        "처음 사용하는 경우 시작하기를 누릅니다.",
        "이미 등록된 사람을 보고 싶으면 등록된 사람 보기를 누릅니다.",
    ])

    add_heading(document, "2. 프로필 등록하기")
    add_image(document, "qa-screenshots/ai-match-ref-register.png", "그림 2. 프로필 등록 화면", 10.0)
    add_numbered(document, [
        "사진 업로드를 눌러 정면 사진을 올립니다.",
        "AI 변환 미리보기가 생성되면 결과 이미지를 확인합니다.",
        "닉네임, 성별, 관심사 태그, 자기소개를 입력합니다.",
        "기본 만남 장소를 선택합니다.",
        "공개 동의 체크박스를 선택한 뒤 등록하기를 누릅니다.",
    ])
    add_bullets(document, [
        "관심사 태그는 최대 6개까지 선택할 수 있습니다.",
        "사진과 소개가 공개 목록에 표시되는 것에 동의해야 등록할 수 있습니다.",
        "PIN은 신청함 확인, 프로필 수정, 삭제에 필요합니다.",
    ])

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "3. 등록된 사람 보기")
    add_image(document, "qa-screenshots/ai-match-ref-people.png", "그림 3. 등록된 사람 목록 화면", 10.0)
    add_numbered(document, [
        "하단 메뉴에서 사람들을 누릅니다.",
        "상단 필터에서 전체, 남자, 여자, 신청 가능을 선택해 목록을 좁힙니다.",
        "프로필 카드에서 사진, 닉네임, 만남 장소, 자기소개, 관심사를 확인합니다.",
        "마음에 드는 프로필의 데이트 신청 버튼을 누릅니다.",
    ])

    add_heading(document, "4. 프로필 상세 확인 및 데이트 신청")
    add_image(document, "qa-screenshots/ai-match-ref-detail.png", "그림 4. 프로필 상세 및 데이트 신청 화면", 10.0)
    add_numbered(document, [
        "상대 프로필의 사진과 소개를 확인합니다.",
        "신청자 닉네임을 입력합니다.",
        "만날 장소를 선택합니다.",
        "짧은 메시지를 작성합니다.",
        "데이트 신청 보내기를 누릅니다.",
    ])

    add_heading(document, "5. 신청함 확인하기")
    add_image(document, "qa-screenshots/ai-match-bottom-check-3.png", "그림 5. 하단 메뉴와 신청함 진입 화면", 10.0)
    add_numbered(document, [
        "하단 메뉴에서 신청함을 누릅니다.",
        "닉네임과 PIN으로 인증합니다.",
        "받은 신청과 보낸 신청을 확인합니다.",
        "받은 신청은 수락 또는 거절할 수 있습니다.",
        "보낸 신청은 대기 중일 때 취소할 수 있습니다.",
    ])

    add_heading(document, "6. 사용 흐름 요약")
    add_table(document, ["단계", "해야 할 일", "화면"], [
        ["1", "/ai-match 접속", "소개 화면"],
        ["2", "시작하기 선택", "프로필 등록"],
        ["3", "사진 업로드 및 정보 입력", "등록 화면"],
        ["4", "등록된 사람 보기", "사람 목록"],
        ["5", "마음에 드는 사람에게 신청", "상세 화면"],
        ["6", "신청함에서 상태 확인", "신청 관리"],
    ])

    add_heading(document, "7. 자주 헷갈리는 부분")
    add_bullets(document, [
        "사진을 올려야 AI 프로필 이미지가 만들어지고 목록에서 더 잘 보입니다.",
        "연락처는 일반 사용자 화면에 바로 공개되지 않습니다.",
        "PIN은 본인 프로필과 신청함을 보호하기 위한 인증 수단입니다.",
        "신청을 보낸 뒤 상대가 수락해야 매칭이 성사됩니다.",
    ])
    document.save(ROOT / "ai-matach_사용설명서.docx")


def build_admin_manual():
    document = Document()
    style_document(document)
    add_title(document, "AI Match 관리자 사용설명서", "대상 화면: /ai-match/admin · 화면 기준: 데스크탑 모드")
    add_note(document, "이 문서는 실제 데스크탑 관리자 화면 캡처를 기준으로 매칭 운영 방법을 설명합니다.")

    add_heading(document, "1. 관리자 로그인")
    add_image(document, "admin-ui.png", "그림 1. 관리자 로그인 화면", 15.8)
    add_numbered(document, [
        "데스크탑 브라우저에서 /ai-match/admin으로 접속합니다.",
        "관리자 아이디와 비밀번호를 입력합니다.",
        "로그인 버튼을 누릅니다.",
    ])
    add_table(document, ["항목", "로컬 개발 기본값"], [["아이디", "0000"], ["비밀번호", "0000"]])

    add_heading(document, "2. 관리자 대시보드 전체 구조")
    add_image(document, "manual-assets/ai-match-admin-dashboard-desktop.png", "그림 2. 관리자 데스크탑 대시보드", 15.8)
    add_table(document, ["영역", "설명"], [
        ["상단 액션", "사용자 화면 이동, 새로고침, 로그아웃"],
        ["운영 요약", "활성 프로필, 전체 신청, 대기중, 성사된 매치 수 확인"],
        ["등록된 사람들", "프로필 검색, 상세 확인, 사진 검수, 관리자 삭제"],
        ["성사된 매치 / 신청 기록", "연락처 확인, 연결 상태 변경, 메모 저장"],
    ])

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "3. 등록된 사람들 관리")
    add_image(document, "manual-assets/ai-match-admin-profiles-desktop.png", "그림 3. 등록된 사람들 관리 화면", 10.0)
    add_numbered(document, [
        "왼쪽의 등록된 사람들 영역을 확인합니다.",
        "검색창에서 닉네임, 전화번호, MBTI, 관심사를 검색합니다.",
        "프로필 카드에서 성별, 관심사, 받은 신청, 보낸 신청, 대기 수, 성사 수를 확인합니다.",
        "사진 검수를 눌러 원본 사진과 AI 변환 사진을 확인합니다.",
        "부적절한 프로필은 관리자 삭제로 비활성화합니다.",
    ])
    add_bullets(document, [
        "삭제된 프로필은 사용자 목록에서 보이지 않습니다.",
        "삭제 후 사용자는 PIN으로 프로필 관리 화면에 접근할 수 없습니다.",
        "삭제 전 닉네임과 사진을 다시 확인해야 합니다.",
    ])

    add_heading(document, "4. 성사된 매치 확인 및 연결 조율")
    add_image(document, "manual-assets/ai-match-admin-matches-desktop.png", "그림 4. 성사된 매치 관리 화면", 10.0)
    add_numbered(document, [
        "오른쪽의 성사된 매치 영역을 확인합니다.",
        "매칭된 두 사람의 닉네임과 프로필 사진을 확인합니다.",
        "상세 보기를 눌러 양쪽 연락처와 사진을 확인합니다.",
        "전화 또는 문자로 만남 장소와 시간을 안내합니다.",
        "처리 상태에 따라 연결 대기중, 연결 완료, 연결 실패 중 하나를 선택합니다.",
    ])
    add_table(document, ["상태", "의미"], [
        ["연결 대기중", "아직 연락 또는 현장 안내가 완료되지 않음"],
        ["연결 완료", "양쪽에게 안내가 끝났고 만남 조율이 완료됨"],
        ["연결 실패", "연락 불가, 취소, 현장 사정 등으로 연결하지 못함"],
    ])

    add_heading(document, "5. 관리자 메모 작성")
    add_bullets(document, [
        "연락 완료 시간",
        "안내한 만남 장소",
        "안내한 약속 시간",
        "특이사항",
        "연결 실패 사유",
    ])
    add_note(document, "예시: 18:20 신청자 전화 완료, 18:25 상대방 문자 발송. 19:00 중앙무대 앞 안내.")

    add_heading(document, "6. 운영 순서 요약")
    add_image(document, "manual-assets/ai-match-admin-requests-desktop.png", "그림 5. 신청 기록 화면", 10.0)
    add_table(document, ["순서", "관리자 작업", "확인 위치"], [
        ["1", "/ai-match/admin 로그인", "로그인 화면"],
        ["2", "대기중/성사된 매치 수 확인", "운영 요약"],
        ["3", "새로 등록된 프로필 확인", "등록된 사람들"],
        ["4", "부적절한 프로필 삭제", "프로필 카드"],
        ["5", "성사된 매치 상세 보기", "성사된 매치"],
        ["6", "양쪽 연락처로 연결 조율", "상세 보기"],
        ["7", "연결 상태 변경", "연결 상태 버튼"],
        ["8", "처리 내용 메모 저장", "관리자 메모"],
    ])

    add_heading(document, "7. 관리자 주의사항")
    add_bullets(document, [
        "연락처는 운영 목적으로만 사용합니다.",
        "사용자 사진은 검수 목적 외에 저장하거나 공유하지 않습니다.",
        "관리자 삭제는 되돌리기 어려우므로 신중하게 사용합니다.",
        "처리 후에는 연결 상태와 메모를 남겨 다음 관리자가 이어서 확인할 수 있게 합니다.",
    ])
    document.save(ROOT / "ai-matach_admin_관리자설명서.docx")


if __name__ == "__main__":
    ensure_manual_assets()
    build_user_manual()
    build_admin_manual()
