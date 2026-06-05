from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "FestFlow_홈페이지_마스터_문서.md"
ASSET_DIR = ROOT / "technical-assets"
OUT_DOCX = ROOT / "FestFlow_홈페이지_마스터_독스.docx"

FONT = Path("C:/Windows/Fonts/malgun.ttf")
BOLD_FONT = Path("C:/Windows/Fonts/malgunbd.ttf")

INK = "111827"
MUTED = "475569"
BLUE = "2563EB"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EFF6FF"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F8FAFC"
CALLOUT_FILL = "EEF2FF"
BORDER = "CBD5E1"
CODE_FILL = "F3F4F6"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def pil_font(size, bold=False):
    return ImageFont.truetype(str(BOLD_FONT if bold else FONT), size)


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in str(text).split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = current + " " + word
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered(draw, box, text, size=27, fill="#111827", bold=False, max_width_ratio=0.86):
    x1, y1, x2, y2 = box
    font = pil_font(size, bold)
    max_width = int((x2 - x1) * max_width_ratio)
    lines = wrap_text(draw, text, font, max_width)
    line_height = size + 9
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height


def rounded_box(draw, box, fill, outline="#93C5FD", width=2, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color="#2563EB", width=5):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def make_home_architecture_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-homepage-master-architecture.png"
    image = Image.new("RGB", (1750, 1080), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((65, 48), "FestFlow 홈페이지 코드 구조", fill="#111827", font=pil_font(48, True))
    draw.text(
        (65, 112),
        "브라우저가 React 라우터를 통해 HomePage를 렌더링하고, api.js를 거쳐 Spring Boot API와 SSE 스트림에 연결됩니다.",
        fill="#475569",
        font=pil_font(25),
    )

    boxes = {
        "browser": (70, 230, 390, 420),
        "main": (505, 230, 825, 420),
        "app": (940, 230, 1260, 420),
        "home": (1375, 230, 1695, 420),
        "api": (505, 610, 825, 800),
        "controller": (940, 610, 1260, 800),
        "service": (1375, 610, 1695, 800),
        "sse": (70, 610, 390, 800),
    }
    labels = {
        "browser": "사용자 브라우저\nURL: /",
        "main": "main.jsx\nBrowserRouter\nRoutes",
        "app": "App.jsx\n공통 레이아웃\n하단 메뉴",
        "home": "HomePage.jsx\n홈 화면 state\n추천/AI/혼잡도",
        "api": "api.js\nfetch 함수\nEventSource 생성",
        "controller": "Controller\n/api/booths\n/api/events\n/api/ai",
        "service": "Service\n조회, 계산,\n상태 변경",
        "sse": "SSE Stream\nbooths/events\n실시간 수신",
    }
    fills = {
        "browser": "#ECFEFF",
        "main": "#EEF2FF",
        "app": "#EFF6FF",
        "home": "#F5F3FF",
        "api": "#F0FDF4",
        "controller": "#FEFCE8",
        "service": "#FFF7ED",
        "sse": "#F0FDFA",
    }
    for key, box in boxes.items():
        rounded_box(draw, box, fills[key])
        draw_centered(draw, box, labels[key], size=25, bold=True)

    arrow(draw, (390, 325), (505, 325))
    arrow(draw, (825, 325), (940, 325))
    arrow(draw, (1260, 325), (1375, 325))
    arrow(draw, (1535, 420), (1535, 610), "#7C3AED")
    arrow(draw, (1375, 705), (1260, 705), "#F97316")
    arrow(draw, (940, 705), (825, 705), "#F97316")
    arrow(draw, (665, 610), (1535, 420), "#16A34A", 4)
    arrow(draw, (505, 705), (390, 705), "#0D9488")
    arrow(draw, (230, 610), (1375, 420), "#0D9488", 4)

    note_font = pil_font(23, True)
    note_text = (
        "읽는 순서: main.jsx에서 라우트 확인 -> App.jsx에서 공통 껍데기 확인 -> "
        "HomePage.jsx에서 state/useEffect/useMemo 확인 -> api.js와 백엔드 Controller/Service 연결 추적"
    )
    note_y = 925
    for note_line in wrap_text(draw, note_text, note_font, 1560):
        draw.text((70, note_y), note_line, fill="#111827", font=note_font)
        note_y += 34
    image.save(path)
    return path


def make_home_data_flow_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-homepage-data-flow.png"
    image = Image.new("RGB", (1750, 1120), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((65, 48), "홈 화면 데이터 흐름", fill="#111827", font=pil_font(48, True))
    draw.text(
        (65, 112),
        "HomePage는 최초 로딩 HTTP 요청과 지속 연결 SSE 스트림을 함께 사용합니다.",
        fill="#475569",
        font=pil_font(25),
    )

    left = [
        ((70, 220, 425, 370), "fetchBooths()\nGET /api/booths", "#ECFEFF", "setBooths"),
        ((70, 420, 425, 570), "fetchEvents()\nGET /api/events", "#EEF2FF", "setEvents"),
        ((70, 620, 425, 770), "fetchTrafficHourly()\nGET /api/analytics/traffic-hourly", "#F0FDF4", "setTraffic"),
        ((70, 820, 425, 970), "fetchAiFestivalGuide()\nGET /api/ai/guide", "#F5F3FF", "setAiGuide"),
    ]
    center = (610, 360, 1030, 650)
    right = [
        ((1225, 260, 1670, 420), "homeCards\n대기 짧은 부스\n다음 공연\n예약 가능 부스", "#FEFCE8"),
        ((1225, 500, 1670, 660), "crowdPercent\n방문량 최신값 / 최대값\n혼잡 단계 변환", "#FFF7ED"),
        ((1225, 740, 1670, 900), "AI 답변 영역\naskChat()\nevidence 버튼", "#FDF2F8"),
    ]

    rounded_box(draw, center, "#F8FAFC", "#94A3B8")
    draw_centered(draw, center, "HomePage state\nbooths / events / traffic\naiGuide / aiAnswer / message", size=27, bold=True)

    for box, label, fill, setter in left:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=23, bold=True)
        arrow(draw, (425, (box[1] + box[3]) // 2), (610, 505), "#2563EB")
        draw.text((462, (box[1] + box[3]) // 2 - 18), setter, fill="#1D4ED8", font=pil_font(19, True))

    for box, label, fill in right:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=23, bold=True)
        arrow(draw, (1030, 505), (1225, (box[1] + box[3]) // 2), "#0D9488")

    stream_box = (610, 805, 1030, 970)
    rounded_box(draw, stream_box, "#E0F2FE", "#38BDF8")
    draw_centered(draw, stream_box, "SSE 실시간\ncreateBoothStream()\ncreateEventStream()", size=23, bold=True)
    arrow(draw, (815, 805), (815, 650), "#0284C7")
    draw.text((930, 740), "실시간 이벤트가 state를 다시 갱신", fill="#0369A1", font=pil_font(20, True))

    image.save(path)
    return path


def make_backend_chain_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-homepage-backend-chain.png"
    image = Image.new("RGB", (1750, 930), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((65, 48), "홈페이지 백엔드 연결 체인", fill="#111827", font=pil_font(48, True))
    draw.text(
        (65, 112),
        "api.js의 함수는 Controller를 호출하고, Controller는 Service로 넘기며, Service는 Repository/외부 API/SSE를 사용합니다.",
        fill="#475569",
        font=pil_font(25),
    )

    rows = [
        ("부스", "fetchBooths", "BoothController", "BoothService", "BoothRepository + 예약 요약"),
        ("공연", "fetchEvents", "EventController", "EventService", "EventRepository + 상태 계산"),
        ("방문량", "fetchTrafficHourly", "AnalyticsController", "AnalyticsService", "GpsLogRepository"),
        ("AI 가이드", "fetchAiFestivalGuide", "AiGuideController", "AiCongestionService", "FestivalSnapshotService"),
        ("AI 질문", "askChat", "ChatController", "ChatService", "근거 검색 + OpenAI/fallback"),
    ]
    x_positions = [70, 365, 660, 995, 1320]
    widths = [230, 235, 255, 255, 360]
    headings = ["기능", "api.js", "Controller", "Service", "데이터/외부 연동"]
    for x, w, h in zip(x_positions, widths, headings):
        rounded_box(draw, (x, 205, x + w, 285), "#DBEAFE", "#93C5FD")
        draw_centered(draw, (x, 205, x + w, 285), h, size=21, bold=True)

    y = 330
    for row in rows:
        boxes = []
        for idx, value in enumerate(row):
            x = x_positions[idx]
            w = widths[idx]
            box = (x, y, x + w, y + 105)
            boxes.append(box)
            rounded_box(draw, box, "#FFFFFF", "#CBD5E1", radius=16)
            draw_centered(draw, box, value, size=19, bold=idx == 0)
        for idx in range(len(boxes) - 1):
            arrow(draw, (boxes[idx][2], y + 52), (boxes[idx + 1][0], y + 52), "#64748B", 3)
        y += 125

    image.save(path)
    return path


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=PAGE_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_dxa):
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                set_cell_width(cell, widths_dxa[idx])


def set_cell_text(cell, text, bold=False, size=8.6, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    for idx, line in enumerate(str(text).split("\n")):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)


def add_table(document, headers, rows, widths_dxa=None, font_size=8.4):
    if widths_dxa is None:
        widths_dxa = [PAGE_WIDTH_DXA // len(headers)] * len(headers)
        widths_dxa[-1] += PAGE_WIDTH_DXA - sum(widths_dxa)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_grid(table, widths_dxa)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, True, font_size + 0.5, INK, WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(cells):
                align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 and idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_text(cells[idx], value, False, font_size, INK, align)
    set_table_grid(table, widths_dxa)
    document.add_paragraph()
    return table


def add_note(document, title, body, fill=CALLOUT_FILL):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_grid(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.2)
    r = p.add_run(title + "\n")
    set_run_font(r, 10.5, True, BLUE)
    r = p.add_run(body)
    set_run_font(r, 9.7, False, INK)
    document.add_paragraph()


def code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_grid(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.05)
    for line_index, line in enumerate(text.strip("\n").splitlines()):
        if line_index:
            p.add_run().add_break()
        safe = line if len(line) <= 112 else line[:109] + "..."
        r = p.add_run(safe)
        set_run_font(r, 8.3, False, INK, name="Consolas")
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    document.add_paragraph()


def add_image(document, path, caption, width_cm=15.7):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=4, line=1.0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line=1.0)
    r = p.add_run(caption)
    set_run_font(r, 8.8, False, MUTED)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 26, INK, 0, 8),
        ("Subtitle", 12, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    section.header.paragraphs[0].text = ""
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("FestFlow 홈페이지 마스터 독스")
    set_run_font(hr, 8.5, False, MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("FestFlow technical reference guide")
    set_run_font(fr, 8.5, False, MUTED)


def add_title_page(document, arch_path):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=4, line=1.0)
    r = p.add_run("FestFlow 홈페이지 마스터 독스")
    set_run_font(r, 28, True, BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=14, line=1.15)
    r = p.add_run("초보자도 코드와 구조를 끝까지 따라갈 수 있도록 풀어쓴 React + Spring Boot 상세 기술 문서")
    set_run_font(r, 12.5, False, MUTED)

    add_table(
        document,
        ["항목", "내용"],
        [
            ["대상", "FestFlow 홈페이지와 방문자 웹앱 전체"],
            ["주요 코드", "frontend/src/main.jsx, App.jsx, pages/HomePage.jsx, api.js"],
            ["백엔드 범위", "Controller, Service, Repository, Entity, SSE, 인증, 배포 설정"],
            ["문서 성격", "초보자용 마스터 독스 + 개발자용 코드 추적 가이드"],
            ["작성 기준", "현재 작업 폴더의 실제 소스 코드"],
        ],
        widths_dxa=[1900, 7460],
        font_size=8.7,
    )

    add_note(
        document,
        "문서 사용 방법",
        "처음 읽는 사람은 1장부터 8장까지로 전체 그림을 잡고, 개발자는 9장 이후의 HomePage, api.js, Controller/Service 연결 흐름을 따라가면 됩니다. 코드를 고칠 때는 마지막의 수정 체크리스트를 반드시 확인하세요.",
    )
    add_image(document, arch_path, "그림 1. FestFlow 홈페이지 코드 구조", 15.6)
    document.add_section(WD_SECTION.NEW_PAGE)


def parse_inline_runs(paragraph, text, size=10.5):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.5, bold=False, color=INK, name="Consolas")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=False, color=INK)


def add_markdown_paragraph(document, text):
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.25)
    parse_inline_runs(p, text, 10.5)


def add_bullet(document, text, numbered=False):
    p = document.add_paragraph(style="List Number" if numbered else "List Bullet")
    parse_inline_runs(p, re.sub(r"^\d+\.\s+", "", text).strip(), 10.0)


def flush_table(document, table_lines):
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        values = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(values)
    if len(rows) < 2:
        return
    headers = rows[0]
    body = rows[2:] if re.fullmatch(r"[-:\s|]+", table_lines[1].strip()) else rows[1:]
    col_count = len(headers)
    if col_count == 2:
        widths = [2500, 6860]
    elif col_count == 3:
        widths = [2200, 3700, 3460]
    elif col_count == 4:
        widths = [1800, 2500, 2800, 2260]
    else:
        base = PAGE_WIDTH_DXA // col_count
        widths = [base] * col_count
        widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    normalized_body = []
    for row in body:
        normalized = row[:col_count] + [""] * max(0, col_count - len(row))
        normalized_body.append(normalized)
    add_table(document, headers, normalized_body, widths_dxa=widths, font_size=7.9 if col_count >= 4 else 8.3)


def markdown_to_docx(document, markdown_text):
    lines = markdown_text.splitlines()
    in_code = False
    code_lines = []
    table_lines = []
    skip_first_title = True

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                flush_table(document, table_lines)
                table_lines = []
                in_code = True
                code_lines = []
            else:
                code_block(document, "\n".join(code_lines))
                in_code = False
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            continue
        else:
            flush_table(document, table_lines)
            table_lines = []

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            p = document.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            p = document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            p = document.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith("#### "):
            p = document.add_heading(stripped[5:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            add_bullet(document, stripped[2:], numbered=False)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_bullet(document, stripped, numbered=True)
            continue

        add_markdown_paragraph(document, stripped)

    flush_table(document, table_lines)


def add_static_toc(document):
    document.add_heading("빠른 목차", level=1)
    toc_items = [
        "1-8장: 프로젝트 구조, 프론트/백엔드 기술 스택, 라우팅과 공통 레이아웃",
        "9-19장: HomePage.jsx 코드 상세 해부, state/useEffect/useMemo/AI 질문 처리",
        "20-24장: 홈 화면과 백엔드 API, Controller/Service, 혼잡도 계산, SSE",
        "25-28장: 주요 페이지, 브라우저 저장소, CSS, i18n",
        "29-34장: 설정, API, DB, 인증, PWA, 배포",
        "35-38장: 오류 해결, 개발 체크리스트, 코드 읽기 연습, 핵심 요약",
    ]
    for item in toc_items:
        add_bullet(document, item)
    add_note(
        document,
        "읽기 팁",
        "이 독스는 처음부터 끝까지 읽어도 되지만, 실제 개발 중에는 목차에서 관련 장으로 바로 이동해도 됩니다. 특히 오류 해결과 수정 체크리스트는 작업 전후로 반복해서 확인하는 용도입니다.",
    )


def add_extra_deep_appendix(document):
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("부록 A. 핵심 코드 상세 해부", level=1)
    add_note(
        document,
        "이 부록의 목적",
        "앞 장들이 전체 구조를 설명했다면, 이 부록은 실제 코드 몇 줄이 무슨 의미인지 초보자 눈높이로 다시 풀어 설명합니다. 코드를 직접 열어놓고 같이 보면 가장 효과적입니다.",
    )

    document.add_heading("A.1 HomePage가 처음 렌더링될 때", level=2)
    code_block(
        document,
        """
useEffect(() => {
  let mounted = true;

  Promise.allSettled([
    fetchBooths(),
    fetchEvents(),
    fetchTrafficHourly(),
  ]).then(([boothResult, eventResult, trafficResult]) => {
    if (!mounted) return;
    ...
  });

  return () => {
    mounted = false;
    streams.forEach((stream) => stream.close());
  };
}, []);
        """,
    )
    add_markdown_paragraph(
        document,
        "이 코드는 홈 화면 생명주기의 핵심입니다. 빈 의존성 배열 `[]` 때문에 화면이 처음 나타날 때 한 번만 실행됩니다. `Promise.allSettled`는 부스, 공연, 방문량 요청을 동시에 보내고, 실패한 요청이 있어도 성공한 요청 결과를 사용할 수 있게 합니다. `return` 안의 함수는 화면이 사라질 때 실행되는 정리 함수이며, SSE 연결을 닫아 메모리 누수와 중복 이벤트 수신을 막습니다.",
    )

    document.add_heading("A.2 api.js에서 URL을 안정적으로 만드는 방식", level=2)
    code_block(
        document,
        """
const API_BASE = (
  import.meta.env.VITE_API_BASE_URL || "http://localhost:8080/api"
)
  .trim()
  .replace(/\\/+$/, "");
        """,
    )
    add_markdown_paragraph(
        document,
        "`import.meta.env.VITE_API_BASE_URL`은 Vite가 제공하는 환경변수 접근 방식입니다. 운영 배포에서는 Railway 백엔드 주소를 여기에 넣고, 로컬에서는 값이 없으므로 `http://localhost:8080/api`가 사용됩니다. 마지막의 `replace(/\\/+$/, \"\")`는 주소 끝의 `/`를 제거해 `/api//booths` 같은 잘못된 URL을 방지합니다.",
    )

    document.add_heading("A.3 백엔드 Controller가 얇은 이유", level=2)
    code_block(
        document,
        """
@GetMapping
public List<BoothResponseDto> getBooths() {
    return boothService.getAllBooths();
}
        """,
    )
    add_markdown_paragraph(
        document,
        "Controller는 HTTP 요청을 받는 입구입니다. 여기서 DB 조회, 정렬, 혼잡도 계산을 직접 하지 않고 `BoothService`로 넘깁니다. 이렇게 하면 API 주소를 담당하는 코드와 실제 도메인 규칙을 담당하는 코드가 분리되어 유지보수가 쉬워집니다.",
    )

    document.add_heading("A.4 Service에서 DTO로 바꾸는 이유", level=2)
    add_markdown_paragraph(
        document,
        "Entity는 DB 테이블과 가깝고, DTO는 API 응답과 가깝습니다. Entity를 그대로 프론트에 보내면 내부 필드나 민감한 정보가 노출될 수 있고, 화면에 필요한 계산값을 붙이기 어렵습니다. 그래서 `BoothService.toDto`처럼 서버 내부 모델을 화면용 응답 모델로 변환합니다.",
    )

    document.add_heading("A.5 SSE 이벤트 이름이 중요한 이유", level=2)
    code_block(
        document,
        """
emitter.send(SseEmitter.event().name(eventName).data(payload));

boothStream.addEventListener("booths", (event) => {
  const next = JSON.parse(event.data);
  setBooths(next);
});
        """,
    )
    add_markdown_paragraph(
        document,
        "백엔드가 `.name(\"booths\")`로 이벤트를 보내면 프론트도 정확히 `addEventListener(\"booths\", ...)`로 들어야 합니다. 경로가 맞아도 이벤트 이름이 다르면 콜백이 실행되지 않습니다. 실시간 갱신 문제를 디버깅할 때는 경로와 이벤트 이름을 모두 확인해야 합니다.",
    )


def build():
    arch = make_home_architecture_diagram()
    data_flow = make_home_data_flow_diagram()
    backend_chain = make_backend_chain_diagram()

    source_text = SOURCE_MD.read_text(encoding="utf-8")

    doc = Document()
    style_document(doc)
    add_title_page(doc, arch)
    add_static_toc(doc)
    add_image(doc, data_flow, "그림 2. 홈 화면 데이터 흐름", 15.6)
    add_image(doc, backend_chain, "그림 3. 홈페이지 백엔드 연결 체인", 15.6)
    doc.add_section(WD_SECTION.NEW_PAGE)
    markdown_to_docx(doc, source_text)
    add_extra_deep_appendix(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
