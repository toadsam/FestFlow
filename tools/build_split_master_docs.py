from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "technical-assets"
OUT_FRONTEND = ROOT / "FestFlow_프론트엔드_완전_마스터_독스.docx"
OUT_BACKEND = ROOT / "FestFlow_백엔드_완전_마스터_독스.docx"

FONT_NAME = "Malgun Gothic"
MONO_FONT = "Consolas"
FONT_PATH = Path("C:/Windows/Fonts/malgun.ttf")
BOLD_FONT_PATH = Path("C:/Windows/Fonts/malgunbd.ttf")

INK = "111827"
MUTED = "475569"
BLUE = "2563EB"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EFF6FF"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F8FAFC"
CALLOUT_FILL = "EEF2FF"
WARNING_FILL = "FFF7ED"
SUCCESS_FILL = "ECFDF5"
BORDER = "CBD5E1"
CODE_FILL = "F3F4F6"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def pil_font(size, bold=False):
    path = BOLD_FONT_PATH if bold else FONT_PATH
    return ImageFont.truetype(str(path), size)


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in str(text).split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def rounded_box(draw, box, fill, outline="#93C5FD", width=2, radius=20):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered(draw, box, text, size=26, fill="#111827", bold=True, max_width_ratio=0.88):
    x1, y1, x2, y2 = box
    font = pil_font(size, bold)
    max_width = int((x2 - x1) * max_width_ratio)
    lines = wrap_text(draw, text, font, max_width)
    line_height = size + 8
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height


def arrow(draw, start, end, color="#2563EB", width=5):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def make_frontend_architecture_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-frontend-master-architecture.png"
    image = Image.new("RGB", (1800, 1120), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), "FestFlow 프론트엔드 전체 구조", fill="#111827", font=pil_font(48, True))
    draw.text(
        (70, 112),
        "브라우저에서 React 앱이 시작되고, 라우터가 페이지를 고르며, 각 페이지는 api.js를 통해 백엔드와 통신합니다.",
        fill="#475569",
        font=pil_font(25),
    )

    boxes = {
        "browser": (70, 240, 370, 410),
        "vite": (480, 240, 780, 410),
        "main": (890, 240, 1190, 410),
        "app": (1300, 240, 1700, 410),
        "pages": (1300, 590, 1700, 760),
        "api": (890, 590, 1190, 760),
        "backend": (480, 590, 780, 760),
        "state": (70, 590, 370, 760),
        "storage": (70, 850, 370, 1010),
        "i18n": (480, 850, 780, 1010),
        "assets": (890, 850, 1190, 1010),
        "pwa": (1300, 850, 1700, 1010),
    }
    labels = {
        "browser": "사용자 브라우저\n주소 입력 / 클릭",
        "vite": "Vite\n개발 서버와 빌드",
        "main": "main.jsx\nReact 시작점\nRoutes 정의",
        "app": "App.jsx\n공통 레이아웃\n하단 메뉴",
        "pages": "pages/*.jsx\n홈, 지도, 공연,\n관리, 운영 화면",
        "api": "api.js\nHTTP fetch\nSSE EventSource",
        "backend": "Spring Boot API\n/api/booths\n/api/events 등",
        "state": "React state\nuseState\nuseEffect\nuseMemo",
        "storage": "localStorage\n즐겨찾기, 토큰,\n언어 설정",
        "i18n": "i18n.js\n화면 텍스트 번역\nMutationObserver",
        "assets": "public/images\n아이콘, PWA,\n부스 이미지",
        "pwa": "manifest\nservice-worker\noffline.html",
    }
    fills = {
        "browser": "#ECFEFF",
        "vite": "#EEF2FF",
        "main": "#EFF6FF",
        "app": "#F5F3FF",
        "pages": "#FFF7ED",
        "api": "#F0FDF4",
        "backend": "#FEFCE8",
        "state": "#FDF2F8",
        "storage": "#F8FAFC",
        "i18n": "#E0F2FE",
        "assets": "#F0FDFA",
        "pwa": "#FAE8FF",
    }
    for key, box in boxes.items():
        rounded_box(draw, box, fills[key])
        draw_centered(draw, box, labels[key], size=24)

    arrow(draw, (370, 325), (480, 325))
    arrow(draw, (780, 325), (890, 325))
    arrow(draw, (1190, 325), (1300, 325))
    arrow(draw, (1500, 410), (1500, 590), "#7C3AED")
    arrow(draw, (1300, 675), (1190, 675), "#16A34A")
    arrow(draw, (890, 675), (780, 675), "#F97316")
    arrow(draw, (370, 675), (1300, 675), "#DB2777", 3)
    arrow(draw, (220, 760), (220, 850), "#64748B", 3)
    arrow(draw, (630, 850), (630, 760), "#0284C7", 3)
    arrow(draw, (1040, 850), (1040, 760), "#0D9488", 3)
    arrow(draw, (1500, 850), (1500, 760), "#A855F7", 3)

    note_font = pil_font(22, True)
    note = "읽는 순서: package.json -> main.jsx -> App.jsx -> 원하는 Page.jsx -> api.js -> 백엔드 Controller 순서로 따라가면 됩니다."
    for i, line in enumerate(wrap_text(draw, note, note_font, 1580)):
        draw.text((70, 1040 + i * 31), line, fill="#111827", font=note_font)
    image.save(path)
    return path


def make_frontend_state_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-frontend-state-flow.png"
    image = Image.new("RGB", (1800, 1080), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), "React 화면이 바뀌는 원리", fill="#111827", font=pil_font(48, True))
    draw.text(
        (70, 112),
        "프론트엔드 초보자가 가장 헷갈리는 흐름은 데이터 요청, state 변경, 화면 재렌더링입니다.",
        fill="#475569",
        font=pil_font(25),
    )

    boxes = [
        ((70, 250, 390, 420), "1. 페이지 진입\nHomePage 실행", "#EEF2FF"),
        ((505, 250, 825, 420), "2. useEffect\n초기 데이터 요청", "#EFF6FF"),
        ((940, 250, 1260, 420), "3. api.js\nfetch / EventSource", "#F0FDF4"),
        ((1375, 250, 1695, 420), "4. 백엔드 응답\nJSON 데이터", "#FEFCE8"),
        ((1375, 600, 1695, 770), "5. setState\nbooths/events 변경", "#FDF2F8"),
        ((940, 600, 1260, 770), "6. useMemo\n추천 카드 계산", "#FFF7ED"),
        ((505, 600, 825, 770), "7. JSX 렌더링\n화면에 표시", "#ECFEFF"),
        ((70, 600, 390, 770), "8. 사용자 클릭\nnavigate / submit", "#F5F3FF"),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=24)
    for i in range(3):
        arrow(draw, (boxes[i][0][2], 335), (boxes[i + 1][0][0], 335))
    arrow(draw, (1535, 420), (1535, 600), "#DB2777")
    arrow(draw, (1375, 685), (1260, 685), "#DB2777")
    arrow(draw, (940, 685), (825, 685), "#DB2777")
    arrow(draw, (505, 685), (390, 685), "#DB2777")
    arrow(draw, (230, 600), (230, 420), "#7C3AED")

    footer = [
        "핵심: React는 DOM을 직접 매번 고치는 방식이 아니라 state가 바뀌면 JSX를 다시 계산해서 화면을 갱신합니다.",
        "SSE는 백엔드가 새 데이터를 밀어 보내는 통로입니다. 그래서 사용자가 새로고침하지 않아도 부스/공연/예약 상태가 바뀔 수 있습니다.",
    ]
    y = 875
    for text in footer:
        for line in wrap_text(draw, text, pil_font(23, True), 1600):
            draw.text((70, y), line, fill="#111827", font=pil_font(23, True))
            y += 35
        y += 8
    image.save(path)
    return path


def make_backend_architecture_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-backend-master-architecture.png"
    image = Image.new("RGB", (1800, 1120), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), "FestFlow 백엔드 전체 구조", fill="#111827", font=pil_font(48, True))
    draw.text(
        (70, 112),
        "Spring Boot 서버는 요청을 받은 뒤 보안 필터, 컨트롤러, 서비스, 저장소, DB 순서로 처리합니다.",
        fill="#475569",
        font=pil_font(25),
    )

    boxes = {
        "client": (70, 250, 360, 420),
        "cors": (465, 250, 755, 420),
        "security": (860, 250, 1150, 420),
        "controller": (1255, 250, 1700, 420),
        "service": (1255, 590, 1700, 760),
        "repo": (860, 590, 1150, 760),
        "db": (465, 590, 755, 760),
        "external": (70, 590, 360, 760),
        "sse": (465, 870, 755, 1010),
        "init": (860, 870, 1150, 1010),
        "upload": (1255, 870, 1700, 1010),
    }
    labels = {
        "client": "프론트엔드\napi.js 요청",
        "cors": "CorsConfig\n허용 출처 확인",
        "security": "SecurityConfig\nJWT / Ops Key\n권한 검사",
        "controller": "Controller\nURL 매핑\nDTO 검증",
        "service": "Service\n비즈니스 규칙\n계산과 저장",
        "repo": "Repository\nJPA 쿼리\n엔티티 조회",
        "db": "DB\nMySQL 또는\nPostgreSQL",
        "external": "외부 연동\nOpenAI / SMS\nS3",
        "sse": "StreamService\nSSE 실시간 방송",
        "init": "DataInitializer\n데모 데이터\n초기 관리자",
        "upload": "UploadStorage\n로컬 파일 또는\nS3 저장",
    }
    fills = {
        "client": "#ECFEFF",
        "cors": "#EEF2FF",
        "security": "#FDF2F8",
        "controller": "#EFF6FF",
        "service": "#F0FDF4",
        "repo": "#FEFCE8",
        "db": "#FFF7ED",
        "external": "#FAE8FF",
        "sse": "#E0F2FE",
        "init": "#F8FAFC",
        "upload": "#F0FDFA",
    }
    for key, box in boxes.items():
        rounded_box(draw, box, fills[key])
        draw_centered(draw, box, labels[key], size=24)

    arrow(draw, (360, 335), (465, 335))
    arrow(draw, (755, 335), (860, 335))
    arrow(draw, (1150, 335), (1255, 335))
    arrow(draw, (1478, 420), (1478, 590), "#16A34A")
    arrow(draw, (1255, 675), (1150, 675), "#16A34A")
    arrow(draw, (860, 675), (755, 675), "#F97316")
    arrow(draw, (1255, 675), (360, 675), "#A855F7", 3)
    arrow(draw, (610, 870), (1478, 760), "#0284C7", 3)
    arrow(draw, (1005, 870), (610, 760), "#64748B", 3)
    arrow(draw, (1478, 870), (1478, 760), "#0D9488", 3)

    note_font = pil_font(22, True)
    note = "읽는 순서: application.properties -> SecurityConfig -> Controller -> Service -> Repository -> Entity -> DTO 순서로 보면 서버 구조가 풀립니다."
    for i, line in enumerate(wrap_text(draw, note, note_font, 1580)):
        draw.text((70, 1040 + i * 31), line, fill="#111827", font=note_font)
    image.save(path)
    return path


def make_backend_security_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-backend-security-flow.png"
    image = Image.new("RGB", (1800, 1040), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), "백엔드 권한 검사 흐름", fill="#111827", font=pil_font(48, True))
    draw.text(
        (70, 112),
        "FestFlow는 공개 API, 관리자 API, 운영 API, 스태프 API, 예약 인증 API를 성격에 맞게 나눕니다.",
        fill="#475569",
        font=pil_font(25),
    )
    rows = [
        ("공개 API", "/api/booths, /api/events, /api/analytics, /api/chat", "대부분 permitAll. 방문자 화면에서 사용"),
        ("관리자 API", "/api/admin/**", "Bearer JWT가 있고 역할이 ADMIN이어야 접근"),
        ("운영 마스터", "/api/ops/master/**", "X-Ops-Key가 마스터 키와 일치해야 접근"),
        ("부스 운영", "/api/ops/booth/**", "마스터 키 또는 부스별 키가 필요"),
        ("스태프", "/api/staff/**", "스태프 로그인 토큰을 각 API에서 검사"),
        ("예약 인증", "/api/reservations/auth/**", "전화번호 인증 코드 발급과 검증"),
    ]
    y = 220
    for name, path_text, rule in rows:
        rounded_box(draw, (70, y, 345, y + 110), "#EEF2FF", "#93C5FD", radius=16)
        draw_centered(draw, (70, y, 345, y + 110), name, size=22)
        rounded_box(draw, (410, y, 950, y + 110), "#F8FAFC", "#CBD5E1", radius=16)
        draw_centered(draw, (410, y, 950, y + 110), path_text, size=20, bold=False)
        rounded_box(draw, (1015, y, 1700, y + 110), "#F0FDF4", "#86EFAC", radius=16)
        draw_centered(draw, (1015, y, 1700, y + 110), rule, size=20, bold=False)
        arrow(draw, (345, y + 55), (410, y + 55), "#64748B", 3)
        arrow(draw, (950, y + 55), (1015, y + 55), "#64748B", 3)
        y += 125
    image.save(path)
    return path


def set_run_font(run, size=None, bold=None, color=None, name=FONT_NAME):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("p. ")
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)
    for r in paragraph.runs:
        set_run_font(r, 8.5, color=MUTED)


def style_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]
    for style_name, size, color, before, after in heading_tokens:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(running_title)
    set_run_font(run, 8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(doc, title, subtitle, label):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=20, after=2, line=1.1)
    run = p.add_run(title)
    set_run_font(run, 24, bold=True, color=INK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.2)
    run = p.add_run(subtitle)
    set_run_font(run, 12.5, color=MUTED)

    meta = [
        ("문서 성격", label),
        ("대상 독자", "프론트/백엔드를 처음 배우는 학생 또는 FestFlow 코드를 처음 넘겨받은 사람"),
        ("작성 기준", "현재 로컬 프로젝트 코드 기준"),
        ("사용 방법", "처음에는 순서대로 읽고, 이후에는 목차처럼 필요한 장만 찾아봅니다."),
    ]
    add_table(doc, ["항목", "내용"], meta, [1850, 7510])
    add_note(
        doc,
        "읽기 전에",
        "이 문서는 코드 이름을 외우게 하려는 문서가 아닙니다. "
        "각 파일이 왜 존재하고, 화면에서 어떤 일이 일어나며, 백엔드 어느 코드와 연결되는지를 "
        "초보자 기준으로 반복해서 설명하는 마스터 문서입니다.",
        fill=SUCCESS_FILL,
    )
    doc.add_page_break()


def add_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_lead:
        lead = p.add_run(bold_lead)
        set_run_font(lead, 10.5, bold=True, color=INK)
        text = text
    run = p.add_run(text)
    set_run_font(run, 10.5, color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run_font(r1, 10.2, bold=True, color=INK)
            r2 = p.add_run(rest)
            set_run_font(r2, 10.2, color=INK)
        else:
            run = p.add_run(str(item))
            set_run_font(run, 10.2, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run_font(r1, 10.2, bold=True, color=INK)
            r2 = p.add_run(rest)
            set_run_font(r2, 10.2, color=INK)
        else:
            run = p.add_run(str(item))
            set_run_font(run, 10.2, color=INK)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grids = tbl.findall(qn("w:tblGrid"))
    for grid in old_grids:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, after=0, line=1.18)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    widths = widths or [int(TABLE_WIDTH_DXA / len(headers))] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], HEADER_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            shade_cell(cells[idx], "FFFFFF")
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=9.0, align=align)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6)
    return table


def add_note(doc, title, body, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.2)
    r1 = p.add_run(f"{title}: ")
    set_run_font(r1, 10.2, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    set_run_font(r2, 10.2, color=INK)
    doc.add_paragraph()


def add_code_block(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2)
        run = p.add_run(caption)
        set_run_font(run, 9.2, bold=True, color=MUTED)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    shade_cell(cell, CODE_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run(textwrap.dedent(code).strip())
    set_run_font(run, 8.2, color=INK, name=MONO_FONT)
    doc.add_paragraph()


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.35))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=8)
    run = p.add_run(caption)
    set_run_font(run, 8.8, bold=True, color=MUTED)


def add_code_walkthrough(doc, title, path, code, explanations):
    doc.add_heading(title, level=3)
    add_paragraph(doc, f"파일 위치: {path}")
    add_code_block(doc, code, caption="대표 코드 조각")
    add_numbered(doc, explanations)


FRONTEND_TERMS = [
    ("컴포넌트", "React에서 화면 조각을 만드는 함수입니다. FestFlow의 HomePage, StageMapPage, CongestionBadge가 모두 컴포넌트입니다.", "컴포넌트를 HTML 파일처럼 생각하면 헷갈립니다. 컴포넌트는 데이터를 받아 JSX를 반환하는 JavaScript 함수입니다."),
    ("JSX", "JavaScript 안에 HTML처럼 보이는 문법을 쓰게 해주는 React 문법입니다. 실제 브라우저에는 JavaScript로 변환되어 들어갑니다.", "class 대신 className을 쓰고, onclick 대신 onClick을 쓰는 이유가 JSX 규칙 때문입니다."),
    ("state", "화면이 기억해야 하는 값입니다. booths, events, loading, message 같은 값이 state입니다.", "state를 직접 바꾸지 않고 setBooths 같은 setter를 써야 React가 화면을 다시 그립니다."),
    ("props", "부모 컴포넌트가 자식 컴포넌트에게 넘겨주는 값입니다. 아이콘 컴포넌트에 className을 넘기는 것도 props입니다.", "props는 자식이 마음대로 고치는 데이터가 아니라 입력값입니다."),
    ("useState", "컴포넌트 내부에 state를 만드는 Hook입니다. const [booths, setBooths] = useState([]) 형태로 씁니다.", "첫 번째 값은 현재 데이터, 두 번째 값은 데이터를 바꾸는 함수입니다."),
    ("useEffect", "화면이 처음 열렸을 때 API 요청을 보내거나, 화면이 닫힐 때 정리 작업을 하는 Hook입니다.", "무한 반복을 막기 위해 의존성 배열을 이해해야 합니다. []는 첫 렌더 이후 한 번 실행한다는 뜻입니다."),
    ("useMemo", "계산 결과를 잠시 기억해서 불필요한 재계산을 줄이는 Hook입니다. 홈 추천 카드와 혼잡 퍼센트 계산에 사용됩니다.", "useMemo는 데이터를 저장하는 용도가 아니라 이미 있는 데이터에서 파생값을 계산하는 용도입니다."),
    ("라우터", "URL에 맞는 페이지 컴포넌트를 선택하는 장치입니다. /stage-map이면 StageMapPage가 렌더링됩니다.", "페이지 이동은 전체 새로고침이 아니라 React Router가 화면만 바꾸는 방식입니다."),
    ("Outlet", "App.jsx 안에서 현재 라우트의 실제 페이지가 끼워지는 자리입니다.", "공통 레이아웃과 페이지 본문을 나누기 위해 필요합니다."),
    ("NavLink", "현재 URL과 메뉴 URL을 비교해서 active 클래스를 붙일 수 있는 링크입니다.", "일반 a 태그보다 SPA 라우팅에 적합합니다."),
    ("lazy와 Suspense", "페이지 코드를 처음부터 모두 불러오지 않고 필요할 때 나눠 불러오는 방식입니다.", "사용자가 특정 페이지를 열 때만 해당 페이지 번들이 로드되므로 초기 로딩이 가벼워집니다."),
    ("fetch", "브라우저가 서버에 HTTP 요청을 보내는 기본 함수입니다. FestFlow는 api.js에서 감싸서 사용합니다.", "fetch가 실패하면 화면 컴포넌트가 아니라 api.js에서 먼저 에러 메시지를 만들도록 되어 있습니다."),
    ("Promise", "아직 끝나지 않은 비동기 작업을 표현하는 객체입니다. API 응답을 기다릴 때 사용합니다.", "then, catch, async/await 모두 Promise를 다루는 방식입니다."),
    ("Promise.allSettled", "여러 요청을 동시에 보내고, 실패한 요청이 있어도 나머지 결과를 받을 수 있게 해줍니다.", "홈 화면은 부스/공연/방문 데이터를 동시에 요청하고 일부 실패 시 기본 데이터를 보여줍니다."),
    ("EventSource", "SSE 연결을 여는 브라우저 API입니다. 서버가 새 이벤트를 보내면 프론트가 즉시 받습니다.", "WebSocket과 달리 서버에서 브라우저로 밀어주는 단방향 실시간 스트림에 가깝습니다."),
    ("localStorage", "브라우저에 작은 문자열 데이터를 저장하는 공간입니다. 새로고침 후에도 남습니다.", "토큰, 언어, 즐겨찾기, 최근 본 부스 같은 값이 저장됩니다."),
    ("환경 변수", "개발/배포 환경마다 달라지는 값을 코드 밖에서 넣는 방식입니다. VITE_API_BASE_URL이 대표적입니다.", "Vite 프론트 환경 변수는 VITE_ 접두사가 있어야 브라우저 코드에서 읽을 수 있습니다."),
    ("PWA", "웹앱을 앱처럼 설치하거나 오프라인 화면을 제공하는 기능 묶음입니다.", "manifest.json, service-worker.js, offline.html이 관련 파일입니다."),
    ("Tailwind CSS", "className에 유틸리티 클래스를 적어 스타일을 빠르게 만드는 CSS 도구입니다.", "FestFlow는 Tailwind와 index.css의 커스텀 클래스를 같이 씁니다."),
    ("Leaflet", "지도 표시를 위한 라이브러리입니다. StageMapPage에서 MapContainer, Marker, Popup이 사용됩니다.", "지도 타일은 OpenStreetMap URL에서 불러옵니다."),
    ("FormData", "이미지나 파일을 업로드할 때 사용하는 브라우저 객체입니다.", "부스 이미지, 분실물 사진, AI 매칭 프로필 사진 업로드에 필요합니다."),
    ("Authorization 헤더", "로그인 토큰을 서버에 보낼 때 쓰는 HTTP 헤더입니다.", "관리자 API는 withAuth가 Authorization: Bearer 토큰을 붙여 보냅니다."),
    ("fallback 데이터", "서버 연결이 실패했을 때 화면이 빈 화면이 되지 않도록 준비한 기본 데이터입니다.", "fallbackBooths, fallbackEvents가 대표 예시입니다."),
    ("MutationObserver", "DOM 변화가 생겼을 때 감지하는 브라우저 API입니다. i18n.js가 새 텍스트도 번역하기 위해 사용합니다.", "React 렌더 후 생기는 텍스트까지 번역하려는 목적입니다."),
    ("controlled input", "input 값이 React state로 관리되는 입력창입니다. value와 onChange가 함께 있습니다.", "HomePage의 AI 질문 입력창, 로그인 폼, 검색창이 이런 패턴입니다."),
    ("navigate", "React Router의 페이지 이동 함수입니다. 버튼 클릭 후 /events 같은 주소로 이동합니다.", "링크가 아닌 버튼에서 화면 이동을 할 때 유용합니다."),
    ("cleanup", "useEffect가 끝날 때 실행되는 정리 함수입니다. SSE 연결과 interval을 닫는 데 사용됩니다.", "정리하지 않으면 페이지를 떠난 뒤에도 네트워크 연결이 남을 수 있습니다."),
]


BACKEND_TERMS = [
    ("Spring Boot", "Java 백엔드를 빠르게 만들 수 있게 해주는 프레임워크입니다. 서버 실행, 웹 요청 처리, 설정, 보안, DB 연결을 묶어줍니다.", "FestFlow의 BackendApplication이 Spring Boot 앱의 시작점입니다."),
    ("Controller", "HTTP URL을 Java 메서드에 연결하는 클래스입니다. /api/booths 요청은 BoothController로 들어갑니다.", "컨트롤러는 일을 직접 다 처리하기보다 Service에 위임합니다."),
    ("Service", "실제 비즈니스 규칙이 들어가는 클래스입니다. 예약 가능 여부, 혼잡도 계산, AI 추천 같은 핵심 로직이 여기에 있습니다.", "Service를 읽으면 이 앱이 어떤 규칙으로 움직이는지 알 수 있습니다."),
    ("Repository", "DB에서 엔티티를 조회하고 저장하는 인터페이스입니다. Spring Data JPA가 구현체를 자동으로 만들어줍니다.", "findById, findAll, save 같은 메서드는 Repository에서 시작합니다."),
    ("Entity", "DB 테이블과 연결되는 Java 클래스입니다. Booth, FestivalEvent, GpsLog가 대표 엔티티입니다.", "엔티티는 서버 내부 저장 모델이고, 프론트에 그대로 보내기보다 DTO로 바꿉니다."),
    ("DTO", "API 요청/응답에 사용하는 데이터 모양입니다. BoothResponseDto, ChatRequestDto 등이 있습니다.", "DTO를 보면 프론트가 받을 JSON 필드를 예측할 수 있습니다."),
    ("JPA", "Java 객체와 DB 테이블을 연결해주는 기술입니다. SQL을 직접 많이 쓰지 않아도 객체 저장이 가능합니다.", "복잡한 성능 최적화가 필요할 때는 Repository 메서드와 쿼리 사용을 확인해야 합니다."),
    ("@RestController", "이 클래스가 JSON API 컨트롤러라는 뜻입니다.", "반환값이 HTML 페이지가 아니라 JSON 응답으로 나갑니다."),
    ("@RequestMapping", "컨트롤러 전체의 공통 URL 앞부분을 지정합니다.", "BoothController의 /api/booths와 @GetMapping이 합쳐져 최종 URL이 됩니다."),
    ("@GetMapping", "GET 요청을 처리합니다. 조회 API에 주로 씁니다.", "목록 조회, 상세 조회, 분석 조회가 대부분 GET입니다."),
    ("@PostMapping", "POST 요청을 처리합니다. 생성, 로그인, 실행 액션에 주로 씁니다.", "예약 생성, 채팅 질문, 로그인, 업로드가 POST입니다."),
    ("@PutMapping", "PUT 요청을 처리합니다. 기존 데이터를 수정할 때 씁니다.", "부스 수정, 공지 수정, 상태 변경에서 사용됩니다."),
    ("@DeleteMapping", "DELETE 요청을 처리합니다. 데이터를 삭제할 때 씁니다.", "관리자 삭제 API들이 이 방식을 사용합니다."),
    ("@Transactional", "메서드 안의 DB 작업을 하나의 트랜잭션으로 묶습니다.", "예약 생성처럼 좌석 차감과 예약 저장이 함께 성공해야 하는 작업에 중요합니다."),
    ("ResponseStatusException", "서버가 특정 HTTP 상태 코드와 메시지를 반환하게 만드는 예외입니다.", "NOT_FOUND, CONFLICT, TOO_MANY_REQUESTS 같은 상황을 명확히 표현합니다."),
    ("Spring Security", "요청 권한을 검사하는 보안 프레임워크입니다.", "SecurityConfig가 어떤 API를 공개하고 어떤 API를 보호할지 정합니다."),
    ("JWT", "로그인 후 발급되는 토큰입니다. 관리자 API는 Authorization 헤더의 JWT를 읽어 ADMIN 권한을 확인합니다.", "세션을 서버에 저장하지 않는 stateless 구조입니다."),
    ("Filter", "컨트롤러에 도달하기 전에 요청을 먼저 검사하는 코드입니다.", "JwtAuthenticationFilter와 OpsKeyAuthenticationFilter가 대표입니다."),
    ("CORS", "프론트 주소와 백엔드 주소가 다를 때 브라우저가 요청을 허용할지 결정하는 정책입니다.", "개발 중 http://localhost:*를 허용하도록 설정되어 있습니다."),
    ("SSE", "Server-Sent Events의 약자입니다. 서버가 브라우저로 실시간 이벤트를 계속 보낼 수 있습니다.", "StreamController와 StreamService가 담당합니다."),
    ("SseEmitter", "Spring에서 SSE 연결 하나를 표현하는 객체입니다.", "StreamService는 연결된 emitter 목록에 이벤트를 보냅니다."),
    ("환경 변수", "DB 주소, JWT secret, OpenAI key처럼 환경마다 다른 값을 코드 밖에서 주입합니다.", "application.properties에서 ${NAME:default} 형태로 읽습니다."),
    ("profile", "local, prod, mysql 같은 실행 환경 구분입니다.", "bootRun은 별도 지정이 없으면 local 프로필을 사용합니다."),
    ("CommandLineRunner", "서버 시작 후 자동으로 실행되는 초기화 코드입니다.", "DataInitializer가 데모 데이터와 관리자 계정을 넣는 데 사용합니다."),
    ("PasswordEncoder", "비밀번호를 평문이 아니라 해시로 저장하기 위한 도구입니다.", "BCryptPasswordEncoder가 등록되어 있습니다."),
    ("Multipart", "파일 업로드 요청 형식입니다.", "이미지 업로드 API들이 multipart/form-data를 사용합니다."),
    ("RestClient", "백엔드에서 외부 HTTP API를 호출하는 Spring 클라이언트입니다.", "ChatService가 OpenAI Responses API를 호출할 때 사용합니다."),
    ("Validation", "요청 DTO 값이 올바른지 검사하는 기능입니다.", "@Valid가 붙은 컨트롤러 파라미터에서 동작합니다."),
    ("Enum", "정해진 값 중 하나만 갖게 하는 Java 타입입니다.", "ReservationStatus, StaffStatus가 상태값을 안전하게 표현합니다."),
    ("Record", "DTO처럼 불변 데이터 묶음을 짧게 정의할 수 있는 Java 문법입니다.", "많은 FestFlow DTO가 record입니다."),
    ("Scheduler", "정해진 주기로 메서드를 실행하는 기능입니다.", "EventService는 30초마다 공연 상태를 SSE로 방송합니다."),
    ("No-show", "예약 후 체크인하지 않은 사용자를 의미합니다.", "ReservationService는 만료 예약을 처리하며 노쇼 횟수를 기록합니다."),
]


FRONTEND_ROUTES = [
    ("/", "HomePage", "홈 화면. 추천 카드, AI 축제 가이드, 실시간 부스/공연 스트림을 보여줍니다."),
    ("/stage-map", "StageMapPage", "지도 화면. 부스 마커, 카테고리 필터, 검색, GPS 전송, AI 안내를 담당합니다."),
    ("/events", "EventPage", "공연 목록 화면. 공연 상태별 필터와 AI 방문 가이드를 표시합니다."),
    ("/events/lineup", "LineupPage", "라인업 중심 공연 화면. 공연 스트림을 받아 최신 상태를 반영합니다."),
    ("/analytics", "AnalyticsPage", "혼잡도 분석 화면. 대시보드, 예측, AI 판단 이력을 표시합니다."),
    ("/booths/:id", "BoothDetailPage", "부스 상세 화면. 메뉴, 혼잡도, 예약, QR 체크인 토큰을 다룹니다."),
    ("/lost-found", "LostFoundPage", "분실물 화면. 분실물 조회, 신고, 상태 확인 기능을 제공합니다."),
    ("/chat", "ChatPage", "AI 챗봇 화면. 질문을 보내고 근거 기반 답변을 표시합니다."),
    ("/staff", "StaffPage", "스태프 화면. 로그인, 담당 구역, 분실물, AI 지원 기능을 제공합니다."),
    ("/more", "MorePage", "더보기/설정 화면. 언어, 부가 기능, 주요 링크를 제공합니다."),
    ("/admin", "AdminPage", "관리자 화면. 로그인 후 부스, 공연, 공지, KPI, 로그를 관리합니다."),
    ("/ops/master", "OpsMasterPage", "운영 마스터 화면. 운영 키로 현장 운영 데이터를 관리합니다."),
    ("/ops/booth/:id", "OpsBoothPage", "부스 운영 화면. 부스별 키로 상태, 메뉴 이미지, 예약을 관리합니다."),
    ("/ai-match", "AiMatchPage", "AI 매칭 사용자 화면. 프로필, 요청, 즐겨찾기, 만남 제안을 다룹니다."),
    ("/ai-match/admin", "AiMatchAdminPage", "AI 매칭 관리자 화면. 요청과 프로필을 관리합니다."),
    ("*", "Navigate", "정의되지 않은 주소는 홈으로 되돌립니다."),
]


FRONTEND_API_GROUPS = {
    "공개 방문자 API": [
        ("fetchBooths", "GET /api/booths", "HomePage, StageMapPage, AdminPage", "부스 목록을 가져옵니다."),
        ("fetchBoothById", "GET /api/booths/{id}", "BoothDetailPage", "부스 상세 정보를 가져옵니다."),
        ("fetchCongestion", "GET /api/booths/{id}/congestion", "BoothDetailPage", "특정 부스 주변 혼잡도를 가져옵니다."),
        ("fetchEvents", "GET /api/events", "HomePage, EventPage, LineupPage", "공연 목록과 상태를 가져옵니다."),
        ("fetchTrafficHourly", "GET /api/analytics/traffic-hourly", "HomePage", "최근 24시간 방문자 데이터를 가져옵니다."),
        ("fetchAnalyticsDashboard", "GET /api/analytics/dashboard", "AnalyticsPage", "구역별 혼잡도 대시보드를 가져옵니다."),
        ("fetchStageCrowd", "GET /api/analytics/stage-crowd", "EventPage/Analytics 관련", "무대 주변 관중 데이터를 가져옵니다."),
        ("sendGps", "POST /api/gps", "StageMapPage", "사용자 GPS를 백엔드에 전송해 혼잡도 계산에 반영합니다."),
    ],
    "AI API": [
        ("fetchAiFestivalGuide", "GET /api/ai/guide", "HomePage", "홈 화면 AI 축제 가이드를 가져옵니다."),
        ("fetchAiVisitorGuide", "GET /api/ai/visitor-guide/{scope}", "StageMapPage, EventPage, AnalyticsPage", "페이지별 AI 방문 안내를 가져옵니다."),
        ("fetchAiCongestionPredictions", "GET /api/ai/congestion/predictions", "AnalyticsPage", "AI 혼잡도 예측 목록을 가져옵니다."),
        ("fetchAiDecisionLogs", "GET /api/ai/decisions", "AnalyticsPage", "AI 판단 이력을 가져옵니다."),
        ("askChat", "POST /api/chat", "HomePage, ChatPage", "질문을 보내고 근거 기반 답변을 받습니다."),
    ],
    "실시간 SSE API": [
        ("createBoothStream", "GET /api/stream/booths", "HomePage, StageMapPage, BoothDetailPage", "부스 상태 변경을 실시간으로 받습니다."),
        ("createEventStream", "GET /api/stream/events", "HomePage, EventPage, LineupPage", "공연 상태 변경을 실시간으로 받습니다."),
        ("createCongestionStream", "GET /api/stream/congestion", "AnalyticsPage", "혼잡도 갱신 이벤트를 실시간으로 받습니다."),
        ("createNoticeStream", "GET /api/stream/notices", "관리/운영 화면", "공지 변경을 실시간으로 받습니다."),
        ("createReservationStream", "GET /api/stream/reservations", "BoothDetailPage, OpsBoothPage", "예약 상태 변경을 실시간으로 받습니다."),
        ("createStaffStream", "GET /api/stream/staff", "StaffPage", "스태프 상태 변경을 실시간으로 받습니다."),
        ("createLostItemStream", "GET /api/stream/lost-items", "LostFoundPage, StaffPage", "분실물 변경을 실시간으로 받습니다."),
    ],
    "관리/운영 API": [
        ("loginAdmin", "POST /api/auth/login", "AdminPage, AiMatchAdminPage", "관리자 로그인 후 JWT를 받습니다."),
        ("createBooth/updateBooth/deleteBooth", "/api/admin/booths", "AdminPage", "관리자가 부스를 생성, 수정, 삭제합니다."),
        ("createEvent/updateEvent/deleteEvent", "/api/admin/events", "AdminPage", "관리자가 공연을 생성, 수정, 삭제합니다."),
        ("createNotice/updateNotice/deleteNotice", "/api/admin/notices", "AdminPage", "관리자가 공지를 생성, 수정, 삭제합니다."),
        ("fetchOpsMasterBootstrap", "GET /api/ops/master/bootstrap", "OpsMasterPage", "운영 마스터 초기 데이터를 가져옵니다."),
        ("fetchOpsBoothBootstrap", "GET /api/ops/booth/{id}/bootstrap", "OpsBoothPage", "부스 운영자 초기 데이터를 가져옵니다."),
    ],
    "예약/스태프/분실물/AI 매칭 API": [
        ("sendReservationAuthCode", "POST /api/reservations/auth/send-code", "BoothDetailPage", "전화번호 인증 코드를 요청합니다."),
        ("verifyReservationAuthCode", "POST /api/reservations/auth/verify-code", "BoothDetailPage", "인증 코드를 검증하고 예약 토큰을 받습니다."),
        ("fetchBoothReservations", "GET /api/booths/{id}/reservations", "BoothDetailPage", "부스 예약 상태와 좌석 정보를 가져옵니다."),
        ("createBoothReservation", "POST /api/booths/{id}/reservations", "BoothDetailPage", "사용자가 부스 예약을 생성합니다."),
        ("loginStaff", "POST /api/staff/auth/login", "StaffPage", "스태프 번호와 PIN으로 로그인합니다."),
        ("fetchLostItems/createLostItem/updateLostItem", "/api/lost-items", "LostFoundPage, StaffPage", "분실물 조회, 생성, 상태 변경을 처리합니다."),
        ("fetchAiMatchProfiles", "GET /api/ai-match/profiles", "AiMatchPage", "AI 매칭 프로필 목록을 가져옵니다."),
        ("createAiMatchProfile", "POST /api/ai-match/profiles", "AiMatchPage", "사진과 폼 데이터를 업로드해 프로필을 생성합니다."),
    ],
}


BACKEND_ENDPOINT_GROUPS = {
    "방문자 공개 API": [
        ("BoothController", "GET /api/booths", "전체 부스 목록을 displayOrder 기준으로 반환합니다."),
        ("BoothController", "GET /api/booths/{id}", "특정 부스 상세 정보를 반환합니다."),
        ("BoothController", "GET /api/booths/{id}/congestion", "GPS 로그를 기반으로 특정 부스 혼잡도를 계산합니다."),
        ("EventController", "GET /api/events", "공연 목록과 현재 상태를 반환합니다."),
        ("NoticeController", "GET /api/notices/active", "활성 공지 목록을 반환합니다."),
        ("GpsController", "POST /api/gps", "사용자 위치 로그를 저장합니다."),
        ("AnalyticsController", "GET /api/analytics/traffic-hourly", "최근 24시간 시간대별 방문 데이터를 반환합니다."),
        ("AnalyticsController", "GET /api/analytics/popular-booths", "최근 1시간 인기 부스 랭킹을 반환합니다."),
        ("AnalyticsController", "GET /api/analytics/congestion-heatmap", "혼잡 히트맵 포인트를 반환합니다."),
        ("AnalyticsController", "GET /api/analytics/dashboard", "구역별 혼잡 대시보드를 반환합니다."),
    ],
    "AI API": [
        ("AiGuideController", "GET /api/ai/guide", "AI 축제 가이드 DTO를 반환합니다."),
        ("AiGuideController", "GET /api/ai/congestion/predictions", "AI 혼잡도 예측 목록을 반환합니다."),
        ("AiGuideController", "GET /api/ai/decisions", "AI 판단 로그를 반환합니다."),
        ("AiGuideController", "GET /api/ai/visitor-guide/{scope}", "페이지별 방문자 가이드를 반환합니다."),
        ("ChatController", "POST /api/chat", "질문을 받아 근거 검색 후 AI 또는 fallback 답변을 반환합니다."),
        ("TranslateController", "POST /api/translate", "번역 요청을 처리합니다."),
        ("TranslateController", "GET /api/translate/metrics", "번역 사용량 지표를 반환합니다."),
    ],
    "관리자 API": [
        ("AuthController", "POST /api/auth/login", "관리자 로그인 후 JWT를 발급합니다."),
        ("AdminBoothController", "POST/PUT/DELETE /api/admin/booths", "부스 생성, 수정, 삭제, 순서 변경, 이미지 업로드를 처리합니다."),
        ("AdminEventController", "POST/PUT/DELETE /api/admin/events", "공연 생성, 수정, 삭제를 처리합니다."),
        ("AdminNoticeController", "GET/POST/PUT/DELETE /api/admin/notices", "관리자 공지 CRUD를 처리합니다."),
        ("AdminDashboardController", "GET /api/admin/dashboard/kpis", "관리자 KPI를 반환합니다."),
        ("AdminDashboardController", "GET /api/admin/audit-logs", "관리자 작업 로그를 반환합니다."),
        ("AdminStaffController", "GET/PUT /api/admin/staff", "스태프 목록과 상태를 관리합니다."),
        ("AdminAiMatchController", "/api/admin/ai-match/**", "AI 매칭 운영 현황과 프로필 관리를 처리합니다."),
    ],
    "운영/스태프 API": [
        ("OpsController", "GET /api/ops/master/bootstrap", "운영 마스터 초기 데이터를 반환합니다."),
        ("OpsController", "/api/ops/master/**", "운영 마스터가 공지, 공연, 부스, AI 브리핑을 관리합니다."),
        ("OpsController", "GET /api/ops/booth/{id}/bootstrap", "부스 운영자 초기 데이터를 반환합니다."),
        ("OpsController", "/api/ops/booth/{id}/reservations/**", "부스별 예약 관리, 체크인, 완료, 테이블 해제를 처리합니다."),
        ("StaffController", "POST /api/staff/auth/login", "스태프 로그인을 처리합니다."),
        ("StaffController", "GET /api/staff/bootstrap", "스태프 화면 초기 데이터를 반환합니다."),
        ("StaffController", "PUT /api/staff/me/status", "내 스태프 상태를 변경합니다."),
        ("StaffController", "/api/staff/ai/**", "스태프용 AI 요약, 분실물 답변, 현장 체크리스트를 생성합니다."),
    ],
    "예약/분실물/업로드/실시간": [
        ("ReservationAuthController", "POST /api/reservations/auth/send-code", "예약 전화번호 인증 코드를 발송합니다."),
        ("ReservationAuthController", "POST /api/reservations/auth/verify-code", "인증 코드를 확인하고 예약 인증 토큰을 발급합니다."),
        ("BoothController", "GET/POST /api/booths/{id}/reservations", "부스 예약 상태 조회와 예약 생성을 처리합니다."),
        ("BoothController", "POST /api/booths/{id}/reservations/{reservationId}/check-in-token", "예약자 체크인 QR 토큰을 발급합니다."),
        ("LostItemController", "/api/lost-items/**", "분실물 조회, 등록, 상태 변경, 주인 확인, 삭제를 처리합니다."),
        ("UploadAssetController", "GET /uploads/{filename}", "로컬 업로드 파일을 반환합니다."),
        ("StreamController", "GET /api/stream/*", "SSE 실시간 스트림 구독을 처리합니다."),
        ("ExportController", "GET /api/export/*.csv", "부스/공연 CSV 다운로드를 제공합니다."),
    ],
}


ENTITY_SUMMARY = [
    ("Booth", "부스", "이름, 위치, 설명, 이미지, 대기시간, 재고, 카테고리, 운영시간, 예약 가능 여부를 저장합니다."),
    ("FestivalEvent", "공연", "공연 제목, 시작/종료 시간, 상태 override, 이미지, 지연 정보, 실시간 메시지를 저장합니다."),
    ("GpsLog", "위치 로그", "방문자가 보낸 위도/경도와 생성 시간을 저장해 혼잡도 계산에 사용합니다."),
    ("Notice", "공지", "제목, 내용, 카테고리, 활성 여부를 저장합니다."),
    ("LostItem", "분실물", "분실물 제목, 설명, 사진, 발견 위치, 상태, 연락 관련 정보를 저장합니다."),
    ("AdminUser", "관리자", "관리자 계정, 암호 해시, 역할을 저장합니다."),
    ("AuditLog", "감사 로그", "관리자가 수행한 주요 작업을 추적합니다."),
    ("StaffMember", "스태프", "스태프 번호, 이름, PIN 해시, 담당 구역, 현재 상태를 저장합니다."),
    ("StaffSession", "스태프 세션", "스태프 로그인 토큰과 만료 시간을 저장합니다."),
    ("BoothReservationTable", "예약 테이블", "부스별 테이블명, 전체 좌석, 남은 좌석, 표시 순서를 저장합니다."),
    ("BoothReservation", "예약", "부스, 테이블, 사용자 키, 좌석 수, 상태, 예약/만료 시간을 저장합니다."),
    ("ReservationCheckInToken", "체크인 토큰", "예약자가 부스에서 보여줄 QR 토큰과 만료/사용 여부를 저장합니다."),
    ("ReservationAuthSession", "예약 인증 세션", "전화번호 인증 후 발급된 예약 사용자 세션을 저장합니다."),
    ("ReservationVerificationCode", "인증 코드", "전화번호 인증 코드, 만료 시간, 시도 횟수를 저장합니다."),
    ("ReservationUserState", "예약 사용자 상태", "노쇼 횟수와 차단 상태를 저장합니다."),
    ("AiMatchProfile", "AI 매칭 프로필", "매칭용 프로필, 사진, 닉네임, PIN, 자기소개 데이터를 저장합니다."),
    ("AiMatchRequest", "AI 매칭 요청", "매칭 요청, 수락/거절/취소/만남 제안 상태를 저장합니다."),
    ("AiMatchFavorite", "AI 매칭 즐겨찾기", "프로필 즐겨찾기 관계를 저장합니다."),
    ("AiMatchPhoneUsage", "전화번호 사용", "AI 매칭 전화번호 중복/사용 제한 정보를 저장합니다."),
]


FRONTEND_FILE_DEEP_DIVES = [
    ("main.jsx", "앱 시작과 라우팅", "처음에는 lazy import 목록을 보고 이 앱에 어떤 화면이 있는지 파악합니다. 그다음 Routes 구조를 따라가며 URL과 페이지 컴포넌트의 관계를 외우지 말고 연결합니다.", ["새 페이지를 추가할 때", "페이지 URL이 잘못 이동될 때", "처음 로딩 fallback 문구를 바꿀 때", "PWA service worker 등록 조건을 확인할 때"]),
    ("App.jsx", "공통 레이아웃과 하단 메뉴", "App은 모든 화면의 공통 껍데기입니다. Outlet이 실제 페이지가 들어가는 자리이고, location.pathname으로 메뉴 표시 여부를 결정합니다.", ["하단 메뉴 항목을 추가할 때", "운영 화면에서 일반 메뉴가 보이는 문제를 고칠 때", "active 메뉴 표시가 잘못될 때", "페이지 전체 레이아웃 여백을 확인할 때"]),
    ("api.js", "백엔드 통신 허브", "프론트에서 백엔드로 나가는 길은 대부분 api.js를 거칩니다. 함수 이름, HTTP method, URL, headers, body, parseJson 에러 메시지를 한 세트로 봐야 합니다.", ["새 백엔드 API를 프론트에 연결할 때", "CORS나 서버 연결 오류를 추적할 때", "관리자 토큰이 API에 붙는지 확인할 때", "SSE 스트림을 새로 만들 때"]),
    ("HomePage.jsx", "첫 화면과 AI 추천", "홈 화면은 state, API, SSE, fallback, useMemo가 모두 들어 있는 종합 예제입니다. 이 파일을 잘 이해하면 다른 페이지도 훨씬 쉬워집니다.", ["홈 추천 카드를 바꿀 때", "AI 질문 입력창을 수정할 때", "서버 실패 시 기본 데이터를 바꿀 때", "실시간 부스/공연 갱신을 확인할 때"]),
    ("StageMapPage.jsx", "지도, 검색, GPS, 실시간 부스", "지도 화면은 UI보다 데이터 가공이 많습니다. 좌표 검증, fallback 좌표, 카테고리 판단, 검색 필터, 지도 viewport 조정이 서로 연결됩니다.", ["마커가 잘못 찍힐 때", "카테고리 필터가 이상할 때", "현재 위치 전송을 고칠 때", "지도 초기 zoom이나 범위를 바꿀 때"]),
    ("BoothDetailPage.jsx", "부스 상세와 예약 UX", "상세 화면은 사용자 예약과 운영자 예약 상태가 만나는 지점입니다. 인증 토큰, 테이블 상태, 내 예약, QR 토큰, 예약 스트림을 함께 봐야 합니다.", ["예약 버튼이 비활성화될 때", "좌석 상태가 늦게 갱신될 때", "QR 토큰 생성 흐름을 바꿀 때", "부스 메뉴 표시를 개선할 때"]),
    ("AnalyticsPage.jsx", "혼잡도 대시보드", "분석 화면은 서버 계산 결과를 시각적으로 해석합니다. 숫자 자체보다 percent, level, trend, recommendation이 화면에 어떻게 쓰이는지 봅니다.", ["혼잡도 그래프나 카드가 비어 있을 때", "AI 예측 영역을 수정할 때", "SSE 혼잡도 갱신을 추가할 때", "분석 기준 시간을 바꿀 때"]),
    ("EventPage.jsx", "공연 일정", "공연 화면은 시간과 상태가 핵심입니다. 시작/종료 시간, statusOverride, liveMessage, delayMinutes를 어떻게 사용자에게 보여주는지 확인합니다.", ["공연 상태 라벨을 바꿀 때", "공연 필터를 추가할 때", "AI 공연 안내를 수정할 때", "공연 이미지 표시를 개선할 때"]),
    ("LineupPage.jsx", "공연 라인업 전용 화면", "LineupPage는 EventPage보다 공연 흐름을 더 강조합니다. 같은 events 데이터를 다른 UI로 보여주는 예로 읽으면 좋습니다.", ["라인업 정렬을 바꿀 때", "라이브 공연 강조 UI를 수정할 때", "공연 스트림 갱신을 확인할 때", "공연 시간 포맷을 바꿀 때"]),
    ("ChatPage.jsx", "AI 챗봇", "ChatPage는 질문 입력, 전송 중 상태, 답변 표시, 근거 표시를 다룹니다. HomePage의 간단 질문창보다 더 독립적인 대화 화면입니다.", ["AI 응답 표시 구조를 바꿀 때", "근거 버튼 이동을 추가할 때", "에러 메시지 UX를 개선할 때", "질문 예시를 추가할 때"]),
    ("LostFoundPage.jsx", "분실물 조회와 신고", "분실물 화면은 목록 조회와 파일 업로드가 섞입니다. FormData, staffToken, status update, claim flow를 구분해서 봅니다.", ["분실물 사진 업로드 오류를 고칠 때", "상태 필터를 바꿀 때", "분실물 상세 표시를 추가할 때", "실시간 분실물 갱신을 확인할 때"]),
    ("StaffPage.jsx", "스태프 현장 콘솔", "스태프 화면은 로그인 토큰, 담당 부스, 분실물, AI 지원 기능을 한 화면에 담습니다. 관리자 화면과 달리 스태프 전용 토큰 흐름을 봐야 합니다.", ["스태프 로그인 문제를 고칠 때", "상태 변경 버튼을 수정할 때", "현장 AI 요약을 바꿀 때", "분실물 처리 권한을 확인할 때"]),
    ("AdminPage.jsx", "관리자 콘솔", "관리자 화면은 JWT 로그인 이후 여러 관리자 API를 호출합니다. 이 파일을 읽을 때는 토큰 저장, 데이터 로딩, CRUD 핸들러, 에러 state를 묶어 봅니다.", ["관리자 로그인 오류를 고칠 때", "부스/공연 CRUD를 수정할 때", "공지 관리 기능을 추가할 때", "감사 로그 표시를 바꿀 때"]),
    ("OpsMasterPage.jsx", "운영 마스터 콘솔", "운영 마스터 화면은 비밀번호 로그인 대신 운영 키로 접근합니다. 현장 운영자가 빠르게 작업하도록 관리자보다 실무 중심으로 설계되어 있습니다.", ["운영 키 입력 UX를 고칠 때", "현장 공지 발행을 수정할 때", "AI 브리핑 버튼을 추가할 때", "운영 데이터 초기 로딩을 확인할 때"]),
    ("OpsBoothPage.jsx", "부스별 운영 콘솔", "부스 운영 화면은 특정 부스만 다루는 제한된 운영 도구입니다. 부스 키, 예약 스트림, 체크인/완료 처리, 테이블 해제 흐름을 봅니다.", ["부스 키 권한 문제를 고칠 때", "예약 상태가 갱신되지 않을 때", "메뉴 이미지 업로드를 수정할 때", "체크인 토큰 처리 오류를 볼 때"]),
    ("AiMatchPage.jsx", "AI 매칭 사용자 기능", "AI 매칭은 폼, 이미지, PIN, 요청 상태, 즐겨찾기가 섞인 복합 화면입니다. 일반 축제 화면보다 사용자 플로우가 길기 때문에 상태 관리가 중요합니다.", ["프로필 생성 폼을 수정할 때", "이미지 미리보기 오류를 고칠 때", "요청 수락/거절 UX를 바꿀 때", "PIN 접근 흐름을 점검할 때"]),
    ("AiMatchAdminPage.jsx", "AI 매칭 관리자 기능", "AI 매칭 관리자 화면은 일반 AdminPage와 별도로 AI 매칭 요청과 프로필 상태를 관리합니다. 관리자 로그인과 현황 조회를 먼저 봅니다.", ["매칭 현황 카드가 비어 있을 때", "연결 상태 변경 버튼을 수정할 때", "프로필 삭제 흐름을 확인할 때", "관리자 메모 저장을 고칠 때"]),
    ("i18n.js", "언어 전환과 DOM 번역", "이 파일은 React context와 DOM 후처리를 함께 사용합니다. 언어 state가 바뀌면 localStorage와 document lang, 텍스트 노드가 함께 바뀝니다.", ["번역 문구를 추가할 때", "placeholder가 번역되지 않을 때", "일부 영역 번역을 막을 때", "영어 전환 후 레이아웃이 깨질 때"]),
    ("index.css", "전역 스타일과 디자인 시스템", "index.css는 앱 전체 UI 톤을 결정합니다. 페이지 전용 클래스와 공통 클래스가 섞여 있으므로 수정 전 영향 범위를 확인해야 합니다.", ["하단 nav가 겹칠 때", "카드 여백을 조정할 때", "지도 높이가 깨질 때", "모바일 반응형 문제를 고칠 때"]),
    ("data/festivalUiData.js", "fallback 데이터", "서버가 실패해도 화면이 비지 않게 하는 기본 데이터입니다. 실제 운영 데이터가 아니라 UX 안전망으로 이해해야 합니다.", ["서버 없이 프론트 화면을 시연할 때", "기본 부스/공연 예시를 바꿀 때", "fallback 이미지나 카테고리를 수정할 때", "서버 실패 메시지를 확인할 때"]),
    ("utils/*.js", "브라우저 저장과 보조 함수", "utils 폴더는 작은 기능을 모아 둔 곳입니다. 인증 토큰, 예약 토큰, 즐겨찾기, 위치 기준값처럼 여러 페이지에서 반복되는 일을 처리합니다.", ["localStorage 키를 바꿀 때", "관리자 토큰 저장 방식을 고칠 때", "예약 인증 값을 초기화할 때", "아주대 중심 좌표를 수정할 때"]),
]


FRONTEND_FLOW_CASES = [
    ("사용자가 홈에 처음 들어오는 흐름", ["브라우저가 / 주소를 엽니다.", "main.jsx의 index route가 HomePage를 선택합니다.", "App.jsx의 Outlet 자리에 HomePage가 들어갑니다.", "HomePage useEffect가 부스/공연/방문 데이터를 요청합니다.", "state가 채워지면 추천 카드와 AI 가이드가 화면에 표시됩니다."], ["main.jsx", "App.jsx", "HomePage.jsx", "api.js"]),
    ("홈 추천 카드가 바뀌는 흐름", ["백엔드 부스 데이터가 fetchBooths 또는 SSE로 들어옵니다.", "setBooths가 실행되어 booths state가 바뀝니다.", "boothSource가 fallback 대신 실제 데이터를 사용합니다.", "homeCards useMemo가 대기시간과 예약 가능 좌석을 다시 계산합니다.", "JSX가 새 카드 제목, 이미지, caption을 렌더링합니다."], ["HomePage.jsx", "api.js", "BoothController", "BoothService"]),
    ("AI 질문을 보내는 흐름", ["사용자가 질문 입력창에 타이핑합니다.", "aiQuestion state가 onChange로 갱신됩니다.", "submit 시 handleAiAsk가 askChat을 호출합니다.", "api.js가 POST /api/chat 요청을 보냅니다.", "응답 answer/evidence가 aiAnswer state에 들어가고 화면에 표시됩니다."], ["HomePage.jsx", "ChatPage.jsx", "api.js", "ChatController", "ChatService"]),
    ("지도에서 현재 위치를 보내는 흐름", ["사용자가 내 위치 버튼을 누릅니다.", "navigator.geolocation.getCurrentPosition이 현재 좌표를 가져옵니다.", "StageMapPage가 currentLocation state를 바꿉니다.", "sendGps가 POST /api/gps 요청을 보냅니다.", "백엔드 GpsLog가 쌓이고 혼잡도 계산 재료가 됩니다."], ["StageMapPage.jsx", "api.js", "GpsController", "GpsService", "GpsLog"]),
    ("부스 상세 예약 상태를 보는 흐름", ["사용자가 /booths/:id에 들어갑니다.", "BoothDetailPage가 id를 읽고 fetchBoothById를 호출합니다.", "예약 인증 토큰이 있으면 fetchBoothReservations도 호출합니다.", "테이블별 예약 가능 좌석과 내 예약 상태가 state에 들어갑니다.", "예약 UI가 가능한 버튼과 남은 시간을 표시합니다."], ["BoothDetailPage.jsx", "api.js", "BoothController", "ReservationService"]),
    ("예약을 만드는 흐름", ["사용자가 전화번호 인증을 완료합니다.", "부스 상세에서 테이블과 좌석 수를 고릅니다.", "createBoothReservation이 예약 토큰과 payload를 보냅니다.", "백엔드가 중복 예약과 좌석 가능 여부를 검사합니다.", "성공 응답과 예약 SSE가 프론트 화면을 갱신합니다."], ["BoothDetailPage.jsx", "reservationAuth.js", "api.js", "ReservationService"]),
    ("공연 상태가 실시간으로 바뀌는 흐름", ["EventService가 30초마다 공연 상태를 계산합니다.", "StreamService.publishEvents가 events 이벤트를 보냅니다.", "프론트 createEventStream 연결이 이벤트를 받습니다.", "setEvents가 실행됩니다.", "EventPage/LineupPage/HomePage의 공연 표시가 바뀝니다."], ["EventService", "StreamService", "api.js", "EventPage.jsx", "LineupPage.jsx"]),
    ("관리자 로그인이 필요한 API 흐름", ["AdminPage에서 username/password를 입력합니다.", "loginAdmin이 POST /api/auth/login을 호출합니다.", "백엔드가 JWT를 반환합니다.", "프론트가 token을 localStorage에 저장합니다.", "관리자 API 호출 때 withAuth가 Authorization 헤더를 붙입니다."], ["AdminPage.jsx", "utils/auth.js", "api.js", "AuthController", "JwtAuthenticationFilter"]),
    ("운영 마스터 키 흐름", ["운영자가 /ops/master에 들어갑니다.", "마스터 키를 입력합니다.", "fetchOpsMasterBootstrap이 X-Ops-Key 헤더를 보냅니다.", "백엔드 필터가 OPS_MASTER 권한을 넣습니다.", "운영 화면이 부스/공지/공연 데이터를 표시합니다."], ["OpsMasterPage.jsx", "api.js", "OpsKeyAuthenticationFilter", "OpsController"]),
    ("분실물 사진 업로드 흐름", ["사용자가 분실물 폼과 파일을 입력합니다.", "FormData에 텍스트 필드와 file이 들어갑니다.", "createLostItem이 multipart 요청을 보냅니다.", "백엔드가 파일을 저장하고 LostItem을 생성합니다.", "lost-items SSE가 목록 화면을 갱신합니다."], ["LostFoundPage.jsx", "StaffPage.jsx", "api.js", "LostItemController", "UploadStorageService"]),
    ("언어 전환 흐름", ["사용자가 언어 토글을 누릅니다.", "LanguageProvider의 setLanguage가 실행됩니다.", "localStorage에 festflow_language가 저장됩니다.", "applyDocumentTranslations가 텍스트와 속성을 번역합니다.", "MutationObserver가 이후 새로 생긴 DOM도 다시 번역합니다."], ["MorePage.jsx", "i18n.js", "localStorage"]),
    ("PWA 오프라인 흐름", ["프로덕션 빌드에서 serviceWorker 등록 조건이 통과합니다.", "브라우저가 public/service-worker.js를 등록합니다.", "정적 파일 일부가 캐시에 들어갑니다.", "네트워크가 끊기면 offline.html 또는 캐시된 자원이 사용됩니다.", "개발 중에는 캐시 때문에 이전 화면이 보일 수 있어 Application 탭에서 정리합니다."], ["main.jsx", "public/service-worker.js", "public/manifest.json", "public/offline.html"]),
]


BACKEND_CONTROLLER_DEEP_DIVES = [
    ("BoothController", "부스 공개 API와 예약 공개 API를 받습니다.", ["GET /api/booths는 전체 부스 목록입니다.", "GET /api/booths/{id}는 상세입니다.", "GET /api/booths/{id}/congestion은 혼잡도입니다.", "예약 조회/생성/체크인 토큰 발급도 같은 부스 맥락에서 처리합니다."]),
    ("EventController", "공연 목록 조회 API를 받습니다.", ["공연 상태는 EventService에서 계산됩니다.", "프론트 EventPage, LineupPage, HomePage가 같은 API를 사용합니다.", "상태 override가 있으면 자동 시간 계산보다 우선합니다."]),
    ("AnalyticsController", "GPS 로그 기반 분석 API를 받습니다.", ["traffic-hourly는 홈과 분석에서 방문 흐름을 보여줍니다.", "popular-booths는 최근 1시간 부스 주변 로그를 봅니다.", "dashboard는 구역별 현재/이전 혼잡도를 비교합니다."]),
    ("AiGuideController", "AI 축제 가이드와 예측 API를 받습니다.", ["guide는 홈 화면 AI 가이드입니다.", "visitor-guide/{scope}는 페이지별 안내입니다.", "decisions는 AI 판단 이력입니다.", "congestion/predictions는 부스별 AI 위험 판단입니다."]),
    ("ChatController", "AI 질문 API를 받습니다.", ["ChatRequestDto의 question을 받습니다.", "ChatService가 근거 검색과 OpenAI/fallback 응답을 담당합니다.", "프론트는 answer와 evidence를 화면에 보여줍니다."]),
    ("GpsController", "사용자 위치 전송 API를 받습니다.", ["StageMapPage의 sendGps가 호출합니다.", "저장된 GpsLog는 혼잡도와 분석의 입력값입니다.", "위치 권한 거부 시 프론트에서 메시지를 보여줍니다."]),
    ("NoticeController", "활성 공지 공개 조회를 받습니다.", ["방문자 화면에서 운영 공지를 보여줄 수 있습니다.", "관리자/운영자는 별도 API로 공지를 생성합니다.", "SSE notices 채널과 연결될 수 있습니다."]),
    ("LostItemController", "분실물 조회/등록/상태 변경을 받습니다.", ["GET은 공개 조회 성격입니다.", "POST는 multipart 파일 업로드가 필요합니다.", "상태 변경과 삭제는 스태프 토큰 확인이 중요합니다."]),
    ("AuthController", "관리자 로그인 API를 받습니다.", ["username/password를 검증합니다.", "성공하면 JWT를 반환합니다.", "프론트 withAuth가 이후 관리자 API에 JWT를 붙입니다."]),
    ("AdminBoothController", "관리자 부스 CRUD를 받습니다.", ["부스 생성/수정/삭제를 처리합니다.", "이미지 업로드와 순서 변경도 포함합니다.", "변경 후 실시간 부스 스트림 방송이 필요할 수 있습니다."]),
    ("AdminEventController", "관리자 공연 CRUD를 받습니다.", ["공연 생성/수정/삭제를 처리합니다.", "상태 override와 지연 정보가 프론트 공연 상태에 영향을 줍니다.", "EventService의 상태 계산과 함께 봐야 합니다."]),
    ("AdminNoticeController", "관리자 공지 CRUD를 받습니다.", ["관리자가 공지를 생성하고 활성화할 수 있습니다.", "활성 공지는 NoticeController 공개 API로 조회됩니다.", "변경 후 StreamService.publishNotices를 고려해야 합니다."]),
    ("AdminDashboardController", "관리자 KPI와 감사 로그 API를 받습니다.", ["관리 콘솔 상단 지표를 제공합니다.", "관리자 작업 이력 추적에 사용됩니다.", "보안상 ADMIN 권한이 필요합니다."]),
    ("AdminStaffController", "관리자 스태프 관리 API를 받습니다.", ["스태프 목록 조회와 상태/정보 수정에 사용됩니다.", "StaffPage와 AdminPage의 역할이 다릅니다.", "스태프 로그인 토큰과 관리자 JWT를 구분해야 합니다."]),
    ("OpsController", "현장 운영 API를 받습니다.", ["운영 마스터와 부스 운영자 API가 한 컨트롤러에 있습니다.", "X-Ops-Key 기반 권한이 필요합니다.", "예약 체크인, 완료, 테이블 해제처럼 현장 작업이 많습니다."]),
    ("StaffController", "스태프 API를 받습니다.", ["스태프 로그인, 로그아웃, bootstrap, 내 상태 변경을 처리합니다.", "AI 현장 지원 API도 포함합니다.", "스태프 토큰은 관리자 JWT와 다릅니다."]),
    ("StreamController", "SSE 구독 API를 받습니다.", ["각 GET 요청은 text/event-stream을 생산합니다.", "컨트롤러는 StreamService.subscribe...를 반환합니다.", "프론트 EventSource URL과 정확히 맞아야 합니다."]),
    ("ReservationAuthController", "예약 전화번호 인증 API를 받습니다.", ["send-code는 인증 코드 발급입니다.", "verify-code는 코드 확인과 예약 인증 토큰 발급입니다.", "예약 생성 전에 프론트가 이 흐름을 거칩니다."]),
    ("UploadAssetController", "업로드 파일 조회 API를 받습니다.", ["local 저장소를 사용할 때 /uploads 파일을 제공합니다.", "S3 public URL을 쓰면 이 컨트롤러 경로와 표시 방식이 달라질 수 있습니다.", "프론트 resolveApiAssetUrl과 함께 봅니다."]),
    ("AiMatchController", "AI 매칭 사용자 API를 받습니다.", ["프로필 생성, 접근, 수정, 삭제를 처리합니다.", "요청 생성, 수락, 거절, 취소, 만남 제안을 처리합니다.", "사진 업로드와 PIN 검증이 중요합니다."]),
    ("AdminAiMatchController", "AI 매칭 관리자 API를 받습니다.", ["운영 현황을 반환합니다.", "요청 연결 상태와 관리자 메모를 수정합니다.", "프로필 삭제와 전화번호 purge를 제공합니다."]),
    ("TranslateController", "번역 API를 받습니다.", ["텍스트 번역 요청을 처리합니다.", "번역 메트릭을 제공합니다.", "프론트 i18n.js의 정적 번역과는 별도 기능입니다."]),
    ("ExportController", "CSV 다운로드 API를 받습니다.", ["부스와 공연 데이터를 CSV로 내려줍니다.", "프론트 downloadBoothCsv/downloadEventCsv와 연결됩니다.", "JSON이 아니라 text/csv 응답입니다."]),
]


BACKEND_SERVICE_DEEP_DIVES = [
    ("BoothService", "부스 조회, 생성, 수정, 삭제, 순서 변경, 혼잡도 계산, 예약 요약 DTO 변환을 담당합니다.", ["displayOrder 정렬", "toDto 변환", "GPS 거리 계산", "예약 테이블 요약"]),
    ("EventService", "공연 목록과 상태 계산, 공연 CRUD, 30초 SSE 방송을 담당합니다.", ["statusOverride 우선", "시간 기반 예정/진행중/종료", "@Scheduled 방송", "EventResponseDto 변환"]),
    ("AnalyticsService", "GPS 로그를 통계 데이터로 바꿉니다.", ["최근 24시간 방문", "최근 60분 인기 부스", "좌표 cell 히트맵", "구역별 혼잡도 dashboard"]),
    ("GpsService", "프론트가 보낸 위치를 GpsLog로 저장합니다.", ["위도/경도 요청 DTO", "혼잡도 계산 입력", "분석 화면 입력", "사용자 개인 식별 최소화 검토 필요"]),
    ("ReservationService", "예약 생성, 조회, 설정, 체크인, 완료, 만료, 노쇼, 좌석 복구를 담당합니다.", ["@Transactional", "findByIdForUpdate", "중복 예약 방지", "StreamService.publishReservations"]),
    ("ReservationAuthService", "전화번호 인증 코드와 예약 사용자 키를 관리합니다.", ["코드 발급", "코드 검증", "세션 토큰 발급", "시도 횟수 제한"]),
    ("ChatService", "질문 의도 분석, 근거 검색, OpenAI 호출, fallback 답변을 담당합니다.", ["retrieveEvidence", "STATIC_KNOWLEDGE", "MAX_EVIDENCE", "warnings/confidence"]),
    ("AiCongestionService", "부스 혼잡 위험 점수와 AI 축제 가이드를 만듭니다.", ["riskScore", "recommendedNow", "avoidNow", "operatorAlerts"]),
    ("PublicAiGuideService", "페이지별 방문자 AI 안내를 만듭니다.", ["scope별 안내", "OpenAI key fallback", "프론트 StageMap/Event/Analytics 연결", "timeout 관리"]),
    ("FestivalSnapshotService", "AI 판단에 필요한 현재 축제 데이터를 모읍니다.", ["부스", "공연", "혼잡도", "예약", "스태프"]),
    ("NoticeService", "공지 조회와 CRUD, 활성 공지 처리를 담당합니다.", ["active notices", "관리자/운영 공지", "SSE notices", "공지 카테고리"]),
    ("LostItemService", "분실물 조회, 생성, 상태 변경, 주인 확인, 삭제를 담당합니다.", ["파일 URL", "상태 라벨", "스태프 권한", "실시간 분실물 방송"]),
    ("StaffService", "스태프 로그인, 세션, bootstrap, 상태 변경, AI 스태프 지원을 담당합니다.", ["staff token", "status update", "zone summary", "field checklist"]),
    ("AuthService", "관리자 로그인과 JWT 발급을 담당합니다.", ["PasswordEncoder", "AdminUserRepository", "JwtService", "LoginResponseDto"]),
    ("JwtService", "JWT 생성과 검증을 담당합니다.", ["secret", "expiration", "claims role", "parse"]),
    ("OpsKeyService", "운영 마스터/부스 키를 검증합니다.", ["master key", "booth keys", "shared booth key", "OpsIdentity"]),
    ("UploadStorageService", "파일을 local 또는 S3에 저장하고 URL을 반환합니다.", ["app.storage.type", "APP_UPLOAD_DIR", "S3 bucket", "public base URL"]),
    ("AiMatchService", "AI 매칭 프로필, 요청, 즐겨찾기, 만남 제안을 처리합니다.", ["PIN 검증", "이미지 미리보기", "전화번호 사용 제한", "요청 상태 전이"]),
    ("AiImageGenerationService", "AI 매칭 이미지 생성/검증 관련 처리를 담당합니다.", ["OpenAI image model", "사진 검증 설정", "원본/웹툰 이미지", "실패 fallback"]),
    ("TranslateService", "번역 요청을 처리합니다.", ["요청 DTO", "외부 번역/AI 가능성", "TranslateMetricsService", "프론트 번역과 차이"]),
    ("AdminDashboardService", "관리자 KPI를 계산합니다.", ["부스 수", "공연 수", "공지/로그", "관리 콘솔 카드"]),
    ("AuditLogService", "관리자 작업 이력을 기록합니다.", ["누가", "무엇을", "언제", "어떤 엔티티를"]),
    ("AdminImportService", "CSV import를 처리합니다.", ["부스 CSV", "공연 CSV", "multipart", "데이터 정규화"]),
    ("OpsAiService", "운영자용 AI 브리핑과 공지 초안을 만듭니다.", ["운영 데이터 요약", "혼잡 완화", "공지 초안", "OpenAI fallback"]),
]


BACKEND_FLOW_CASES = [
    ("GET /api/booths 흐름", ["프론트 fetchBooths가 요청합니다.", "SecurityConfig에서 공개 API로 통과합니다.", "BoothController.getAllBooths가 호출됩니다.", "BoothService.getAllBooths가 Repository를 조회합니다.", "BoothResponseDto 목록이 JSON으로 반환됩니다."], ["api.js", "BoothController", "BoothService", "BoothRepository", "BoothResponseDto"]),
    ("GET /api/booths/{id}/congestion 흐름", ["프론트가 특정 부스 혼잡도를 요청합니다.", "BoothService가 부스를 찾습니다.", "최근 15분 GpsLog를 가져옵니다.", "부스 반경 80m 안의 로그에 시간 가중치를 줍니다.", "CongestionResponseDto를 반환합니다."], ["BoothController", "BoothService", "GpsLogRepository", "CongestionResponseDto"]),
    ("POST /api/gps 흐름", ["StageMapPage가 현재 좌표를 보냅니다.", "GpsController가 요청 DTO를 받습니다.", "GpsService가 GpsLog 엔티티를 저장합니다.", "이 로그가 나중에 분석과 혼잡도 계산에 쓰입니다.", "프론트는 위치 반영 메시지를 표시합니다."], ["StageMapPage", "GpsController", "GpsService", "GpsLog"]),
    ("관리자 로그인 흐름", ["AdminPage가 username/password를 보냅니다.", "AuthController가 AuthService를 호출합니다.", "AuthService가 AdminUser와 PasswordEncoder로 검증합니다.", "JwtService가 role=ADMIN 토큰을 만듭니다.", "프론트가 localStorage에 토큰을 저장합니다."], ["AdminPage", "AuthController", "AuthService", "AdminUser", "JwtService"]),
    ("관리자 부스 수정 흐름", ["프론트 updateBooth가 Authorization 헤더와 payload를 보냅니다.", "JwtAuthenticationFilter가 ADMIN 권한을 넣습니다.", "AdminBoothController가 요청을 받습니다.", "BoothService.updateBooth가 엔티티를 수정합니다.", "필요 시 StreamService가 부스 변경을 방송합니다."], ["api.js", "JwtAuthenticationFilter", "AdminBoothController", "BoothService", "StreamService"]),
    ("운영 부스 체크인 흐름", ["OpsBoothPage가 부스 키와 예약 ID를 보냅니다.", "OpsKeyAuthenticationFilter가 OPS_BOOTH 권한을 넣습니다.", "OpsController가 check-in 메서드를 호출합니다.", "ReservationService.checkIn이 상태를 CHECKED_IN으로 바꿉니다.", "예약 SSE가 상세/운영 화면을 갱신합니다."], ["OpsBoothPage", "OpsKeyAuthenticationFilter", "OpsController", "ReservationService", "StreamService"]),
    ("예약 만료 흐름", ["예약 조회나 생성 시 expireStaleReservations가 실행됩니다.", "만료 시간이 지난 RESERVED 예약을 찾습니다.", "상태를 EXPIRED로 바꿉니다.", "좌석을 복구하고 노쇼를 기록합니다.", "예약 SSE로 변경을 방송합니다."], ["ReservationService", "BoothReservationRepository", "ReservationUserState", "StreamService"]),
    ("AI 홈 가이드 흐름", ["프론트 fetchAiFestivalGuide가 요청합니다.", "AiGuideController가 AiCongestionService.guide를 호출합니다.", "FestivalSnapshotService가 현재 축제 데이터를 모읍니다.", "AiCongestionService가 riskScore와 추천/회피 목록을 만듭니다.", "AiFestivalGuideDto가 홈 화면에 표시됩니다."], ["HomePage", "AiGuideController", "AiCongestionService", "FestivalSnapshotService", "AiFestivalGuideDto"]),
    ("챗봇 질문 흐름", ["프론트 askChat이 question을 보냅니다.", "ChatController가 ChatService.answer를 호출합니다.", "ChatService가 질문 의도에 따라 근거 후보를 모읍니다.", "OpenAI key가 있으면 Responses API를 호출하고 없으면 fallback을 만듭니다.", "answer/evidence/warnings가 프론트에 표시됩니다."], ["ChatPage", "ChatController", "ChatService", "OpenAI", "ChatResponseDto"]),
    ("분실물 등록 흐름", ["프론트가 FormData로 텍스트와 파일을 보냅니다.", "LostItemController가 multipart 요청을 받습니다.", "UploadStorageService가 파일을 저장합니다.", "LostItemService가 LostItem을 저장합니다.", "StreamService.publishLostItems가 목록 갱신을 방송합니다."], ["LostFoundPage", "LostItemController", "UploadStorageService", "LostItemService", "StreamService"]),
    ("SSE 구독 흐름", ["프론트가 new EventSource를 생성합니다.", "StreamController가 text/event-stream 요청을 받습니다.", "StreamService가 SseEmitter를 목록에 추가합니다.", "Service 변경 작업이 publish 메서드를 호출합니다.", "프론트 addEventListener가 JSON payload를 파싱해 state를 갱신합니다."], ["api.js", "StreamController", "StreamService", "Page.jsx"]),
    ("AI 매칭 프로필 생성 흐름", ["프론트가 프로필 폼과 사진을 FormData로 보냅니다.", "AiMatchController가 multipart 요청을 받습니다.", "AiMatchService가 전화번호/PIN/프로필 규칙을 검증합니다.", "UploadStorageService와 AI 이미지 서비스가 사진 처리를 돕습니다.", "AiMatchProfileResponseDto가 프론트 목록에 반영됩니다."], ["AiMatchPage", "AiMatchController", "AiMatchService", "AiMatchProfile", "UploadStorageService"]),
]


def add_beginner_terms(doc, terms, heading):
    doc.add_heading(heading, level=1)
    add_note(
        doc,
        "이 장의 목표",
        "처음 보는 용어 때문에 코드 읽기가 막히지 않도록 핵심 단어를 FestFlow 코드와 연결해서 설명합니다. "
        "용어를 외우기보다 '이 단어가 실제 어느 파일에서 어떤 역할을 하는지'를 이해하는 것이 중요합니다.",
    )
    for term, meaning, caution in terms:
        doc.add_heading(term, level=3)
        add_paragraph(doc, meaning, bold_lead="뜻: ")
        add_paragraph(doc, caution, bold_lead="초보자 포인트: ")


def add_frontend_deep_file_guide(doc):
    doc.add_heading("16. 프론트 파일별 초심자 상세 해설", level=1)
    add_note(
        doc,
        "이 장의 사용법",
        "아래 내용은 파일을 열었을 때 무엇을 먼저 보고, 어떤 상황에서 이 파일을 수정하는지 알려주는 안내서입니다. "
        "처음에는 모든 파일을 완벽히 이해하려고 하지 말고, 파일의 책임을 한 문장으로 말할 수 있게 되는 것을 목표로 합니다.",
    )
    for filename, role, reading, use_cases in FRONTEND_FILE_DEEP_DIVES:
        doc.add_heading(filename, level=2)
        add_paragraph(doc, role, bold_lead="파일의 핵심 책임: ")
        add_paragraph(doc, reading, bold_lead="읽는 방법: ")
        add_bullets(doc, [(case, "") for case in use_cases])
        add_note(
            doc,
            "초보자 확인 질문",
            f"{filename}을 읽은 뒤에는 '이 파일이 직접 화면을 그리는가, 서버와 통신하는가, 공통 설정을 제공하는가, 데이터를 저장하는가'를 스스로 구분해 봅니다.",
            fill=LIGHT_FILL,
        )


def add_frontend_flow_cases(doc):
    doc.add_heading("17. 실제 사용자 흐름으로 프론트와 백엔드 연결 추적하기", level=1)
    add_paragraph(
        doc,
        "코드를 파일 단위로만 읽으면 흐름이 끊겨 보입니다. 실제 앱은 사용자의 행동 하나가 여러 파일을 지나갑니다. "
        "이 장은 '사용자가 무엇을 한다 -> 어떤 프론트 코드가 실행된다 -> 어떤 API가 호출된다 -> 어떤 백엔드 코드와 연결된다'를 사례별로 추적합니다.",
    )
    for title, steps, files in FRONTEND_FLOW_CASES:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)
        add_paragraph(doc, ", ".join(files), bold_lead="함께 열어볼 파일/클래스: ")
        add_note(
            doc,
            "연습 방법",
            "각 단계를 실제 코드에서 검색해 보세요. 함수 이름으로 검색하고, 그 함수가 다시 어떤 함수를 호출하는지 한 단계씩 내려가면 전체 흐름이 보입니다.",
            fill=SUCCESS_FILL,
        )


def add_frontend_state_and_event_catalog(doc):
    doc.add_heading("18. 프론트 state, 이벤트, 데이터 흐름 카탈로그", level=1)
    add_paragraph(
        doc,
        "React 화면을 고칠 때는 HTML처럼 보이는 JSX만 바꾸면 안 됩니다. 화면에 보이는 값은 대부분 state에서 오고, state는 API 응답이나 사용자 이벤트로 바뀝니다. "
        "이 표는 주요 state를 어디서 보고 어떻게 추적할지 정리합니다.",
    )
    rows = [
        ("booths", "HomePage, StageMapPage, AdminPage", "fetchBooths 또는 createBoothStream", "부스 카드, 지도 마커, 관리자 목록"),
        ("events", "HomePage, EventPage, LineupPage", "fetchEvents 또는 createEventStream", "공연 카드, 라인업, 다음 공연 추천"),
        ("traffic", "HomePage", "fetchTrafficHourly", "홈 혼잡 퍼센트"),
        ("aiGuide", "HomePage, StageMapPage, EventPage, AnalyticsPage", "fetchAiFestivalGuide 또는 fetchAiVisitorGuide", "AI 안내 카드와 행동 제안"),
        ("aiQuestion", "HomePage, ChatPage", "input onChange", "AI 질문 입력창"),
        ("aiAnswer/messages", "HomePage, ChatPage", "askChat 응답", "AI 답변과 근거"),
        ("activeCategory", "StageMapPage", "카테고리 버튼 클릭", "지도 부스 필터"),
        ("query", "StageMapPage, LostFoundPage", "검색 입력 onChange 또는 URL query", "검색 결과 필터"),
        ("currentLocation", "StageMapPage", "navigator.geolocation", "내 위치 마커와 지도 중심"),
        ("geoMessage", "StageMapPage", "GPS 성공/실패 콜백", "위치 전송 안내 문구"),
        ("reservationState", "BoothDetailPage, OpsBoothPage", "fetchBoothReservations 또는 reservation stream", "테이블, 내 예약, 체크인 상태"),
        ("authToken/reservationToken", "BoothDetailPage", "reservationAuth utils", "전화번호 인증 후 예약 가능 여부"),
        ("adminToken", "AdminPage, AiMatchAdminPage", "loginAdmin과 utils/auth", "관리자 API 접근"),
        ("opsKey", "OpsMasterPage, OpsBoothPage", "운영 키 입력", "운영 API 접근"),
        ("staffToken", "StaffPage", "loginStaff", "스태프 API 접근"),
        ("lostItems", "LostFoundPage, StaffPage", "fetchLostItems 또는 lost item stream", "분실물 목록과 상태"),
        ("language", "LanguageProvider", "언어 토글 또는 URL query", "한국어/영어 표시"),
        ("favorites/recents/memos", "여러 부스 화면", "utils/storage localStorage", "즐겨찾기, 최근 본 부스, 메모"),
        ("profiles/requests", "AiMatchPage, AiMatchAdminPage", "AI 매칭 API", "프로필 목록과 매칭 요청"),
        ("loading flags", "대부분의 페이지", "요청 시작/종료", "로딩 스피너, 버튼 비활성화"),
    ]
    add_table(doc, ["state", "주요 위치", "바뀌는 원인", "화면 영향"], rows, [1700, 2500, 2600, 2560])

    doc.add_heading("19. 이벤트 핸들러를 읽는 법", level=1)
    add_paragraph(
        doc,
        "프론트에서 '사용자가 뭘 눌렀을 때 무슨 일이 일어나는지'는 이벤트 핸들러를 보면 알 수 있습니다. "
        "onClick, onSubmit, onChange, addEventListener는 모두 사용자의 행동이나 외부 이벤트와 연결됩니다.",
    )
    handler_rows = [
        ("onClick", "버튼 클릭", "navigate('/events'), setSearchOpen(...), handleLocate()", "버튼이 단순 링크인지, state 변경인지, API 호출인지 구분합니다."),
        ("onSubmit", "폼 제출", "AI 질문, 로그인, 분실물 신고", "event.preventDefault가 있는지 확인합니다."),
        ("onChange", "입력값 변경", "검색어, 로그인 폼, 질문 입력", "value와 state가 연결된 controlled input인지 봅니다."),
        ("addEventListener", "SSE 이벤트 수신", "boothStream.addEventListener('booths', ...)", "이벤트 이름과 JSON.parse를 확인합니다."),
        ("setInterval", "주기적 새로고침", "StageMapPage 부스/AI 가이드 refresh", "cleanup에서 clearInterval이 있는지 확인합니다."),
        ("navigator.geolocation", "위치 권한 요청", "handleLocate", "성공/실패 콜백과 timeout 옵션을 봅니다."),
    ]
    add_table(doc, ["핸들러", "뜻", "FestFlow 예시", "읽을 때 질문"], handler_rows, [1600, 1900, 3100, 2760])


def add_frontend_practice_appendix(doc):
    doc.add_heading("20. 프론트 실습 과제와 정답 방향", level=1)
    tasks = [
        ("홈 화면에 '가장 한산한 부스' 배지를 추가하기", "HomePage의 homeCards 또는 별도 useMemo에서 estimatedWaitMinutes가 가장 낮은 부스를 찾고 JSX에 배지를 표시합니다."),
        ("지도 카테고리에 '굿즈'를 추가하기", "festivalUiData의 카테고리, StageMapPage의 displayCategory, pinTone, CategoryIcon을 함께 수정합니다."),
        ("AI 챗봇 근거를 2개에서 3개로 늘리기", "HomePage 또는 ChatPage에서 evidence.slice 제한을 찾고 UI 공간이 충분한지 확인합니다."),
        ("관리자 부스 이미지 업로드 후 목록 즉시 갱신하기", "uploadBoothImage 호출 성공 후 fetchBooths를 다시 호출하거나 해당 booth state를 교체합니다."),
        ("예약 인증 토큰 초기화 버튼 만들기", "reservationAuth util의 저장/삭제 함수를 확인하고 BoothDetailPage에 버튼을 추가합니다."),
        ("공연 시작까지 남은 시간을 표시하기", "EventPage에서 startTime을 Date로 변환하고 현재 시간과 차이를 계산합니다."),
        ("분실물 목록에 '최근 등록순' 정렬 추가하기", "LostFoundPage에서 filtered list를 useMemo로 정렬하고 상태 필터와 검색 필터 후에 적용합니다."),
        ("영어 번역에 새 문구 추가하기", "i18n.js의 KO_TO_EN에 정확한 원문을 넣고 placeholder 번역까지 확인합니다."),
        ("API 서버 주소를 배포 주소로 바꾸기", ".env에 VITE_API_BASE_URL을 넣고 api.js의 getApiBase로 실제 값을 확인합니다."),
        ("SSE 연결 실패 시 폴링 fallback 추가하기", "EventSource 생성 catch에서 setInterval 기반 fetch를 사용하고 cleanup에서 clearInterval을 호출합니다."),
    ]
    for title, direction in tasks:
        doc.add_heading(title, level=2)
        add_paragraph(doc, direction, bold_lead="정답 방향: ")
        add_bullets(doc, [
            "수정 전 관련 state와 API 함수를 먼저 찾습니다.",
            "성공 상태뿐 아니라 서버 실패, 빈 데이터, 로딩 중 상태를 함께 고려합니다.",
            "모바일 화면에서 텍스트가 넘치지 않는지 확인합니다.",
            "수정 후 npm run build로 문법과 번들 오류를 확인합니다.",
        ])


def add_backend_deep_controller_guide(doc):
    doc.add_heading("24. 컨트롤러별 초심자 상세 해설", level=1)
    add_note(
        doc,
        "컨트롤러를 읽는 공식",
        "컨트롤러 파일을 열면 먼저 @RequestMapping의 공통 경로를 보고, 각 @GetMapping/@PostMapping이 최종 URL을 어떻게 만드는지 계산합니다. "
        "그다음 생성자 주입된 Service 이름을 보고 실제 로직이 어디로 넘어가는지 따라갑니다.",
    )
    for name, summary, points in BACKEND_CONTROLLER_DEEP_DIVES:
        doc.add_heading(name, level=2)
        add_paragraph(doc, summary, bold_lead="역할: ")
        add_bullets(doc, points)
        add_note(
            doc,
            "읽는 질문",
            f"{name}에서 직접 DB를 만지는지, 아니면 Service에 위임하는지 확인하세요. 좋은 컨트롤러는 요청을 받고 검증한 뒤 핵심 일을 Service에 맡깁니다.",
            fill=LIGHT_FILL,
        )


def add_backend_deep_service_guide(doc):
    doc.add_heading("25. 서비스별 초심자 상세 해설", level=1)
    add_paragraph(
        doc,
        "Service는 백엔드의 실제 두뇌입니다. 컨트롤러가 URL을 받는 입구라면, 서비스는 규칙을 결정하고 Repository를 통해 데이터를 바꾸는 곳입니다. "
        "어떤 버그가 '화면은 뜨는데 결과가 이상하다' 형태라면 대부분 Service를 읽어야 합니다.",
    )
    for name, summary, keywords in BACKEND_SERVICE_DEEP_DIVES:
        doc.add_heading(name, level=2)
        add_paragraph(doc, summary, bold_lead="핵심 책임: ")
        add_bullets(doc, [(keyword, "") for keyword in keywords])
        add_note(
            doc,
            "수정 전 체크",
            f"{name}를 수정할 때는 같은 데이터를 쓰는 Controller, DTO, Repository, 프론트 api.js 함수까지 함께 검색해야 합니다.",
            fill=SUCCESS_FILL,
        )


def add_backend_flow_cases(doc):
    doc.add_heading("26. 실제 요청 흐름으로 백엔드 연결 추적하기", level=1)
    add_paragraph(
        doc,
        "백엔드 코드는 컨트롤러, 서비스, 저장소, 엔티티로 나뉘어 있어 처음에는 흩어져 보입니다. "
        "실제 HTTP 요청 하나를 끝까지 따라가면 역할 분리가 자연스럽게 이해됩니다.",
    )
    for title, steps, files in BACKEND_FLOW_CASES:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)
        add_paragraph(doc, ", ".join(files), bold_lead="함께 열어볼 파일/클래스: ")
        add_note(
            doc,
            "추적 팁",
            "프론트 api.js 함수 이름으로 시작해 URL을 찾고, 백엔드에서 그 URL의 @Mapping을 검색한 뒤 Service 메서드로 내려갑니다.",
            fill=SUCCESS_FILL,
        )


def add_backend_error_and_status_catalog(doc):
    doc.add_heading("27. HTTP 상태 코드와 백엔드 예외 읽기", level=1)
    add_paragraph(
        doc,
        "백엔드가 실패할 때는 단순히 '에러'가 아니라 상태 코드와 이유를 봐야 합니다. 프론트 parseJson은 실패 응답을 Error로 바꾸고, "
        "백엔드 Service는 ResponseStatusException으로 의미 있는 상태 코드를 던집니다.",
    )
    rows = [
        ("200 OK", "조회/수정/생성이 정상 처리됨", "프론트가 JSON을 state에 반영합니다."),
        ("400 Bad Request", "요청 형식이나 값이 잘못됨", "DTO validation, 필수 값 누락을 확인합니다."),
        ("401 Unauthorized", "인증 정보가 없거나 유효하지 않음", "토큰, 스태프 세션, 예약 인증 토큰을 확인합니다."),
        ("403 Forbidden", "권한이 부족함", "SecurityConfig, JWT role, X-Ops-Key를 확인합니다."),
        ("404 Not Found", "대상 엔티티를 찾지 못함", "id가 실제 DB에 있는지 확인합니다."),
        ("409 Conflict", "상태 충돌", "예약 중복, 이미 점유된 테이블, 다른 부스 테이블 선택 같은 규칙 위반입니다."),
        ("429 Too Many Requests", "요청 제한 또는 차단", "예약 노쇼 차단, rate limit, 인증 시도 제한을 확인합니다."),
        ("500 Internal Server Error", "서버 내부 오류", "로그, 외부 API, DB 예외, null 처리 누락을 확인합니다."),
    ]
    add_table(doc, ["상태 코드", "의미", "FestFlow에서 볼 곳"], rows, [1800, 3100, 4460])

    doc.add_heading("28. 상태 전이 표", level=1)
    add_table(
        doc,
        ["도메인", "상태", "다음 상태", "전이 조건"],
        [
            ("예약", "RESERVED", "CHECKED_IN", "부스 운영자가 체크인하거나 QR 토큰으로 체크인합니다."),
            ("예약", "RESERVED", "EXPIRED", "만료 시간까지 체크인하지 않았습니다."),
            ("예약", "RESERVED", "CANCELLED", "운영자가 테이블을 해제하거나 예약이 취소됩니다."),
            ("예약", "CHECKED_IN", "COMPLETED", "이용 완료 처리합니다."),
            ("공연", "예정", "진행중", "현재 시간이 startTime 이상 endTime 이하입니다."),
            ("공연", "진행중", "종료", "현재 시간이 endTime 이후입니다."),
            ("공연", "임의 상태", "statusOverride", "관리자나 운영자가 override를 설정했습니다."),
            ("분실물", "보관 중", "주인 확인", "사용자가 claim을 요청했습니다."),
            ("분실물", "주인 확인", "반환 완료", "스태프가 반환 처리를 완료했습니다."),
            ("스태프", "AVAILABLE", "BUSY/URGENT 등", "스태프가 내 상태를 변경합니다."),
            ("AI 매칭 요청", "PENDING", "ACCEPTED/REJECTED/CANCELLED", "상대방 또는 요청자가 상태를 바꿉니다."),
            ("AI 매칭 요청", "ACCEPTED", "MEETUP_PROPOSED/CONFIRMED", "만남 제안과 확인 흐름이 진행됩니다."),
        ],
        [1600, 1800, 2300, 3660],
    )


def add_backend_practice_appendix(doc):
    doc.add_heading("29. 백엔드 실습 과제와 정답 방향", level=1)
    tasks = [
        ("부스 응답에 새 필드 추가하기", "Booth 엔티티, BoothResponseDto, BoothService.toDto, BoothUpsertRequestDto, 프론트 표시 위치를 함께 수정합니다."),
        ("공지 카테고리 필터 API 만들기", "NoticeController에 query parameter를 받고 NoticeService에서 category 조건을 적용합니다."),
        ("혼잡도 기준 변경하기", "BoothService의 convertLevel과 BOOTH_RADIUS_METERS, timeWeight를 검토합니다."),
        ("예약 최대 유지 시간을 바꾸기", "Booth.maxReservationMinutes, ReservationService.sanitizeReservationMinutes, 운영 예약 설정 API를 확인합니다."),
        ("관리자 작업 로그 추가하기", "Admin 컨트롤러 작업 후 AuditLogService를 호출하고 AdminDashboardController에서 표시되는지 봅니다."),
        ("새 SSE 이벤트 추가하기", "StreamService 목록, StreamController 구독 엔드포인트, publish 호출 위치, 프론트 EventSource를 모두 추가합니다."),
        ("OpenAI API timeout 조정하기", "ChatService requestFactory와 fetchWithTimeout 프론트 timeout을 함께 확인합니다."),
        ("S3 업로드로 전환하기", "application.properties의 app.storage.type과 S3 설정을 넣고 UploadStorageService URL 반환을 확인합니다."),
        ("운영 부스 키 정책 바꾸기", "OpsKeyService와 application.properties의 app.ops.booth-keys/shared-booth-key를 함께 수정합니다."),
        ("예약 노쇼 차단 시간을 조정하기", "ReservationUserState의 registerNoShow/isBlocked 로직과 ReservationService 만료 처리를 확인합니다."),
    ]
    for title, direction in tasks:
        doc.add_heading(title, level=2)
        add_paragraph(doc, direction, bold_lead="정답 방향: ")
        add_bullets(doc, [
            "Controller, Service, Repository, Entity, DTO 중 무엇이 필요한지 먼저 나눕니다.",
            "프론트 api.js와 페이지 state가 응답 변경을 받아들일 수 있는지 확인합니다.",
            "권한이 필요한 API라면 SecurityConfig와 필터를 확인합니다.",
            "상태 변경이 있으면 SSE publish가 필요한지 검토합니다.",
            "테스트 또는 최소한 수동 API 호출로 성공/실패 케이스를 확인합니다.",
        ])


def add_frontend_doc():
    front_arch = make_frontend_architecture_diagram()
    state_flow = make_frontend_state_diagram()
    doc = Document()
    style_document(doc, "FestFlow 프론트엔드 완전 마스터 독스")
    add_title(
        doc,
        "FestFlow 프론트엔드 완전 마스터 독스",
        "React, Vite, 라우팅, 페이지 상태, API 연결, 실시간 스트림, 화면 수정 방법을 초보자 기준으로 끝까지 설명합니다.",
        "프론트엔드 전용 기술 설명서",
    )

    doc.add_heading("1. 이 문서를 읽는 방법", level=1)
    add_paragraph(
        doc,
        "프론트엔드는 사용자가 직접 보는 화면입니다. 버튼, 지도, 카드, 검색창, 로그인 폼, AI 질문창처럼 눈에 보이는 대부분의 UI가 "
        "frontend 폴더 안에 있습니다. 이 문서는 '화면이 어디서 시작되는지', '페이지가 어떻게 선택되는지', "
        "'서버 데이터가 어떻게 화면에 들어오는지', '사용자가 클릭하면 어떤 함수가 실행되는지'를 순서대로 설명합니다.",
    )
    add_bullets(
        doc,
        [
            ("처음 읽을 때: ", "1장부터 8장까지 순서대로 읽으면 큰 구조가 잡힙니다."),
            ("코드를 고칠 때: ", "수정하려는 화면 장을 먼저 보고, 그 장의 API 연결표와 수정 레시피를 확인합니다."),
            ("오류를 찾을 때: ", "디버깅 장에서 증상을 고른 뒤 API, state, 라우터, CSS 순서로 추적합니다."),
            ("백엔드와 연결할 때: ", "api.js 장을 보고 백엔드 문서의 컨트롤러 장으로 이어서 봅니다."),
        ],
    )
    add_image(doc, front_arch, "그림 1. 프론트엔드 파일과 실행 흐름")
    add_image(doc, state_flow, "그림 2. React state 변경과 화면 갱신 흐름")

    add_beginner_terms(doc, FRONTEND_TERMS, "2. 프론트엔드 초보자 개념 사전")

    doc.add_heading("3. frontend 폴더 전체 지도", level=1)
    add_paragraph(
        doc,
        "frontend 폴더는 Vite 기반 React 앱입니다. package.json은 실행 방법과 라이브러리를 정의하고, src 폴더는 실제 화면 코드, "
        "public 폴더는 브라우저가 직접 가져가는 정적 파일을 담습니다.",
    )
    add_table(
        doc,
        ["경로", "역할", "처음 볼 때 확인할 것"],
        [
            ("frontend/package.json", "프론트 실행 스크립트와 의존성", "dev, build, preview 스크립트와 React/Vite/Leaflet 의존성"),
            ("frontend/index.html", "React가 붙을 HTML 뼈대", "id=\"root\"가 있는지 확인"),
            ("frontend/src/main.jsx", "React 앱 시작점과 전체 라우트 정의", "BrowserRouter, Routes, lazy 페이지 목록"),
            ("frontend/src/App.jsx", "공통 화면 껍데기", "Outlet, 하단 메뉴, 운영/AI 매칭 화면의 메뉴 숨김 조건"),
            ("frontend/src/api.js", "백엔드와 통신하는 모든 함수", "API_BASE, fetchWithTimeout, parseJson, EventSource 함수"),
            ("frontend/src/pages", "각 URL에 대응하는 페이지", "HomePage, StageMapPage, EventPage 등"),
            ("frontend/src/components", "여러 페이지가 재사용하는 UI 조각", "CongestionBadge, UxIcons"),
            ("frontend/src/data", "서버 실패 시 쓰는 기본 축제 데이터", "fallbackBooths, fallbackEvents"),
            ("frontend/src/utils", "저장소, 인증 토큰, 위치 같은 보조 함수", "auth, reservationAuth, storage"),
            ("frontend/src/index.css", "전체 디자인 시스템과 커스텀 CSS", "app-shell, uni-card, map 관련 클래스"),
            ("frontend/public", "이미지, PWA, 오프라인 파일", "manifest.json, service-worker.js, images 폴더"),
        ],
        [2450, 3550, 3360],
    )

    doc.add_heading("4. package.json과 실행 명령", level=1)
    add_paragraph(
        doc,
        "package.json은 프론트 프로젝트의 사용 설명서와 같습니다. 어떤 명령으로 개발 서버를 켜는지, 어떤 라이브러리를 쓰는지, "
        "빌드는 어떤 도구가 담당하는지 여기서 확인합니다.",
    )
    add_table(
        doc,
        ["항목", "내용", "왜 중요한가"],
        [
            ("dev/start", "vite", "개발 서버를 실행합니다. 코드 수정 후 브라우저에서 바로 확인할 때 사용합니다."),
            ("build", "vite build", "배포용 정적 파일을 만듭니다. 문법 오류와 번들 문제를 잡을 수 있습니다."),
            ("preview", "vite preview", "빌드 결과물을 로컬에서 미리 봅니다."),
            ("react/react-dom", "React 18", "컴포넌트와 렌더링의 핵심 라이브러리입니다."),
            ("react-router-dom", "6.x", "URL별 페이지 전환을 담당합니다."),
            ("leaflet/react-leaflet", "지도 라이브러리", "StageMapPage의 캠퍼스 지도와 마커를 만듭니다."),
            ("qrcode", "QR 생성", "예약 체크인 QR 기능에서 사용합니다."),
            ("tailwindcss", "CSS 유틸리티", "빠른 스타일링과 반응형 UI 구현에 사용합니다."),
        ],
        [1850, 2350, 5160],
    )
    add_code_walkthrough(
        doc,
        "package.json 대표 코드 해설",
        "frontend/package.json",
        """
        "scripts": {
          "dev": "vite",
          "start": "vite",
          "build": "vite build",
          "preview": "vite preview"
        }
        """,
        [
            ("dev/start는 개발용입니다. ", "브라우저에서 화면을 보며 코드를 고칠 때 사용합니다."),
            ("build는 배포용 검사입니다. ", "최종 산출물을 만들며 import 오류나 빌드 오류를 잡습니다."),
            ("preview는 빌드 결과 확인입니다. ", "개발 서버가 아니라 실제 빌드된 파일이 잘 보이는지 확인합니다."),
        ],
    )

    doc.add_heading("5. main.jsx: 앱이 시작되는 첫 번째 파일", level=1)
    add_paragraph(
        doc,
        "main.jsx는 React 앱의 입구입니다. 브라우저가 index.html을 열면 root라는 빈 영역이 있고, ReactDOM.createRoot가 "
        "그 영역에 React 앱을 붙입니다. 그 안에서 LanguageProvider가 언어 상태를 제공하고, BrowserRouter가 URL을 관리하고, "
        "Routes가 실제 페이지를 고릅니다.",
    )
    add_table(doc, ["URL", "렌더링 컴포넌트", "설명"], FRONTEND_ROUTES, [2100, 2300, 4960])
    add_code_walkthrough(
        doc,
        "라우트 구조 대표 코드 해설",
        "frontend/src/main.jsx",
        """
        <Route path="/" element={<App />}>
          <Route index element={lazyElement(HomePage)} />
          <Route path="stage-map" element={lazyElement(StageMapPage)} />
          <Route path="events" element={lazyElement(EventPage)} />
          <Route path="booths/:id" element={lazyElement(BoothDetailPage)} />
          <Route path="*" element={<Navigate to="/" replace />} />
        </Route>
        """,
        [
            ("path=\"/\"의 element가 App입니다. ", "모든 주요 페이지는 App이라는 공통 레이아웃 안에 들어갑니다."),
            ("index 라우트는 / 주소입니다. ", "주소가 정확히 /이면 HomePage가 Outlet 자리에 들어갑니다."),
            ("booths/:id의 :id는 변수입니다. ", "/booths/3처럼 들어오면 BoothDetailPage가 id=3을 읽어 상세 정보를 요청합니다."),
            ("* 라우트는 예외 처리입니다. ", "정의되지 않은 주소로 들어오면 홈으로 돌려보냅니다."),
        ],
    )
    add_note(
        doc,
        "lazyElement가 필요한 이유",
        "페이지가 많아지면 모든 코드를 첫 로딩에 가져오면 느려집니다. lazy와 Suspense를 쓰면 사용자가 실제로 이동한 페이지 코드만 나눠 불러올 수 있습니다.",
    )

    doc.add_heading("6. App.jsx: 모든 페이지를 감싸는 공통 레이아웃", level=1)
    add_paragraph(
        doc,
        "App.jsx는 직접적인 페이지 내용보다 '공통 껍데기'에 가깝습니다. 모든 페이지가 들어가는 main 영역을 만들고, 일반 방문자 화면에서는 "
        "하단 메뉴를 보여주며, 운영 패널이나 AI 매칭 화면처럼 전용 UI가 필요한 곳에서는 하단 메뉴를 숨깁니다.",
    )
    add_code_walkthrough(
        doc,
        "App.jsx 메뉴 표시 조건 해설",
        "frontend/src/App.jsx",
        """
        const isOpsPanelRoute = location.pathname.startsWith("/ops");
        const isAiMatchRoute = location.pathname.startsWith("/ai-match");

        {!isOpsPanelRoute && !isAiMatchRoute && (
          <nav className="festival-bottom-nav">...</nav>
        )}
        """,
        [
            ("location.pathname은 현재 주소입니다. ", "/stage-map, /ops/master 같은 문자열을 확인합니다."),
            ("운영 패널은 별도 하단 메뉴를 씁니다. ", "그래서 일반 방문자 메뉴가 보이면 화면 목적이 섞입니다."),
            ("AI 매칭 화면도 메뉴를 숨깁니다. ", "전용 플로우가 있고 일반 축제 탭과 섞이면 사용자가 헷갈립니다."),
        ],
    )
    add_table(
        doc,
        ["메뉴", "이동 주소", "활성화 기준", "사용자 관점"],
        [
            ("홈", "/", "정확히 /", "처음 들어오는 추천/안내 화면"),
            ("지도", "/stage-map", "/stage-map으로 시작", "부스 위치와 실시간 대기 상태 확인"),
            ("공연", "/events", "/events로 시작", "공연 일정과 라인업 확인"),
            ("분석", "/analytics", "/analytics로 시작", "혼잡도와 AI 예측 확인"),
            ("더보기", "/more", "/more, /chat, /lost-found, /staff, /admin 포함", "부가 기능과 설정으로 이동"),
        ],
        [1400, 1700, 2700, 3560],
    )

    doc.add_heading("7. api.js: 프론트와 백엔드의 계약서", level=1)
    add_paragraph(
        doc,
        "api.js는 프론트엔드에서 가장 중요한 연결 파일입니다. 페이지 컴포넌트가 직접 fetch URL을 매번 쓰지 않고, "
        "fetchBooths, askChat, createBoothStream 같은 의미 있는 함수 이름으로 백엔드를 호출하게 만듭니다. "
        "이 파일을 이해하면 화면과 서버가 어떻게 이어지는지 거의 모두 추적할 수 있습니다.",
    )
    add_code_walkthrough(
        doc,
        "API_BASE와 환경 변수 해설",
        "frontend/src/api.js",
        """
        const API_BASE = (
          import.meta.env.VITE_API_BASE_URL || "http://localhost:8080/api"
        )
          .trim()
          .replace(/\\/+$/, "");
        """,
        [
            ("VITE_API_BASE_URL이 있으면 그 값을 사용합니다. ", "배포 환경에서는 백엔드 주소가 로컬과 다를 수 있기 때문입니다."),
            ("없으면 localhost:8080/api를 사용합니다. ", "Spring Boot 기본 포트 8080과 application.properties의 /api 경로에 맞춘 기본값입니다."),
            ("replace(/\\/+$/, \"\")는 끝 슬래시 제거입니다. ", "주소가 /api/처럼 끝나도 /api로 정리해 중복 슬래시를 막습니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "parseJson 오류 처리 해설",
        "frontend/src/api.js",
        """
        async function parseJson(response, errorMessage) {
          if (!response.ok) {
            const error = new Error(errorMessage);
            error.status = response.status;
            throw error;
          }
          const contentType = response.headers.get("content-type") || "";
          if (!contentType.includes("application/json")) {
            return response.text();
          }
          return response.json();
        }
        """,
        [
            ("response.ok가 false면 실패입니다. ", "404, 409, 500 같은 상태 코드는 여기서 에러로 바뀝니다."),
            ("error.status를 붙입니다. ", "화면에서 상태 코드별 처리를 하고 싶을 때 사용할 수 있습니다."),
            ("JSON이 아니면 text로 받습니다. ", "CSV 다운로드처럼 JSON이 아닌 응답도 처리할 여지를 둡니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "SSE 연결 함수 해설",
        "frontend/src/api.js",
        """
        export function createBoothStream() {
          return new EventSource(`${API_BASE}/stream/booths`);
        }

        export function createEventStream() {
          return new EventSource(`${API_BASE}/stream/events`);
        }
        """,
        [
            ("EventSource는 SSE 연결을 엽니다. ", "fetch처럼 한 번 받고 끝나는 요청이 아니라 계속 열려 있는 통로입니다."),
            ("이벤트 이름은 백엔드 StreamService가 정합니다. ", "booths 스트림은 addEventListener(\"booths\", ...)로 받습니다."),
            ("페이지 cleanup에서 close해야 합니다. ", "닫지 않으면 페이지를 이동해도 연결이 계속 남을 수 있습니다."),
        ],
    )
    for group, rows in FRONTEND_API_GROUPS.items():
        doc.add_heading(group, level=2)
        add_table(doc, ["프론트 함수", "백엔드 주소", "사용 화면", "역할"], rows, [2300, 2450, 2300, 2310])

    doc.add_heading("8. HomePage.jsx 완전 해설", level=1)
    add_paragraph(
        doc,
        "HomePage는 FestFlow의 첫 화면입니다. 이 파일은 초보자에게 React 흐름을 이해하기 가장 좋은 예시입니다. "
        "왜냐하면 state, useEffect, useMemo, API 호출, SSE 연결, 폼 제출, navigate가 모두 들어 있기 때문입니다.",
    )
    add_code_walkthrough(
        doc,
        "HomePage state 선언 해설",
        "frontend/src/pages/HomePage.jsx",
        """
        const [booths, setBooths] = useState([]);
        const [events, setEvents] = useState([]);
        const [traffic, setTraffic] = useState([]);
        const [aiGuide, setAiGuide] = useState(null);
        const [aiQuestion, setAiQuestion] = useState("");
        const [aiAnswer, setAiAnswer] = useState(null);
        """,
        [
            ("booths/events/traffic는 서버에서 가져온 데이터입니다. ", "홈 추천 카드와 혼잡도 표시의 재료가 됩니다."),
            ("aiGuide는 홈 상단 AI 안내 카드입니다. ", "서버가 실패하면 DEFAULT_AI_GUIDE로 대체됩니다."),
            ("aiQuestion은 입력창 값입니다. ", "controlled input이라 사용자가 타이핑할 때마다 setAiQuestion이 실행됩니다."),
            ("aiAnswer는 질문 결과입니다. ", "pending 상태를 먼저 넣고, 응답이 오면 실제 답변으로 바꿉니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "초기 데이터 로딩 해설",
        "frontend/src/pages/HomePage.jsx",
        """
        Promise.allSettled([
          fetchBooths(),
          fetchEvents(),
          fetchTrafficHourly(),
        ]).then(([boothResult, eventResult, trafficResult]) => {
          if (boothResult.status === "fulfilled") setBooths(boothResult.value || []);
          if (eventResult.status === "fulfilled") setEvents(eventResult.value || []);
          if (trafficResult.status === "fulfilled") setTraffic(trafficResult.value || []);
        });
        """,
        [
            ("세 요청을 동시에 보냅니다. ", "부스, 공연, 방문 데이터는 서로 기다릴 필요가 없기 때문입니다."),
            ("allSettled는 일부 실패를 허용합니다. ", "예를 들어 방문 데이터만 실패해도 부스와 공연은 표시할 수 있습니다."),
            ("fulfilled인 결과만 state에 반영합니다. ", "실패한 요청은 fallback 데이터나 메시지로 처리됩니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "홈 추천 카드 계산 해설",
        "frontend/src/pages/HomePage.jsx",
        """
        const homeCards = useMemo(() => {
          const sortedBooths = [...boothSource].sort(
            (a, b) => (Number(a.estimatedWaitMinutes) || 0) - (Number(b.estimatedWaitMinutes) || 0),
          );
          const firstBooth = sortedBooths[0];
          return [
            { type: "booth", title: firstBooth?.name },
            { type: "event", title: nextEvent?.title },
            { type: "booth", title: reservable?.name },
          ];
        }, [boothSource, eventSource]);
        """,
        [
            ("boothSource는 실제 데이터 또는 fallback입니다. ", "서버 실패 시에도 홈 카드가 비지 않게 합니다."),
            ("대기시간이 짧은 부스를 먼저 찾습니다. ", "사용자에게 지금 가기 좋은 부스를 추천하기 위한 계산입니다."),
            ("다음 공연과 예약 가능 부스를 함께 넣습니다. ", "홈 카드 3개가 각각 다른 행동을 유도하도록 설계되어 있습니다."),
            ("의존성은 boothSource와 eventSource입니다. ", "둘 중 하나가 바뀔 때만 추천 카드를 다시 계산합니다."),
        ],
    )
    add_note(
        doc,
        "HomePage를 수정하는 감각",
        "새로운 홈 카드가 필요하면 JSX만 고치지 말고 homeCards 계산, 이미지 선택, navigate 경로, fallback 상황을 함께 확인해야 합니다.",
        fill=WARNING_FILL,
    )

    page_chapters = [
        ("StageMapPage", "지도 화면", [
            "Leaflet 지도와 OpenStreetMap 타일을 사용합니다.",
            "fetchBooths로 부스 데이터를 가져오고 createBoothStream으로 실시간 갱신을 받습니다.",
            "navigator.geolocation으로 현재 위치를 얻고 sendGps로 백엔드에 보냅니다.",
            "검색어와 카테고리를 state로 관리하고 useMemo로 필터링된 부스 목록을 계산합니다.",
            "부스 좌표가 이상하면 아주대 중심 좌표 기준 fallback 위치를 만들어 지도에 표시합니다.",
        ]),
        ("EventPage", "공연 목록 화면", [
            "fetchEvents로 공연 목록을 불러오고 createEventStream으로 상태 변경을 받습니다.",
            "fetchAiVisitorGuide(\"events\")를 호출해 공연 관람 관련 AI 안내를 보여줍니다.",
            "공연 상태, 시작 시간, 지연 정보, 이미지 포커스를 화면 요소로 변환합니다.",
        ]),
        ("LineupPage", "라인업 화면", [
            "공연을 시간 순으로 보여주는 화면입니다.",
            "EventPage와 마찬가지로 createEventStream을 사용해 공연 상태 변화를 반영합니다.",
            "공연 중심 UI라 홈/지도보다 데이터 종류는 적지만 상태 표시가 중요합니다.",
        ]),
        ("BoothDetailPage", "부스 상세와 예약 화면", [
            "URL의 :id를 읽어 fetchBoothById와 fetchCongestion을 호출합니다.",
            "전화번호 인증 토큰이 있으면 fetchBoothReservations로 내 예약 상태까지 가져옵니다.",
            "createBoothReservation으로 예약을 만들고 createBoothReservationCheckInToken으로 QR 토큰을 발급합니다.",
            "createReservationStream을 통해 운영자가 체크인하거나 테이블을 해제한 변화를 즉시 받습니다.",
        ]),
        ("AnalyticsPage", "혼잡도 분석 화면", [
            "fetchAnalyticsDashboard로 구역별 현재 혼잡도를 가져옵니다.",
            "fetchAiCongestionPredictions와 fetchAiDecisionLogs로 AI 판단 근거를 표시합니다.",
            "createCongestionStream으로 실시간 혼잡도 갱신을 받습니다.",
        ]),
        ("ChatPage", "AI 챗봇 화면", [
            "사용자 질문을 askChat으로 백엔드에 보냅니다.",
            "응답에는 answer, confidence, evidence, warnings가 포함될 수 있습니다.",
            "evidence는 답변 근거를 보여주고 사용자가 관련 화면으로 이동할 수 있게 합니다.",
        ]),
        ("LostFoundPage", "분실물 화면", [
            "fetchLostItems로 분실물 목록을 가져옵니다.",
            "createLostItemStream으로 새 분실물이나 상태 변경을 실시간 반영합니다.",
            "createLostItem은 FormData를 사용해 사진 파일과 폼 데이터를 함께 보냅니다.",
        ]),
        ("StaffPage", "스태프 화면", [
            "loginStaff로 스태프 토큰을 받고 fetchStaffBootstrap으로 초기 데이터를 가져옵니다.",
            "createStaffStream과 createLostItemStream을 함께 사용합니다.",
            "AI 현장 요약, 분실물 답변 초안, 체크리스트 생성 API를 호출합니다.",
        ]),
        ("AdminPage", "관리자 화면", [
            "loginAdmin으로 JWT를 받고 localStorage에 저장합니다.",
            "withAuth가 Authorization 헤더를 붙여 관리자 API를 호출합니다.",
            "부스, 공연, 공지, KPI, 감사 로그, 스태프 데이터를 한 화면에서 다룹니다.",
        ]),
        ("OpsMasterPage", "운영 마스터 화면", [
            "운영 마스터 키를 사용해 fetchOpsMasterBootstrap을 호출합니다.",
            "관리자 로그인 없이 현장 운영자가 공지, 부스 상태, 공연 상태를 빠르게 바꿀 수 있게 설계되었습니다.",
            "백엔드에서는 OPS_MASTER 권한으로 검사됩니다.",
        ]),
        ("OpsBoothPage", "부스 운영 화면", [
            "부스별 운영 키를 사용해 해당 부스의 상태와 예약만 관리합니다.",
            "예약 스트림을 받아 테이블 상태, 체크인, 완료 처리를 최신으로 유지합니다.",
            "운영자가 메뉴 이미지를 업로드할 수 있어 FormData 호출도 포함합니다.",
        ]),
        ("AiMatchPage", "AI 매칭 사용자 화면", [
            "프로필 목록 조회, 전화번호 중복 확인, 사진 미리보기, 프로필 생성, 요청 수락/거절을 처리합니다.",
            "이미지 업로드와 PIN 기반 접근 제어가 섞여 있어 폼 상태 관리가 중요합니다.",
        ]),
        ("AiMatchAdminPage", "AI 매칭 관리자 화면", [
            "관리자 로그인 후 AI 매칭 요청과 프로필을 운영 관점에서 관리합니다.",
            "fetchAdminAiMatchOverview로 전체 현황을 불러오고 상태 변경 API를 호출합니다.",
        ]),
        ("MorePage", "더보기와 설정 화면", [
            "언어 설정, 부가 화면 진입, 주요 관리 링크를 제공합니다.",
            "LanguageProvider와 연결된 언어 상태를 바꾸는 UI가 포함됩니다.",
        ]),
    ]
    doc.add_heading("9. 페이지별 완전 해설", level=1)
    for page, label, points in page_chapters:
        doc.add_heading(f"{page}: {label}", level=2)
        add_paragraph(doc, f"{page}는 {label}을 담당하는 페이지 컴포넌트입니다. 이 페이지를 읽을 때는 import한 API 함수, state 목록, useEffect 로딩 흐름, JSX 이벤트 핸들러 순서로 보면 됩니다.")
        add_bullets(doc, points)
        add_note(
            doc,
            "읽는 순서",
            f"{page}.jsx를 열면 먼저 import를 보고 어떤 API와 유틸을 쓰는지 확인합니다. 그다음 useState 선언을 읽고, useEffect에서 데이터를 어디서 가져오는지 본 뒤, return JSX에서 사용자가 보는 화면과 연결합니다.",
        )

    doc.add_heading("10. i18n.js: 언어 전환 구조", level=1)
    add_paragraph(
        doc,
        "i18n.js는 단순히 문자열 하나를 바꾸는 파일이 아닙니다. LanguageProvider가 현재 언어를 저장하고, translateText가 한국어를 영어로 바꾸며, "
        "MutationObserver가 새로 렌더링된 DOM 텍스트까지 감지해 번역을 적용합니다.",
    )
    add_code_walkthrough(
        doc,
        "언어 저장과 DOM 번역 해설",
        "frontend/src/i18n.js",
        """
        const [language, setLanguage] = useState(() => {
          const queryLanguage = new URLSearchParams(window.location.search).get("lang");
          if (queryLanguage === "en" || queryLanguage === "ko") return queryLanguage;
          return window.localStorage.getItem(LANGUAGE_STORAGE_KEY) === "en" ? "en" : "ko";
        });

        useEffect(() => {
          window.localStorage.setItem(LANGUAGE_STORAGE_KEY, language);
          applyDocumentTranslations(language);
        }, [language]);
        """,
        [
            ("처음 언어는 URL 또는 localStorage에서 읽습니다. ", "?lang=en이면 영어로 시작할 수 있습니다."),
            ("language가 바뀌면 localStorage에 저장합니다. ", "새로고침 후에도 선택 언어가 유지됩니다."),
            ("applyDocumentTranslations가 실제 DOM을 바꿉니다. ", "React state와 DOM 후처리가 함께 쓰이는 구조입니다."),
        ],
    )

    doc.add_heading("11. CSS와 디자인 시스템 읽는 법", level=1)
    add_paragraph(
        doc,
        "FestFlow는 Tailwind 클래스와 index.css의 커스텀 클래스를 같이 씁니다. className에 text-sm, flex, grid 같은 Tailwind가 보이고, "
        "uni-page, uni-card, festival-bottom-nav 같은 앱 전용 클래스도 보입니다.",
    )
    add_table(
        doc,
        ["클래스/영역", "역할", "수정 시 주의"],
        [
            ("app-shell / festival-shell", "앱 전체 외곽 레이아웃", "여기를 바꾸면 모든 페이지에 영향이 갑니다."),
            ("festival-main", "페이지 본문이 들어가는 영역", "하단 nav와 겹치지 않도록 padding을 확인합니다."),
            ("festival-bottom-nav", "모바일 하단 메뉴", "아이콘, 글자, active 상태가 모두 연결됩니다."),
            ("uni-page", "페이지 기본 폭과 여백", "새 페이지를 만들 때 기본 래퍼로 쓰면 디자인이 맞습니다."),
            ("uni-card", "반복되는 카드 스타일", "카드 내부에 카드가 과하게 중첩되지 않도록 주의합니다."),
            ("map-page / real-campus-map", "지도 화면 레이아웃", "지도 높이와 marker 레이어가 깨지지 않게 확인합니다."),
            ("ai-guide-card", "AI 안내 카드", "로딩, 답변, 근거 버튼 상태를 함께 확인합니다."),
        ],
        [2350, 3200, 3810],
    )

    doc.add_heading("12. localStorage와 토큰", level=1)
    add_paragraph(
        doc,
        "프론트가 새로고침 후에도 기억해야 하는 값은 localStorage에 저장됩니다. 단, 보안상 민감한 토큰을 localStorage에 저장하면 XSS에 취약할 수 있으므로 "
        "실서비스에서는 보안 검토가 필요합니다. 현재 FestFlow는 관리자 토큰, 예약 인증 토큰, 언어, 즐겨찾기 등을 저장합니다.",
    )
    add_table(
        doc,
        ["파일", "저장 키", "역할"],
        [
            ("utils/auth.js", "festflow_access_token, festflow_admin_name", "관리자 JWT와 관리자 이름 저장"),
            ("utils/reservationAuth.js", "festflow_reservation_auth_token, festflow_reservation_phone", "예약 전화번호 인증 토큰과 전화번호 저장"),
            ("utils/reservation.js", "festflow_reservation_user_key", "예약 사용자 식별 키 저장"),
            ("utils/storage.js", "festflow_favorites, festflow_recents, festflow_memos", "즐겨찾기, 최근 본 부스, 메모 저장"),
            ("i18n.js", "festflow_language", "선택 언어 저장"),
        ],
        [2450, 3350, 3560],
    )

    doc.add_heading("13. 새 기능을 추가하는 실전 레시피", level=1)
    recipes = [
        ("새 페이지 추가", ["pages 폴더에 NewPage.jsx를 만듭니다.", "main.jsx에서 lazy import를 추가합니다.", "Routes 안에 path와 element를 추가합니다.", "App.jsx 하단 메뉴에 넣을지 결정합니다.", "필요한 API 함수가 있으면 api.js에 추가합니다."]),
        ("새 API 호출 추가", ["백엔드 주소와 HTTP method를 확인합니다.", "api.js에 의미 있는 함수 이름을 만듭니다.", "parseJson에 표시할 사용자용 에러 메시지를 넣습니다.", "토큰이 필요하면 withAuth를 사용합니다.", "페이지에서 useEffect 또는 이벤트 핸들러로 호출합니다."]),
        ("홈 카드 하나 추가", ["HomePage의 homeCards useMemo를 찾습니다.", "카드 객체에 id, type, tag, title, caption, image를 맞춰 넣습니다.", "JSX에서 렌더링 개수 제한이 있는지 확인합니다.", "클릭 시 이동할 navigate 경로를 정합니다.", "fallback 데이터일 때도 카드가 깨지지 않는지 봅니다."]),
        ("지도 필터 추가", ["StageMapPage의 displayCategory 또는 categoryMatches를 확인합니다.", "festivalUiData의 mapCategories를 필요하면 수정합니다.", "검색어 필터와 카테고리 필터가 동시에 적용되는지 봅니다.", "아이콘과 pinTone도 새 카테고리에 맞춥니다.", "모바일 지도에서 마커가 과도하게 많지 않은지 확인합니다."]),
        ("관리자 버튼 추가", ["AdminPage에서 로그인 토큰이 필요한지 확인합니다.", "api.js에 관리자 API 함수를 withAuth로 만듭니다.", "버튼 클릭 핸들러에서 로딩과 에러 state를 둡니다.", "성공 후 목록을 다시 불러오거나 state를 직접 갱신합니다.", "백엔드 SecurityConfig에서 /api/admin/** 권한을 확인합니다."]),
        ("SSE 스트림 추가", ["백엔드 StreamController와 StreamService에 채널을 만듭니다.", "api.js에 createSomethingStream 함수를 추가합니다.", "페이지 useEffect에서 addEventListener를 등록합니다.", "cleanup에서 stream.close()를 호출합니다.", "초기 fetch와 SSE 갱신이 같은 state 구조를 쓰게 맞춥니다."]),
        ("이미지 업로드 추가", ["input type=file을 만듭니다.", "FormData를 생성하고 file을 append합니다.", "api.js에서 Content-Type을 직접 지정하지 않습니다.", "백엔드 컨트롤러는 multipart/form-data를 받게 합니다.", "응답 이미지 URL은 resolveApiAssetUrl로 표시합니다."]),
        ("문구 번역 추가", ["i18n.js의 KO_TO_EN에 원문과 영어를 추가합니다.", "placeholder, aria-label도 번역 대상인지 확인합니다.", "data-i18n-skip이 필요한 영역은 번역에서 제외합니다.", "문구가 동적으로 생성되면 부분 치환 규칙이 필요한지 봅니다.", "한국어/영어 전환 후 레이아웃 깨짐을 확인합니다."]),
        ("PWA 수정", ["public/manifest.json의 앱 이름과 아이콘을 확인합니다.", "service-worker.js 캐시 대상 파일을 확인합니다.", "offline.html을 수정해 오프라인 화면을 바꿉니다.", "개발 중에는 서비스워커 캐시 때문에 예전 파일이 보일 수 있습니다.", "브라우저 Application 탭에서 캐시를 지워 확인합니다."]),
        ("CSS 수정", ["수정하려는 className을 JSX에서 찾습니다.", "index.css에서 해당 클래스 정의를 찾습니다.", "공통 클래스인지 페이지 전용 클래스인지 확인합니다.", "모바일 폭과 데스크톱 폭을 모두 확인합니다.", "하단 nav와 겹치는 padding을 확인합니다."]),
    ]
    for title, steps in recipes:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)

    doc.add_heading("14. 프론트엔드 디버깅 플레이북", level=1)
    debug_rows = [
        ("화면이 빈 페이지", "main.jsx 라우트, lazy import, 브라우저 콘솔 오류", "import 경로 오타나 컴포넌트 export 문제를 먼저 봅니다."),
        ("API 데이터가 안 보임", "Network 탭, api.js API_BASE, 백엔드 서버 실행 여부", "VITE_API_BASE_URL과 CORS를 확인합니다."),
        ("로그인은 되는데 관리자 API 실패", "localStorage 토큰, withAuth, Authorization 헤더", "토큰 만료 또는 SecurityConfig 권한 문제일 수 있습니다."),
        ("실시간 갱신이 안 됨", "EventSource 연결, stream event name, 백엔드 StreamService publish", "프론트 addEventListener 이름과 백엔드 send eventName이 같아야 합니다."),
        ("지도 마커가 안 보임", "좌표 유효성, Leaflet CSS import, 지도 컨테이너 높이", "leaflet/dist/leaflet.css가 main.jsx에서 import되는지 확인합니다."),
        ("이미지가 깨짐", "resolveApiAssetUrl, 업로드 URL, public 경로", "서버 업로드는 /uploads 경로, 정적 이미지는 /images 경로를 구분합니다."),
        ("언어 전환이 이상함", "i18n.js KO_TO_EN, MutationObserver, data-i18n-skip", "동적 문구가 사전에 없으면 부분 번역만 적용될 수 있습니다."),
        ("예약 버튼이 비활성", "예약 인증 토큰, reservationEnabled, 좌석 상태", "백엔드 예약 상태와 프론트 검증 조건을 함께 확인합니다."),
        ("빌드 실패", "npm run build 로그, import 경로, JSX 문법", "개발 서버에서 보이던 경고가 빌드에서 오류가 될 수 있습니다."),
        ("스타일이 갑자기 깨짐", "공통 CSS 클래스 변경 여부", "uni-card, festival-main, bottom nav 같은 공통 클래스 변경은 전역 영향이 큽니다."),
    ]
    add_table(doc, ["증상", "먼저 볼 곳", "해결 방향"], debug_rows, [2200, 3300, 3860])

    add_frontend_deep_file_guide(doc)
    add_frontend_flow_cases(doc)
    add_frontend_state_and_event_catalog(doc)
    add_frontend_practice_appendix(doc)

    doc.add_heading("21. 프론트엔드 마스터 체크리스트", level=1)
    add_bullets(
        doc,
        [
            "main.jsx의 모든 라우트를 보고 URL과 페이지 컴포넌트를 설명할 수 있다.",
            "App.jsx의 Outlet과 하단 메뉴 표시 조건을 설명할 수 있다.",
            "api.js의 API_BASE, parseJson, withAuth, EventSource 함수를 설명할 수 있다.",
            "HomePage의 state, useEffect, useMemo, handleAiAsk를 순서대로 설명할 수 있다.",
            "StageMapPage의 지도 표시, 필터, GPS 전송, SSE 갱신 흐름을 설명할 수 있다.",
            "관리자/운영/스태프 화면에서 토큰과 키가 어떻게 API에 전달되는지 설명할 수 있다.",
            "localStorage에 저장되는 값과 위험성을 설명할 수 있다.",
            "새 화면, 새 API, 새 SSE 스트림을 추가하는 순서를 설명할 수 있다.",
            "브라우저 DevTools Network와 Console을 사용해 문제를 추적할 수 있다.",
        ],
    )
    doc.add_page_break()
    doc.add_heading("부록 A. 프론트 코드 읽기 순서 요약", level=1)
    add_numbered(
        doc,
        [
            "package.json에서 실행 명령과 라이브러리를 확인합니다.",
            "main.jsx에서 URL과 페이지 매핑을 확인합니다.",
            "App.jsx에서 공통 레이아웃과 메뉴를 확인합니다.",
            "관심 있는 Page.jsx에서 import, state, useEffect, 이벤트 핸들러, JSX 순서로 읽습니다.",
            "api.js에서 Page.jsx가 호출한 함수의 URL과 method를 확인합니다.",
            "백엔드 문서에서 같은 URL의 Controller와 Service를 찾아 이어서 읽습니다.",
            "index.css에서 화면 클래스의 실제 디자인을 확인합니다.",
        ],
    )
    doc.save(OUT_FRONTEND)
    return OUT_FRONTEND


def add_backend_doc():
    back_arch = make_backend_architecture_diagram()
    security_flow = make_backend_security_diagram()
    doc = Document()
    style_document(doc, "FestFlow 백엔드 완전 마스터 독스")
    add_title(
        doc,
        "FestFlow 백엔드 완전 마스터 독스",
        "Spring Boot, 보안, 컨트롤러, 서비스, DB 엔티티, 예약, 실시간 스트림, AI 연동, 배포 설정을 초보자 기준으로 설명합니다.",
        "백엔드 전용 기술 설명서",
    )

    doc.add_heading("1. 백엔드를 어떻게 이해해야 하는가", level=1)
    add_paragraph(
        doc,
        "백엔드는 사용자가 직접 보는 화면은 아니지만, 프론트엔드가 필요로 하는 데이터를 만들고 저장하고 검사하는 서버입니다. "
        "FestFlow 백엔드는 Spring Boot로 만들어졌고, /api로 시작하는 HTTP 요청을 받아 JSON을 반환합니다. "
        "부스 목록, 공연 상태, 혼잡도, 예약, 분실물, 관리자, 운영자, 스태프, AI 매칭 기능이 모두 백엔드에 연결되어 있습니다.",
    )
    add_bullets(
        doc,
        [
            ("요청 관점: ", "프론트의 api.js 함수가 HTTP 요청을 보내면 Controller가 받습니다."),
            ("규칙 관점: ", "Controller는 Service에 일을 맡기고 Service가 실제 비즈니스 규칙을 실행합니다."),
            ("저장 관점: ", "Service는 Repository를 통해 Entity를 DB에 저장하거나 조회합니다."),
            ("보안 관점: ", "SecurityConfig와 필터가 관리자, 운영자, 공개 API 접근을 구분합니다."),
            ("실시간 관점: ", "StreamService가 SSE 연결을 유지하고 데이터 변경을 브라우저로 방송합니다."),
        ],
    )
    add_image(doc, back_arch, "그림 1. 백엔드 요청 처리 전체 구조")
    add_image(doc, security_flow, "그림 2. API 성격별 권한 검사 흐름")

    add_beginner_terms(doc, BACKEND_TERMS, "2. 백엔드 초보자 개념 사전")

    doc.add_heading("3. backend 폴더 전체 지도", level=1)
    add_table(
        doc,
        ["경로", "역할", "읽을 때 확인할 것"],
        [
            ("backend/build.gradle", "Java/Spring 의존성과 빌드 설정", "Spring Boot 버전, Java 17, JPA, Security, SMS, S3, JWT"),
            ("backend/src/main/resources/application.properties", "서버 설정 기본값", "DB, 포트, CORS, JWT, 업로드, OpenAI, SMS 환경 변수"),
            ("BackendApplication.java", "Spring Boot 시작점", "main 메서드와 애플리케이션 실행"),
            ("controller", "HTTP API 진입점", "URL, method, 요청 DTO, 응답 DTO"),
            ("service", "핵심 비즈니스 로직", "예약 규칙, 혼잡도 계산, AI 판단, 업로드, 인증"),
            ("repository", "DB 접근 계층", "findBy..., save, delete, existsBy..."),
            ("entity", "DB 테이블 모델", "필드, 관계, 상태 변경 메서드"),
            ("dto", "API 요청/응답 모델", "프론트가 보내고 받는 JSON 구조"),
            ("security", "JWT, 운영 키, rate limit 필터", "권한을 어떻게 SecurityContext에 넣는지"),
            ("config", "보안과 CORS 설정", "허용 origin, 필터 체인, 비밀번호 인코더"),
            ("init", "초기 데이터 생성", "데모 부스, 공연, 관리자, 스태프, GPS 로그"),
            ("src/test", "테스트", "예약 서비스와 부스 컨트롤러 테스트"),
        ],
        [2850, 3100, 3410],
    )

    doc.add_heading("4. build.gradle: 백엔드가 쓰는 기술", level=1)
    add_paragraph(
        doc,
        "build.gradle은 백엔드 프로젝트의 기술 목록입니다. 어떤 프레임워크를 쓰는지, 어떤 외부 서비스를 연결할 수 있는지, "
        "테스트와 빌드가 어떤 Java 버전으로 돌아가는지 확인할 수 있습니다.",
    )
    add_table(
        doc,
        ["의존성", "역할", "FestFlow에서 쓰이는 곳"],
        [
            ("spring-boot-starter-web", "HTTP API 서버", "모든 Controller와 JSON 응답"),
            ("spring-boot-starter-data-jpa", "DB ORM", "Entity와 Repository"),
            ("spring-boot-starter-validation", "요청 값 검증", "@Valid DTO 검증"),
            ("spring-boot-starter-security", "권한 검사", "SecurityConfig, JWT, Ops key"),
            ("software.amazon.awssdk:s3", "S3 파일 저장", "UploadStorageService"),
            ("net.nurigo:sdk", "Solapi/Nurigo SMS", "SMS 인증/AI 매칭 알림"),
            ("twilio", "Twilio SMS", "전화번호 인증 provider 선택"),
            ("jjwt", "JWT 생성/검증", "관리자 로그인 토큰"),
            ("mysql-connector-j", "MySQL 연결", "로컬 또는 MySQL 배포 DB"),
            ("postgresql", "PostgreSQL 연결", "Postgres 배포 DB"),
            ("spring-boot-starter-test", "테스트", "ReservationServiceTest, BoothControllerTest"),
        ],
        [2900, 2800, 3660],
    )

    doc.add_heading("5. application.properties: 서버 설정 읽기", level=1)
    add_paragraph(
        doc,
        "application.properties는 백엔드의 환경 설정입니다. ${환경변수:기본값} 형태는 환경 변수가 있으면 그 값을 쓰고, 없으면 기본값을 쓴다는 뜻입니다. "
        "개발 환경과 배포 환경에서 DB 주소, JWT secret, OpenAI key, CORS origin이 달라질 수 있으므로 이 파일을 반드시 이해해야 합니다.",
    )
    add_table(
        doc,
        ["설정", "기본값/의미", "주의할 점"],
        [
            ("spring.datasource.url", "MySQL localhost festival_db", "배포에서는 실제 DB URL 환경 변수를 넣어야 합니다."),
            ("spring.jpa.hibernate.ddl-auto", "update", "개발에는 편하지만 운영 DB 스키마 변경은 신중해야 합니다."),
            ("server.port", "8080", "프론트 API_BASE 기본값과 맞물립니다."),
            ("app.jwt.secret", "빈 값", "운영에서는 반드시 강한 secret을 넣어야 합니다."),
            ("app.cors.allowed-origins", "localhost 계열", "배포 프론트 도메인을 추가해야 브라우저 요청이 허용됩니다."),
            ("app.ops.master-key", "0000", "운영 마스터 키입니다. 실사용에서는 바꿔야 합니다."),
            ("app.ops.booth-keys", "1:1111,...", "부스별 운영 키입니다. 실사용에서는 고유 값으로 바꿉니다."),
            ("app.sms.provider", "none", "SMS 인증 provider를 선택합니다. none이면 실제 문자 발송이 안 될 수 있습니다."),
            ("app.openai.api-key", "빈 값", "없으면 AI 기능은 fallback 답변을 사용합니다."),
            ("app.storage.type", "local", "local 또는 S3 저장 전략을 고릅니다."),
        ],
        [2850, 3150, 3360],
    )
    add_note(
        doc,
        "초보자가 자주 놓치는 부분",
        "프론트의 VITE_API_BASE_URL과 백엔드의 server.port는 직접 연결됩니다. 백엔드가 8080이 아닌 포트에서 실행되면 프론트 API_BASE도 같이 바꿔야 합니다.",
        fill=WARNING_FILL,
    )

    doc.add_heading("6. 요청 처리 생명주기", level=1)
    add_numbered(
        doc,
        [
            ("브라우저가 요청을 보냅니다. ", "예: fetchBooths()가 GET http://localhost:8080/api/booths를 호출합니다."),
            ("CORS 설정이 출처를 확인합니다. ", "프론트 주소가 허용되지 않으면 브라우저가 응답을 막습니다."),
            ("보안 필터가 권한을 확인합니다. ", "관리자 JWT, 운영 키, 공개 API 여부를 검사합니다."),
            ("Controller가 URL과 method를 매칭합니다. ", "BoothController의 getAllBooths 메서드가 호출됩니다."),
            ("Controller가 Service를 호출합니다. ", "BoothService.getAllBooths가 실제 조회를 담당합니다."),
            ("Service가 Repository를 사용합니다. ", "DB의 Booth 엔티티를 가져옵니다."),
            ("Service가 DTO로 변환합니다. ", "프론트에 필요한 필드만 BoothResponseDto로 만듭니다."),
            ("Controller가 JSON을 반환합니다. ", "프론트 parseJson이 JSON을 읽어 화면 state에 넣습니다."),
        ],
    )

    doc.add_heading("7. SecurityConfig와 보안 필터", level=1)
    add_paragraph(
        doc,
        "SecurityConfig는 백엔드 API의 문지기입니다. 어떤 API는 누구나 볼 수 있고, 어떤 API는 관리자만, 어떤 API는 운영 키가 있어야 접근할 수 있습니다. "
        "FestFlow는 세션 기반 로그인보다 토큰/키 기반 stateless 구조에 가깝습니다.",
    )
    add_code_walkthrough(
        doc,
        "SecurityConfig 권한 규칙 해설",
        "backend/src/main/java/com/festflow/backend/config/SecurityConfig.java",
        """
        .authorizeHttpRequests(auth -> auth
            .requestMatchers("/api/ops/master/**").hasRole("OPS_MASTER")
            .requestMatchers("/api/ops/booth/**").hasAnyRole("OPS_MASTER", "OPS_BOOTH")
            .requestMatchers("/api/admin/**").hasRole("ADMIN")
            .requestMatchers("/api/auth/**", "/api/**", "/uploads/**").permitAll()
            .anyRequest().permitAll()
        )
        .addFilterBefore(opsKeyAuthenticationFilter, UsernamePasswordAuthenticationFilter.class)
        .addFilterBefore(jwtAuthenticationFilter, UsernamePasswordAuthenticationFilter.class);
        """,
        [
            ("운영 마스터 API는 OPS_MASTER 역할이 필요합니다. ", "OpsKeyAuthenticationFilter가 키를 보고 역할을 넣어줍니다."),
            ("부스 운영 API는 OPS_MASTER 또는 OPS_BOOTH가 필요합니다. ", "마스터는 모든 부스, 부스 키는 해당 부스 작업에 사용됩니다."),
            ("관리자 API는 ADMIN 역할이 필요합니다. ", "JwtAuthenticationFilter가 Bearer 토큰을 파싱해 역할을 넣습니다."),
            ("/api/** permitAll이 뒤에 있습니다. ", "앞에서 더 구체적인 admin/ops 규칙이 먼저 평가되므로 공개 API는 허용됩니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "JWT 필터 해설",
        "backend/src/main/java/com/festflow/backend/security/JwtAuthenticationFilter.java",
        """
        String authHeader = request.getHeader("Authorization");
        if (authHeader != null && authHeader.startsWith("Bearer ")) {
          String token = authHeader.substring(7);
          Claims claims = jwtService.parse(token);
          String username = claims.getSubject();
          String rawRole = String.valueOf(claims.get("role"));
          SecurityContextHolder.getContext().setAuthentication(authentication);
        }
        """,
        [
            ("Authorization 헤더를 읽습니다. ", "프론트 api.js의 withAuth가 이 헤더를 붙입니다."),
            ("Bearer 뒤의 토큰만 잘라냅니다. ", "JWT 본문을 JwtService가 검증하고 claims를 읽습니다."),
            ("role을 ROLE_ 접두사 형태로 바꿉니다. ", "Spring Security의 hasRole 검사가 인식할 수 있게 합니다."),
            ("SecurityContext에 인증 정보를 넣습니다. ", "이후 SecurityConfig의 hasRole 조건이 통과될 수 있습니다."),
        ],
    )

    doc.add_heading("8. 컨트롤러와 엔드포인트 전체 목록", level=1)
    add_paragraph(
        doc,
        "컨트롤러는 프론트의 api.js와 직접 맞닿는 계층입니다. 어떤 API가 어디에 있는지 찾을 때는 먼저 이 표를 봅니다. "
        "그다음 해당 컨트롤러가 어떤 Service를 호출하는지 따라가면 됩니다.",
    )
    for group, rows in BACKEND_ENDPOINT_GROUPS.items():
        doc.add_heading(group, level=2)
        add_table(doc, ["컨트롤러", "엔드포인트", "역할"], rows, [2450, 3050, 3860])

    doc.add_heading("9. Entity와 DB 테이블 이해", level=1)
    add_paragraph(
        doc,
        "Entity는 DB에 저장되는 데이터의 실제 모양입니다. 프론트가 받는 DTO와는 다릅니다. 예를 들어 Booth 엔티티에는 내부 관리 필드가 있고, "
        "BoothResponseDto는 화면에 필요한 형태로 가공된 응답입니다.",
    )
    add_table(doc, ["엔티티", "도메인", "저장 내용"], ENTITY_SUMMARY, [2450, 1900, 5010])
    add_note(
        doc,
        "Entity와 DTO를 나누는 이유",
        "DB 모델을 그대로 프론트에 노출하면 보안, 호환성, 화면 요구사항이 뒤섞입니다. DTO를 쓰면 내부 저장 구조를 바꿔도 API 응답 모양을 비교적 안정적으로 유지할 수 있습니다.",
    )

    doc.add_heading("10. Booth 모듈 완전 해설", level=1)
    add_paragraph(
        doc,
        "Booth 모듈은 FestFlow의 중심 도메인입니다. 부스 목록, 부스 상세, 운영 상태, 대기시간, 재고, 메뉴 이미지, 예약 가능 좌석, 혼잡도 계산이 모두 이 모듈과 연결됩니다.",
    )
    add_code_walkthrough(
        doc,
        "부스 목록 조회 흐름",
        "backend/src/main/java/com/festflow/backend/service/BoothService.java",
        """
        public List<BoothResponseDto> getAllBooths() {
            return boothRepository.findAll().stream()
                    .sorted(Comparator.comparing(Booth::getDisplayOrder).thenComparing(Booth::getId))
                    .map(this::toDto)
                    .toList();
        }
        """,
        [
            ("Repository에서 모든 Booth를 가져옵니다. ", "DB 조회는 boothRepository.findAll이 담당합니다."),
            ("displayOrder 기준으로 정렬합니다. ", "관리자가 정한 표시 순서를 화면에 반영합니다."),
            ("id로 한 번 더 정렬합니다. ", "displayOrder가 같을 때 순서를 안정적으로 만듭니다."),
            ("toDto로 변환합니다. ", "엔티티를 그대로 보내지 않고 프론트 응답 DTO로 바꿉니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "혼잡도 계산 흐름",
        "backend/src/main/java/com/festflow/backend/service/BoothService.java",
        """
        LocalDateTime threshold = LocalDateTime.now().minusMinutes(15);
        List<GpsLog> recentLogs = gpsLogRepository.findByCreatedAtAfter(threshold);

        double weightedScore = recentLogs.stream()
            .filter(log -> distanceInMeters(booth.getLatitude(), booth.getLongitude(),
                                            log.getLatitude(), log.getLongitude()) <= BOOTH_RADIUS_METERS)
            .mapToDouble(log -> timeWeight(log.getCreatedAt(), now))
            .sum();
        """,
        [
            ("최근 15분 GPS 로그만 봅니다. ", "오래된 위치는 현재 혼잡도를 잘 설명하지 못합니다."),
            ("부스 반경 80m 안의 로그만 셉니다. ", "해당 부스 주변 사람만 혼잡도에 반영하기 위함입니다."),
            ("timeWeight로 시간 가중치를 줍니다. ", "방금 들어온 로그는 크게, 15분 가까이 된 로그는 작게 반영합니다."),
            ("weightedScore를 level로 바꿉니다. ", "프론트는 숫자뿐 아니라 여유/보통/혼잡 같은 라벨을 표시합니다."),
        ],
    )

    doc.add_heading("11. Event 모듈 완전 해설", level=1)
    add_paragraph(
        doc,
        "Event 모듈은 공연 목록과 공연 상태를 담당합니다. 공연 상태는 관리자가 직접 override할 수도 있고, 시작/종료 시간에 따라 자동 계산될 수도 있습니다.",
    )
    add_code_walkthrough(
        doc,
        "공연 상태 계산 해설",
        "backend/src/main/java/com/festflow/backend/service/EventService.java",
        """
        if (event.getStatusOverride() != null && !event.getStatusOverride().isBlank()) {
            return event.getStatusOverride();
        }

        if (now.isBefore(event.getStartTime())) {
            status = "예정";
        } else if (now.isAfter(event.getEndTime())) {
            status = "종료";
        } else {
            status = "진행중";
        }
        """,
        [
            ("statusOverride가 있으면 우선합니다. ", "지연, 취소, 수동 상태 변경을 반영하기 위해서입니다."),
            ("현재 시간이 시작 전이면 예정입니다. ", "공연 전 대기 상태입니다."),
            ("현재 시간이 종료 후면 종료입니다. ", "끝난 공연을 구분합니다."),
            ("그 사이면 진행중입니다. ", "프론트에서 라이브 상태로 보여줄 수 있습니다."),
        ],
    )
    add_note(
        doc,
        "30초마다 공연 스트림 방송",
        "EventService의 @Scheduled(fixedDelay = 30000) 메서드는 getAllEvents 결과를 StreamService.publishEvents로 보냅니다. 그래서 프론트 EventSource가 공연 상태 변화를 받을 수 있습니다.",
    )

    doc.add_heading("12. Analytics와 GPS 모듈", level=1)
    add_paragraph(
        doc,
        "혼잡도 분석은 GpsLog를 기반으로 합니다. 프론트 StageMapPage가 sendGps로 위치를 보내면 GpsService가 로그를 저장하고, AnalyticsService가 시간대별 방문, 인기 부스, 히트맵, 구역별 혼잡도를 계산합니다.",
    )
    add_table(
        doc,
        ["메서드", "계산 기준", "프론트 사용"],
        [
            ("trafficHourly", "최근 24시간 GPS 로그를 시간 단위로 그룹화", "HomePage 방문 흐름 카드"),
            ("popularBooths", "최근 60분 GPS 로그가 부스 100m 안에 있는 횟수", "분석 화면 인기 부스"),
            ("congestionHeatmap", "GPS 좌표를 0.001 단위 cell로 묶어 강도 계산", "히트맵 포인트"),
            ("stageCrowd", "무대 구역 반경 안의 최근 GPS 로그 수와 capacityHint 비교", "무대 관중 혼잡도"),
            ("dashboard", "현재/이전 시간창을 비교해 구역별 percent와 trend 계산", "AnalyticsPage 메인 대시보드"),
        ],
        [2600, 4650, 2110],
    )

    doc.add_heading("13. AI 가이드와 챗봇 모듈", level=1)
    add_paragraph(
        doc,
        "FestFlow의 AI 기능은 단순히 질문을 OpenAI로 보내는 구조가 아닙니다. 먼저 부스, 공연, 분실물, 공지, 혼잡도 같은 축제 데이터를 모아 근거를 만들고, "
        "API 키가 있으면 OpenAI 응답을 사용하고, 없거나 실패하면 fallback 답변을 제공합니다.",
    )
    add_code_walkthrough(
        doc,
        "AI 혼잡도 riskScore 해설",
        "backend/src/main/java/com/festflow/backend/service/AiCongestionService.java",
        """
        int riskScore = 0;
        riskScore += Math.min(30, crowdCount * 5);
        riskScore += Math.min(20, waitMinutes / 2);
        riskScore += Math.min(20, (int) activeReservations * 5);
        if (availableSeats <= 0 && Boolean.TRUE.equals(booth.reservationEnabled())) {
            riskScore += 12;
        }
        riskScore = Math.max(0, Math.min(100, riskScore));
        """,
        [
            ("crowdCount는 주변 사람 수입니다. ", "사람이 많을수록 위험 점수가 올라갑니다."),
            ("waitMinutes는 운영자가 입력한 예상 대기시간입니다. ", "기다리는 시간이 길수록 방문 추천이 약해집니다."),
            ("activeReservations는 예약/체크인 상태입니다. ", "예약이 많으면 자리 여유가 줄어듭니다."),
            ("availableSeats가 0이면 추가 점수를 줍니다. ", "예약 가능한 부스인데 좌석이 없으면 혼잡 위험이 큽니다."),
            ("마지막에 0에서 100 사이로 제한합니다. ", "프론트가 안정적으로 게이지나 라벨을 만들 수 있습니다."),
        ],
    )
    add_code_walkthrough(
        doc,
        "ChatService 근거 검색 흐름",
        "backend/src/main/java/com/festflow/backend/service/ChatService.java",
        """
        RetrievalResult retrieval = retrieveEvidence(question);
        String confidence = resolveConfidence(retrieval);

        if (apiKey == null || apiKey.isBlank()) {
            return new ChatResponseDto(buildFallbackAnswer(question, retrieval),
                                       confidence, retrieval.evidence(), warnings);
        }
        """,
        [
            ("먼저 retrieveEvidence로 근거를 찾습니다. ", "부스, 공연, 분실물, 공지, 정적 지식을 질문 의도에 맞춰 검색합니다."),
            ("confidence를 계산합니다. ", "근거가 충분한지 프론트에 알려줄 수 있습니다."),
            ("OpenAI 키가 없으면 fallback을 씁니다. ", "개발 환경에서도 챗봇 기능이 완전히 죽지 않습니다."),
            ("응답에는 evidence와 warnings가 포함됩니다. ", "프론트가 근거 버튼과 경고 문구를 보여줄 수 있습니다."),
        ],
    )

    doc.add_heading("14. Reservation 모듈 완전 해설", level=1)
    add_paragraph(
        doc,
        "예약 모듈은 가장 조심해서 읽어야 하는 부분입니다. 예약은 좌석 차감, 중복 예약 방지, 인증 토큰 확인, 체크인, 만료, 노쇼 기록, SSE 방송이 모두 연결됩니다. "
        "그래서 @Transactional이 중요합니다.",
    )
    add_code_walkthrough(
        doc,
        "예약 생성 핵심 흐름",
        "backend/src/main/java/com/festflow/backend/service/ReservationService.java",
        """
        String userKey = reservationAuthService.requireUserKey(authToken);
        if (userState.isBlocked(now)) {
            throw new ResponseStatusException(TOO_MANY_REQUESTS, "Reservation is temporarily blocked");
        }
        if (activeReservation.isPresent()) {
            throw new ResponseStatusException(CONFLICT, "Only one active reservation is allowed");
        }
        BoothReservationTable table = boothReservationTableRepository.findByIdForUpdate(requestDto.tableId())
            .orElseThrow(() -> new ResponseStatusException(NOT_FOUND, "Table not found."));
        """,
        [
            ("예약 인증 토큰으로 userKey를 확인합니다. ", "전화번호 인증이 끝난 사용자만 예약할 수 있게 합니다."),
            ("차단된 사용자인지 검사합니다. ", "반복 노쇼 사용자는 일시적으로 예약이 막힐 수 있습니다."),
            ("활성 예약이 이미 있는지 검사합니다. ", "한 사람이 여러 부스를 동시에 잡는 것을 막습니다."),
            ("findByIdForUpdate로 테이블을 잠급니다. ", "동시에 두 사람이 같은 좌석을 잡는 문제를 줄이기 위한 DB 락입니다."),
        ],
    )
    add_table(
        doc,
        ["예약 상태", "의미", "좌석에 미치는 영향"],
        [
            ("RESERVED", "예약됨. 아직 체크인 전", "테이블을 점유하므로 다른 사용자가 예약할 수 없습니다."),
            ("CHECKED_IN", "현장 체크인 완료", "사용 중인 좌석으로 간주합니다."),
            ("COMPLETED", "이용 완료", "좌석을 다시 복구합니다."),
            ("CANCELLED", "취소됨", "좌석을 다시 복구합니다."),
            ("EXPIRED", "만료됨", "좌석을 복구하고 노쇼 기록에 반영합니다."),
        ],
        [1900, 3200, 4260],
    )

    doc.add_heading("15. StreamService와 SSE 실시간 구조", level=1)
    add_paragraph(
        doc,
        "SSE는 FestFlow가 실시간처럼 보이게 만드는 핵심입니다. 프론트가 EventSource로 /api/stream/booths에 연결하면 백엔드는 SseEmitter를 목록에 저장합니다. "
        "나중에 부스나 예약이 바뀌면 StreamService가 연결된 브라우저들에게 이벤트를 보냅니다.",
    )
    add_code_walkthrough(
        doc,
        "StreamService 구조 해설",
        "backend/src/main/java/com/festflow/backend/service/stream/StreamService.java",
        """
        private final List<SseEmitter> boothEmitters = new CopyOnWriteArrayList<>();

        public SseEmitter subscribeBooths() {
            return createEmitter(boothEmitters);
        }

        public void publishBooths(Object payload) {
            send(boothEmitters, "booths", payload);
        }
        """,
        [
            ("boothEmitters는 연결 목록입니다. ", "현재 부스 스트림을 듣고 있는 브라우저들이 들어 있습니다."),
            ("subscribeBooths는 새 연결을 등록합니다. ", "StreamController의 /api/stream/booths가 이 메서드를 호출합니다."),
            ("publishBooths는 이벤트를 방송합니다. ", "이벤트 이름 booths와 payload를 모든 연결에 보냅니다."),
            ("CopyOnWriteArrayList를 씁니다. ", "연결이 추가/삭제되는 동안 순회해도 비교적 안전합니다."),
        ],
    )

    doc.add_heading("16. Admin, Ops, Staff 모듈 차이", level=1)
    add_table(
        doc,
        ["모듈", "사용자", "인증 방식", "대표 기능"],
        [
            ("Admin", "시스템 관리자", "JWT Bearer token / ROLE_ADMIN", "부스, 공연, 공지, 스태프, AI 매칭, 감사 로그 관리"),
            ("Ops Master", "현장 총괄 운영자", "X-Ops-Key master key / ROLE_OPS_MASTER", "현장 공지, 공연 상태, 부스 상태, AI 브리핑"),
            ("Ops Booth", "개별 부스 운영자", "X-Ops-Key booth key / ROLE_OPS_BOOTH", "내 부스 상태, 메뉴 이미지, 예약 체크인/완료"),
            ("Staff", "현장 스태프", "스태프 로그인 토큰", "상태 변경, 분실물, AI 현장 지원"),
        ],
        [1800, 1900, 3000, 2660],
    )
    add_note(
        doc,
        "왜 Admin과 Ops를 나눴나",
        "Admin은 시스템 관리 권한이고 Ops는 축제 현장에서 빠르게 쓰는 운영 권한입니다. 현장 운영자에게 관리자 비밀번호를 공유하지 않고도 필요한 작업만 열어줄 수 있습니다.",
    )

    doc.add_heading("17. LostItem, Upload, SMS, AI Match", level=1)
    add_table(
        doc,
        ["모듈", "핵심 파일", "설명"],
        [
            ("분실물", "LostItemController, LostItemService, LostItem", "분실물 등록, 사진 업로드, 상태 변경, 주인 확인, 삭제를 처리합니다."),
            ("업로드", "UploadStorageService, UploadAssetController", "local 저장 또는 S3 저장을 선택하고 /uploads URL을 제공합니다."),
            ("SMS", "SmsSender, TwilioSmsSender, AligoSmsSender, SolapiSmsSender", "전화번호 인증과 AI 매칭 알림에 사용할 수 있는 문자 발송 계층입니다."),
            ("AI 매칭", "AiMatchController, AiMatchService, AiMatchProfile, AiMatchRequest", "프로필 생성, 이미지 미리보기, 요청/수락/거절/만남 제안 흐름을 처리합니다."),
            ("번역", "TranslateController, TranslateService, TranslateMetricsService", "번역 요청과 사용량 메트릭을 처리합니다."),
        ],
        [1800, 3300, 4260],
    )

    doc.add_heading("18. DataInitializer: 데모 데이터가 생기는 이유", level=1)
    add_paragraph(
        doc,
        "DataInitializer는 서버가 시작될 때 DB가 비어 있으면 데모 부스, 공연, 공지, 스태프, 예약 테이블, GPS 로그를 넣습니다. "
        "프론트 개발자가 백엔드만 켜도 화면에 데이터가 보이게 하기 위한 장치입니다.",
    )
    add_bullets(
        doc,
        [
            "초기 관리자 계정은 app.init.admin.username과 app.init.admin.password 설정을 사용합니다.",
            "DB에 부스가 없으면 seedBooths 결과를 저장합니다.",
            "공연이 없으면 현재 시간 기준의 데모 공연을 저장합니다.",
            "스태프가 없으면 부스와 연결된 데모 스태프를 생성합니다.",
            "예약 테이블과 예약 데이터를 넣어 예약 화면을 바로 테스트할 수 있게 합니다.",
            "GPS 로그를 넣어 혼잡도/분석 화면이 비어 있지 않게 합니다.",
        ],
    )
    add_note(
        doc,
        "운영 환경 주의",
        "데모 데이터 초기화는 개발에는 편리하지만 운영 DB에서는 원치 않는 데이터가 생기지 않도록 프로필과 초기화 조건을 반드시 확인해야 합니다.",
        fill=WARNING_FILL,
    )

    doc.add_heading("19. 테스트 구조", level=1)
    add_paragraph(
        doc,
        "backend/src/test에는 예약 서비스와 부스 컨트롤러 테스트가 있습니다. 백엔드에서 특히 예약 로직처럼 실수하면 실제 사용자에게 피해가 갈 수 있는 영역은 테스트가 중요합니다.",
    )
    add_table(
        doc,
        ["테스트 파일", "대상", "의미"],
        [
            ("ReservationServiceTest", "예약 생성/상태 변경 로직", "좌석 차감, 중복 예약, 완료/만료 같은 규칙을 검증합니다."),
            ("ReservationAuthServiceTest", "예약 전화번호 인증", "코드 발급, 검증, 토큰 발급 흐름을 검증합니다."),
            ("BoothControllerTest", "부스 API 컨트롤러", "HTTP 요청과 응답 DTO가 정상인지 확인합니다."),
        ],
        [3000, 2800, 3560],
    )

    doc.add_heading("20. 백엔드 수정 실전 레시피", level=1)
    recipes = [
        ("새 공개 조회 API 추가", ["Controller에 @GetMapping을 추가합니다.", "Service에 조회 로직을 만듭니다.", "Repository 메서드가 필요하면 추가합니다.", "응답 DTO를 만듭니다.", "프론트 api.js에 fetch 함수를 추가합니다."]),
        ("새 관리자 API 추가", ["Controller 경로를 /api/admin/** 아래에 둡니다.", "SecurityConfig에서 /api/admin/**는 이미 ADMIN 권한입니다.", "프론트 api.js에서는 withAuth를 사용합니다.", "AuditLog가 필요한 작업이면 기록을 남깁니다.", "테스트 또는 수동 API 확인을 진행합니다."]),
        ("새 Entity 필드 추가", ["Entity에 필드를 추가합니다.", "DTO 요청/응답에 필요한 필드를 추가합니다.", "Service의 toDto와 update 로직을 수정합니다.", "프론트 타입/사용 위치를 수정합니다.", "DB ddl-auto나 마이그레이션 전략을 확인합니다."]),
        ("예약 규칙 변경", ["ReservationService의 createReservation 또는 상태 변경 메서드를 찾습니다.", "좌석 복구와 StreamService.publishReservations 호출을 놓치지 않습니다.", "@Transactional 범위가 필요한지 확인합니다.", "ReservationServiceTest를 추가/수정합니다.", "프론트 BoothDetailPage와 OpsBoothPage 상태 표시를 확인합니다."]),
        ("새 SSE 채널 추가", ["StreamService에 emitter 목록, subscribe, publish 메서드를 추가합니다.", "StreamController에 GET /api/stream/... 엔드포인트를 추가합니다.", "Service에서 데이터 변경 시 publish를 호출합니다.", "프론트 api.js에 create...Stream 함수를 만듭니다.", "페이지 useEffect cleanup에서 close합니다."]),
        ("OpenAI 프롬프트 수정", ["ChatService 또는 PublicAiGuideService의 instructions/input 생성 부분을 찾습니다.", "응답 실패 시 fallback이 여전히 자연스러운지 확인합니다.", "max_output_tokens와 timeout을 확인합니다.", "API key 없는 개발 환경도 동작하는지 봅니다.", "프론트 evidence/warnings 표시와 맞춥니다."]),
        ("파일 업로드 저장소 변경", ["application.properties의 app.storage.type을 확인합니다.", "S3 관련 bucket, region, endpoint, public-base-url을 설정합니다.", "UploadStorageService의 저장 URL 반환 형태를 확인합니다.", "프론트 resolveApiAssetUrl이 표시할 수 있는 URL인지 확인합니다.", "업로드 용량 제한도 확인합니다."]),
        ("CORS 문제 해결", ["브라우저 콘솔에서 CORS 오류 메시지를 확인합니다.", "app.cors.allowed-origins에 프론트 주소를 추가합니다.", "와일드카드와 배포 도메인 형식을 확인합니다.", "백엔드 재시작 후 다시 요청합니다.", "프론트 API_BASE가 실제 백엔드 주소인지 확인합니다."]),
        ("환경 변수 추가", ["application.properties에 ${ENV_NAME:default}를 추가합니다.", "해당 값을 @Value 또는 @ConfigurationProperties로 주입합니다.", "application-secrets.example.properties에도 예시를 적습니다.", "배포 환경 변수에 값을 등록합니다.", "민감 정보는 코드에 직접 쓰지 않습니다."]),
        ("테스트 추가", ["변경한 Service의 핵심 성공/실패 케이스를 정합니다.", "Repository나 외부 API 의존성을 어떻게 준비할지 결정합니다.", "상태 코드 예외도 테스트합니다.", "Gradle test를 실행합니다.", "프론트와 연결되는 응답 DTO 변경도 확인합니다."]),
    ]
    for title, steps in recipes:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)

    doc.add_heading("21. 백엔드 디버깅 플레이북", level=1)
    add_table(
        doc,
        ["증상", "먼저 볼 곳", "해결 방향"],
        [
            ("서버가 안 켜짐", "bootrun.log, application.properties, DB 연결", "DB URL/계정/포트와 Java 17 설정을 확인합니다."),
            ("프론트에서 Failed to fetch", "백엔드 실행 여부, CORS, API_BASE", "서버 포트와 프론트 VITE_API_BASE_URL을 맞춥니다."),
            ("관리자 API 403", "JWT 토큰, SecurityConfig, JwtAuthenticationFilter", "로그인 토큰의 role claim이 ADMIN인지 확인합니다."),
            ("운영 API 403", "X-Ops-Key, OpsKeyAuthenticationFilter, app.ops 설정", "마스터/부스 키가 올바른지 확인합니다."),
            ("예약 충돌 409", "ReservationService createReservation", "이미 활성 예약이 있거나 테이블이 점유되었는지 봅니다."),
            ("예약 429", "ReservationUserState", "노쇼로 임시 차단된 사용자인지 확인합니다."),
            ("SSE가 안 옴", "StreamController, StreamService, publish 호출", "구독 엔드포인트와 이벤트 이름이 프론트와 같은지 봅니다."),
            ("AI 답변이 fallback", "OPENAI_API_KEY, ChatService RestClient 오류", "API key와 외부 네트워크, timeout을 확인합니다."),
            ("이미지 업로드 실패", "multipart limit, UploadStorageService, storage type", "파일 크기와 저장 경로/S3 설정을 확인합니다."),
            ("DB 데이터가 이상함", "DataInitializer, ddl-auto, seed 조건", "데모 데이터가 다시 들어갔는지 확인합니다."),
        ],
        [2300, 3400, 3660],
    )

    add_backend_deep_controller_guide(doc)
    add_backend_deep_service_guide(doc)
    add_backend_flow_cases(doc)
    add_backend_error_and_status_catalog(doc)
    add_backend_practice_appendix(doc)

    doc.add_heading("30. 배포와 운영 체크리스트", level=1)
    add_bullets(
        doc,
        [
            "APP_JWT_SECRET은 빈 값으로 운영하지 않습니다.",
            "APP_INIT_ADMIN_PASSWORD 기본값을 운영에서 그대로 쓰지 않습니다.",
            "APP_OPS_MASTER_KEY와 부스 키를 기본값 0000/1111로 두지 않습니다.",
            "APP_CORS_ALLOWED_ORIGINS에 실제 프론트 배포 도메인을 넣습니다.",
            "SPRING_DATASOURCE_URL, USERNAME, PASSWORD를 운영 DB로 설정합니다.",
            "OPENAI_API_KEY가 없을 때 fallback 동작이 괜찮은지 확인합니다.",
            "SMS provider가 none이면 전화번호 인증/알림 UX를 별도로 확인합니다.",
            "업로드 저장소가 local이면 배포 환경에서 파일이 유지되는지 확인합니다.",
            "JPA ddl-auto=update를 운영에서 계속 쓸지 팀 정책을 정합니다.",
            "관리자/운영/스태프 권한을 실제 사용자 역할에 맞게 분리합니다.",
        ],
    )

    doc.add_heading("31. 백엔드 마스터 체크리스트", level=1)
    add_bullets(
        doc,
        [
            "SecurityConfig의 권한 규칙을 설명할 수 있다.",
            "api.js의 함수 하나를 백엔드 Controller와 Service까지 추적할 수 있다.",
            "Entity와 DTO의 차이를 설명할 수 있다.",
            "BoothService의 부스 조회와 혼잡도 계산을 설명할 수 있다.",
            "EventService의 상태 계산과 30초 SSE 방송을 설명할 수 있다.",
            "ReservationService의 인증, 중복 예약 방지, 좌석 차감, 체크인, 만료 처리를 설명할 수 있다.",
            "StreamService의 subscribe/publish 구조를 설명할 수 있다.",
            "ChatService가 근거 검색, OpenAI 호출, fallback을 어떻게 처리하는지 설명할 수 있다.",
            "application.properties의 주요 환경 변수를 운영 관점에서 설명할 수 있다.",
            "새 API를 추가할 때 Controller, Service, Repository, DTO, 프론트 api.js를 함께 수정할 수 있다.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("부록 A. 백엔드 코드 읽기 순서 요약", level=1)
    add_numbered(
        doc,
        [
            "application.properties에서 포트, DB, JWT, CORS, 외부 API 설정을 확인합니다.",
            "SecurityConfig에서 해당 API가 공개인지 보호 API인지 확인합니다.",
            "Controller에서 URL과 HTTP method를 찾습니다.",
            "Controller가 호출하는 Service 메서드를 찾습니다.",
            "Service에서 비즈니스 규칙, 예외, 트랜잭션, publish 호출을 확인합니다.",
            "Repository 메서드와 Entity 관계를 확인합니다.",
            "응답 DTO가 프론트 api.js와 페이지 state에서 어떻게 쓰이는지 확인합니다.",
        ],
    )
    doc.save(OUT_BACKEND)
    return OUT_BACKEND


def build():
    frontend = add_frontend_doc()
    backend = add_backend_doc()
    print(f"created={frontend}")
    print(f"created={backend}")


if __name__ == "__main__":
    build()
