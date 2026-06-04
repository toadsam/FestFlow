from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "technical-assets"
OUT_DOCX = ROOT / "FestFlow_전체_기술설명서.docx"
FONT = Path("C:/Windows/Fonts/malgun.ttf")
BOLD_FONT = Path("C:/Windows/Fonts/malgunbd.ttf")

INK = "111827"
MUTED = "475569"
BLUE = "2563EB"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F8FAFC"
CALLOUT_FILL = "EEF2FF"
BORDER = "CBD5E1"


def pil_font(size, bold=False):
    return ImageFont.truetype(str(BOLD_FONT if bold else FONT), size)


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in str(text).split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = current + " " + word
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered(draw, box, text, size=28, fill="#111827", bold=False, max_width_ratio=0.86):
    x1, y1, x2, y2 = box
    font = pil_font(size, bold)
    max_width = int((x2 - x1) * max_width_ratio)
    lines = wrap_text(draw, text, font, max_width)
    line_height = size + 9
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height


def rounded_box(draw, box, fill, outline="#93C5FD", width=2, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color="#2563EB", width=5):
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def make_architecture_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-overall-architecture.png"
    image = Image.new("RGB", (1700, 1050), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((60, 44), "FestFlow 전체 아키텍처", fill="#111827", font=pil_font(46, True))
    draw.text(
        (60, 104),
        "브라우저 기반 React 앱이 Spring Boot API를 호출하고, 백엔드는 DB/업로드 저장소/AI/SMS/SSE를 통해 축제 운영 데이터를 처리합니다.",
        fill="#475569",
        font=pil_font(24),
    )

    boxes = {
        "client": (60, 210, 390, 430),
        "frontend": (500, 210, 850, 430),
        "api": (980, 210, 1320, 430),
        "admin": (60, 590, 390, 790),
        "ops": (500, 590, 850, 790),
        "service": (980, 590, 1320, 790),
        "db": (1430, 190, 1645, 350),
        "storage": (1430, 395, 1645, 555),
        "external": (1430, 600, 1645, 800),
        "sse": (980, 845, 1320, 990),
    }
    labels = {
        "client": "사용자 브라우저\n모바일/데스크탑\nPWA 설치 가능",
        "frontend": "React + Vite\n라우트/페이지\napi.js 호출 계층",
        "api": "Spring Boot 3 API\nController\nSecurity Filter",
        "admin": "관리자/스태프\nJWT, 운영 키,\n스태프 토큰",
        "ops": "운영 기능\n부스/공연/공지\n예약/분실물/AI 보조",
        "service": "Service 계층\n검증, 상태 전이,\n도메인 규칙",
        "db": "MySQL\nJPA Entity\nRepository",
        "storage": "업로드 저장소\nlocal 또는 S3\n이미지/메뉴/프로필",
        "external": "외부 연동\nOpenAI\nSMS, 지도 타일",
        "sse": "SSE 실시간 스트림\n혼잡도, 공지, 공연,\n예약, 스태프, 분실물",
    }
    fills = {
        "client": "#ECFEFF",
        "frontend": "#EEF2FF",
        "api": "#EFF6FF",
        "admin": "#FDF2F8",
        "ops": "#F0FDF4",
        "service": "#F5F3FF",
        "db": "#FEFCE8",
        "storage": "#FFF7ED",
        "external": "#FDF2F8",
        "sse": "#F0FDFA",
    }
    for key, box in boxes.items():
        rounded_box(draw, box, fills[key])
        draw_centered(draw, box, labels[key], size=25, bold=key in {"frontend", "api", "service"})

    arrow(draw, (390, 320), (500, 320))
    arrow(draw, (850, 320), (980, 320))
    arrow(draw, (1150, 430), (1150, 590))
    arrow(draw, (1320, 300), (1430, 270))
    arrow(draw, (1320, 380), (1430, 475), "#F97316")
    arrow(draw, (1320, 670), (1430, 700), "#DB2777")
    arrow(draw, (1150, 790), (1150, 845), "#0D9488")
    arrow(draw, (980, 915), (850, 725), "#0D9488")
    arrow(draw, (500, 700), (390, 700), "#16A34A")
    arrow(draw, (500, 700), (850, 700), "#16A34A")

    draw.text(
        (70, 960),
        "핵심 경계: 화면은 사용자 입력과 표시를 담당하고, 백엔드 Service가 상태 변경과 검증을 책임지며, DB/저장소/외부 API는 결과와 보조 데이터를 제공합니다.",
        fill="#111827",
        font=pil_font(24, True),
    )
    image.save(path)
    return path


def make_backend_flow_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-backend-flow.png"
    image = Image.new("RGB", (1700, 980), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((60, 44), "백엔드 요청 처리 흐름", fill="#111827", font=pil_font(44, True))
    draw.text((60, 100), "HTTP 요청은 인증 필터를 통과한 뒤 Controller, Service, Repository, Entity 순서로 처리됩니다.", fill="#475569", font=pil_font(24))

    boxes = [
        ((80, 210, 360, 380), "HTTP 요청\n브라우저/api.js", "#ECFEFF"),
        ((470, 210, 750, 380), "Security\nJWT / OPS KEY\nStaff Token", "#FDF2F8"),
        ((860, 210, 1140, 380), "Controller\n요청 매핑\nDTO 수신", "#EFF6FF"),
        ((1250, 210, 1530, 380), "Service\n트랜잭션\n검증/상태 변경", "#F0FDF4"),
        ((470, 560, 750, 730), "Repository\nJPA Query\nLock/조회", "#FEFCE8"),
        ((860, 560, 1140, 730), "Entity\n도메인 상태\n테이블 매핑", "#FFF7ED"),
        ((1250, 560, 1530, 730), "SSE/외부 API\nOpenAI, SMS,\nUploadStorage", "#F5F3FF"),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=25, bold=True)
    arrow(draw, (360, 295), (470, 295))
    arrow(draw, (750, 295), (860, 295))
    arrow(draw, (1140, 295), (1250, 295))
    arrow(draw, (1390, 380), (1390, 560))
    arrow(draw, (1250, 645), (1140, 645))
    arrow(draw, (860, 645), (750, 645))
    arrow(draw, (610, 560), (1390, 380), "#64748B", 3)
    draw.text((80, 850), "예: 예약 생성은 인증 토큰 확인 -> 테이블 행 잠금 -> 중복 예약/좌석 검증 -> 예약 생성 -> SSE 발행 순서로 처리됩니다.", fill="#111827", font=pil_font(25, True))
    image.save(path)
    return path


def make_realtime_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "festflow-realtime-streams.png"
    image = Image.new("RGB", (1700, 860), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((60, 44), "실시간 반영 구조", fill="#111827", font=pil_font(44, True))
    draw.text((60, 100), "상태를 바꾸는 서비스는 StreamService에 이벤트를 발행하고, 프론트엔드는 EventSource로 구독합니다.", fill="#475569", font=pil_font(24))
    stream_box = (660, 285, 1040, 505)
    rounded_box(draw, stream_box, "#F0FDFA")
    draw_centered(draw, stream_box, "StreamService\nSseEmitter 목록 관리\n죽은 연결 제거", size=27, bold=True)

    left = [
        ((80, 190, 390, 320), "GpsService\n혼잡도 갱신", "#ECFEFF", "congestion"),
        ((80, 365, 390, 495), "Event/Notice/Booth\n운영 데이터 변경", "#EEF2FF", "events/notices/booths"),
        ((80, 540, 390, 670), "Reservation/Staff\n예약/스태프/분실물", "#F0FDF4", "reservations/staff/lost-items"),
    ]
    right = [
        ((1280, 190, 1600, 320), "홈/지도/공연\n상태 즉시 갱신", "#EFF6FF"),
        ((1280, 365, 1600, 495), "운영 콘솔\n대시보드 반영", "#F5F3FF"),
        ((1280, 540, 1600, 670), "스태프/예약 화면\n현장 상태 반영", "#FFF7ED"),
    ]
    for box, label, fill, event_name in left:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=23, bold=True)
        arrow(draw, (390, (box[1] + box[3]) // 2), (660, 395), "#2563EB")
        draw.text((430, (box[1] + box[3]) // 2 - 24), event_name, fill="#1D4ED8", font=pil_font(19, True))
    for box, label, fill in right:
        rounded_box(draw, box, fill)
        draw_centered(draw, box, label, size=23, bold=True)
        arrow(draw, (1040, 395), (1280, (box[1] + box[3]) // 2), "#0D9488")
    draw.text((85, 765), "구독 엔드포인트: /api/stream/congestion, /events, /notices, /booths, /staff, /lost-items, /reservations", fill="#111827", font=pil_font(24, True))
    image.save(path)
    return path


def set_run_font(run, size=None, bold=None, color=None, name="Malgun Gothic"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=9.0, color=INK):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def set_column_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Title", 24, INK, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    run = p.add_run("FestFlow 전체 기술설명서")
    set_run_font(run, 27, True, BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10, line=1.15)
    run = p.add_run("대학교 축제 관리 웹앱의 프론트엔드, 백엔드, DB, 실시간 처리, 인증, 배포 구조 상세 문서")
    set_run_font(run, 11.5, False, MUTED)

    meta = [
        ("대상 프로젝트", "FestFlow"),
        ("문서 범위", "React/Vite 프론트엔드, Spring Boot 백엔드, MySQL/JPA, PWA, SSE, 운영/예약/AI/분실물/스태프/AI Match 기능"),
        ("작성 기준", "현재 작업 폴더의 소스 코드와 설정 파일"),
        ("스타일 프리셋", "compact_reference_guide, 한글 렌더링을 위해 Malgun Gothic 적용"),
    ]
    add_table(document, ["항목", "내용"], meta, widths=[3.2, 12.4])
    add_note(
        document,
        "문서 사용 방법",
        "처음 보는 사람은 1~4장을 먼저 읽어 전체 구조를 잡고, 개발자는 API/DB/서비스 섹션을 기능별로 찾아보면 됩니다. 운영자는 인증, 예약, SSE, 배포/환경변수 섹션을 우선 확인하면 됩니다.",
    )


def heading(document, text, level=1):
    return document.add_heading(text, level=level)


def paragraph(document, text, bold_prefix=None):
    p = document.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 10.5, True, INK)
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    set_run_font(r, 10.5, False, INK)
    return p


def bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(str(item))
        set_run_font(r, 10, False, INK)


def numbers(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        r = p.add_run(str(item))
        set_run_font(r, 10, False, INK)


def add_note(document, title, body, fill=CALLOUT_FILL):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.2)
    r = p.add_run(title + "\n")
    set_run_font(r, 10.5, True, BLUE)
    r = p.add_run(body)
    set_run_font(r, 9.7, False, INK)
    document.add_paragraph()


def add_table(document, headers, rows, widths=None, font_size=8.8):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, True, font_size + 0.3, INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(cells):
                set_cell_text(cells[idx], value, False, font_size, INK)
    if widths:
        set_column_widths(table, widths)
    document.add_paragraph()
    return table


def add_image(document, path, caption, width_cm=15.7):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line=1.0)
    r = p.add_run(caption)
    set_run_font(r, 9, False, MUTED)


def code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.05)
    for line in text.strip("\n").splitlines():
        r = p.add_run(line + "\n")
        set_run_font(r, 8.7, False, INK, name="Consolas")
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    document.add_paragraph()


def add_page_break(document):
    document.add_section(WD_SECTION.NEW_PAGE)


def section_project_overview(document):
    heading(document, "1. 프로젝트 개요")
    paragraph(
        document,
        "FestFlow는 대학교 축제 현장의 부스, 공연, 공지, 혼잡도, 예약, 분실물, 스태프 운영, AI 안내 기능을 하나의 웹앱으로 묶은 운영형 서비스입니다. 방문자는 모바일 중심 화면을 사용하고, 운영자는 관리자/운영 콘솔/스태프 콘솔을 통해 현장 데이터를 관리합니다.",
    )
    add_table(
        document,
        ["사용자군", "주요 화면", "핵심 작업", "인증 방식"],
        [
            ["방문자", "/", "/stage-map, /events, /booths/:id, /lost-found, /chat, /ai-match", "부스 탐색, 공연 확인, 예약, 분실물 확인, AI 안내, AI Match", "대부분 공개 API, 예약은 휴대폰 인증 토큰"],
            ["관리자", "/admin, /ai-match/admin", "부스/공연/공지 CRUD, CSV import, 감사 로그, AI Match 운영", "JWT Bearer 토큰"],
            ["통합 운영자", "/ops/master", "전체 부스/공연/공지/AI 브리핑/혼잡 완화 공지 처리", "X-OPS-KEY, OPS_MASTER"],
            ["부스 운영자", "/ops/booth/:id", "부스 상태, 메뉴 이미지, 예약 설정, 체크인/완료 처리", "X-OPS-KEY, OPS_BOOTH"],
            ["스태프", "/staff", "로그인, 상태/위치/업무 갱신, AI 현장 보조, 분실물 지원", "X-Staff-Token"],
        ],
        widths=[2.3, 3.9, 6.6, 3.1],
    )
    add_note(
        document,
        "핵심 설계 방향",
        "서비스는 방문자 편의 기능과 운영자 도구가 같은 데이터 모델을 공유하도록 설계되어 있습니다. 예를 들어 부스 대기 시간, 재고, 예약 상태는 방문자 화면에서 안내 정보가 되고, 운영 화면에서는 실시간 조정 대상이 됩니다.",
    )

    heading(document, "1.1 기술 스택", 2)
    add_table(
        document,
        ["계층", "기술", "역할"],
        [
            ["Frontend", "React 18, Vite, React Router, Tailwind CSS, Leaflet/React Leaflet", "페이지 라우팅, 모바일/데스크탑 UI, 지도, PWA 셸, API 호출"],
            ["Backend", "Java 17, Spring Boot 3.3.5, Spring Web, Spring Security, Spring Data JPA", "REST API, 인증/권한, 도메인 서비스, DB 접근"],
            ["Database", "MySQL 8 기본, PostgreSQL runtime 의존성 포함", "축제 운영 데이터와 예약/스태프/AI Match 상태 저장"],
            ["Realtime", "Server-Sent Events, SseEmitter", "혼잡도, 공지, 공연, 부스, 스태프, 분실물, 예약 상태 push"],
            ["AI/External", "OpenAI Responses API, OpenAI Images API, Google Translate endpoint, SMS provider", "방문자 AI 가이드, 운영 AI 보조, 이미지 변환, 번역, 예약 인증"],
            ["Deploy", "Railway backend, Vercel frontend", "백엔드 API와 프론트 정적 앱 배포"],
        ],
        widths=[2.6, 5.0, 8.0],
    )


def section_architecture(document, arch, backend_flow, realtime):
    heading(document, "2. 전체 아키텍처")
    add_image(document, arch, "그림 1. FestFlow 전체 아키텍처")
    paragraph(
        document,
        "프론트엔드는 단일 React 앱이지만 실제 역할은 공개 방문자 화면, 관리자 화면, 운영 콘솔, 스태프 콘솔로 나뉩니다. 백엔드는 URL prefix와 인증 필터를 통해 권한 경계를 나누며, 서비스 계층에서 실제 상태 변경 규칙을 처리합니다.",
    )
    add_image(document, backend_flow, "그림 2. 백엔드 요청 처리 흐름")
    add_table(
        document,
        ["구성 요소", "대표 파일", "설명"],
        [
            ["라우팅", "frontend/src/main.jsx", "React Router가 모든 화면 경로를 등록하고 lazy loading으로 페이지를 분리합니다."],
            ["API 클라이언트", "frontend/src/api.js", "VITE_API_BASE_URL을 기준으로 모든 REST/SSE 호출을 중앙화합니다."],
            ["보안", "SecurityConfig.java, JwtAuthenticationFilter.java, OpsKeyAuthenticationFilter.java", "/api/admin/**, /api/ops/** 권한을 필터에서 나눕니다."],
            ["도메인 서비스", "backend/service/*.java", "CRUD보다 중요한 검증, 상태 전이, 권한 확인, SSE 발행을 담당합니다."],
            ["저장소", "backend/repository/*.java", "JPA Repository 기반 DB 접근. 예약 테이블 등 일부는 pessimistic lock을 사용합니다."],
            ["실시간", "StreamService.java, StreamController.java", "SseEmitter 리스트를 관리하고 변경 이벤트를 각 채널에 발행합니다."],
        ],
        widths=[3.0, 5.2, 7.2],
    )
    add_image(document, realtime, "그림 3. SSE 실시간 반영 구조")


def section_frontend(document):
    heading(document, "3. 프론트엔드 구조")
    paragraph(
        document,
        "프론트엔드는 `frontend/src/main.jsx`에서 페이지를 lazy import하고, `App.jsx`가 공통 shell과 하단/상단 내비게이션 노출 조건을 제어합니다. 일반 축제 화면과 운영 화면은 같은 앱 안에 있지만 라우트 scope에 따라 UI가 다르게 동작합니다.",
    )
    add_table(
        document,
        ["경로", "페이지 파일", "역할"],
        [
            ["/", "HomePage.jsx", "축제 홈, 주요 부스/공지/요약 진입점"],
            ["/stage-map", "StageMapPage.jsx", "지도 기반 부스/혼잡도/공연 위치 확인"],
            ["/events", "EventPage.jsx", "공연 목록과 상태 확인"],
            ["/events/lineup", "LineupPage.jsx", "공연 라인업 전용 보기"],
            ["/analytics", "AnalyticsPage.jsx", "혼잡도 분석, AI 예측/가이드"],
            ["/booths/:id", "BoothDetailPage.jsx", "부스 상세, 지도, 예약 상태"],
            ["/lost-found", "LostFoundPage.jsx", "분실물 목록, 등록/상태 확인"],
            ["/chat", "ChatPage.jsx", "질문 기반 챗봇 응답"],
            ["/staff", "StaffPage.jsx", "스태프 로그인, 상태 갱신, AI 보조"],
            ["/more", "MorePage.jsx", "부가 기능 진입"],
            ["/admin", "AdminPage.jsx", "관리자 대시보드와 CRUD"],
            ["/ops/master", "OpsMasterPage.jsx", "통합 운영 콘솔"],
            ["/ops/booth/:id", "OpsBoothPage.jsx", "부스별 운영 콘솔, 예약/체크인"],
            ["/ai-match", "AiMatchPage.jsx", "AI Match 방문자 화면"],
            ["/ai-match/admin", "AiMatchAdminPage.jsx", "AI Match 관리자 화면"],
        ],
        widths=[3.0, 4.2, 8.2],
    )
    heading(document, "3.1 API 호출 계층", 2)
    paragraph(
        document,
        "`frontend/src/api.js`는 API base URL, 인증 헤더, JSON 파싱, timeout 처리, SSE 생성 함수를 한 파일에 모읍니다. 이 구조는 페이지 컴포넌트가 직접 URL 문자열을 조합하는 일을 줄이고, 백엔드 경로 변경 시 수정 지점을 제한합니다.",
    )
    add_table(
        document,
        ["함수 묶음", "대표 함수", "설명"],
        [
            ["관리자", "loginAdmin, fetchAdminDashboardKpis, fetchAuditLogs", "JWT 토큰 기반 관리자 호출"],
            ["부스/공연/공지", "fetchBooths, createBooth, updateEvent, createNotice", "공개 조회와 관리자 CRUD"],
            ["운영 콘솔", "fetchOpsMasterBootstrap, updateOpsBoothLiveStatus", "X-OPS-KEY 기반 운영 API"],
            ["예약", "sendReservationAuthCode, createBoothReservation, checkInOpsBoothReservation", "휴대폰 인증 토큰과 운영 키를 나누어 사용"],
            ["스태프", "loginStaff, fetchStaffBootstrap, updateMyStaffStatus", "스태프 세션 토큰 기반 호출"],
            ["분실물", "fetchLostItems, createLostItem, claimLostItem", "권한에 따라 연락처 마스킹과 수정 가능 범위가 달라짐"],
            ["AI/번역", "fetchAiVisitorGuide, fetchOpsMasterAiBriefing, translateText", "OpenAI/번역 기능의 프론트 진입점"],
            ["SSE", "createCongestionStream, createReservationStream", "EventSource 구독 객체 생성"],
        ],
        widths=[2.7, 5.2, 7.6],
    )
    heading(document, "3.2 PWA와 정적 배포", 2)
    bullets(
        document,
        [
            "`manifest.json`은 앱 이름, 아이콘, 시작 URL, standalone 표시 방식을 정의합니다.",
            "`service-worker.js`는 앱 shell과 build asset을 캐시하고, API/SSE/uploads/OSM 지도 타일은 캐시하지 않습니다.",
            "`vercel.json`은 파일 시스템 우선 처리 후 나머지 경로를 `index.html`로 보내 React Router 새로고침을 지원합니다.",
            "프로덕션에서만 service worker를 등록하므로 로컬 개발 중 캐시 혼선을 줄입니다.",
        ],
    )


def section_backend(document):
    heading(document, "4. 백엔드 패키지 구조")
    add_table(
        document,
        ["패키지", "역할", "예시"],
        [
            ["controller", "REST API 입구. 요청 경로, 메서드, 권한 prefix가 명확히 드러남", "BoothController, LostItemController, OpsController"],
            ["controller.admin", "관리자 전용 API. SecurityConfig에서 ROLE_ADMIN 필요", "AdminBoothController, AdminNoticeController"],
            ["controller.ops", "운영 키 기반 master/booth API", "OpsController"],
            ["controller.staff", "스태프 로그인과 AI 보조 API", "StaffController"],
            ["service", "도메인 규칙, 트랜잭션, 상태 변경, 외부 API 호출", "ReservationService, AiMatchService"],
            ["service.analytics", "혼잡도/방문량/대시보드 분석", "AnalyticsService"],
            ["service.stream", "SSE 구독/발행", "StreamService"],
            ["entity", "JPA 테이블 매핑과 상태 변경 메서드", "Booth, BoothReservation, StaffMember"],
            ["repository", "JPA Repository와 custom query", "BoothReservationRepository"],
            ["security", "JWT, 운영 키, rate limit filter", "JwtService, OpsKeyAuthenticationFilter"],
            ["init", "데모 데이터와 레거시 schema 보정", "DataInitializer, AiMatchSchemaInitializer"],
        ],
        widths=[3.1, 7.2, 5.0],
    )
    heading(document, "4.1 인증 경계", 2)
    add_table(
        document,
        ["API 영역", "인증 수단", "권한", "설명"],
        [
            ["/api/admin/**", "Authorization: Bearer <JWT>", "ROLE_ADMIN", "관리자 로그인 후 발급된 JWT로 접근"],
            ["/api/ops/master/**", "X-OPS-KEY", "OPS_MASTER", "통합 운영 콘솔 전용"],
            ["/api/ops/booth/**", "X-OPS-KEY", "OPS_MASTER 또는 OPS_BOOTH", "부스 운영 콘솔 전용"],
            ["/api/staff/**", "X-Staff-Token", "스태프 세션", "스태프 상태/AI 보조/분실물 운영"],
            ["/api/reservations/auth/**", "전화번호 인증", "예약 사용자", "예약용 사용자 계정/세션 토큰 발급"],
            ["/api/ai-match/**", "PIN 기반 프로필 인증", "프로필 소유자", "공개 API이지만 개인 작업은 닉네임+PIN 확인"],
            ["/api/** 공개 조회", "없음", "방문자", "부스, 공연, 공지, 혼잡도, 일부 AI 안내"],
        ],
        widths=[3.5, 4.0, 3.2, 7.8],
    )


def section_feature_modules(document):
    heading(document, "5. 기능 모듈별 상세")
    modules = [
        ("부스/지도/혼잡도", "BoothService, GpsService, AnalyticsService", "부스 CRUD, 좌표 기반 지도 표시, 최근 15분 GPS 로그와 시간 가중치를 이용한 혼잡도 산정", "혼잡도 단계는 여유/보통/혼잡/매우혼잡이며, 80m 반경 내 GPS 로그를 기준으로 합니다."),
        ("공연/라인업", "EventService", "공연 일정, 상태, 이미지, 실시간 메시지, 지연 시간 관리", "공연 상태는 시간 기반 상태와 관리자 override/liveMessage를 함께 사용합니다."),
        ("공지", "NoticeService, AdminActionService", "활성 공지 조회, 관리자 CRUD, 혼잡 완화/공연 시작 공지 자동 발행", "SSE notices 채널로 홈과 운영 화면에 반영됩니다."),
        ("관리자 대시보드", "AdminDashboardService, AuditLogService", "KPI, 감사 로그, 부스/공연/공지 관리", "관리 작업은 감사 로그로 추적됩니다."),
        ("운영 콘솔", "OpsController, OpsAiService", "통합 운영과 부스별 운영을 분리하고, 운영 키로 보호", "혼잡 완화 공지, AI 브리핑, 부스별 예약 체크인 처리를 제공합니다."),
        ("예약", "ReservationAuthService, ReservationService", "전화번호 인증, 예약 생성, QR 체크인 토큰, 노쇼 패널티, 테이블 상태 관리", "동시성 방지를 위해 테이블 조회에 write lock을 사용합니다."),
        ("스태프", "StaffService", "스태프 로그인, 상태/업무/위치 갱신, 관리자 수정", "데모 계정과 DB 기반 계정을 모두 지원합니다."),
        ("분실물", "LostItemService", "분실물 등록, 상태 변경, 찾는 사람 claim, 연락처 마스킹", "공개 조회와 운영자 조회에서 연락처 노출 수준이 달라집니다."),
        ("AI 방문 안내", "PublicAiGuideService, AiCongestionService", "혼잡도, 공연, 지도 데이터를 기반으로 방문자 행동 추천", "OpenAI 키가 없거나 실패하면 fallback 로직을 반환합니다."),
        ("번역", "TranslateService, TranslateMetricsService", "한국어/영어 preset, Google Translate endpoint, fallback 번역", "외부 호출 실패 시에도 서비스가 끊기지 않도록 fallback을 사용합니다."),
        ("AI Match", "AiMatchService, AiImageGenerationService", "AI 프로필 이미지 변환, 프로필/신청/즐겨찾기/관리자 연결 상태", "전화번호별 이미지 변환 횟수 제한과 삭제 번호 재가입 차단이 있습니다."),
    ]
    add_table(document, ["모듈", "대표 구현", "기능", "중요 규칙"], modules, widths=[2.8, 4.3, 6.0, 6.0], font_size=8.1)

    heading(document, "5.1 혼잡도 계산", 2)
    numbers(
        document,
        [
            "GpsService가 사용자의 latitude/longitude를 받아 gps_logs에 저장합니다.",
            "BoothService는 최근 15분 로그만 조회합니다.",
            "각 부스 좌표에서 80m 이내 로그를 필터링합니다.",
            "오래된 로그는 낮은 가중치를 받도록 timeWeight를 적용합니다.",
            "가중 점수를 반올림해 nearbyUserCount로 만들고 단계로 변환합니다.",
        ],
    )
    add_table(
        document,
        ["가중 인원", "혼잡도 단계", "해석"],
        [
            ["0~2", "여유", "방문 부담 낮음"],
            ["3~6", "보통", "약간의 대기 가능"],
            ["7~11", "혼잡", "우회 또는 대기 시간 고려"],
            ["12 이상", "매우혼잡", "운영자 개입 또는 방문 지연 권장"],
        ],
        widths=[3.0, 3.0, 8.0],
    )

    heading(document, "5.2 예약 상태 전이", 2)
    add_table(
        document,
        ["상태", "생성/변경 조건", "좌석 영향", "후속 처리"],
        [
            ["RESERVED", "사용자가 인증 토큰으로 테이블 예약 생성", "선택 좌석 수만큼 availableSeats 감소", "만료 전 QR 체크인 가능"],
            ["CHECKED_IN", "부스 운영자가 직접 체크인하거나 QR 토큰 확인", "좌석 점유 유지", "완료 또는 테이블 release 가능"],
            ["COMPLETED", "체크인된 예약을 완료 처리", "좌석 복구", "예약 이력으로 남음"],
            ["CANCELLED", "운영자가 reserved 상태 테이블을 release", "좌석 복구", "노쇼 패널티 없음"],
            ["EXPIRED", "예약 만료 시간 초과", "좌석 복구", "사용자 noShowCount 증가 및 조건부 차단"],
        ],
        widths=[2.6, 6.0, 4.0, 4.8],
    )


def section_api(document):
    heading(document, "6. API 명세 요약")
    paragraph(
        document,
        "아래 표는 Controller 기준으로 정리한 API 표입니다. 실제 프론트 호출은 `frontend/src/api.js`에 대부분 매핑되어 있습니다.",
    )
    api_rows = [
        ["공개 부스", "GET", "/api/booths", "없음", "부스 목록"],
        ["공개 부스", "GET", "/api/booths/{id}", "없음", "부스 상세"],
        ["공개 부스", "GET", "/api/booths/{id}/congestion", "없음", "부스 혼잡도"],
        ["예약", "GET", "/api/booths/{id}/reservations", "예약 토큰 선택", "부스 예약 상태"],
        ["예약", "POST", "/api/booths/{id}/reservations", "예약 토큰", "예약 생성"],
        ["예약", "POST", "/api/booths/{id}/reservations/{reservationId}/check-in-token", "예약 토큰", "QR 체크인 토큰 발급"],
        ["공연", "GET", "/api/events", "없음", "공연 목록"],
        ["GPS", "POST", "/api/gps", "없음", "위치 로그 저장"],
        ["채팅", "POST", "/api/chat", "없음", "챗봇 응답"],
        ["공지", "GET", "/api/notices/active", "없음", "활성 공지"],
        ["분실물", "GET", "/api/lost-items", "선택: 관리자/스태프 토큰", "목록 조회, 권한 없으면 연락처 마스킹"],
        ["분실물", "POST", "/api/lost-items", "선택: 관리자/스태프 토큰", "분실물 등록"],
        ["분실물", "PUT", "/api/lost-items/{id}/status", "관리자/스태프 권장", "상태 변경"],
        ["분실물", "PUT", "/api/lost-items/{id}/claim", "없음", "소유자 claim 요청"],
        ["분실물", "DELETE", "/api/lost-items/{id}", "관리자/스태프", "삭제"],
        ["분석", "GET", "/api/analytics/traffic-hourly", "없음", "시간대별 방문량"],
        ["분석", "GET", "/api/analytics/popular-booths", "없음", "인기 부스"],
        ["분석", "GET", "/api/analytics/congestion-heatmap", "없음", "히트맵 포인트"],
        ["분석", "GET", "/api/analytics/stage-crowd", "없음", "무대 혼잡도"],
        ["분석", "GET", "/api/analytics/dashboard", "없음", "혼잡도 대시보드"],
        ["AI", "GET", "/api/ai/guide", "없음", "축제 AI 가이드"],
        ["AI", "GET", "/api/ai/congestion/predictions", "없음", "혼잡도 예측"],
        ["AI", "GET", "/api/ai/decisions", "없음", "AI 판단 이력"],
        ["AI", "GET", "/api/ai/visitor-guide/{scope}", "없음", "페이지별 방문자 추천"],
        ["번역", "POST", "/api/translate", "없음", "텍스트 번역"],
        ["번역", "GET", "/api/translate/metrics", "없음", "번역 지표"],
    ]
    add_table(document, ["그룹", "메서드", "경로", "권한", "역할"], api_rows, widths=[2.2, 1.6, 5.5, 3.0, 5.6], font_size=7.7)

    heading(document, "6.1 관리자 API", 2)
    admin_rows = [
        ["POST", "/api/auth/login", "없음", "관리자 로그인, JWT 발급"],
        ["GET", "/api/admin/dashboard/kpis", "ADMIN", "관리자 KPI"],
        ["GET", "/api/admin/audit-logs", "ADMIN", "감사 로그"],
        ["POST/PUT/DELETE", "/api/admin/booths, /api/admin/booths/{id}", "ADMIN", "부스 CRUD"],
        ["PUT", "/api/admin/booths/{id}/live-status", "ADMIN", "부스 실시간 상태 변경"],
        ["POST", "/api/admin/booths/{id}/image", "ADMIN", "부스 이미지 업로드"],
        ["PUT", "/api/admin/booths/reorder", "ADMIN", "부스 순서 저장"],
        ["POST/PUT/DELETE", "/api/admin/events, /api/admin/events/{id}", "ADMIN", "공연 CRUD"],
        ["POST/PUT/DELETE", "/api/admin/notices, /api/admin/notices/{id}", "ADMIN", "공지 CRUD"],
        ["POST", "/api/admin/import/booths", "ADMIN", "부스 CSV import"],
        ["POST", "/api/admin/import/events", "ADMIN", "공연 CSV import"],
        ["POST", "/api/admin/actions/congestion-relief-notice", "ADMIN", "혼잡 완화 공지 발행"],
        ["POST", "/api/admin/actions/events/{eventId}/start-notice", "ADMIN", "공연 시작 공지 발행"],
        ["GET/PUT", "/api/admin/staff, /api/admin/staff/{id}", "ADMIN", "스태프 목록/수정"],
        ["GET/PUT/DELETE/POST", "/api/admin/ai-match/**", "ADMIN", "AI Match 운영"],
    ]
    add_table(document, ["메서드", "경로", "권한", "역할"], admin_rows, widths=[2.5, 7.0, 2.4, 6.0], font_size=7.6)

    heading(document, "6.2 운영/스태프/예약 API", 2)
    ops_rows = [
        ["OPS_MASTER", "GET", "/api/ops/master/bootstrap", "통합 운영 초기 데이터"],
        ["OPS_MASTER", "POST/PUT/DELETE", "/api/ops/master/notices, /events, /booths", "운영 콘솔 CRUD"],
        ["OPS_MASTER", "POST", "/api/ops/master/ai/briefing", "운영 AI 브리핑"],
        ["OPS_MASTER", "POST", "/api/ops/master/ai/notice-draft", "공지 초안 생성"],
        ["OPS_BOOTH", "GET", "/api/ops/booth/{id}/bootstrap", "부스 운영 초기 데이터"],
        ["OPS_BOOTH", "PUT", "/api/ops/booth/{id}/live-status", "부스 상태 저장"],
        ["OPS_BOOTH", "POST", "/api/ops/booth/{id}/menu-image", "메뉴 이미지 업로드"],
        ["OPS_BOOTH", "PUT", "/api/ops/booth/{id}/reservations/config", "예약 테이블/시간 설정"],
        ["OPS_BOOTH", "POST", "/api/ops/booth/{id}/reservations/{reservationId}/check-in", "예약 체크인"],
        ["OPS_BOOTH", "POST", "/api/ops/booth/{id}/reservations/{reservationId}/complete", "테이블 비우기"],
        ["OPS_BOOTH", "POST", "/api/ops/booth/{id}/reservations/check-in/by-token", "QR 토큰 체크인"],
        ["STAFF", "POST", "/api/staff/auth/login", "스태프 로그인"],
        ["STAFF", "GET", "/api/staff/bootstrap", "스태프 대시보드"],
        ["STAFF", "PUT", "/api/staff/me/status", "스태프 상태/위치 갱신"],
        ["STAFF", "POST", "/api/staff/ai/*", "스태프 AI 보조"],
        ["예약 인증", "POST", "/api/reservations/auth/send-code", "휴대폰 인증번호 발송"],
        ["예약 인증", "POST", "/api/reservations/auth/verify-code", "예약 세션 토큰 발급"],
    ]
    add_table(document, ["권한/영역", "메서드", "경로", "역할"], ops_rows, widths=[2.4, 2.1, 7.1, 6.0], font_size=7.5)

    heading(document, "6.3 SSE API", 2)
    add_table(
        document,
        ["경로", "이벤트 이름", "대표 발행 시점"],
        [
            ["/api/stream/congestion", "congestion", "GPS 로그 저장 후 혼잡도 갱신"],
            ["/api/stream/events", "events", "공연 생성/수정/삭제 또는 운영 공지"],
            ["/api/stream/notices", "notices", "공지 생성/수정/삭제, 자동 공지"],
            ["/api/stream/booths", "booths", "부스 상태/이미지/순서/삭제 변경"],
            ["/api/stream/staff", "staff", "스태프 상태/위치/업무 갱신"],
            ["/api/stream/lost-items", "lost-items", "분실물 등록/수정/상태 변경"],
            ["/api/stream/reservations", "reservations", "예약 생성/만료/체크인/완료/설정 변경"],
        ],
        widths=[5.0, 3.2, 8.0],
    )


def section_database(document):
    heading(document, "7. DB/Entity 설계")
    paragraph(
        document,
        "DB는 Spring Data JPA Entity 중심으로 관리됩니다. `spring.jpa.hibernate.ddl-auto`는 로컬에서 update, prod에서 validate 기본값을 사용하므로 운영 배포 전 schema 정합성 확인이 중요합니다.",
    )
    entity_rows = [
        ["admin_users", "AdminUser", "관리자 계정", "username unique, password hash, role"],
        ["audit_logs", "AuditLog", "관리자 작업 감사 로그", "adminUsername, action, targetType, targetId, details"],
        ["booths", "Booth", "부스 기본/운영/예약 가능 정보", "좌표, 설명, 이미지, 대기시간, 재고, 운영시간, reservationEnabled"],
        ["events", "FestivalEvent", "공연 일정과 상태", "startTime, endTime, status, statusOverride, liveMessage, delayMinutes"],
        ["gps_logs", "GpsLog", "혼잡도 산정용 위치 로그", "latitude, longitude, createdAt"],
        ["notices", "Notice", "운영 공지", "title, content, category, active, timestamps"],
        ["lost_items", "LostItem", "분실물", "title, description, category, location, contacts, claim fields, status"],
        ["staff_members", "StaffMember", "스태프 계정/현장 상태", "staffNo, pinHash, team, status, task, assignedBoothId, location"],
        ["staff_sessions", "StaffSession", "스태프 로그인 세션", "token unique, expiresAt, lastSeenAt"],
        ["booth_reservation_tables", "BoothReservationTable", "부스별 테이블/좌석", "booth_id, tableName, totalSeats, availableSeats"],
        ["booth_reservations", "BoothReservation", "예약 이력과 상태", "booth_id, table_id, userKey, status, reservedAt, expiresAt"],
        ["reservation_checkin_tokens", "ReservationCheckInToken", "QR 체크인 토큰", "reservation_id, token, expiresAt, usedAt"],
        ["reservation_user_accounts", "ReservationUserAccount", "예약 사용자 전화번호 계정", "phoneNumber unique, lastVerifiedAt"],
        ["reservation_auth_sessions", "ReservationAuthSession", "예약 인증 세션", "user_account_id, token, expiresAt"],
        ["reservation_verification_codes", "ReservationVerificationCode", "인증번호", "phoneNumber, code, expiresAt, failedAttempts"],
        ["reservation_user_states", "ReservationUserState", "노쇼 패널티", "userKey unique, noShowCount, blockedUntil"],
        ["ai_match_profiles", "AiMatchProfile", "AI Match 프로필", "nickname, gender, intro, pinHash, phoneNumber, image URLs, status"],
        ["ai_match_requests", "AiMatchRequest", "AI Match 신청/약속/관리자 상태", "profile_id, requester_profile_id, status, connectionStatus, adminNote"],
        ["ai_match_favorites", "AiMatchFavorite", "AI Match 즐겨찾기", "requester_profile_id + favorite_profile_id unique"],
        ["ai_match_phone_usages", "AiMatchPhoneUsage", "AI Match 전화번호 사용량", "phone_number unique, image conversion count, blocked"],
    ]
    add_table(document, ["테이블", "Entity", "역할", "주요 필드"], entity_rows, widths=[3.8, 3.9, 4.7, 7.0], font_size=7.3)

    heading(document, "7.1 주요 관계", 2)
    bullets(
        document,
        [
            "Booth 1:N BoothReservationTable, Booth 1:N BoothReservation 구조입니다.",
            "BoothReservation은 Booth와 BoothReservationTable을 모두 참조하고, RESERVED/CHECKED_IN 상태는 좌석 점유 상태로 간주합니다.",
            "ReservationCheckInToken은 BoothReservation에 연결되고, 60초 만료/1회 사용 규칙을 가집니다.",
            "StaffMember는 assignedBoothId로 부스와 느슨하게 연결됩니다.",
            "AiMatchRequest는 profile과 requesterProfile 두 개의 AiMatchProfile 참조를 가지며, requesterProfile은 레거시/삭제 케이스를 고려해 optional 처리됩니다.",
            "AiMatchFavorite은 requester_profile_id와 favorite_profile_id 조합 unique constraint로 중복 즐겨찾기를 막습니다.",
        ],
    )


def section_security(document):
    heading(document, "8. 인증과 보안")
    add_table(
        document,
        ["주제", "구현", "주의점"],
        [
            ["관리자 JWT", "AuthService 로그인 성공 시 JwtService가 role claim을 포함한 JWT 발급", "APP_JWT_SECRET은 32 UTF-8 bytes 이상이어야 하며 운영에서 반드시 별도 설정"],
            ["관리자 권한", "SecurityConfig에서 /api/admin/**는 ROLE_ADMIN 필요", "프론트는 localStorage 토큰을 Authorization Bearer로 전송"],
            ["운영 키", "OpsKeyAuthenticationFilter가 master/booth key 검증", "master key는 전체 운영 권한이므로 노출 금지"],
            ["스태프 토큰", "StaffService가 세션 또는 stateless demo token을 처리", "토큰 만료와 logout 처리 확인 필요"],
            ["예약 인증", "ReservationAuthService가 전화번호 인증 후 예약 사용자 key/token 발급", "예약 API는 토큰으로 사용자 중복 예약과 노쇼 상태를 판단"],
            ["PIN", "AI Match 프로필 소유 확인용. BCrypt로 hash 저장", "PIN 원문 저장 금지, 삭제 프로필은 pinHash를 null로 비활성화"],
            ["파일 업로드", "UploadStorageService가 10MB 이하 JPG/PNG/WEBP/GIF만 허용", "로컬 저장소 경로 traversal 방지, S3 사용 시 public base URL 관리"],
            ["개인정보", "분실물 연락처와 AI Match 전화번호는 일반 방문자 화면에 노출하지 않음", "관리자/운영 목적 외 사용 금지"],
            ["CORS", "CorsConfig가 app.cors.allowed-origins를 allowedOriginPatterns로 적용", "운영 도메인을 정확히 설정해야 함"],
        ],
        widths=[2.8, 7.0, 7.0],
        font_size=8.0,
    )
    add_note(
        document,
        "운영 보안 메모",
        "문서에는 로컬 개발 기본값을 그대로 비밀번호로 재사용하지 않도록 명시해야 합니다. 실제 배포에서는 관리자 계정, JWT secret, OPS key, SMS/OpenAI/S3 키를 모두 환경변수로 분리하고 저장소에 커밋하지 않는 것이 원칙입니다.",
    )


def section_ai_external(document):
    heading(document, "9. AI와 외부 연동")
    add_table(
        document,
        ["기능", "서비스", "외부 API", "fallback"],
        [
            ["방문자 AI 가이드", "PublicAiGuideService", "OpenAI Responses API", "혼잡도/공연/부스 데이터를 기반으로 자체 추천 반환"],
            ["혼잡도 예측", "AiCongestionService", "외부 호출 없음", "룰 기반 점수 계산"],
            ["운영 AI 브리핑", "OpsAiService", "OpenAI Responses API", "현재 운영 snapshot 기반 자체 브리핑"],
            ["AI Match 이미지", "AiImageGenerationService", "OpenAI Images Edits API + Responses API 사진 검증", "키 누락 시 503 또는 변환 실패 메시지"],
            ["번역", "TranslateService", "Google Translate public endpoint", "preset 또는 [EN]/[KO] fallback"],
            ["예약 인증 SMS", "SmsSender 구현체", "Solapi, Aligo, Twilio, Noop", "provider none/noop 사용 가능"],
        ],
        widths=[3.2, 4.2, 6.2, 5.0],
        font_size=8.0,
    )
    heading(document, "9.1 AI 혼잡도 위험 점수", 2)
    paragraph(
        document,
        "AiCongestionService는 단순 GPS 인원만 보지 않고, 대기시간, 활성 예약, 체크인 예약, 테이블 수, 예약 가능 좌석, 재고, 30분 내 공연 시작 여부를 종합해 riskScore를 만듭니다.",
    )
    add_table(
        document,
        ["입력 신호", "점수 영향", "운영 의미"],
        [
            ["주변 감지 인원", "최대 30점", "현재 현장 밀도"],
            ["예상 대기 시간", "최대 20점", "운영자가 입력한 체감 대기"],
            ["활성 예약/체크인", "최대 20점 + 테이블 비율 가중", "테이블 점유와 회전 지연"],
            ["예약 가능 좌석", "0석/3석 이하일 때 가중", "예약 불가 또는 마감 임박"],
            ["재고", "10개 이하 또는 소진 시 가중", "품절/운영 지연 가능성"],
            ["공연 임박", "추가 가중", "무대 주변 이동량 증가 예측"],
        ],
        widths=[4.0, 3.5, 8.0],
    )


def section_config_deploy(document):
    heading(document, "10. 실행, 설정, 배포")
    heading(document, "10.1 로컬 실행", 2)
    code_block(
        document,
        """
# 백엔드
cd backend
./gradlew bootRun

# 프론트엔드
cd frontend
npm install
npm run dev
        """,
    )
    add_table(
        document,
        ["항목", "기본값/설명"],
        [
            ["백엔드 포트", "server.port=${PORT:8080}"],
            ["프론트 dev 서버", "Vite 기본 포트, 일반적으로 http://localhost:5173"],
            ["API Base", "frontend/src/api.js에서 VITE_API_BASE_URL 또는 http://localhost:8080/api"],
            ["DB", "local profile은 MySQL festival_db를 바라봄"],
            ["JPA", "local update, prod validate"],
            ["업로드", "APP_UPLOAD_DIR, APP_STORAGE_TYPE=local 또는 s3"],
        ],
        widths=[3.4, 10.5],
    )
    heading(document, "10.2 핵심 환경변수", 2)
    env_rows = [
        ["SPRING_DATASOURCE_URL / USERNAME / PASSWORD", "DB 연결 정보", "prod 필수"],
        ["SPRING_PROFILES_ACTIVE", "local/prod 등 Spring profile", "배포 시 prod 권장"],
        ["APP_JWT_SECRET", "JWT 서명 secret", "32 bytes 이상, prod 필수"],
        ["APP_CORS_ALLOWED_ORIGINS", "프론트 허용 origin", "Vercel 도메인 포함"],
        ["APP_INIT_ADMIN_USERNAME / PASSWORD", "초기 관리자 계정", "운영에서는 별도 강한 값"],
        ["APP_OPS_MASTER_KEY / APP_OPS_BOOTH_KEYS", "운영 콘솔 키", "노출 금지"],
        ["APP_SMS_PROVIDER", "none/twilio/aligo/solapi", "SMS 인증 provider"],
        ["OPENAI_API_KEY", "OpenAI 기능용 API key", "AI 기능 사용 시 필요"],
        ["OPENAI_MODEL / OPENAI_IMAGE_MODEL", "텍스트/이미지 모델", "기본값 사용 가능"],
        ["APP_STORAGE_TYPE", "local 또는 s3", "S3 사용 시 bucket/region 필요"],
        ["AWS_S3_BUCKET / AWS_REGION / AWS_S3_PUBLIC_BASE_URL", "S3 업로드 설정", "s3 저장 시 필요"],
        ["VITE_API_BASE_URL", "프론트가 호출할 백엔드 API base", "Vercel에서 설정"],
    ]
    add_table(document, ["환경변수", "역할", "운영 메모"], env_rows, widths=[6.0, 5.5, 5.0], font_size=7.7)
    heading(document, "10.3 배포", 2)
    bullets(
        document,
        [
            "백엔드는 Railway에서 backend 폴더를 서비스 루트로 설정하고, 환경변수와 DB를 연결합니다.",
            "프론트엔드는 Vercel에서 frontend 폴더를 Root Directory로 설정하고, Build Command는 `npm run build`, Output Directory는 `dist`를 사용합니다.",
            "React Router fallback은 `frontend/vercel.json`이 처리합니다.",
            "백엔드 prod profile은 JPA validate가 기본이므로 schema가 맞지 않으면 기동 실패가 정상입니다.",
        ],
    )


def section_maintenance(document):
    heading(document, "11. 검증과 유지보수")
    add_table(
        document,
        ["검증 대상", "명령", "확인할 내용"],
        [
            ["백엔드 전체 빌드", "cd backend && ./gradlew clean build", "컴파일, 단위 테스트, Spring 설정"],
            ["프론트 빌드", "cd frontend && npm run build", "Vite build, import 오류, 번들 생성"],
            ["API 수동 점검", "브라우저 Network 또는 curl/Postman", "권한 헤더, payload, 응답 DTO"],
            ["실시간 SSE", "EventSource 구독 또는 프론트 화면", "연결 유지, 이벤트명, payload"],
            ["문서 렌더 QA", "render_docx.py", "표 잘림, 한글 렌더링, 이미지/페이지 여백"],
        ],
        widths=[3.2, 5.8, 7.2],
    )
    heading(document, "11.1 현재 테스트 구성과 추가 권장", 2)
    paragraph(
        document,
        "현재 repository에는 예약/부스 일부 테스트가 존재하지만, 전체 기능 대비 테스트가 충분하다고 보기는 어렵습니다. 특히 운영 콘솔, 스태프, 분실물, AI Match, SSE, 외부 API fallback은 회귀 위험이 큽니다.",
    )
    add_table(
        document,
        ["영역", "권장 테스트"],
        [
            ["예약", "동일 사용자 중복 예약 차단, 좌석 잠금, 만료/노쇼/차단, QR token 1회 사용"],
            ["관리자", "JWT 없는 접근 거부, CRUD 후 감사 로그/SSE 발행 확인"],
            ["운영 키", "master/booth 권한 분리, 잘못된 booth key 접근 차단"],
            ["스태프", "로그인/만료/logout, 상태/위치 갱신, 관리자 수정"],
            ["분실물", "공개 조회 연락처 마스킹, claim 상태, returned 이후 claim 차단"],
            ["AI", "OpenAI key 누락 시 fallback 또는 명확한 에러, JSON parse 실패 fallback"],
            ["AI Match", "전화번호 변환 제한, 삭제 번호 재가입 차단, PIN 인증, 상태 전이"],
            ["SSE", "각 서비스 상태 변경 후 올바른 채널로 이벤트 발행"],
        ],
        widths=[3.2, 12.5],
    )
    heading(document, "11.2 변경 시 체크리스트", 2)
    bullets(
        document,
        [
            "새 API를 추가하면 Controller, Service, DTO, api.js, 권한 설정을 함께 확인합니다.",
            "Entity 필드를 추가하면 prod validate 배포 전 migration 또는 schema 변경 절차가 필요합니다.",
            "운영자가 보는 실시간 값이면 상태 변경 후 StreamService 발행 여부를 확인합니다.",
            "사용자 개인정보가 DTO에 포함되는 경우 공개/관리자/스태프 응답을 분리합니다.",
            "예약/좌석 관련 변경은 동시성 lock, 좌석 복구, 만료 처리, 노쇼 패널티를 같이 검증합니다.",
            "AI 기능은 외부 API 실패 시 화면이 멈추지 않는 fallback을 유지합니다.",
            "PWA service worker 변경 후 cache name을 올려 오래된 shell이 남지 않게 합니다.",
        ],
    )


def section_appendix(document):
    heading(document, "12. 부록: 코드 읽는 순서")
    add_table(
        document,
        ["목적", "읽을 파일 순서"],
        [
            ["전체 라우트 파악", "frontend/src/main.jsx -> frontend/src/App.jsx -> 각 pages/*.jsx"],
            ["API 호출 파악", "frontend/src/api.js -> backend/controller/*Controller.java -> backend/service/*.java"],
            ["DB 구조 파악", "backend/entity/*.java -> backend/repository/*.java"],
            ["예약 로직 파악", "ReservationAuthController -> ReservationAuthService -> ReservationService -> BoothReservation* Entity"],
            ["운영 콘솔 파악", "OpsMasterPage/OpsBoothPage -> OpsController -> BoothService/NoticeService/EventService/ReservationService/OpsAiService"],
            ["스태프 로직 파악", "StaffPage -> StaffController -> StaffService -> StaffMember/StaffSession"],
            ["분실물 로직 파악", "LostFoundPage -> LostItemController -> LostItemService -> LostItem"],
            ["AI Match 파악", "AiMatchPage/AiMatchAdminPage -> AiMatchController/AdminAiMatchController -> AiMatchService -> AiImageGenerationService"],
            ["실시간 흐름 파악", "api.js EventSource 함수 -> StreamController -> StreamService -> 각 service publish 호출"],
            ["배포 설정 파악", "README.md -> backend/application*.properties -> frontend/vercel.json -> frontend/public/service-worker.js"],
        ],
        widths=[4.0, 12.0],
        font_size=8.0,
    )
    add_note(
        document,
        "최종 요약",
        "FestFlow의 핵심은 '현장 상태 데이터'를 여러 화면에서 공유하는 구조입니다. 방문자 화면은 정보를 소비하고, 운영/관리/스태프 화면은 같은 데이터를 갱신하며, 백엔드 서비스 계층이 권한과 상태 전이를 통제합니다.",
    )


def build():
    arch = make_architecture_diagram()
    backend_flow = make_backend_flow_diagram()
    realtime = make_realtime_diagram()

    doc = Document()
    style_document(doc)
    add_title(doc)
    section_project_overview(doc)
    add_page_break(doc)
    section_architecture(doc, arch, backend_flow, realtime)
    add_page_break(doc)
    section_frontend(doc)
    add_page_break(doc)
    section_backend(doc)
    section_feature_modules(doc)
    add_page_break(doc)
    section_api(doc)
    add_page_break(doc)
    section_database(doc)
    section_security(doc)
    add_page_break(doc)
    section_ai_external(doc)
    section_config_deploy(doc)
    add_page_break(doc)
    section_maintenance(doc)
    section_appendix(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
