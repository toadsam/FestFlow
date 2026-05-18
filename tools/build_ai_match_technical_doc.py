from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "technical-assets"
OUT_DOCX = ROOT / "ai-matach_기술문서.docx"
FONT = Path("C:/Windows/Fonts/malgun.ttf")
BOLD_FONT = Path("C:/Windows/Fonts/malgunbd.ttf")


def font(size, bold=False):
    return ImageFont.truetype(str(BOLD_FONT if bold else FONT), size)


def rounded_box(draw, box, fill, outline="#93c5fd", width=2):
    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fill="#111827", size=26, bold=False, max_width=20):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=max_width) or [""])
    fnt = font(size, bold)
    line_h = size + 8
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=fill, font=fnt)
        y += line_h


def arrow(draw, start, end, color="#2563eb"):
    draw.line([start, end], fill=color, width=5)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(points, fill=color)


def make_architecture_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "ai-match-architecture.png"
    image = Image.new("RGB", (1500, 850), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.text((55, 35), "AI Match 전체 구조", fill="#111827", font=font(38, True))
    draw.text(
        (55, 88),
        "사용자 화면과 관리자 화면은 같은 백엔드 서비스를 사용하고, 사진/DB/AI 변환이 뒤에서 함께 동작합니다.",
        fill="#475569",
        font=font(22),
    )

    boxes = {
        "browser": (60, 170, 360, 350),
        "react": (455, 170, 755, 350),
        "api": (850, 170, 1150, 350),
        "service": (455, 470, 755, 650),
        "db": (850, 470, 1150, 650),
        "storage": (1195, 470, 1450, 650),
        "openai": (1195, 170, 1450, 350),
    }
    labels = {
        "browser": "브라우저\n/mobile or desktop",
        "react": "React 화면\nAiMatchPage\nAiMatchAdminPage",
        "api": "Spring Boot API\nController",
        "service": "AiMatchService\n검증/상태 변경",
        "db": "MySQL\nprofiles\nrequests",
        "storage": "업로드 저장소\nlocal 또는 S3",
        "openai": "OpenAI 이미지 API\n웹툰 프로필 변환",
    }
    fills = {
        "browser": "#ecfeff",
        "react": "#eef2ff",
        "api": "#eff6ff",
        "service": "#f0fdf4",
        "db": "#fefce8",
        "storage": "#fff7ed",
        "openai": "#fdf2f8",
    }
    for key, box in boxes.items():
        rounded_box(draw, box, fills[key])
        center_text(draw, box, labels[key], max_width=16)

    arrow(draw, (360, 260), (455, 260))
    arrow(draw, (755, 260), (850, 260))
    arrow(draw, (1000, 350), (1000, 470))
    arrow(draw, (850, 560), (755, 560))
    arrow(draw, (1150, 560), (1195, 560))
    arrow(draw, (1195, 260), (1150, 260))
    arrow(draw, (1150, 260), (1195, 260), "#db2777")
    arrow(draw, (605, 350), (605, 470))

    draw.text((70, 720), "핵심 생각법: 화면은 입력을 받고, API는 요청을 받고, Service는 규칙을 검사하고, DB/저장소는 결과를 보관합니다.", fill="#111827", font=font(24, True))
    image.save(path)
    return path


def make_status_diagram():
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "ai-match-status-flow.png"
    image = Image.new("RGB", (1500, 760), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((55, 35), "데이트 신청 상태 흐름", fill="#111827", font=font(38, True))
    draw.text((55, 88), "신청 한 건은 아래 상태 중 하나로 이동합니다. 관리자는 성사된 상태에서 연결 상태를 별도로 관리합니다.", fill="#475569", font=font(22))

    pending = (90, 250, 330, 390)
    accepted = (520, 160, 760, 300)
    rejected = (520, 360, 760, 500)
    canceled = (520, 560, 760, 700)
    proposed = (930, 160, 1170, 300)
    confirmed = (930, 360, 1170, 500)
    admin = (1210, 250, 1450, 500)

    for box, label, fill in [
        (pending, "PENDING\n신청 대기", "#eff6ff"),
        (accepted, "ACCEPTED\n상대가 수락", "#ecfdf5"),
        (rejected, "REJECTED\n상대가 거절", "#fef2f2"),
        (canceled, "CANCELED\n신청자가 취소", "#f8fafc"),
        (proposed, "PROPOSED\n약속 제안", "#f5f3ff"),
        (confirmed, "CONFIRMED\n약속 확정", "#f0fdfa"),
        (admin, "관리자 연결 상태\nWAITING\nCOMPLETED\nFAILED", "#fff7ed"),
    ]:
        rounded_box(draw, box, fill)
        center_text(draw, box, label, max_width=13)

    arrow(draw, (330, 320), (520, 230))
    arrow(draw, (330, 320), (520, 430), "#ef4444")
    arrow(draw, (330, 320), (520, 630), "#64748b")
    arrow(draw, (760, 230), (930, 230), "#7c3aed")
    arrow(draw, (930, 430), (760, 230), "#7c3aed")
    arrow(draw, (1170, 230), (1210, 320), "#f97316")
    arrow(draw, (1170, 430), (1210, 430), "#f97316")
    image.save(path)
    return path


def set_doc_style(document):
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 25, "111827"),
        ("Heading 1", 17, "2563EB"),
        ("Heading 2", 13, "111827"),
        ("Heading 3", 11, "334155"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Match 기술 문서")
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(37, 99, 235)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("초보자도 이해할 수 있게 풀어쓴 구조, API, DB, 보안, 운영 설명서")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(71, 85, 105)
    document.add_paragraph()


def heading(document, text, level=1):
    p = document.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def paragraph(document, text, bold_prefix=None):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(10.5)
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10.5)


def bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(10)


def numbers(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(10)


def note(document, title_text, body):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title_text + "\n")
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(29, 78, 216)
    r = p.add_run(body)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    document.add_paragraph()


def table(document, headers, rows, widths=None):
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_shading(tbl.rows[0].cells[idx], "DBEAFE")
        set_cell_text(tbl.rows[0].cells[idx], header, True)
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, False, 9.2)
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    document.add_paragraph()


def image(document, path, caption, width_cm=15.8):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100, 116, 139)


def code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line in text.strip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(9)
    document.add_paragraph()


def build():
    arch = make_architecture_diagram()
    status = make_status_diagram()
    doc = Document()
    set_doc_style(doc)
    title(doc)
    note(
        doc,
        "문서 목적",
        "이 문서는 AI Match 기능을 처음 보는 사람이 코드와 기능 구조를 함께 이해할 수 있도록 만든 기술 설명서입니다. "
        "전문 용어는 쉬운 말로 풀고, 실제 파일명과 API 경로를 함께 적었습니다.",
    )

    heading(doc, "1. 한 문장으로 이해하기")
    paragraph(
        doc,
        "AI Match는 축제 방문자가 사진과 자기소개로 프로필을 만들고, 다른 사람에게 데이트 신청을 보내며, 관리자가 성사된 매칭을 확인해 연락을 조율하는 기능입니다.",
    )
    bullets(doc, [
        "사용자는 /ai-match 화면에서 프로필을 만들고 신청을 보냅니다.",
        "관리자는 /ai-match/admin 화면에서 프로필과 신청 기록을 확인합니다.",
        "백엔드는 Spring Boot API로 요청을 받고 MySQL에 저장합니다.",
        "사진은 업로드 저장소에 저장되고, OpenAI 이미지 API로 웹툰 스타일 이미지가 생성됩니다.",
    ])

    heading(doc, "2. 초보자 용어 사전")
    table(doc, ["용어", "쉬운 뜻", "이 프로젝트에서의 예"], [
        ["프론트엔드", "사용자가 브라우저에서 직접 보는 화면", "React로 만든 /ai-match 화면"],
        ["백엔드", "화면 뒤에서 데이터 저장, 검증, 인증을 처리하는 서버", "Spring Boot 서버"],
        ["API", "프론트엔드가 백엔드에게 일을 부탁하는 주소", "POST /api/ai-match/profiles"],
        ["DTO", "API로 주고받는 데이터 모양을 정한 객체", "AiMatchProfileResponseDto"],
        ["Entity", "DB 테이블과 거의 1:1로 연결되는 Java 클래스", "AiMatchProfile, AiMatchRequest"],
        ["Repository", "DB에서 찾기/저장하기를 쉽게 해주는 객체", "AiMatchProfileRepository"],
        ["Service", "실제 규칙과 핵심 로직이 들어가는 곳", "AiMatchService"],
        ["JWT", "관리자가 로그인 후 API를 호출할 때 쓰는 인증 토큰", "Authorization: Bearer ..."],
        ["BCrypt", "PIN 같은 비밀번호를 안전하게 해시로 바꾸는 방식", "pinHash 저장"],
        ["Multipart", "파일과 텍스트를 함께 보내는 API 요청 방식", "사진 업로드"],
        ["S3", "운영 환경에서 이미지 파일을 저장할 수 있는 클라우드 저장소", "APP_STORAGE_TYPE=s3"],
    ], [3.0, 6.5, 5.5])

    heading(doc, "3. 프로젝트 파일 지도")
    paragraph(doc, "처음 코드를 볼 때는 모든 파일을 한꺼번에 보려고 하면 어렵습니다. 아래 순서대로 보면 기능 흐름이 잘 보입니다.")
    table(doc, ["위치", "파일", "무엇을 보면 되나"], [
        ["frontend/src/main.jsx", "라우팅 설정", "/ai-match와 /ai-match/admin이 어떤 컴포넌트로 연결되는지 확인"],
        ["frontend/src/pages", "AiMatchPage.jsx", "사용자 화면의 탭, 폼, 버튼 이벤트"],
        ["frontend/src/pages", "AiMatchAdminPage.jsx", "관리자 로그인, 대시보드, 연결 상태 변경"],
        ["frontend/src", "api.js", "프론트가 호출하는 API 주소와 요청 방식"],
        ["backend/controller", "AiMatchController.java", "사용자 API 주소"],
        ["backend/controller/admin", "AdminAiMatchController.java", "관리자 API 주소"],
        ["backend/service", "AiMatchService.java", "입력 검증, PIN 인증, 신청 상태 변경"],
        ["backend/service", "AiImageGenerationService.java", "OpenAI 이미지 변환 호출"],
        ["backend/service", "UploadStorageService.java", "로컬/S3 이미지 저장"],
        ["backend/entity", "AiMatchProfile.java", "프로필 DB 컬럼과 삭제 처리"],
        ["backend/entity", "AiMatchRequest.java", "신청 DB 컬럼과 상태 변경"],
        ["backend/repository", "AiMatchProfileRepository.java", "프로필 조회 메서드"],
        ["backend/repository", "AiMatchRequestRepository.java", "신청 조회 메서드"],
    ], [4.0, 4.2, 6.6])

    heading(doc, "4. 전체 구조")
    image(doc, arch, "그림 1. AI Match 전체 아키텍처", 15.8)
    paragraph(doc, "가장 중요한 흐름은 브라우저 -> React 화면 -> Spring Boot API -> Service -> DB/저장소 순서입니다.")
    table(doc, ["구성 요소", "역할", "대표 파일"], [
        ["React 화면", "사용자가 보는 화면과 버튼, 입력값을 처리합니다.", "frontend/src/pages/AiMatchPage.jsx, AiMatchAdminPage.jsx"],
        ["API 함수", "React에서 백엔드로 요청을 보냅니다.", "frontend/src/api.js"],
        ["Controller", "HTTP 요청 주소를 받고 Service로 넘깁니다.", "AiMatchController.java, AdminAiMatchController.java"],
        ["Service", "검증, 인증, 상태 변경, 저장 규칙을 처리합니다.", "AiMatchService.java"],
        ["Entity", "DB 테이블 구조와 상태 변경 메서드를 정의합니다.", "AiMatchProfile.java, AiMatchRequest.java"],
        ["Repository", "DB 조회와 저장을 담당합니다.", "AiMatchProfileRepository.java, AiMatchRequestRepository.java"],
        ["Image Service", "OpenAI 이미지 변환을 호출합니다.", "AiImageGenerationService.java"],
        ["Storage Service", "사진 파일을 로컬 또는 S3에 저장합니다.", "UploadStorageService.java"],
    ], [3.0, 6.2, 6.0])

    heading(doc, "5. 프론트엔드 화면 구조")
    paragraph(doc, "프론트엔드는 Vite + React로 작성되어 있고, 라우팅은 React Router가 담당합니다.")
    code_block(doc, """
/ai-match        -> AiMatchPage.jsx
/ai-match/admin  -> AiMatchAdminPage.jsx
""")
    heading(doc, "5.1 사용자 화면 AiMatchPage", 2)
    paragraph(doc, "사용자 화면은 하나의 컴포넌트 안에서 activeScreen 값에 따라 화면을 바꿉니다. 쉽게 말하면 '현재 어떤 탭을 보여줄지'를 상태값으로 기억합니다.")
    table(doc, ["화면 상태", "사용자가 보는 화면", "주요 기능"], [
        ["intro", "AI 소개팅 부스 첫 화면", "시작하기, 등록된 사람 보기"],
        ["register", "프로필 등록 화면", "사진 업로드, AI 변환, 닉네임/PIN/전화번호/소개 입력"],
        ["people", "등록된 사람 목록", "필터, 검색, 관심 표시, 데이트 신청 진입"],
        ["requests", "신청함", "받은 신청, 보낸 신청, 프로필 수정/삭제"],
        ["detail", "프로필 상세", "상대 정보 확인 후 데이트 신청"],
    ])
    heading(doc, "5.2 사용자 화면의 중요한 상태값", 2)
    table(doc, ["상태값", "쉬운 설명", "왜 필요한가"], [
        ["activeScreen", "현재 보여줄 화면", "intro/register/people/requests 전환"],
        ["profiles", "목록에 보여줄 프로필들", "사람들 화면 표시"],
        ["accessProfile", "PIN 인증이 끝난 내 프로필", "신청함과 신청 보내기 권한 판단"],
        ["accessRequests", "내가 받은 신청", "수락/거절 화면 표시"],
        ["accessSentRequests", "내가 보낸 신청", "신청 상태와 취소 버튼 표시"],
        ["selectedProfile", "상세로 보고 있는 상대 프로필", "데이트 신청 폼 표시"],
        ["previewUrl/originalImageUrl/generatedImageUrl", "업로드 이미지와 AI 변환 이미지", "등록 전 미리보기와 저장"],
        ["converting", "AI 변환 중인지 여부", "변환 중 버튼 비활성화"],
        ["submitting", "서버 요청 처리 중인지 여부", "중복 클릭 방지"],
        ["errorMessage/successMessage", "사용자에게 보여줄 안내 문구", "실패/성공 피드백"],
    ], [4.0, 5.0, 5.8])

    heading(doc, "5.3 관리자 화면 AiMatchAdminPage", 2)
    paragraph(doc, "관리자 화면은 로그인 후 운영 대시보드를 보여줍니다. 관리자 API는 JWT 토큰이 있어야 접근할 수 있습니다.")
    bullets(doc, [
        "로그인 전: 아이디와 비밀번호 입력 화면",
        "로그인 후: 운영 요약, 등록된 사람들, 성사된 매치, 신청 기록",
        "2초마다 조용히 새로고침하여 최신 상태를 반영",
        "연결 상태와 관리자 메모를 서버에 저장",
    ])

    heading(doc, "6. 백엔드 API 구조")
    paragraph(doc, "사용자 API는 /api/ai-match로 시작하고, 관리자 API는 /api/admin/ai-match로 시작합니다.")
    heading(doc, "6.1 사용자 API", 2)
    table(doc, ["메서드", "주소", "하는 일"], [
        ["POST", "/api/ai-match/image-preview", "사진을 업로드하고 AI 변환 미리보기를 만듭니다."],
        ["POST", "/api/ai-match/profiles", "닉네임, PIN, 전화번호, 소개, 이미지 URL을 저장해 프로필을 만듭니다."],
        ["POST", "/api/ai-match/profiles/access", "닉네임 + PIN으로 본인 프로필과 신청함을 엽니다."],
        ["PUT", "/api/ai-match/profiles/{profileId}", "본인 인증 후 프로필을 수정합니다."],
        ["POST", "/api/ai-match/profiles/{profileId}/delete", "본인 인증 후 프로필을 삭제 상태로 바꿉니다."],
        ["POST", "/api/ai-match/profiles/{profileId}/requests", "상대 프로필에 데이트 신청을 보냅니다."],
        ["POST", "/api/ai-match/requests/{requestId}/accept", "받은 신청을 수락합니다."],
        ["POST", "/api/ai-match/requests/{requestId}/reject", "받은 신청을 거절합니다."],
        ["POST", "/api/ai-match/requests/{requestId}/cancel", "보낸 신청을 취소합니다."],
        ["POST", "/api/ai-match/requests/{requestId}/meetup/propose", "성사된 매치에 약속 시간/장소를 제안합니다."],
        ["POST", "/api/ai-match/requests/{requestId}/meetup/confirm", "상대가 제안한 약속을 확정합니다."],
    ], [2.0, 5.8, 7.2])
    heading(doc, "6.2 관리자 API", 2)
    table(doc, ["메서드", "주소", "하는 일"], [
        ["GET", "/api/admin/ai-match/overview", "관리자 대시보드 전체 데이터를 가져옵니다."],
        ["PUT", "/api/admin/ai-match/requests/{requestId}/connection-status", "연결 대기/완료/실패 상태를 저장합니다."],
        ["PUT", "/api/admin/ai-match/requests/{requestId}/admin-note", "관리자 메모를 저장합니다."],
        ["DELETE", "/api/admin/ai-match/profiles/{profileId}", "관리자 권한으로 프로필을 삭제 처리합니다."],
    ], [2.0, 6.0, 7.0])

    heading(doc, "6.3 API 요청 예시", 2)
    paragraph(doc, "아래 예시는 실제 개발자가 브라우저 개발자도구 Network 탭이나 API 테스트 도구에서 이해해야 하는 요청 모양입니다.")
    heading(doc, "프로필 접근 요청", 3)
    code_block(doc, """
POST /api/ai-match/profiles/access
Content-Type: application/json

{
  "nickname": "꾸에엥",
  "pin": "1234"
}
""")
    heading(doc, "데이트 신청 요청", 3)
    code_block(doc, """
POST /api/ai-match/profiles/5/requests
Content-Type: application/json

{
  "requesterNickname": "꾸에엥",
  "requesterPin": "1234",
  "meetPlace": "중앙무대 앞",
  "message": "같이 공연 보러 갈래요?"
}
""")
    heading(doc, "관리자 연결 상태 변경 요청", 3)
    code_block(doc, """
PUT /api/admin/ai-match/requests/10/connection-status
Authorization: Bearer <관리자 JWT>
Content-Type: application/json

{
  "connectionStatus": "COMPLETED"
}
""")

    heading(doc, "7. 데이터베이스 구조")
    paragraph(doc, "AI Match에서 핵심 테이블은 ai_match_profiles와 ai_match_requests 두 개입니다.")
    heading(doc, "7.1 ai_match_profiles", 2)
    table(doc, ["컬럼", "쉬운 설명", "비고"], [
        ["id", "프로필 고유 번호", "자동 증가"],
        ["nickname", "사용자가 입력한 닉네임", "ACTIVE 상태에서 중복 불가"],
        ["gender", "성별", "남자/여자/비공개 등 화면 값"],
        ["intro", "자기소개, MBTI, 태그가 함께 들어갈 수 있음", "최대 500자"],
        ["pinHash", "PIN을 암호화한 값", "원본 PIN은 저장하지 않음"],
        ["phoneNumber", "관리자 연락 조율용 전화번호", "일반 목록에는 노출하지 않음"],
        ["meetPlace", "기본 만남 장소", "프로필/신청 기본값"],
        ["originalImageUrl", "원본 사진 저장 주소", "/uploads/... 또는 S3 URL"],
        ["generatedImageUrl", "AI 변환 이미지 저장 주소", "목록에 주로 표시"],
        ["consent", "프로필 공개 동의 여부", "true여야 등록 가능"],
        ["status", "프로필 상태", "ACTIVE 또는 DELETED"],
        ["createdAt", "등록 시간", "생성 시 자동 입력"],
    ], [3.0, 7.0, 5.0])
    heading(doc, "7.2 ai_match_requests", 2)
    table(doc, ["컬럼", "쉬운 설명", "비고"], [
        ["id", "신청 고유 번호", "자동 증가"],
        ["profile_id", "신청을 받은 사람의 프로필 ID", "대상자"],
        ["requester_profile_id", "신청을 보낸 사람의 프로필 ID", "신청자"],
        ["requesterNickname", "신청자 닉네임", "표시용으로 함께 저장"],
        ["meetPlace", "신청자가 선택한 만남 장소", "신청 메시지와 함께 저장"],
        ["message", "신청자가 보낸 짧은 메시지", "최대 500자"],
        ["status", "신청 처리 상태", "PENDING/ACCEPTED/REJECTED/CANCELED/PROPOSED/CONFIRMED"],
        ["connectionStatus", "관리자 연결 상태", "WAITING/COMPLETED/FAILED"],
        ["adminNote", "관리자 운영 메모", "최대 1000자"],
        ["meetupPlace", "약속 제안 장소", "PROPOSED 이후 사용"],
        ["meetupAt", "약속 제안 시간", "과거 시간 불가"],
        ["updatedAt", "마지막 변경 시간", "상태 변경 시 갱신"],
    ], [3.0, 7.0, 5.0])

    heading(doc, "7.3 두 테이블 관계", 2)
    paragraph(doc, "신청 테이블은 프로필 테이블을 두 번 참조합니다. 하나는 신청을 받은 사람이고, 다른 하나는 신청을 보낸 사람입니다.")
    table(doc, ["관계", "DB 컬럼", "뜻"], [
        ["신청 대상자", "ai_match_requests.profile_id", "데이트 신청을 받은 프로필"],
        ["신청자", "ai_match_requests.requester_profile_id", "데이트 신청을 보낸 프로필"],
        ["표시용 닉네임", "ai_match_requests.requesterNickname", "신청자 프로필이 나중에 삭제되어도 기록에 표시할 이름"],
    ])

    heading(doc, "8. 상태값 이해하기")
    image(doc, status, "그림 2. 데이트 신청 상태 흐름", 15.8)
    table(doc, ["상태", "뜻", "누가 바꾸는가"], [
        ["PENDING", "신청을 보냈고 아직 상대가 응답하지 않음", "신청 생성 시 자동"],
        ["ACCEPTED", "상대가 신청을 수락함", "신청 받은 사용자"],
        ["REJECTED", "상대가 신청을 거절함", "신청 받은 사용자"],
        ["CANCELED", "신청자가 대기 중인 신청을 취소함", "신청 보낸 사용자 또는 관리자 삭제 로직"],
        ["PROPOSED", "성사된 뒤 누군가 약속 시간/장소를 제안함", "매칭 참여자"],
        ["CONFIRMED", "상대가 약속 제안을 확정함", "제안자가 아닌 상대"],
    ])
    note(doc, "관리자 상태는 별도입니다", "신청 status가 ACCEPTED, PROPOSED, CONFIRMED 중 하나이면 '성사된 매치'로 보고, 관리자는 connectionStatus를 WAITING, COMPLETED, FAILED 중 하나로 관리합니다.")

    heading(doc, "9. 주요 기능이 실제로 동작하는 순서")
    heading(doc, "9.1 프로필 등록 흐름", 2)
    numbers(doc, [
        "사용자가 사진을 선택합니다.",
        "React가 /api/ai-match/image-preview로 multipart/form-data 요청을 보냅니다.",
        "UploadStorageService가 원본 사진을 저장합니다.",
        "AiImageGenerationService가 OpenAI 이미지 API를 호출해 웹툰 스타일 이미지를 만듭니다.",
        "생성된 이미지를 다시 업로드 저장소에 저장합니다.",
        "사용자가 닉네임, PIN, 전화번호, 소개, 동의 여부를 입력하고 등록합니다.",
        "AiMatchService가 입력값을 검증하고 PIN을 BCrypt로 암호화합니다.",
        "ai_match_profiles 테이블에 프로필이 저장됩니다.",
    ])
    heading(doc, "9.2 신청 보내기 흐름", 2)
    numbers(doc, [
        "사용자가 사람 목록에서 상대 프로필을 선택합니다.",
        "신청자 본인은 닉네임 + PIN으로 인증되어 있어야 합니다.",
        "자기 자신에게 신청하는지 검사합니다.",
        "같은 상대에게 이미 PENDING 신청이 있는지 검사합니다.",
        "문제가 없으면 ai_match_requests 테이블에 PENDING 상태로 저장합니다.",
    ])
    heading(doc, "9.3 신청 수락/거절/취소 흐름", 2)
    bullets(doc, [
        "수락: 신청 받은 사람의 닉네임 + PIN이 맞고 상태가 PENDING이면 ACCEPTED로 변경됩니다.",
        "거절: 신청 받은 사람의 닉네임 + PIN이 맞고 상태가 PENDING이면 REJECTED로 변경됩니다.",
        "취소: 신청 보낸 사람의 닉네임 + PIN이 맞고 상태가 PENDING이면 CANCELED로 변경됩니다.",
        "이미 처리된 신청은 다시 처리하지 못하도록 CONFLICT 오류를 냅니다.",
    ])
    heading(doc, "9.4 관리자 운영 흐름", 2)
    numbers(doc, [
        "관리자가 /api/auth/login으로 로그인합니다.",
        "서버가 JWT 토큰을 발급합니다.",
        "React가 토큰을 저장하고 /api/admin/ai-match/overview를 호출합니다.",
        "관리자는 성사된 매치를 확인하고 양쪽 연락처를 봅니다.",
        "연락이 끝나면 connectionStatus를 COMPLETED로 저장합니다.",
        "운영 기록이 필요하면 adminNote에 메모를 남깁니다.",
    ])

    heading(doc, "10. 검증 규칙")
    paragraph(doc, "백엔드는 프론트에서 이미 검사한 값도 다시 검사합니다. 사용자가 브라우저 화면을 우회해서 직접 API를 호출할 수 있기 때문입니다.")
    table(doc, ["항목", "규칙", "실패 시 의미"], [
        ["동의 여부", "consent가 true여야 등록 가능", "공개 동의 없이 프로필 생성 방지"],
        ["닉네임", "필수, 최대 40자, ACTIVE 상태 중복 불가", "같은 이름 혼동 방지"],
        ["성별", "필수, 최대 20자", "프로필 필터 기준"],
        ["자기소개", "필수, 최대 500자", "목록과 상세에 표시"],
        ["PIN", "필수, 4~6자리 숫자", "너무 약하거나 잘못된 입력 방지"],
        ["전화번호", "필수 등록, 숫자 기준 8~15자리", "관리자 연결 조율용"],
        ["만남 장소", "필수, 최대 120자", "신청/약속 안내에 사용"],
        ["이미지", "등록 시 AI 변환 이미지 필요", "사진 없는 프로필 방지"],
        ["중복 신청", "같은 신청자 -> 같은 대상자 PENDING 1개만 허용", "중복 신청 방지"],
        ["자기 자신 신청", "신청자와 대상자가 같으면 실패", "자기 자신에게 신청 방지"],
        ["과거 약속 시간", "현재보다 과거 시간이면 실패", "이미 지난 약속 방지"],
    ], [3.2, 6.0, 5.6])

    heading(doc, "11. Service 메서드 매핑")
    paragraph(doc, "Controller는 주소를 받는 입구이고, 실제 일은 대부분 AiMatchService 메서드에서 처리합니다.")
    table(doc, ["사용 상황", "Controller", "Service 메서드"], [
        ["AI 이미지 미리보기", "createImagePreview", "createImagePreview"],
        ["프로필 등록", "createProfile", "createProfile"],
        ["신청함 열기", "accessProfile", "accessProfile"],
        ["프로필 수정", "updateProfile", "updateProfile"],
        ["프로필 삭제", "deleteProfile", "deleteProfile"],
        ["신청 보내기", "createRequest", "createRequest"],
        ["신청 수락", "acceptRequest", "acceptRequest"],
        ["신청 거절", "rejectRequest", "rejectRequest"],
        ["신청 취소", "cancelRequest", "cancelRequest"],
        ["약속 제안", "proposeMeetup", "proposeMeetup"],
        ["약속 확정", "confirmMeetup", "confirmMeetup"],
        ["관리자 대시보드", "getOverview", "getAdminOverview"],
        ["관리자 연결 상태 변경", "updateConnectionStatus", "updateConnectionStatus"],
        ["관리자 메모 저장", "updateAdminNote", "updateAdminNote"],
        ["관리자 프로필 삭제", "deleteProfile", "deleteProfileByAdmin"],
    ], [4.0, 5.0, 5.0])

    heading(doc, "12. 인증과 보안")
    heading(doc, "12.1 사용자 PIN", 2)
    paragraph(doc, "사용자 PIN은 프로필 소유자 확인용입니다. 서버는 원본 PIN을 저장하지 않고 BCrypt로 암호화한 pinHash만 저장합니다.")
    bullets(doc, [
        "PIN 형식: 4~6자리 숫자",
        "사용 위치: 신청함 열기, 프로필 수정, 프로필 삭제, 신청 수락/거절/취소",
        "삭제된 프로필은 pinHash가 null이 되어 더 이상 PIN 로그인이 되지 않습니다.",
    ])
    heading(doc, "12.2 관리자 JWT", 2)
    paragraph(doc, "관리자는 아이디/비밀번호로 로그인하고 JWT 토큰을 받습니다. /api/admin/** 요청은 ADMIN 권한이 있어야 통과합니다.")
    code_block(doc, """
Authorization: Bearer <JWT_TOKEN>
""")
    heading(doc, "12.3 사진과 개인정보", 2)
    bullets(doc, [
        "일반 사용자 목록에는 전화번호가 표시되지 않습니다.",
        "관리자 대시보드에는 매칭 연결 조율을 위해 전화번호가 표시됩니다.",
        "사진은 원본과 AI 변환본이 모두 저장됩니다.",
        "업로드 파일은 10MB 이하 JPG, PNG, WEBP, GIF만 허용합니다.",
    ])

    heading(doc, "13. 환경변수와 실행 설정")
    paragraph(doc, "로컬 실행과 배포 환경은 application.properties, application-local.properties, application-prod.properties에서 제어합니다.")
    table(doc, ["설정", "무엇을 의미하나", "예시/기본값"], [
        ["VITE_API_BASE_URL", "프론트가 호출할 백엔드 API 주소", "http://localhost:8080/api"],
        ["SPRING_DATASOURCE_URL", "MySQL 연결 주소", "jdbc:mysql://localhost:3306/festival_db"],
        ["APP_JWT_SECRET", "JWT 서명 비밀키", "32바이트 이상 필요"],
        ["APP_UPLOAD_DIR", "로컬 업로드 저장 폴더", "uploads"],
        ["APP_STORAGE_TYPE", "사진 저장 방식", "local 또는 s3"],
        ["OPENAI_API_KEY", "OpenAI 이미지 변환용 API 키", "없으면 이미지 변환 실패"],
        ["OPENAI_IMAGE_MODEL", "이미지 변환 모델", "gpt-image-1.5"],
        ["APP_CORS_ALLOWED_ORIGINS", "프론트 접근 허용 주소", "localhost 또는 배포 도메인"],
    ], [3.8, 6.5, 4.8])

    heading(doc, "14. 로컬 실행 흐름")
    paragraph(doc, "개발할 때는 보통 백엔드와 프론트엔드를 따로 실행합니다. 백엔드는 8080, 프론트는 5173 포트를 사용합니다.")
    code_block(doc, """
# 1. 백엔드 실행
cd backend
./gradlew bootRun

# 2. 프론트엔드 실행
cd frontend
npm install
npm run dev

# 3. 접속
사용자 화면: http://localhost:5173/ai-match
관리자 화면: http://localhost:5173/ai-match/admin
""")

    heading(doc, "15. 초보자용 코드 읽는 순서")
    numbers(doc, [
        "frontend/src/main.jsx에서 /ai-match와 /ai-match/admin 라우트를 확인합니다.",
        "frontend/src/pages/AiMatchPage.jsx에서 사용자 화면 상태와 버튼 동작을 봅니다.",
        "frontend/src/pages/AiMatchAdminPage.jsx에서 관리자 화면과 상태 변경 버튼을 봅니다.",
        "frontend/src/api.js에서 어떤 API 주소로 요청을 보내는지 확인합니다.",
        "AiMatchController.java와 AdminAiMatchController.java에서 API 주소와 메서드를 확인합니다.",
        "AiMatchService.java에서 실제 검증과 저장 로직을 읽습니다.",
        "AiMatchProfile.java와 AiMatchRequest.java에서 DB 컬럼과 상태 변경 메서드를 확인합니다.",
        "AiImageGenerationService.java와 UploadStorageService.java에서 사진 처리 방식을 확인합니다.",
    ])

    heading(doc, "16. 자주 나는 오류와 원인")
    table(doc, ["증상", "가능한 원인", "확인할 것"], [
        ["AI 이미지 변환 실패", "OPENAI_API_KEY가 없거나 OpenAI 호출 실패", "백엔드 환경변수와 서버 로그"],
        ["서버 연결 실패", "백엔드가 꺼졌거나 API 주소가 다름", "VITE_API_BASE_URL, localhost:8080"],
        ["관리자 로그인 실패", "계정 정보 또는 JWT 설정 문제", "APP_INIT_ADMIN_USERNAME/PASSWORD, APP_JWT_SECRET"],
        ["PIN 인증 실패", "닉네임/PIN 불일치 또는 프로필 삭제 상태", "ACTIVE 상태, pinHash 존재 여부"],
        ["같은 사람에게 신청 불가", "이미 PENDING 신청이 있음", "ai_match_requests 상태"],
        ["업로드 실패", "파일 크기 또는 확장자 제한 초과", "10MB 이하 JPG/PNG/WEBP/GIF"],
        ["배포 후 CORS 오류", "프론트 도메인이 허용되지 않음", "APP_CORS_ALLOWED_ORIGINS"],
    ], [4.0, 6.0, 5.0])

    heading(doc, "17. 유지보수 체크리스트")
    bullets(doc, [
        "새 API를 추가하면 frontend/src/api.js에도 함수가 필요한지 확인합니다.",
        "DB 컬럼을 추가하면 Entity, DTO, Service 변환 메서드를 함께 확인합니다.",
        "사용자에게 보이면 안 되는 개인정보가 DTO에 포함되지 않았는지 확인합니다.",
        "관리자 기능은 /api/admin/** 아래에 두고 JWT 권한을 적용합니다.",
        "상태값을 추가하면 프론트 라벨, 백엔드 검증, 관리자 필터를 함께 수정합니다.",
        "이미지 처리 로직을 바꾸면 로컬 저장과 S3 저장 모두에서 동작하는지 확인합니다.",
    ])

    heading(doc, "18. 핵심 요약")
    note(
        doc,
        "기억할 5가지",
        "1) 사용자는 React 화면에서 프로필과 신청을 만든다.\n"
        "2) 백엔드는 AiMatchService에서 입력값과 권한을 검사한다.\n"
        "3) 프로필은 ai_match_profiles, 신청은 ai_match_requests에 저장된다.\n"
        "4) PIN은 원문 저장이 아니라 BCrypt 해시로 저장된다.\n"
        "5) 관리자는 JWT로 로그인하고 성사된 매치의 연결 상태와 메모를 관리한다.",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
